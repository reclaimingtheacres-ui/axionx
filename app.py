from flask import Flask, render_template, request, redirect, url_for, flash, session, send_from_directory, send_file, jsonify, Response, abort, make_response
from functools import wraps
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
import sqlite3
import csv
import json
import os
import re
import hmac
import io
import uuid
import mimetypes
import traceback
import threading
import secrets
import pytesseract
from PIL import Image, ImageFilter, ImageEnhance
from datetime import date, datetime
import pytz
from azure.storage.blob import BlobServiceClient, ContentSettings

_melbourne = pytz.timezone("Australia/Melbourne")

ARCHIVED_STATUSES = ("Archived - Invoiced", "Cold Stored")

from security import throttle_check, throttle_fail, throttle_success
from datetime import timedelta as _td

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)

is_prod = os.environ.get("ENV", "").lower() in ("prod", "production") or os.environ.get("FLASK_ENV", "").lower() == "production"

app.secret_key = os.environ.get("SESSION_SECRET", "axion-dev-secret")
app.config.update(
    PERMANENT_SESSION_LIFETIME=_td(hours=8),
    SESSION_COOKIE_SECURE=True,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="None",
)

DB_PATH = os.path.abspath(os.getenv("DB_PATH", "axion.db"))

UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "webp", "pdf", "doc", "docx", "xls", "xlsx", "csv", "heic", "heif"}
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


@app.after_request
def add_security_headers(resp):
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["X-Frame-Options"] = "DENY"
    resp.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    resp.headers["Content-Security-Policy"] = (
        "default-src 'self' https: data:; "
        "img-src 'self' https: data:; "
        "script-src 'self' 'unsafe-inline' https:; "
        "style-src 'self' 'unsafe-inline' https:;"
    )
    return resp


@app.errorhandler(404)
def handle_404(e):
    if _is_mobile_request():
        return render_template("m/error.html", error_message="Page not found", path=request.path), 404
    return render_template("error_500.html", error_message="Page not found", path=request.path), 404


@app.errorhandler(405)
def handle_405(e):
    if request.method == "POST" and "/api/" not in request.path:
        return redirect(request.path, code=303)
    if "/api/" in request.path:
        return jsonify({"error": "Method not allowed"}), 405
    if _is_mobile_request():
        return render_template("m/error.html", error_message="Method not allowed", path=request.path), 405
    return render_template("error_500.html", error_message="Method not allowed", path=request.path), 405


@app.route("/health")
def health():
    return "ok", 200


@app.errorhandler(Exception)
def handle_exception(e):
    tb = traceback.format_exc()
    import sys
    print(f"[UNHANDLED EXCEPTION] {request.method} {request.path}\n{tb}", file=sys.stderr, flush=True)
    print(f"[UNHANDLED EXCEPTION] {request.method} {request.path}\n{tb}", flush=True)
    if _is_mobile_request():
        return render_template("m/error.html", error_message="An unexpected error occurred. Please try again or contact support.", path=request.path), 500
    return render_template("error_500.html", error_message="An unexpected error occurred. Please try again or contact support.", path=request.path), 500


# ── Azure Blob Storage ─────────────────────────────────────────────────────────
_AZURE_CONN_STR  = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
_AZURE_CONTAINER = os.getenv("AZURE_UPLOADS_CONTAINER", "axion-uploads")

try:
    _blob_service      = BlobServiceClient.from_connection_string(_AZURE_CONN_STR) if _AZURE_CONN_STR else None
    _uploads_container = _blob_service.get_container_client(_AZURE_CONTAINER) if _blob_service else None
except Exception as _e:
    print(f"[Azure] Client init failed: {_e}")
    _blob_service = _uploads_container = None


def upload_to_blob(file_storage, blob_name: str) -> int:
    """Upload a Werkzeug FileStorage to Azure Blob or local uploads folder. Returns file size in bytes."""
    data = file_storage.read()
    ct   = file_storage.mimetype or mimetypes.guess_type(blob_name)[0] or "application/octet-stream"
    if _uploads_container:
        _uploads_container.upload_blob(
            name=blob_name,
            data=data,
            overwrite=True,
            content_settings=ContentSettings(content_type=ct),
        )
    else:
        dest = os.path.join(UPLOAD_FOLDER, blob_name)
        with open(dest, "wb") as fh:
            fh.write(data)
    return len(data)


def delete_blob_safely(blob_name: str):
    """Delete a blob, silently ignore errors."""
    try:
        if _uploads_container:
            _uploads_container.delete_blob(blob_name)
    except Exception:
        pass


def _save_bytes_to_storage(data: bytes, blob_name: str,
                            content_type: str = "application/pdf") -> int:
    """Save raw bytes to Azure Blob Storage or local uploads folder. Returns size."""
    if _uploads_container:
        _uploads_container.upload_blob(
            name=blob_name, data=data, overwrite=True,
            content_settings=ContentSettings(content_type=content_type),
        )
    else:
        dest = os.path.join(UPLOAD_FOLDER, blob_name)
        with open(dest, "wb") as fh:
            fh.write(data)
    return len(data)
# ──────────────────────────────────────────────────────────────────────────────


@app.route("/favicon.ico")
def favicon():
    return send_from_directory("static", "favicon.ico", mimetype="image/png")

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


_db_initialized = False
_db_init_lock = threading.Lock()


def _raw_db():
    os.makedirs(os.path.dirname(DB_PATH) or ".", exist_ok=True)
    last_err = None
    for attempt in range(5):
        try:
            conn = sqlite3.connect(DB_PATH, timeout=30)
            conn.row_factory = sqlite3.Row
            conn.execute("PRAGMA journal_mode=WAL")
            conn.execute("PRAGMA synchronous=NORMAL")
            conn.execute("PRAGMA busy_timeout=60000")
            return conn
        except sqlite3.OperationalError as e:
            last_err = e
            if attempt < 4:
                import time
                time.sleep(1 + attempt)
            else:
                raise


def _lazy_init():
    global _db_initialized
    if _db_initialized:
        return
    with _db_init_lock:
        if not _db_initialized:
            init_db()
            _migrate_update_builder()
            _db_initialized = True


def db():
    _lazy_init()
    return _raw_db()


def now_ts():
    melb = pytz.timezone("Australia/Melbourne")
    return datetime.now(melb).strftime("%Y-%m-%dT%H:%M:%S")


def add_column_if_missing(cur_or_conn, table, col, coltype):
    import sqlite3 as _sq3
    cur = cur_or_conn.cursor() if isinstance(cur_or_conn, _sq3.Connection) else cur_or_conn
    cur.execute(f"PRAGMA table_info({table})")
    cols = [r["name"] for r in cur.fetchall()]
    if col not in cols:
        cur.execute(f"ALTER TABLE {table} ADD COLUMN {col} {coltype}")


def init_db():
    conn = _raw_db()
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        full_name TEXT NOT NULL,
        email TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL,
        role TEXT NOT NULL,
        active INTEGER NOT NULL DEFAULT 1,
        created_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS clients (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        phone TEXT,
        email TEXT,
        address TEXT,
        notes TEXT,
        created_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS customers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        first_name TEXT NOT NULL,
        last_name TEXT NOT NULL,
        company TEXT,
        email TEXT,
        dob TEXT,
        address TEXT,
        notes TEXT,
        id_image_filename TEXT,
        id_image_path TEXT,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL DEFAULT ''
    )
    """)

    for col, definition in [
        ("updated_at", "TEXT NOT NULL DEFAULT ''"),
    ]:
        try:
            cur.execute(f"ALTER TABLE clients ADD COLUMN {col} {definition}")
        except Exception:
            pass

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        doc_type TEXT NOT NULL,
        title TEXT,
        original_filename TEXT NOT NULL,
        stored_filename TEXT NOT NULL,
        mime_type TEXT,
        file_size INTEGER,
        uploaded_by_user_id INTEGER,
        uploaded_at TEXT NOT NULL,
        notes TEXT,
        FOREIGN KEY(job_id) REFERENCES jobs(id),
        FOREIGN KEY(uploaded_by_user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS contact_phone_numbers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        entity_type TEXT NOT NULL,
        entity_id INTEGER NOT NULL,
        label TEXT NOT NULL,
        phone_number TEXT NOT NULL,
        created_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS contact_emails (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        entity_type TEXT NOT NULL,
        entity_id INTEGER NOT NULL,
        label TEXT,
        email TEXT NOT NULL,
        created_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS jobs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,

        internal_job_number TEXT NOT NULL,
        client_reference TEXT,
        display_ref TEXT NOT NULL,

        client_id INTEGER,
        customer_id INTEGER,
        assigned_user_id INTEGER,

        job_type TEXT NOT NULL,
        visit_type TEXT NOT NULL,
        status TEXT NOT NULL,
        priority TEXT NOT NULL,

        job_address TEXT,
        description TEXT,

        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL,

        FOREIGN KEY(client_id) REFERENCES clients(id),
        FOREIGN KEY(customer_id) REFERENCES customers(id),
        FOREIGN KEY(assigned_user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,

        item_type TEXT NOT NULL,
        description TEXT,

        reg TEXT,
        vin TEXT,
        make TEXT,
        model TEXT,
        year TEXT,

        property_address TEXT,
        lot_details TEXT,

        serial_number TEXT,
        identifier TEXT,

        engine_number TEXT,

        notes TEXT,

        created_at TEXT NOT NULL,
        FOREIGN KEY(job_id) REFERENCES jobs(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_assets (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        asset_type TEXT NOT NULL DEFAULT 'Other',
        description TEXT,
        rego TEXT,
        vin TEXT,
        make TEXT,
        model TEXT,
        year TEXT,
        address TEXT,
        serial TEXT,
        notes TEXT,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL,
        FOREIGN KEY(job_id) REFERENCES jobs(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS interactions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        event_type TEXT NOT NULL,
        narrative TEXT NOT NULL,
        occurred_at TEXT NOT NULL,
        created_at TEXT NOT NULL,
        photo_path TEXT,
        FOREIGN KEY(job_id) REFERENCES jobs(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS cue_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        visit_type TEXT NOT NULL,
        due_date TEXT NOT NULL,
        time_window_start TEXT,
        time_window_end TEXT,
        priority TEXT NOT NULL DEFAULT 'Normal',
        status TEXT NOT NULL DEFAULT 'Pending',
        assigned_user_id INTEGER,
        instructions TEXT,
        created_by_user_id INTEGER,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL,
        completed_at TEXT,
        FOREIGN KEY(job_id) REFERENCES jobs(id),
        FOREIGN KEY(assigned_user_id) REFERENCES users(id),
        FOREIGN KEY(created_by_user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS password_reset_tokens (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        token TEXT NOT NULL UNIQUE,
        expires_at TEXT NOT NULL,
        used INTEGER NOT NULL DEFAULT 0,
        created_at TEXT NOT NULL,
        FOREIGN KEY(user_id) REFERENCES users(id)
    );
    """
    )
    cur.execute("""
    CREATE TABLE IF NOT EXISTS audit_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        actor_user_id INTEGER,
        entity_type TEXT NOT NULL,
        entity_id INTEGER,
        action TEXT NOT NULL,
        message TEXT NOT NULL,
        meta_json TEXT,
        created_at TEXT NOT NULL,
        FOREIGN KEY(actor_user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_field_notes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        created_by_user_id INTEGER,
        note_text TEXT,
        created_at TEXT NOT NULL,
        FOREIGN KEY(job_id) REFERENCES jobs(id),
        FOREIGN KEY(created_by_user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_note_files (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_field_note_id INTEGER NOT NULL,
        filename TEXT NOT NULL,
        filepath TEXT NOT NULL,
        uploaded_at TEXT NOT NULL,
        FOREIGN KEY(job_field_note_id) REFERENCES job_field_notes(id)
    )
    """)
    cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_jnf_note_file ON job_note_files(job_field_note_id, filename)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_jnf_note_id ON job_note_files(job_field_note_id)")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS booking_types (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE,
        active INTEGER NOT NULL DEFAULT 1
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS schedules (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        booking_type_id INTEGER NOT NULL,
        scheduled_for TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'Booked',
        notes TEXT,
        created_by_user_id INTEGER,
        created_at TEXT NOT NULL,
        FOREIGN KEY(job_id) REFERENCES jobs(id),
        FOREIGN KEY(booking_type_id) REFERENCES booking_types(id),
        FOREIGN KEY(created_by_user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS system_settings (
        id INTEGER PRIMARY KEY CHECK (id = 1),
        job_prefix TEXT NOT NULL,
        job_sequence INTEGER NOT NULL,
        auto_prefix_enabled INTEGER NOT NULL DEFAULT 1,
        updated_at TEXT NOT NULL
    )
    """)

    cur.execute("SELECT COUNT(*) c FROM system_settings")
    if cur.fetchone()["c"] == 0:
        current_prefix = datetime.now().strftime("%y%m")
        cur.execute("""
            INSERT INTO system_settings (id, job_prefix, job_sequence, auto_prefix_enabled, updated_at)
            VALUES (1, ?, 0, 1, ?)
        """, (current_prefix, now_ts()))

    for col, coltype in [
        ("lender_name",       "TEXT"),
        ("account_number",    "TEXT"),
        ("regulation_type",   "TEXT"),
        ("arrears_cents",     "INTEGER"),
        ("costs_cents",       "INTEGER"),
        ("mmp_cents",         "INTEGER"),
        ("job_due_date",      "TEXT"),
        ("payment_frequency", "TEXT"),
    ]:
        add_column_if_missing(cur, "jobs", col, coltype)

    add_column_if_missing(cur, "interactions", "photo_path", "TEXT")
    add_column_if_missing(cur, "system_settings", "email_signature", "TEXT")
    add_column_if_missing(cur, "schedules", "assigned_to_user_id", "INTEGER")
    add_column_if_missing(cur, "jobs", "lat", "REAL")
    add_column_if_missing(cur, "jobs", "lng", "REAL")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS agent_locations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL UNIQUE,
        lat REAL NOT NULL,
        lng REAL NOT NULL,
        accuracy REAL,
        updated_at TEXT NOT NULL,
        FOREIGN KEY(user_id) REFERENCES users(id)
    )
    """)
    add_column_if_missing(cur, "jobs", "bill_to_client_id", "INTEGER")
    add_column_if_missing(cur, "jobs", "client_job_number", "TEXT")
    add_column_if_missing(cur, "customers", "role", "TEXT")
    add_column_if_missing(cur, "jobs", "deliver_to", "TEXT")
    add_column_if_missing(cur, "clients", "nickname", "TEXT")

    _default_booking_types = [
        "New Visit", "Re-attend", "Urgent New Visit",
        "Update Required", "Urgent Update Required",
    ]
    cur.execute("SELECT name FROM booking_types")
    existing_bt = {r["name"] for r in cur.fetchall()}
    for bt_name in _default_booking_types:
        if bt_name not in existing_bt:
            cur.execute("INSERT INTO booking_types (name) VALUES (?)", (bt_name,))

    for col, coltype in [
        ("lender_name",    "TEXT"),
        ("account_number", "TEXT"),
        ("regulation_type","TEXT"),
        ("engine_number",  "TEXT"),
        ("deliver_to",     "TEXT"),
        ("colour",         "TEXT"),
    ]:
        add_column_if_missing(cur, "job_items", col, coltype)

    add_column_if_missing(cur, "jobs", "costs2_cents", "INTEGER")
    add_column_if_missing(cur, "jobs", "geoop_source_description", "TEXT")
    add_column_if_missing(cur, "jobs", "geoop_job_id", "TEXT")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_payments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        payment_date TEXT NOT NULL,
        amount_cents INTEGER NOT NULL,
        note TEXT,
        recorded_by_user_id INTEGER,
        created_at TEXT NOT NULL,
        FOREIGN KEY(job_id) REFERENCES jobs(id)
    )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS job_types (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            active INTEGER NOT NULL DEFAULT 1
        )
    """)
    _default_job_types = ["Repo/Collect", "Collect Only", "Field Call", "Process Serve"]
    cur.execute("SELECT name FROM job_types")
    existing_jt = {r["name"] for r in cur.fetchall()}
    for jt_name in _default_job_types:
        if jt_name not in existing_jt:
            cur.execute("INSERT INTO job_types (name) VALUES (?)", (jt_name,))

    cur.execute("""
        CREATE TABLE IF NOT EXISTS job_customers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            job_id INTEGER NOT NULL,
            customer_id INTEGER NOT NULL,
            role TEXT NOT NULL DEFAULT 'Primary',
            sort_order INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL,
            UNIQUE(job_id, customer_id)
        )
    """)

    cur.execute("""
        INSERT OR IGNORE INTO job_customers (job_id, customer_id, role, sort_order, created_at)
        SELECT id, customer_id, 'Primary', 0, created_at
        FROM jobs
        WHERE customer_id IS NOT NULL
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS tow_operators (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_name TEXT NOT NULL,
            phone TEXT,
            address TEXT,
            active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            contact_name TEXT,
            mobile TEXT,
            other_phone TEXT,
            email TEXT,
            suburb TEXT,
            state TEXT,
            postcode TEXT,
            notes TEXT,
            created_by_user_id INTEGER
        )
    """)
    for col, defn in [
        ("contact_name",       "TEXT"),
        ("mobile",             "TEXT"),
        ("other_phone",        "TEXT"),
        ("email",              "TEXT"),
        ("suburb",             "TEXT"),
        ("state",              "TEXT"),
        ("postcode",           "TEXT"),
        ("notes",              "TEXT"),
        ("created_by_user_id", "INTEGER"),
    ]:
        add_column_if_missing(cur, "tow_operators", col, defn)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS auction_yards (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            address TEXT,
            active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            contact_name TEXT,
            mobile TEXT,
            other_phone TEXT,
            email TEXT,
            suburb TEXT,
            state TEXT,
            postcode TEXT,
            notes TEXT,
            created_by_user_id INTEGER
        )
    """)
    for col, defn in [
        ("contact_name",       "TEXT"),
        ("mobile",             "TEXT"),
        ("phone",              "TEXT"),
        ("other_phone",        "TEXT"),
        ("email",              "TEXT"),
        ("suburb",             "TEXT"),
        ("state",              "TEXT"),
        ("postcode",           "TEXT"),
        ("notes",              "TEXT"),
        ("created_by_user_id", "INTEGER"),
    ]:
        add_column_if_missing(cur, "auction_yards", col, defn)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS form_templates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            field_list TEXT NOT NULL,
            created_by INTEGER,
            created_at TEXT NOT NULL,
            active INTEGER NOT NULL DEFAULT 1
        )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS pending_uploads (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        original_filename TEXT NOT NULL,
        storage_key TEXT NOT NULL,
        content_type TEXT,
        uploaded_by_user_id INTEGER,
        uploaded_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS document_extractions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        pending_upload_id INTEGER,
        status TEXT NOT NULL DEFAULT 'success',
        provider_used TEXT NOT NULL DEFAULT 'rule_based',
        extracted_json TEXT NOT NULL,
        extracted_text TEXT,
        created_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS schedule_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        schedule_id INTEGER NOT NULL,
        job_id INTEGER NOT NULL,
        action TEXT NOT NULL,
        old_scheduled_for TEXT,
        new_scheduled_for TEXT,
        old_status TEXT,
        new_status TEXT,
        changed_by_user_id INTEGER,
        created_at TEXT NOT NULL,
        notes TEXT,
        FOREIGN KEY(schedule_id) REFERENCES schedules(id),
        FOREIGN KEY(job_id) REFERENCES jobs(id),
        FOREIGN KEY(changed_by_user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_lifecycle_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        action TEXT NOT NULL,
        from_status TEXT,
        to_status TEXT,
        performed_by_user_id INTEGER,
        performed_at TEXT NOT NULL,
        notes TEXT,
        batch_id TEXT,
        FOREIGN KEY(job_id) REFERENCES jobs(id),
        FOREIGN KEY(performed_by_user_id) REFERENCES users(id)
    )
    """)

    for col, coltype in [
        ("archived_at",            "TEXT"),
        ("archived_by_user_id",    "INTEGER"),
        ("cold_stored_at",         "TEXT"),
        ("cold_stored_by_user_id", "INTEGER"),
        ("cold_storage_ref",       "TEXT"),
        ("lifecycle_status",       "TEXT DEFAULT 'active'"),
    ]:
        add_column_if_missing(cur, "jobs", col, coltype)

    for col, coltype in [
        ("archive_after_days",        "INTEGER DEFAULT 90"),
        ("cold_store_after_years",    "INTEGER DEFAULT 3"),
        ("archive_mode",              "TEXT DEFAULT 'manual'"),
        ("cold_storage_mode",         "TEXT DEFAULT 'manual'"),
        ("allow_restore_to_active",   "INTEGER DEFAULT 1"),
        ("allow_permanent_delete",    "INTEGER DEFAULT 0"),
        ("archive_exclude_client_ids","TEXT"),
    ]:
        add_column_if_missing(cur, "system_settings", col, coltype)

    conn.commit()
    conn.close()


def _migrate_update_builder():
    conn = _raw_db()
    cur = conn.cursor()
    add_column_if_missing(cur, "jobs", "is_regional", "INTEGER")
    add_column_if_missing(cur, "jobs", "confirmed_skip", "INTEGER")
    add_column_if_missing(cur, "system_settings", "openai_api_key",          "TEXT")
    add_column_if_missing(cur, "system_settings", "ai_use_own_key",          "INTEGER")
    add_column_if_missing(cur, "system_settings", "lpr_patrol_mode_enabled", "INTEGER DEFAULT 1")
    add_column_if_missing(cur, "cue_items", "cue_link", "TEXT")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_updates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id INTEGER NOT NULL,
        created_by_user_id INTEGER,
        status TEXT NOT NULL DEFAULT 'draft',
        attend_date TEXT,
        attend_time TEXT,
        is_first_attendance INTEGER NOT NULL DEFAULT 0,
        property_description TEXT,
        security_sighted INTEGER NOT NULL DEFAULT 0,
        security_make_model TEXT,
        security_reg TEXT,
        security_location TEXT,
        calling_card INTEGER NOT NULL DEFAULT 0,
        neighbour_outcome TEXT,
        call_made INTEGER NOT NULL DEFAULT 0,
        call_outcome TEXT,
        voicemail_left INTEGER NOT NULL DEFAULT 0,
        sms_sent INTEGER NOT NULL DEFAULT 0,
        customer_mobile TEXT,
        points_of_contact INTEGER NOT NULL DEFAULT 0,
        eta_next_date TEXT,
        generated_narrative TEXT,
        final_narrative TEXT,
        narrative_edited INTEGER NOT NULL DEFAULT 0,
        structured_inputs_json TEXT,
        ai_model_used TEXT,
        ai_tokens_used INTEGER,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL,
        FOREIGN KEY(job_id) REFERENCES jobs(id),
        FOREIGN KEY(created_by_user_id) REFERENCES users(id)
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS ai_usage_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        job_id INTEGER,
        feature TEXT NOT NULL,
        model TEXT,
        tokens_used INTEGER,
        key_source TEXT NOT NULL DEFAULT 'replit',
        created_at TEXT NOT NULL,
        FOREIGN KEY(user_id) REFERENCES users(id),
        FOREIGN KEY(job_id) REFERENCES jobs(id)
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS user_mobile_settings (
        user_id INTEGER PRIMARY KEY,
        list_sort TEXT NOT NULL DEFAULT 'visit_date',
        list_dir TEXT NOT NULL DEFAULT 'asc',
        distance_unit TEXT NOT NULL DEFAULT 'km',
        gps_foreground INTEGER NOT NULL DEFAULT 1,
        gps_bg INTEGER NOT NULL DEFAULT 0,
        gps_interval_mins INTEGER NOT NULL DEFAULT 5,
        updated_at TEXT,
        job_scope TEXT NOT NULL DEFAULT 'mine',
        show_completed TEXT NOT NULL DEFAULT 'week',
        quick_status TEXT NOT NULL DEFAULT ''
    )
    """)
    add_column_if_missing(cur, "user_mobile_settings", "job_scope",              "TEXT NOT NULL DEFAULT 'mine'")
    add_column_if_missing(cur, "user_mobile_settings", "show_completed",         "TEXT NOT NULL DEFAULT 'week'")
    add_column_if_missing(cur, "user_mobile_settings", "quick_status",           "TEXT NOT NULL DEFAULT ''")
    add_column_if_missing(cur, "user_mobile_settings", "mobile_default_view",    "TEXT NOT NULL DEFAULT 'schedule'")
    add_column_if_missing(cur, "user_mobile_settings", "show_status_on_visits",  "INTEGER NOT NULL DEFAULT 1")
    add_column_if_missing(cur, "job_field_notes",       "note_type",             "TEXT NOT NULL DEFAULT 'text'")
    add_column_if_missing(cur, "job_field_notes",       "audio_filename",        "TEXT")
    add_column_if_missing(cur, "job_field_notes",       "updated_at",            "TEXT")
    add_column_if_missing(cur, "job_field_notes",       "updated_by_user_id",    "INTEGER")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS mobile_auth_tokens (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        token TEXT NOT NULL UNIQUE,
        user_id INTEGER NOT NULL,
        device_name TEXT,
        created_at TEXT NOT NULL,
        expires_at TEXT NOT NULL,
        revoked_at TEXT,
        FOREIGN KEY(user_id) REFERENCES users(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS repo_lock_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id  INTEGER NOT NULL,
        item_id INTEGER NOT NULL,
        client_name               TEXT,
        client_reference          TEXT,
        swpi_ref                  TEXT,
        finance_company           TEXT,
        repo_date                 TEXT,
        start_time                TEXT,
        end_time                  TEXT,
        customer_name             TEXT,
        account_number            TEXT,
        repo_address              TEXT,
        contact_number            TEXT,
        description               TEXT,
        registration              TEXT,
        rego_expiry               TEXT,
        registered                TEXT,
        insured                   TEXT,
        insured_with              TEXT,
        vin                       TEXT,
        engine_number             TEXT,
        speedometer               TEXT,
        person_present            TEXT,
        keys_obtained             TEXT,
        how_many_keys             TEXT,
        vol_surrender             TEXT,
        form_13                   TEXT,
        form_13_signed_by         TEXT,
        repossessed_from          TEXT,
        lien_paid                 TEXT,
        security_drivable         TEXT,
        police_notified           TEXT,
        station_officer           TEXT,
        personal_effects_removed  TEXT,
        removed_by_who            TEXT,
        personal_effects_list     TEXT,
        tyres                     TEXT,
        body                      TEXT,
        duco                      TEXT,
        interior                  TEXT,
        engine_condition          TEXT,
        transmission              TEXT,
        fuel_level                TEXT,
        any_damage                TEXT,
        damage_list               TEXT,
        tow_company_id            INTEGER,
        tow_company_name          TEXT,
        tow_costs                 TEXT,
        deliver_to                TEXT,
        delivery_address          TEXT,
        expected_delivery_date    TEXT,
        customers_intention       TEXT,
        other_info                TEXT,
        agent_name                TEXT,
        agent_user_id             INTEGER,
        created_by_user_id        INTEGER,
        created_at                TEXT NOT NULL,
        updated_by_user_id        INTEGER,
        updated_at                TEXT NOT NULL,
        FOREIGN KEY(job_id)  REFERENCES jobs(id),
        FOREIGN KEY(item_id) REFERENCES job_items(id)
    )
    """)

    add_column_if_missing(cur, "repo_lock_records", "status",           "TEXT NOT NULL DEFAULT 'Draft'")
    add_column_if_missing(cur, "repo_lock_records", "submitted_at",     "TEXT")
    add_column_if_missing(cur, "repo_lock_records", "agent_signature",  "TEXT")
    add_column_if_missing(cur, "repo_lock_records", "customer_signature","TEXT")
    add_column_if_missing(cur, "repo_lock_records", "tow_signature",    "TEXT")
    add_column_if_missing(cur, "repo_lock_records", "agent_signed_at",  "TEXT")
    add_column_if_missing(cur, "repo_lock_records", "customer_signed_at","TEXT")
    add_column_if_missing(cur, "repo_lock_records", "tow_signed_at",    "TEXT")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS repo_lock_queue (
        id                    INTEGER PRIMARY KEY AUTOINCREMENT,
        job_id                INTEGER NOT NULL,
        item_id               INTEGER NOT NULL,
        repo_lock_id          INTEGER NOT NULL,
        status                TEXT NOT NULL DEFAULT 'Pending',
        submission_count      INTEGER NOT NULL DEFAULT 1,
        submitted_at          TEXT NOT NULL,
        submitted_by_user_id  INTEGER,
        reviewed_by_user_id   INTEGER,
        reviewed_at           TEXT,
        notes                 TEXT,
        created_at            TEXT NOT NULL,
        updated_at            TEXT NOT NULL,
        FOREIGN KEY(job_id)          REFERENCES jobs(id),
        FOREIGN KEY(repo_lock_id)    REFERENCES repo_lock_records(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_update_photos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        job_update_id INTEGER NOT NULL,
        job_id INTEGER NOT NULL,
        filename TEXT NOT NULL,
        filepath TEXT NOT NULL,
        tag TEXT NOT NULL DEFAULT 'general',
        uploaded_at TEXT NOT NULL,
        FOREIGN KEY(job_update_id) REFERENCES job_updates(id),
        FOREIGN KEY(job_id) REFERENCES jobs(id)
    )
    """)
    add_column_if_missing(cur, "job_updates", "photos_count", "INTEGER NOT NULL DEFAULT 0")
    add_column_if_missing(cur, "job_updates", "agent_notes", "TEXT DEFAULT ''")

    conn.commit()
    conn.close()


@app.context_processor
def inject_globals():
    pending_drafts = 0
    uid = session.get("user_id")
    if uid:
        try:
            conn = db()
            row = conn.execute(
                "SELECT COUNT(*) c FROM job_updates WHERE created_by_user_id=? AND status='draft'",
                (uid,)
            ).fetchone()
            pending_drafts = row["c"] if row else 0
            conn.close()
        except Exception:
            pending_drafts = 0
    return {
        "GOOGLE_MAPS_API_KEY": os.environ.get("GOOGLE_MAPS_API_KEY", ""),
        "google_maps_api_key": os.getenv("GOOGLE_MAPS_API_KEY", ""),
        "pending_draft_count": pending_drafts,
        "now_iso": now_ts(),
    }


DRAFT_LOCKOUT_THRESHOLD = 5

def _agent_draft_count(uid):
    try:
        conn = db()
        row = conn.execute(
            "SELECT COUNT(*) c FROM job_updates WHERE created_by_user_id=? AND status='draft'",
            (uid,)
        ).fetchone()
        conn.close()
        return row["c"] if row else 0
    except Exception:
        return 0

def _check_agent_lockout():
    role = session.get("role")
    if role not in ("agent",):
        return None
    uid = session.get("user_id")
    if not uid:
        return None
    count = _agent_draft_count(uid)
    if count >= DRAFT_LOCKOUT_THRESHOLD:
        return redirect(url_for("my_drafts"))
    return None


# -------- Helpers --------


def _write_schedule_history(cur, schedule_id, job_id, action,
                            old_scheduled_for=None, new_scheduled_for=None,
                            old_status=None, new_status=None,
                            changed_by_user_id=None, notes=None):
    cur.execute("""
        INSERT INTO schedule_history
            (schedule_id, job_id, action, old_scheduled_for, new_scheduled_for,
             old_status, new_status, changed_by_user_id, created_at, notes)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (schedule_id, job_id, action, old_scheduled_for, new_scheduled_for,
          old_status, new_status, changed_by_user_id, now_ts(), notes))


def normalise_registration(reg_text: str) -> str:
    if not reg_text:
        return ""
    return re.sub(r"[^A-Za-z0-9]", "", reg_text).upper()


_LPR_TMP = "/tmp/axionx_lpr"
os.makedirs(_LPR_TMP, exist_ok=True)


def extract_plate_from_image(image_path: str) -> str:
    """Run OCR on an uploaded plate image and return the normalised plate text."""
    try:
        img = Image.open(image_path).convert("L")
        img = img.resize((img.width * 2, img.height * 2), Image.LANCZOS)
        img = ImageEnhance.Contrast(img).enhance(2.5)
        img = img.filter(ImageFilter.SHARPEN)
        custom_config = r"--oem 3 --psm 7 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
        raw = pytesseract.image_to_string(img, config=custom_config)
        return normalise_registration(raw.strip())
    except Exception:
        return ""


# ── Document auto-fill helpers ────────────────────────────────────────────────

PENDING_UPLOAD_DIR = os.path.join(UPLOAD_FOLDER, "pending")
os.makedirs(PENDING_UPLOAD_DIR, exist_ok=True)


def _save_pending_upload(file_storage):
    original = secure_filename(file_storage.filename or "upload")
    ext = os.path.splitext(original)[1].lower()
    key = f"{uuid.uuid4().hex}{ext}"
    path = os.path.join(PENDING_UPLOAD_DIR, key)
    file_storage.save(path)
    return original, file_storage.mimetype or "application/octet-stream", key, path


def _extract_text_docx(path):
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def _extract_text_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    out = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            out.append(t)
    return "\n".join(out)


def _find_antiword():
    import shutil, subprocess as _sp, os as _os
    found = shutil.which("antiword")
    if found:
        return found
    try:
        r = _sp.run(["bash", "-lc", "which antiword"], capture_output=True, text=True, timeout=5)
        candidate = r.stdout.strip()
        if candidate and _os.path.isfile(candidate):
            return candidate
    except Exception:
        pass
    return None


def _extract_text_doc_olefile(path):
    import olefile, re as _re, struct
    ole = olefile.OleFileIO(path)
    try:
        if not ole.exists("WordDocument"):
            raise RuntimeError("Not a valid Word .doc file")
        wd = ole.openstream("WordDocument").read()
        if len(wd) < 24:
            raise RuntimeError("WordDocument stream too short")
        magic = struct.unpack_from("<H", wd, 0)[0]
        flags = struct.unpack_from("<H", wd, 10)[0] if len(wd) > 12 else 0
        is_complex = bool(flags & 0x0004)
        ccpText = struct.unpack_from("<I", wd, 76)[0] if len(wd) > 80 else 0
        table_name = "1Table" if (flags & 0x0200) else "0Table"
        clx_data = None
        if ole.exists(table_name):
            clx_data = ole.openstream(table_name).read()
        text_pieces = []
        if clx_data and ccpText > 0:
            i = 0
            while i < len(clx_data):
                if clx_data[i] == 0x02:
                    if i + 4 > len(clx_data):
                        break
                    grpprl_len = struct.unpack_from("<H", clx_data, i + 1)[0]
                    i += 3 + grpprl_len
                elif clx_data[i] == 0x01:
                    i += 1
                    if i + 4 > len(clx_data):
                        break
                    n = struct.unpack_from("<I", clx_data, i)[0]
                    i += 4
                    cps = []
                    for j in range(n + 1):
                        if i + 4 > len(clx_data):
                            break
                        cps.append(struct.unpack_from("<I", clx_data, i)[0])
                        i += 4
                    for j in range(min(n, len(cps) - 1)):
                        if i + 8 > len(clx_data):
                            break
                        fc_compressed = struct.unpack_from("<I", clx_data, i)[0]
                        i += 8
                        char_count = cps[j + 1] - cps[j]
                        if char_count <= 0 or char_count > 500000:
                            continue
                        is_ansi = bool(fc_compressed & 0x40000000)
                        fc_offset = fc_compressed & 0x3FFFFFFF
                        if is_ansi:
                            start = fc_offset // 2
                            end = start + char_count
                            if end <= len(wd):
                                chunk = wd[start:end]
                                text_pieces.append(chunk.decode("cp1252", errors="replace"))
                        else:
                            start = fc_offset
                            end = start + char_count * 2
                            if end <= len(wd):
                                chunk = wd[start:end]
                                text_pieces.append(chunk.decode("utf-16-le", errors="replace"))
                    break
                else:
                    i += 1
        text = "".join(text_pieces)
        if len(text.strip()) < 20:
            text_bytes = bytearray()
            for b in wd:
                if 32 <= b < 127 or b in (9, 10, 13):
                    text_bytes.append(b)
                else:
                    text_bytes.append(32)
            text = text_bytes.decode("ascii", errors="replace")
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\x07", "\t")
        text = _re.sub(r'[^\x09\x0a\x20-\x7e\xc0-\xff]', ' ', text)
        text = _re.sub(r'INCLUDEPICTURE\s+"[^"]*"[^\\]*(?:\\[^\n]*)*', '', text)
        text = _re.sub(r'HYPERLINK\s+"[^"]*"[^\\]*(?:\\[^\n]*)*', '', text)
        text = _re.sub(r'\\[ot]\s+"[^"]*"', '', text)
        text = _re.sub(r'\\[*]\s+\w+', '', text)
        text = _re.sub(r'[ \t]{3,}', '  ', text)
        text = _re.sub(r'\n{3,}', '\n\n', text)
        lines = []
        skip_binary = True
        for line in text.split('\n'):
            stripped = line.strip()
            if not stripped or len(stripped) <= 1:
                continue
            if skip_binary and _re.match(r'^[^a-zA-Z0-9]*$', stripped):
                continue
            if skip_binary and all(len(w) <= 2 for w in stripped.split()):
                continue
            has_alpha = sum(1 for c in stripped if c.isalpha())
            if has_alpha < 3 and len(stripped) > 10:
                continue
            skip_binary = False
            lines.append(stripped)
        result = '\n'.join(lines)
        return result.strip() if result.strip() else None
    finally:
        ole.close()


def _extract_text_doc(path):
    import subprocess, logging
    _log = logging.getLogger(__name__)
    antiword = _find_antiword()
    if antiword:
        try:
            result = subprocess.run(
                [antiword, "-w", "0", path],
                capture_output=True, text=True, timeout=15
            )
            if result.returncode == 0:
                text = result.stdout.strip()
                if text and len(text) > 10:
                    return text
            else:
                _log.debug("antiword failed (rc=%d): %s", result.returncode, result.stderr.strip()[:200])
        except FileNotFoundError:
            _log.debug("antiword binary not found at %s", antiword)
        except Exception as e:
            _log.debug("antiword error: %s", e)
    try:
        text = _extract_text_doc_olefile(path)
        if text:
            return text
    except Exception as e:
        _log.debug("olefile .doc extraction failed: %s", e)
    raise RuntimeError(
        "Could not read .doc file. Try saving it as .docx format and uploading again."
    )


def _normalise_phone(s):
    if not s:
        return None
    cleaned = "".join(ch for ch in s if ch.isdigit() or ch == "+")
    return cleaned or None


_AUCTION_YARDS = {
    "slattery", "pickles", "grays", "manheim", "adesa", "fowles", "lloyds",
    "pickles auctions", "slattery auctions", "grays online",
}

_UNREGULATED_TYPES = {
    "chattel mortgage", "chattel", "lease", "novated lease",
    "hire purchase", "hire", "commercial lease", "fleet lease",
    "commercial", "fleet", "finance lease",
}

_REGULATED_TYPES = {"consumer", "regulated", "personal loan", "consumer loan"}


def _parse_instruction_text(text):
    """
    Parse extracted text from an instruction document and return a structured dict.
    Returns confidence per field: 'extracted' = cleanly parsed, 'low' = uncertain.
    Confidence 'matched' is added by the route layer after DB lookups.
    """
    def find_after(patterns):
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                val = (m.group(1) or "").strip()
                return val if val else None
        return None

    def find_block(patterns, max_chars=1200):
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE | re.DOTALL)
            if m:
                val = (m.group(1) or "").strip()
                val = re.sub(r'\n{3,}', '\n\n', val)
                return val[:max_chars].strip() if val else None
        return None

    # ── ACS (Australian Collection Services) worksheet detection ─────────
    is_acs = bool(
        re.search(r"auscollect\.com\.au", text, re.IGNORECASE)
        or re.search(r"WORKSHEET\s*-\s*REPOSSESSION\s*/\s*COLLECTION", text, re.IGNORECASE)
        or (re.search(r"Australian\s*Collection\s*Services", text, re.IGNORECASE)
            and re.search(r"Job\s*No\.?", text, re.IGNORECASE))
    )

    if is_acs:
        acs_job_no = find_after([r"Job\s*No\.?\s*:?\s*([A-Za-z0-9\-\/]+)"])
        acs_client_line = find_after([
            r"(?:^|\n)\s*Client\s*:?\s*([^\n]{3,80}?)(?:\s*Contract\s*#|\s{2,}|\s*$)",
        ])
        acs_contract = find_after([r"Contract\s*#\s*([A-Za-z0-9\-\/]+)"])
        acs_con_type = find_after([r"Con\s*Type\s*:?\s*([A-Za-z][A-Za-z ]{2,30}?)(?:\s*\n|\s{2,}|$)"])
        acs_job_type = find_after([
            r"Job\s*Type\s*:?\s*([^\n]{3,40}?)(?:\s+Con\s*Type|\s{2,}|\s*$)",
        ])

        acs_customer = find_after([
            r"Customer\s*:?\s*([A-Z][A-Z0-9 ,.'&\-]{2,80}?)(?:\s{2,}|\s*DOB|\s*$)",
        ])
        acs_dob = find_after([r"DOB\s*:?\s*([0-9]{1,2}[\/\-][0-9]{1,2}[\/\-][0-9]{2,4})"])

        acs_address = find_after([r"(?:^|\n)\s*Address\s*:?\s*([^\n]{3,120})"])
        acs_suburb = find_after([r"(?:^|\n)\s*Suburb\s*:?\s*([^\n]{3,80})"])
        acs_full_addr = None
        if acs_address:
            acs_full_addr = acs_address.strip().rstrip(",")
            if acs_suburb:
                acs_full_addr += ", " + acs_suburb.strip()

        acs_mobile = find_after([r"(?:^|\n)\s*Mob\s*:?\s+([+0-9\(\)\s\-]{8,20})"])
        acs_ph_home = find_after([r"Ph\s*Hm\s*:?\s*([+0-9\(\)\s\-]{8,20})"])
        acs_ph_bus = find_after([r"Ph\s*Bus\s*:?\s*([+0-9\(\)\s\-]{8,20})"])
        if not acs_mobile and not acs_ph_home and not acs_ph_bus:
            _gtor_phone = find_after([
                r"G/tor\s*:?[\s\S]*?\n\s*(0[0-9\s\-]{8,14})\s*\n",
                r"\n\s*(04[0-9]{8})\s*\n",
            ])
            if _gtor_phone:
                acs_mobile = _gtor_phone
        acs_email_raw = find_after([r"Email\s*:?\s*([^\s@]+@[^\s]+)"])
        acs_email = None
        if acs_email_raw and "auscollect" not in acs_email_raw.lower():
            acs_email = acs_email_raw
        acs_dl = find_after([r"D[\/\\]?L\s*:?\s*([A-Za-z0-9]{5,15})"])

        acs_deliver = find_after([r"Deliver\s*(?:to|vehicle\s*to)\s*:?\s*([A-Za-z0-9 ,.'&\-]{3,60}?)(?:\s*\n|$)"])
        acs_make_model = find_after([r"Make\s*/\s*Model\s*:?\s*([A-Za-z0-9][A-Za-z0-9 \-]{1,40}?)(?:\s{2,}|Year|\s*$)"])
        acs_year = find_after([r"Year\s*:?\s*([12][0-9]{3})"])
        acs_colour = find_after([r"Colou?r\s*:?\s*([A-Za-z][A-Za-z ]{1,20}?)(?:\s{2,}|Regn|\s*$)"])

        acs_regn_raw = find_after([r"Regn\s*:?\s*([^\n]{2,30}?)(?:\s{2,}|\s*\n|\s*$)"])
        acs_rego = None
        acs_reg_note = None
        _REG_STATUS_TOKENS = {"unreg", "unregistered", "expired", "cancelled", "suspended", "stolen", "written off", "defected"}
        if acs_regn_raw:
            regn_parts = re.split(r'\s*-\s*', acs_regn_raw.strip(), maxsplit=1)
            if len(regn_parts) > 1 and regn_parts[1].strip().lower() in _REG_STATUS_TOKENS:
                acs_rego = regn_parts[0].strip()
                acs_reg_note = regn_parts[1].strip()
            else:
                acs_rego = acs_regn_raw.strip()

        acs_vin = find_after([
            r"VIN\s*#?\s*:?\s*([A-HJ-NPR-Z0-9]{11,17})",
            r"VIN\s*/?\s*Chassis\s*:?\s*([A-HJ-NPR-Z0-9]{11,17})",
        ])
        acs_engine = find_after([r"Engine\s*:?\s+([A-Za-z0-9\-\/]{3,20}?)(?:\s{2,}|\s*\n|\s*$)"])
        if acs_engine and acs_engine.strip().lower() in ('reg', 'regn', 'vin', 'color', 'colour'):
            acs_engine = None
        acs_reg_exp = find_after([r"Reg\s*Exp\s*:?\s*([^\n]{3,40}?)(?:\s{2,}|\s*\n|\s*$)"])

        acs_arrears = find_after([r"(?:^|\n)\s*Arrears\s*\$?\s*([0-9][0-9,]*\.?\d{0,2})"])
        acs_install_due = find_after([r"Install\s*due\s*\$?\s*([0-9][0-9,]*\.?\d{0,2})"])
        acs_next_due = find_after([r"Next\s*Due\s*:?\s*([0-9]{1,2}[\/\-][0-9]{1,2}[\/\-][0-9]{2,4})"])
        acs_total_collect = find_after([r"Total\s*collect\s*\$?\s*([0-9][0-9,]*\.?\d{0,2})"])

        acs_make = acs_model = None
        if acs_make_model:
            parts = acs_make_model.strip().split(None, 1)
            acs_make = parts[0].upper()
            acs_model = parts[1].upper() if len(parts) > 1 else None

        acs_lender = acs_client_line.strip().rstrip(".,") if acs_client_line else None
        acs_account_name = find_after([r"Account\s*Name\s*:?\s+([^\n]{3,60}?)(?:\s*\n|$)"])
        if acs_account_name and acs_lender:
            if len(acs_account_name.strip()) > len(acs_lender):
                acs_lender = acs_account_name.strip().rstrip(".,")
        elif acs_account_name and not acs_lender:
            acs_lender = acs_account_name.strip().rstrip(".,")


        acs_reg_type = None
        if acs_con_type:
            ct = acs_con_type.lower().strip()
            if any(k in ct for k in _UNREGULATED_TYPES) or "unregulated" in ct:
                acs_reg_type = "UNREGULATED"
            elif any(k in ct for k in _REGULATED_TYPES) or "regulated" in ct:
                acs_reg_type = "REGULATED"

        acs_instructions_parts = []
        if acs_job_type:
            acs_instructions_parts.append(f"Job Type: {acs_job_type.strip()}")
        if acs_reg_type:
            acs_instructions_parts.append(f"Contract Type: {acs_reg_type}")
        if acs_reg_note:
            acs_instructions_parts.append(f"Registration Status: {acs_reg_note}")
        if acs_reg_exp:
            acs_instructions_parts.append(f"Rego Expiry: {acs_reg_exp.strip()}")

        spec_instr = find_block([
            r"Special\s*Instructions?\s*:?\s*\n+([\s\S]+?)(?=\n\s*Security\b|\n\s*FINANCIALS\b|\Z)",
        ])
        if not spec_instr:
            spec_instr = find_after([r"Special\s*Instructions?\s*:?\s*([^\n]{10,})"])

        if spec_instr:
            lines = spec_instr.strip().split('\n')
            formatted = []
            for line in lines:
                line = line.strip()
                if line:
                    formatted.append(line)
            acs_instructions_parts.append("")
            acs_instructions_parts.append("Special Instructions:")
            acs_instructions_parts.extend(formatted)

        if acs_deliver:
            acs_instructions_parts.append("")
            acs_instructions_parts.append(f"Deliver to: {acs_deliver.strip()}")

        if acs_dl:
            acs_instructions_parts.append(f"Driver's Licence: {acs_dl}")

        acs_instructions = "\n".join(acs_instructions_parts) if acs_instructions_parts else None

        acs_desc_parts = [p for p in [acs_year, acs_make, acs_model, acs_colour] if p]
        acs_auto_desc = " ".join(acs_desc_parts) if acs_desc_parts else None

        acs_phone = _normalise_phone(acs_mobile) or _normalise_phone(acs_ph_bus)

        acs_confidence = {}
        acs_field_map = {
            "lender_name": acs_lender, "client_reference": acs_contract,
            "account_number": acs_contract, "regulated_type": acs_reg_type,
            "customer_name": acs_customer, "customer_phone": acs_phone,
            "customer_address": acs_full_addr, "deliver_to": acs_deliver,
            "rego": acs_rego, "vin": acs_vin, "engine_number": acs_engine,
            "year": acs_year, "make": acs_make, "model": acs_model,
            "colour": acs_colour, "instructions": acs_instructions,
            "arrears": acs_arrears, "our_ref": acs_job_no,
        }
        for k, v in acs_field_map.items():
            acs_confidence[k] = "extracted" if v else None
        if acs_lender and acs_account_name and acs_lender.lower() != acs_account_name.strip().lower():
            acs_confidence["lender_name"] = "review"

        return {
            "client_reference":  acs_contract,
            "contract_number":   acs_contract,
            "account_number":    acs_contract,
            "our_ref":           acs_job_no,
            "regulated_type":    acs_reg_type,
            "lender_name":       acs_lender,
            "from_name":         "Australian Collection Services",
            "deliver_to":        acs_deliver.strip().rstrip(".,") if acs_deliver else None,
            "costs":             None,
            "instalment_amount": acs_install_due,
            "total_collect":     acs_total_collect,
            "security_type":     "Vehicle",
            "customer": {
                "full_name": acs_customer.strip() if acs_customer else None,
                "company":   None,
                "dob":       acs_dob,
                "email":     acs_email,
                "mobile":    acs_phone,
            },
            "job_address_full": acs_full_addr,
            "asset_address":    None,
            "security": {
                "year":          acs_year,
                "make":          acs_make,
                "model":         acs_model,
                "rego":          acs_rego,
                "vin":           acs_vin,
                "engine_number": acs_engine,
                "colour":        acs_colour,
                "description":   acs_auto_desc,
                "reg_expiry":    acs_reg_exp,
                "reg_note":      acs_reg_note,
            },
            "financials": {
                "arrears":   (acs_arrears or "").replace(",", "") or None,
                "due_date":  acs_next_due,
            },
            "instructions_raw": acs_instructions,
            "_confidence":      acs_confidence,
            "_filled_count":    sum(1 for v in acs_confidence.values() if v),
            "_source":          "acs_worksheet",
        }

    # ── Wise Group document detection ──────────────────────────────────────
    is_wise = bool(re.search(r"WISE\s*GROUP\s*CASE\s*NUMBER", text, re.IGNORECASE))

    if is_wise:
        wise_case = find_after([r"WISE\s*GROUP\s*CASE\s*NUMBER[:\s]*([A-Za-z0-9\-\/]+)"])
        wise_client = find_after([r"CLIENT[:\s]+([^\n]{3,80}?)(?:\s*\n|$|\s{2,})"])
        wise_financier = find_after([r"FINANCIER[:\s]+([^\n]{3,60}?)(?:\s*\n|,\s*|$)"])
        wise_debtor1 = find_after([
            r"NAME\s*OF\s*DEBTOR\s*1[:\s]*([A-Za-z][A-Za-z\s,.']{2,60}?)\s*(?:\(D\.?O\.?B|$|\n)",
        ])
        wise_dob = find_after([
            r"NAME\s*OF\s*DEBTOR\s*1[^(]*\(\s*D\.?O\.?B\.?\s*([0-9]{1,2}[\/\-][0-9]{1,2}[\/\-][0-9]{2,4})",
        ])
        wise_addr = find_after([r"Given\s*address[:\s]*([^\n]{5,120})"])
        wise_mobile = find_after([
            r"Phone\s*numbers?[^M]*M[:\s]*([+0-9\(\)\s\-]{8,20})",
            r"\bM[:\s]*([+0-9]{8,12})\b",
        ])
        wise_email = find_after([r"Email[:\s]*([^\s@]+@[^\s\n]+)"])
        wise_deliver = find_after([r"DELIVER\s*VEHICLE\s*TO[:\s]*([A-Za-z0-9 ,.'&\-]{3,60}?)(?:\s*\n|$)"])
        wise_account = find_after([r"ACCOUNT\s*NUMBER[:\s]*([A-Za-z0-9\-\/]{4,20})"])
        wise_loan_type = find_after([r"LOAN\s*TYPE[:\s]*([^\n,]{3,40}?)(?:\s*\n|,|$)"])
        wise_arrears = find_after([r"ARREARS[:\s]*\$?\s*([0-9,]+\.?\d{0,2})"])
        wise_costs = find_after([r"COSTS[:\s]*\$?\s*([0-9,]+\.?\d{0,2})"])
        wise_instalment = find_after([r"INSTALMENT\s*AMOUNT[:\s]*([0-9,]+\.?\d{0,2})"])
        wise_freq = find_after([r"INSTALMENT\s*FREQUENCY[:\s]*([A-Za-z]{3,12})"])
        wise_make_model = find_after([r"MAKE[\/\\]MODEL[:\s]*([A-Za-z0-9][A-Za-z0-9 ]{1,40}?)(?:\s*\n|$|\s{2,})"])
        wise_year_v = find_after([r"(?:^|\n)\s*YEAR[:\s]*([12][0-9]{3})"])
        wise_rego = find_after([r"REGISTRATION\s*#[:\s]*([A-Za-z0-9]{2,9})"])
        wise_colour = find_after([r"COLOUR[:\s]*([A-Za-z][A-Za-z ]{1,20}?)(?:\s*\n|$|\s{2,})"])
        wise_engine = find_after([r"ENGINE\s*#[:\s]*([A-Za-z0-9\-]{5,20})"])
        wise_vin_v = find_after([r"VIN[\/\\]CHASSIS\s*#[:\s]*([A-HJ-NPR-Z0-9]{11,17})"])

        wise_make = wise_model = None
        if wise_make_model:
            parts = wise_make_model.strip().split(None, 1)
            wise_make = parts[0].title()
            wise_model = parts[1].title() if len(parts) > 1 else None

        wise_lender = wise_financier or wise_client
        wise_reg_type = None
        if wise_loan_type:
            lt = wise_loan_type.lower().strip()
            if any(k in lt for k in _UNREGULATED_TYPES):
                wise_reg_type = "UNREGULATED"
            elif any(k in lt for k in _REGULATED_TYPES):
                wise_reg_type = "REGULATED"

        return {
            "wise_case_number":  wise_case,
            "client_reference":  wise_case,
            "contract_number":   None,
            "account_number":    wise_account,
            "our_ref":           wise_case,
            "regulated_type":    wise_reg_type,
            "lender_name":       wise_lender,
            "from_name":         wise_client,
            "deliver_to":        wise_deliver.strip().rstrip(".,") if wise_deliver else None,
            "costs":             (wise_costs or "").replace(",", "") or None,
            "instalment_amount": wise_instalment,
            "payment_frequency": wise_freq,
            "security_type":     "Vehicle",
            "customer": {
                "full_name": wise_debtor1.strip() if wise_debtor1 else None,
                "company":   None,
                "dob":       wise_dob,
                "email":     wise_email,
                "mobile":    _normalise_phone(wise_mobile),
            },
            "job_address_full": wise_addr.strip(" ,") if wise_addr else None,
            "asset_address":    None,
            "security": {
                "year":          wise_year_v,
                "make":          wise_make,
                "model":         wise_model,
                "rego":          wise_rego,
                "vin":           wise_vin_v,
                "engine_number": wise_engine,
                "colour":        wise_colour,
                "description":   " ".join(filter(None, [wise_year_v, wise_make, wise_model])) or None,
            },
            "financials": {
                "arrears":   (wise_arrears or "").replace(",", "") or None,
                "due_date":  find_after([r"NEXT\s*DUE\s*DATE[:\s]*([0-9]{1,2}[\/\-][0-9]{1,2}[\/\-][0-9]{2,4})"]),
            },
            "instructions_raw": None,
            "_confidence": {k: "extracted" for k in [
                "lender_name", "client_reference", "account_number", "customer_name",
                "customer_phone", "customer_address", "deliver_to", "rego", "vin",
                "year", "make", "model", "colour", "engine_number", "regulated_type",
                "arrears", "costs",
            ] if locals().get(f"wise_{k.replace('_name','').replace('customer_','').replace('lender_','lender')}") or True},
            "_filled_count": sum(1 for v in [wise_case, wise_lender, wise_debtor1, wise_rego,
                                              wise_vin_v, wise_year_v, wise_make, wise_account] if v),
            "_source": "wise_group",
        }

    # ── Sender / client (FROM: line at top of document) ─────────────────
    from_name = find_after([
        r"^FROM[:\s]+([^\n]{3,60})",
        r"\nFROM[:\s]+([^\n]{3,60})",
        r"(?:Sent\s*(?:by|from)|Instructing\s*(?:Party|Client))[:\s]+([^\n]{3,60})",
    ])
    if from_name:
        from_name = from_name.strip().rstrip(".,")
        # If it's an auction yard, discard (it's a Deliver To, not a client)
        if any(yard in from_name.lower() for yard in _AUCTION_YARDS):
            from_name = None

    # ── Reference fields ─────────────────────────────────────────────────
    contract = find_after([r"Contract\s*(?:No\.?|Number)[:\s]*([A-Za-z0-9\-\/]+)"])
    account  = find_after([
        r"(?:^|\n)\s*A[\/\\]?C\s*(?:No\.?|Number)?[:\s]*([A-Za-z0-9\-\/]{4,20})",
        r"(?:^|\n|\s)Account\s*(?:No\.?|Number)[:\s]*([A-Za-z0-9\-\/]+)",
        r"Customer\s*(?:Account|No\.?)\s*(?:No\.?|Number)?[:\s]*([A-Za-z0-9\-\/]+)",
    ])
    client_ref = find_after([
        r"Authority\s*Ref(?:erence)?[:\s#]*([A-Za-z0-9\-\/]+)",
        r"Client\s*Reference[:\s]*([A-Za-z0-9\-\/]+)",
        r"Your\s*Ref(?:erence)?[:\s]*([A-Za-z0-9\-\/]+)",
    ])
    our_ref = find_after([
        r"Our\s*Ref(?:erence)?[:\s]*([A-Za-z0-9\-\/]+)",
        r"Instruction\s*(?:No\.?|Number)[:\s]*([A-Za-z0-9\-\/]+)",
        r"(?:Job|File)\s*(?:No\.?|Number)[:\s#]*([A-Za-z0-9\-\/]+)",
    ])

    # ── Lender / sender name for client lookup ────────────────────────────
    lender_raw = find_after([
        r"(?:Lender|Finance\s*Company|Financier|Credit\s*Provider|Secured\s*Party)[:\s]+"
        r"([A-Za-z0-9 ,.'&\-]{3,60}?)(?:\s*\n|\s{3,}|,\s*ABN|$)",
    ])
    # Prefer from_name if no explicit lender label
    if not lender_raw and from_name:
        lender_raw = from_name

    deliver_to = None
    lender     = None
    if lender_raw:
        lender_clean = lender_raw.strip().rstrip(".,")
        if any(yard in lender_clean.lower() for yard in _AUCTION_YARDS):
            deliver_to = lender_clean
        else:
            lender = lender_clean

    explicit_deliver = find_after([
        r"Deliver(?:\s*Vehicle)?\s*To[:\s]*([A-Za-z0-9 ,.'&\-]{3,60}?)(?:\s*\n|$)",
        r"Deliver\s*To\s*\|?\s*([A-Za-z0-9 ,.'&\-]{3,60}?)(?:\s*\n|$)",
        r"Release\s*To[:\s]*([A-Za-z0-9 ,.'&\-]{3,60}?)(?:\s*\n|$)",
        r"(?:Yard|Auction)[:\s]*([A-Za-z0-9 ,.'&\-]{3,40}?)(?:\s*\n|$)",
    ])
    if explicit_deliver:
        deliver_to = explicit_deliver.strip().rstrip(".,")

    # ── Regulation type ───────────────────────────────────────────────────
    contract_type = find_after([
        r"(?:Contract|Product|Loan|Finance|Agreement)\s*Type[:\s]*([A-Za-z][A-Za-z ]{2,35}?)(?:\s*\n|$|\s{2,})",
        r"(?:Type\s*of\s*(?:Finance|Contract|Agreement))[:\s]*([A-Za-z][A-Za-z ]{2,35}?)(?:\s*\n|$)",
        r"Regulation[:\s]*([A-Za-z][A-Za-z ]{2,35}?)(?:\s*\n|$|\s{2,})",
    ])
    reg_type = find_after([
        r"Contract\s*Type[:\s]*(REGULATED|UNREGULATED)",
        r"Regulation[:\s]*(REGULATED|UNREGULATED)",
        r"\b(REGULATED|UNREGULATED)\b",
    ])
    if not reg_type and contract_type:
        ct = contract_type.lower().strip()
        if any(k in ct for k in _UNREGULATED_TYPES):
            reg_type = "UNREGULATED"
        elif any(k in ct for k in _REGULATED_TYPES):
            reg_type = "REGULATED"

    # ── Customer: try to split company vs individual ──────────────────────
    cust_name = find_after([
        r"Customer\s*(?:Name)?[:\s]*([A-Za-z][A-Za-z ,.'&\-]{2,60}?)(?:\s*\n|$)",
        r"Debtor\s*(?:Name)?[:\s]*([A-Za-z][A-Za-z ,.'&\-]{2,60}?)(?:\s*\n|$)",
        r"Borrower(?:'s)?\s*(?:Name)?[:\s]*([A-Za-z][A-Za-z ,.'&\-]{2,60}?)(?:\s*\n|$)",
        r"(?:^|\n)Name[:\s]*([A-Za-z][A-Za-z ,.'&\-]{2,60}?)(?:\s*\n|$)",
    ])
    cust_company = None
    # Heuristic: if name contains company indicators, treat as company
    _COMPANY_WORDS = ("pty", "ltd", "pty ltd", "limited", "inc", "llc", "liquidat", "admin", "trust", "group")
    if cust_name and any(kw in cust_name.lower() for kw in _COMPANY_WORDS):
        cust_company = cust_name.strip()
        # Try to also find an individual contact name after the company line
        contact_name = find_after([
            r"Contact(?:\s*Person)?[:\s]*([A-Za-z][A-Za-z ,.']{2,40}?)(?:\s*\n|$)",
            r"Attention[:\s]*([A-Za-z][A-Za-z ,.']{2,40}?)(?:\s*\n|$)",
        ])
        if contact_name:
            cust_name = contact_name.strip()
        else:
            cust_name = None  # Just a company, no individual

    dob   = find_after([
        r"D\.?O\.?B\.?[:\s]*([0-9]{1,2}[\/\-][0-9]{1,2}[\/\-][0-9]{2,4})",
        r"Date\s*of\s*Birth[:\s]*([0-9]{1,2}[\/\-][0-9]{1,2}[\/\-][0-9]{2,4})",
    ])
    email = find_after([r"Email[:\s]*([^\s@]+@[^\s]+)"])
    vin   = find_after([
        r"V\.?I\.?N\.?\s*#?[:\s]*([A-HJ-NPR-Z0-9]{11,17})",
        r"\bVIN\b\s*#?[:\s]*([A-HJ-NPR-Z0-9]{11,17})",
        r"VIN[\/]Chassis[:\s]*([A-HJ-NPR-Z0-9]{11,17})",
        r"Chassis\s*(?:No\.?|Number)[:\s]*([A-HJ-NPR-Z0-9]{11,17})",
    ])
    rego  = find_after([
        r"Registration\s*(?:Plate|Number)?[:\s]+([A-Za-z0-9]{2,9})\b",
        r"(?<![A-Za-z])Rego[:\s#]+([A-Za-z0-9]{2,9})\b",
        r"(?<![A-Za-z])Reg[:\s#]+([A-Za-z0-9]{2,9})\b",
        r"\bPlate\b[:\s]+([A-Za-z0-9]{2,9})\b",
    ])
    engine_no = find_after([
        r"Engine\s*[#\s]*(?:No\.?|Number)?[:\s]*([A-Za-z0-9\-\/]+)",
    ])
    colour = find_after([r"Colou?r[:\s]*([A-Za-z][A-Za-z ]{1,20}?)(?:\s*\n|$|\s{2,})"])
    year   = find_after([r"\bYear\b[:\s]*([12][0-9]{3})", r"\b((?:19|20)[0-9]{2})\b"])
    make   = find_after([r"\bMake\b[:\s]*([A-Za-z][A-Za-z0-9\- ]{1,25}?)(?:\s*\n|$|\s{2,})"])
    model  = find_after([r"\bModel\b[:\s]*([A-Za-z0-9][A-Za-z0-9\- ]{1,35}?)(?:\s*\n|$|\s{2,})"])
    if not make and not model:
        make_model_combined = find_after([
            r"Make\s*/\s*Model[:\s]*([A-Za-z0-9][A-Za-z0-9 \-]{2,40}?)(?:\s*\n|\s{2,}|$)",
        ])
        if make_model_combined:
            mm_parts = make_model_combined.strip().split(None, 1)
            make = mm_parts[0]
            model = mm_parts[1] if len(mm_parts) > 1 else None

    # ── Security type ─────────────────────────────────────────────────────
    sec_type_raw = find_after([
        r"Security\s*Type[:\s]*([A-Za-z][A-Za-z ]{2,20}?)(?:\s*\n|$|\s{2,})",
        r"(?:^|\n)Security[:\s]*([A-Za-z][A-Za-z ]{2,20}?)(?:\s*\n|$|\s{2,})",
        r"Asset\s*Type[:\s]*([A-Za-z][A-Za-z ]{2,20}?)(?:\s*\n|$)",
    ])
    _SEC_TYPES = {
        "vehicle": "Vehicle", "motor vehicle": "Vehicle", "motor": "Vehicle",
        "car": "Vehicle", "truck": "Vehicle", "ute": "Vehicle",
        "property": "Property", "real property": "Property", "land": "Property",
        "equipment": "Equipment", "machinery": "Equipment", "plant": "Equipment",
        "other": "Other",
    }
    security_type = "Vehicle"  # default
    if sec_type_raw:
        stl = sec_type_raw.lower().strip()
        for k, v in _SEC_TYPES.items():
            if k in stl:
                security_type = v
                break

    # ── Asset / last-known address (separate from customer address) ───────
    asset_address = find_after([
        r"(?:Last\s*Known\s*Location|Last\s*Known\s*Address)[:\s]*([^\n]{5,100})",
        r"Asset\s*Address[:\s]*([^\n]{5,100})",
        r"(?:Located|Location)\s*(?:at|At)[:\s]*([^\n]{5,100})",
        r"(?:Vehicle|Asset)\s*(?:Located|Location)[:\s]*([^\n]{5,100})",
        r"(?:Property|Site)\s*Address[:\s]*([^\n]{5,100})",
    ])
    if asset_address:
        asset_address = asset_address.strip(" ,")

    # ── Customer address (home / service / registered) ────────────────────
    addr = find_after([
        r"Home\s*Address[:\s]*([^\n]{5,120})",
        r"(?:Service|Customer|Registered)\s*Address[:\s]*([^\n]{5,120})",
        r"(?:^|\n)Address[:\s]*([^\n]{5,120})",
    ])
    addr = addr.strip(" ,") if addr else None

    # ── Phone ─────────────────────────────────────────────────────────────
    phone = find_after([
        r"Mobile\s*(?:Phone|No\.?)?[:\s]*([+0-9\(\)\s\-]{8,20})",
        r"(?:Contact\s*(?:No\.?|Number)|Phone|Tel(?:ephone)?)[:\s]*([+0-9\(\)\s\-]{8,20})",
    ])
    phone = _normalise_phone(phone)

    # ── Financial ─────────────────────────────────────────────────────────
    arrears  = find_after([r"Arrears?[:\s]*\$?\s*([0-9,]+\.?\d{0,2})"])
    due_date = find_after([
        r"(?:Next\s*Due\s*Date|Due\s*Date|Payment\s*Due)[:\s]*"
        r"([0-9]{1,2}[A-Za-z]{3}[0-9]{2,4}|[0-9]{1,2}[\/\-][0-9]{1,2}[\/\-][0-9]{2,4})",
    ])
    costs = find_after([
        r"(?:Recovery|Repossession|Repo|Our)\s*(?:Fee|Cost)s?[:\s]*\$?\s*([0-9,]+(?:\.\d{2})?)",
        r"(?:Total\s*)?(?:Costs?|Fee)[s]?[:\s]*\$?\s*([0-9,]+(?:\.\d{2})?)",
        r"(?:Charge|Commission)[:\s]*\$?\s*([0-9,]+(?:\.\d{2})?)",
    ])
    if costs:
        costs = costs.replace(",", "")

    # ── Instructions narrative ────────────────────────────────────────────
    instructions = find_block([
        r"(?:Special\s*)?Instructions?[:\s]*\n+([\s\S]+?)(?=\n{2,}[A-Z][A-Z\s]+:|\Z)",
        r"Background[:\s]*\n+([\s\S]+?)(?=\n{2,}[A-Z][A-Z\s]+:|\Z)",
        r"Notes?[:\s]*\n+([\s\S]+?)(?=\n{2,}[A-Z][A-Z\s]+:|\Z)",
        r"Additional\s*(?:Notes?|Information)[:\s]*\n+([\s\S]+?)(?=\n{2,}[A-Z][A-Z\s]+:|\Z)",
        r"(?:Comments?|Remarks?)[:\s]*\n+([\s\S]+?)(?=\n{2,}[A-Z][A-Z\s]+:|\Z)",
    ])
    if not instructions:
        instructions = find_block([
            r"(?:Special\s*)?Instructions?[:\s]*\n+([\s\S]+)",
            r"Background[:\s]*\n+([\s\S]+)",
        ])
    if not instructions:
        instructions = find_after([
            r"Special\s*Instructions?[:\s]+([^\n]{15,})",
            r"Notes?[:\s]+([^\n]{15,})",
            r"Instructions?[:\s]+([^\n]{25,})",
        ])

    # ── Build auto description from vehicle parts ─────────────────────────
    desc_parts = [p for p in [year, make, model] if p]
    if colour:
        desc_parts.append(colour)
    auto_description = " ".join(desc_parts) if desc_parts else None

    # ── Confidence per field ──────────────────────────────────────────────
    def _conf(val):
        return "extracted" if val else None

    confidence = {
        "lender_name":      _conf(lender),
        "from_name":        _conf(from_name),
        "client_reference": _conf(client_ref or contract),
        "our_ref":          _conf(our_ref),
        "account_number":   _conf(account),
        "regulated_type":   _conf(reg_type),
        "customer_name":    _conf(cust_name or cust_company),
        "customer_company": _conf(cust_company),
        "customer_phone":   _conf(phone),
        "customer_address": _conf(addr),
        "deliver_to":       _conf(deliver_to),
        "rego":             _conf(rego),
        "vin":              _conf(vin),
        "engine_number":    _conf(engine_no),
        "year":             _conf(year),
        "make":             _conf(make),
        "model":            _conf(model),
        "colour":           _conf(colour),
        "asset_address":    _conf(asset_address),
        "instructions":     _conf(instructions),
        "costs":            _conf(costs),
        "arrears":          _conf(arrears),
    }
    filled_count = sum(1 for v in confidence.values() if v)

    return {
        "client_reference":  client_ref or contract,
        "contract_number":   contract,
        "account_number":    account,
        "our_ref":           our_ref,
        "regulated_type":    reg_type,
        "lender_name":       lender,
        "from_name":         from_name,
        "deliver_to":        deliver_to,
        "costs":             costs,
        "security_type":     security_type,
        "customer": {
            "full_name":  cust_name,
            "company":    cust_company,
            "dob":        dob,
            "email":      email,
            "mobile":     phone,
        },
        "job_address_full": addr,
        "asset_address":    asset_address,
        "security": {
            "year":          year,
            "make":          make,
            "model":         model,
            "rego":          rego,
            "vin":           vin,
            "engine_number": engine_no,
            "colour":        colour,
            "description":   auto_description,
        },
        "financials": {
            "arrears":   arrears,
            "due_date":  due_date,
        },
        "instructions_raw": instructions,
        "_confidence":      confidence,
        "_filled_count":    filled_count,
    }

# ─────────────────────────────────────────────────────────────────────────────


def money_to_cents(s: str) -> int:
    s = (s or "").strip()
    if not s:
        return 0
    s = s.replace("$", "").replace(",", "").strip()
    if not re.match(r"^\d+(\.\d{1,2})?$", s):
        return 0
    if "." in s:
        dollars, cents = s.split(".", 1)
        cents = (cents + "00")[:2]
    else:
        dollars, cents = s, "00"
    return int(dollars) * 100 + int(cents)


def cents_to_money(cents) -> str:
    return f"${int(cents or 0) / 100:,.2f}"


def format_ddmmyyyy(d):
    if d is None:
        return ""
    if isinstance(d, str):
        try:
            d = datetime.strptime(d[:10], "%Y-%m-%d").date()
        except Exception:
            return d
    if isinstance(d, datetime):
        d = d.date()
    return d.strftime("%d/%m/%Y")


def parse_interaction_datetime(date_str: str, time_str: str) -> str:
    date_str = (date_str or "").strip()
    time_str = (time_str or "").strip()
    if not date_str:
        return datetime.now().isoformat(timespec="seconds")
    if time_str:
        combined = f"{date_str} {time_str}"
        for fmt in ("%d/%m/%Y %H:%M", "%d/%m/%Y %I:%M %p", "%d/%m/%Y %I:%M%p"):
            try:
                return datetime.strptime(combined, fmt).isoformat(timespec="seconds")
            except Exception:
                continue
    try:
        return datetime.strptime(date_str, "%d/%m/%Y").isoformat(timespec="seconds")
    except Exception:
        return datetime.now().isoformat(timespec="seconds")


_AUTO_ADVANCE_TYPES = {
    "attendance", "repo attempt", "card left", "neighbour interview", "note"
}

def maybe_auto_advance_status(cur, job_id: int, current_status: str,
                               event_type: str, role: str) -> bool:
    if role not in ("admin", "both"):
        return False
    if event_type.lower() not in _AUTO_ADVANCE_TYPES:
        return False
    if current_status != "New":
        return False
    cur.execute("UPDATE jobs SET status = 'Active' WHERE id = ?", (job_id,))
    return True


def format_interaction_dt(s: str) -> str:
    if not s:
        return ""
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M"):
        try:
            dt = datetime.strptime(s[:19], fmt)
            return dt.strftime("%-d/%-m/%Y %-I:%M %p").lower()
        except Exception:
            continue
    return s


def calc_total_due_now(arrears, costs, mmp, due_date, costs2=0):
    arrears = float(arrears or 0)
    costs   = float(costs or 0)
    costs2  = float(costs2 or 0)
    mmp     = float(mmp or 0)
    if isinstance(due_date, str):
        try:
            due_date = datetime.strptime(due_date[:10], "%Y-%m-%d").date()
        except Exception:
            due_date = None
    if isinstance(due_date, datetime):
        due_date = due_date.date()
    today = datetime.now(_melbourne).date()
    total = arrears + costs + costs2
    include_mmp = bool(due_date and due_date < today)
    if include_mmp:
        total += mmp
    return round(total, 2), include_mmp


def _add_months(d, n):
    import calendar
    month = d.month + n
    year  = d.year + (month - 1) // 12
    month = (month - 1) % 12 + 1
    day   = min(d.day, calendar.monthrange(year, month)[1])
    return d.replace(year=year, month=month, day=day)


def advance_due_date_display(due_date_str, payment_frequency):
    """Return the next upcoming due date (string YYYY-MM-DD) by advancing
    past today according to the payment frequency.  Used for display only.
    Defaults to monthly if no frequency is set."""
    if not due_date_str:
        return due_date_str
    try:
        due = datetime.strptime(due_date_str[:10], "%Y-%m-%d").date()
    except Exception:
        return due_date_str
    today = datetime.now(_melbourne).date()
    if due >= today:
        return due_date_str
    freq = (payment_frequency or "monthly").lower()
    from datetime import timedelta as _td
    max_iter = 1000
    i = 0
    while due < today and i < max_iter:
        if "fortnight" in freq:
            due += _td(days=14)
        elif "week" in freq:
            due += _td(days=7)
        else:
            due = _add_months(due, 1)
        i += 1
    return due.isoformat()


app.jinja_env.globals.update(
    cents_to_money=cents_to_money,
    format_ddmmyyyy=format_ddmmyyyy,
    format_interaction_dt=format_interaction_dt,
)


@app.template_filter("strip_ai_prefix")
def strip_ai_prefix(text):
    """Strip the [AI Update] marker from stored field notes before display.
    The AI assistance feature remains available; this only affects the label."""
    if not text:
        return text
    if text.startswith("[AI Update]\n"):
        return text[len("[AI Update]\n"):]
    return text


@app.template_filter("fmt_queue_dt")
def fmt_queue_dt(ts_str):
    """Format a cue_items created_at string as '05Mar26 09:07'."""
    if not ts_str:
        return "—"
    try:
        dt = datetime.strptime(ts_str[:19], "%Y-%m-%dT%H:%M:%S")
        return dt.strftime("%d%b%y %H:%M")
    except Exception:
        return ts_str[:16]


def users_count():
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) c FROM users")
    c = cur.fetchone()["c"]
    conn.close()
    return c


def generate_internal_job_number():
    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM system_settings WHERE id = 1")
    settings = cur.fetchone()
    if not settings:
        conn.close()
        return datetime.now().strftime("%y%m") + "001"

    current_prefix = settings["job_prefix"] or datetime.now().strftime("%y%m")
    auto_enabled = settings["auto_prefix_enabled"]

    if auto_enabled:
        actual_prefix = datetime.now().strftime("%y%m")
        if actual_prefix != current_prefix:
            current_prefix = actual_prefix
            cur.execute("""
                UPDATE system_settings
                SET job_prefix = ?, job_sequence = 0, updated_at = ?
                WHERE id = 1
            """, (actual_prefix, now_ts()))
            conn.commit()

    new_sequence = settings["job_sequence"] + 1
    padded = str(new_sequence).zfill(3)

    cur.execute("""
        UPDATE system_settings
        SET job_sequence = ?, updated_at = ?
        WHERE id = 1
    """, (new_sequence, now_ts()))

    conn.commit()
    conn.close()

    return f"{current_prefix}{padded}"


def audit(entity_type, entity_id, action, message, meta=None):
    conn = db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO audit_log (actor_user_id, entity_type, entity_id, action, message, meta_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        session.get("user_id"),
        entity_type,
        entity_id,
        action,
        message,
        json.dumps(meta) if meta else None,
        now_ts()
    ))
    conn.commit()
    conn.close()


def _log_lifecycle(cur, job_id, action, from_status, to_status, user_id, notes=None, batch_id=None):
    cur.execute("""
        INSERT INTO job_lifecycle_log
            (job_id, action, from_status, to_status, performed_by_user_id, performed_at, notes, batch_id)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (job_id, action, from_status, to_status, user_id, now_ts(), notes, batch_id))


# -------- Auth helpers --------
def _is_mobile_request():
    ua = (request.headers.get("User-Agent") or "").lower()
    if request.path.startswith("/m/") or request.path == "/m":
        return True
    if "axionx/" in ua:
        return True
    if any(tok in ua for tok in ("iphone", "ipad", "ipod", "android")) and "mobile" in ua:
        return True
    if any(tok in ua for tok in ("iphone", "ipad", "ipod")) and "applewebkit" in ua:
        return True
    return False

def _login_redirect():
    if _is_mobile_request():
        next_path = request.path
        return redirect(url_for("m_login") + f"?next={next_path}")
    return redirect(url_for("login"))

def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get("user_id"):
            return _login_redirect()
        return f(*args, **kwargs)
    return wrapper


def admin_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get("user_id"):
            return _login_redirect()
        if session.get("role") not in ("admin", "both"):
            flash("Admin access required.", "danger")
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return wrapper


import hmac as _hmac

def _geoop_pw():
    return os.environ.get("GEOOP_PASSWORD", "")


def _geoop_is_unlocked():
    gpw = _geoop_pw()
    if not gpw:
        return True
    return session.get("geoop_unlocked") == session.get("user_id")


def geoop_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get("user_id"):
            return _login_redirect()
        if session.get("role") not in ("admin", "both"):
            flash("Admin access required.", "danger")
            return redirect(url_for("index"))
        if not _geoop_is_unlocked():
            return redirect(url_for("geoop_import_page"))
        return f(*args, **kwargs)
    return wrapper


@app.post("/admin/geoop-unlock")
@admin_required
def geoop_unlock_post():
    gpw = _geoop_pw()
    if not gpw:
        return redirect(url_for("geoop_import_page"))
    pw = request.form.get("password", "")
    if _hmac.compare_digest(pw, gpw):
        session["geoop_unlocked"] = session["user_id"]
        flash("GeoOp Import unlocked.", "success")
    else:
        flash("Incorrect password.", "danger")
    return redirect(url_for("geoop_import_page"))


# -------- Login / Logout --------

import smtplib as _smtplib
import os as _os
from email.mime.multipart import MIMEMultipart as _MIMEMultipart
from email.mime.text import MIMEText as _MIMEText
from email.mime.base import MIMEBase as _MIMEBase
from email import encoders as _encoders


def send_reset_email(to_addr, reset_link):
    host = _os.environ.get("SMTP_HOST", "smtp.gmail.com")
    port = int(_os.environ.get("SMTP_PORT", "587"))
    user = _os.environ.get("SMTP_USER", "")
    pswd = _os.environ.get("SMTP_PASS", "")
    frm  = _os.environ.get("SMTP_FROM", user)
    if not user or not pswd:
        raise RuntimeError("SMTP credentials not configured.")
    msg = _MIMEMultipart("alternative")
    msg["Subject"] = "Axion — Password Reset"
    msg["From"]    = frm
    msg["To"]      = to_addr
    txt = f"Click the link below to reset your Axion password:\n\n{reset_link}\n\nThis link expires in 1 hour."
    htm = f"""<div style="font-family:sans-serif;max-width:480px;margin:0 auto">
  <h2 style="color:#0f172a">Reset your Axion password</h2>
  <p>Click the button below to choose a new password. This link expires in <strong>1 hour</strong>.</p>
  <p><a href="{reset_link}" style="display:inline-block;background:#3b82f6;color:#fff;
     padding:12px 24px;border-radius:8px;text-decoration:none;font-weight:600">Reset Password</a></p>
  <p style="color:#6b7280;font-size:13px">If you did not request this, you can safely ignore this email.</p>
  <hr style="border:none;border-top:1px solid #e5e7eb">
  <p style="color:#9ca3af;font-size:12px">Axion Field Operations Management</p>
</div>"""
    msg.attach(_MIMEText(txt, "plain"))
    msg.attach(_MIMEText(htm, "html"))
    with _smtplib.SMTP(host, port, timeout=10) as s:
        s.ehlo()
        s.starttls()
        s.login(user, pswd)
        s.sendmail(frm, to_addr, msg.as_string())


def send_email(to_list, subject, body_txt, body_html=None, cc_list=None, attachments=None):
    """Generic SMTP helper.
    to_list/cc_list: lists of email strings.
    attachments: list of (filename, bytes_data, mime_type_str) tuples.
    """
    import os as _os
    host = _os.environ.get("SMTP_HOST", "smtp.gmail.com")
    port = int(_os.environ.get("SMTP_PORT", "587"))
    user = _os.environ.get("SMTP_USER", "")
    pswd = _os.environ.get("SMTP_PASS", "")
    frm  = _os.environ.get("SMTP_FROM", user)
    if not user or not pswd:
        raise RuntimeError("SMTP credentials not configured.")
    to_list = [e for e in (to_list or []) if e and "@" in e]
    if not to_list:
        raise ValueError("No valid recipient email addresses.")
    cc_list = [e for e in (cc_list or []) if e and "@" in e]

    msg = _MIMEMultipart("mixed")
    msg["Subject"] = subject
    msg["From"]    = frm
    msg["To"]      = ", ".join(to_list)
    if cc_list:
        msg["Cc"] = ", ".join(cc_list)

    alt = _MIMEMultipart("alternative")
    alt.attach(_MIMEText(body_txt, "plain"))
    if body_html:
        alt.attach(_MIMEText(body_html, "html"))
    msg.attach(alt)

    for fname, fdata, fmime in (attachments or []):
        maintype, subtype = (fmime or "application/octet-stream").split("/", 1)
        part = _MIMEBase(maintype, subtype)
        part.set_payload(fdata)
        _encoders.encode_base64(part)
        part.add_header("Content-Disposition", "attachment", filename=fname)
        msg.attach(part)

    all_recipients = to_list + cc_list
    with _smtplib.SMTP(host, port, timeout=10) as s:
        s.ehlo()
        s.starttls()
        s.login(user, pswd)
        s.sendmail(frm, all_recipients, msg.as_string())


@app.get("/forgot-password")
def forgot_password():
    return render_template("forgot_password.html")


@app.post("/forgot-password")
def forgot_password_post():
    import secrets as _tok
    from datetime import timedelta as _td
    email = request.form.get("email", "").strip().lower()
    generic_msg = "If that email is registered you\u2019ll receive a reset link shortly."
    if not email:
        flash(generic_msg, "success")
        return redirect(url_for("forgot_password"))
    conn = db()
    cur  = conn.cursor()
    cur.execute("SELECT id FROM users WHERE email = ? AND active = 1", (email,))
    user = cur.fetchone()
    if user:
        token   = _tok.token_urlsafe(32)
        expires = (datetime.now() + _td(hours=1)).isoformat(timespec="seconds")
        created = datetime.now().isoformat(timespec="seconds")
        cur.execute("""INSERT INTO password_reset_tokens (user_id, token, expires_at, used, created_at)
                       VALUES (?, ?, ?, 0, ?)""", (user["id"], token, expires, created))
        conn.commit()
        link = request.host_url.rstrip("/") + url_for("reset_password", token=token)
        try:
            send_reset_email(email, link)
        except Exception as e:
            import sys
            print(f"[SMTP ERROR] {e}", file=sys.stderr)
    conn.close()
    flash(generic_msg, "success")
    return redirect(url_for("forgot_password"))


@app.get("/reset-password/<token>")
def reset_password(token):
    conn = db()
    cur  = conn.cursor()
    cur.execute("""SELECT * FROM password_reset_tokens
                   WHERE token = ? AND used = 0 AND expires_at > ?""",
                (token, datetime.now().isoformat(timespec="seconds")))
    row = cur.fetchone()
    conn.close()
    if not row:
        flash("That reset link is invalid or has expired.", "danger")
        return redirect(url_for("forgot_password"))
    return render_template("reset_password.html", token=token)


@app.post("/reset-password/<token>")
def reset_password_post(token):
    conn = db()
    cur  = conn.cursor()
    cur.execute("""SELECT * FROM password_reset_tokens
                   WHERE token = ? AND used = 0 AND expires_at > ?""",
                (token, datetime.now().isoformat(timespec="seconds")))
    row = cur.fetchone()
    if not row:
        conn.close()
        flash("That reset link is invalid or has expired.", "danger")
        return redirect(url_for("forgot_password"))
    new_pw  = request.form.get("new_password", "")
    confirm = request.form.get("confirm_password", "")
    if len(new_pw) < 8:
        conn.close()
        flash("Password must be at least 8 characters.", "danger")
        return redirect(url_for("reset_password", token=token))
    if new_pw != confirm:
        conn.close()
        flash("Passwords do not match.", "danger")
        return redirect(url_for("reset_password", token=token))
    hashed = generate_password_hash(new_pw)
    cur.execute("UPDATE users SET password = ? WHERE id = ?", (hashed, row["user_id"]))
    cur.execute("UPDATE password_reset_tokens SET used = 1 WHERE id = ?", (row["id"],))
    conn.commit()
    conn.close()
    flash("Password updated. Please sign in.", "success")
    return _login_redirect()


@app.get("/login")
def login():
    if session.get("user_id"):
        return redirect(url_for("index"))
    if _is_mobile_request():
        return redirect(url_for("m_login"))
    return render_template("login.html")


def _client_ip():
    xff = request.headers.get("X-Forwarded-For", "")
    if xff:
        return xff.split(",")[0].strip()
    return request.remote_addr or ""

def _audit(event: str, ip: str, ok: bool):
    print(f"[BREAKGLASS] {datetime.utcnow().isoformat()}Z event={event} ok={ok} ip={ip}")

@app.route("/dev/break-glass")
def break_glass():
    if os.environ.get("DEV_BREAKGLASS_ENABLED", "false").lower() != "true":
        abort(404)

    ip = _client_ip()

    allow = [x.strip() for x in os.environ.get("DEV_BREAKGLASS_IPS", "").split(",") if x.strip()]
    if allow and ip not in allow:
        _audit("ip_block", ip, False)
        abort(404)

    provided = request.args.get("token", "")
    expected = os.environ.get("DEV_ACCESS_TOKEN", "")
    if not expected or not hmac.compare_digest(provided, expected):
        _audit("token_fail", ip, False)
        abort(404)

    session.clear()
    session["user_id"] = "breakglass"
    session["role"] = "admin"
    session["breakglass"] = True

    _audit("login_success", ip, True)
    return redirect("/")


@app.post("/login")
def login_post():
    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "").strip()

    ip = request.headers.get("X-Forwarded-For", request.remote_addr or "").split(",")[0].strip()
    ip_key = f"ip:{ip}"

    conn = db()
    cur = conn.cursor()

    allowed, locked_until = throttle_check(conn, ip_key)
    if not allowed:
        conn.close()
        flash(f"Too many failed attempts. Try again after {locked_until} UTC.", "danger")
        return redirect(url_for("login"))

    cur.execute("SELECT * FROM users WHERE email = ? AND active = 1", (email,))
    user = cur.fetchone()

    if not user or not check_password_hash(user["password"], password):
        throttle_fail(conn, ip_key)
        conn.commit()
        conn.close()
        audit("auth", None, "login_failed", f"Failed login attempt for '{email}'", {"ip": ip})
        flash("Invalid email or password.", "danger")
        return redirect(url_for("login"))

    throttle_success(conn, ip_key)
    conn.commit()
    conn.close()

    session.permanent = True
    session["user_id"] = user["id"]
    session["user_name"] = user["full_name"]
    session["role"] = user["role"]
    audit("auth", user["id"], "login_success", f"Login: {user['full_name']}", {"ip": ip})
    next_url = request.args.get("next", "").strip()
    return redirect(next_url if next_url and next_url.startswith("/") else url_for("jobs_list"), code=303)


@app.get("/logout")
def logout():
    reason = request.args.get("reason", "")
    user_id = session.get("user_id")
    user_name = session.get("user_name", "Unknown")
    mobile = _is_mobile_request()
    if reason == "timeout" and user_id:
        audit("user", user_id, "logout", f"Session auto-expired due to inactivity: {user_name}", {})
    session.clear()
    if mobile:
        return redirect(url_for("m_login"))
    if reason == "timeout":
        flash("Your session expired due to inactivity. Please sign in again.", "warning")
    return redirect(url_for("login"))


# -------- Home --------
@app.get("/")
def index():
    if session.get("user_id"):
        if _is_mobile_request():
            return redirect(url_for("m_today"))
        return redirect(url_for("jobs_list"))
    return _login_redirect()



def auto_queue_schedule_alerts(cur, admin_user_id):
    """Auto-create queue items for jobs with overdue/today/tomorrow schedules.
    Also auto-completes stale overdue/today cue items when the schedule has
    been rescheduled to a future date, and skips jobs that already have a
    pending Agent Note Review cue item."""
    import datetime as _dt
    _mel_now = datetime.now(_melbourne)
    today = _mel_now.date().isoformat()
    tomorrow = (_mel_now.date() + _dt.timedelta(days=1)).isoformat()
    now_str = _mel_now.isoformat(timespec="seconds")

    cur.execute("""
        UPDATE cue_items SET status='Completed', completed_at=?, updated_at=?
        WHERE visit_type = 'Urgent: Schedule Overdue'
          AND status IN ('Pending', 'In Progress')
          AND job_id NOT IN (
              SELECT s.job_id FROM schedules s
              WHERE date(s.scheduled_for, 'localtime') < ?
                AND s.status NOT IN ('Cancelled', 'Completed')
          )
    """, (now_str, now_str, today))

    cur.execute("""
        UPDATE cue_items SET status='Completed', completed_at=?, updated_at=?
        WHERE visit_type = 'Schedule Due Today'
          AND status IN ('Pending', 'In Progress')
          AND job_id NOT IN (
              SELECT s.job_id FROM schedules s
              WHERE date(s.scheduled_for, 'localtime') = ?
                AND s.status NOT IN ('Cancelled', 'Completed')
          )
    """, (now_str, now_str, today))

    note_job_ids = set()
    cur.execute("""
        SELECT job_id FROM cue_items
        WHERE visit_type = 'Agent Note Review'
          AND status = 'Pending'
    """)
    for r in cur.fetchall():
        note_job_ids.add(r["job_id"])

    cur.execute("""
        SELECT j.id, j.display_ref, s.scheduled_for,
               date(s.scheduled_for,'localtime') AS sched_date
        FROM jobs j
        JOIN (
            SELECT job_id, MIN(scheduled_for) AS scheduled_for
            FROM schedules
            WHERE date(scheduled_for,'localtime') <= ?
              AND status NOT IN ('Cancelled', 'Completed')
            GROUP BY job_id
        ) s ON s.job_id = j.id
        WHERE j.status NOT IN ('Completed', 'Invoiced', 'New', 'Archived - Invoiced', 'Cold Stored')
    """, (tomorrow,))
    candidates = cur.fetchall()

    visit_map = {
        "past":     ("Urgent: Schedule Overdue",  "Urgent"),
        "today":    ("Schedule Due Today",         "High"),
        "tomorrow": ("Schedule Due Tomorrow",      "Normal"),
    }

    for row in candidates:
        if row["id"] in note_job_ids:
            continue

        sched_date = row["sched_date"]
        if sched_date < today:
            bucket = "past"
        elif sched_date == today:
            bucket = "today"
        else:
            bucket = "tomorrow"

        visit_type, priority = visit_map[bucket]

        cur.execute("""
            SELECT id FROM cue_items
            WHERE job_id = ? AND visit_type = ?
              AND status IN ('Pending','In Progress')
        """, (row["id"], visit_type))
        if cur.fetchone():
            continue

        cur.execute("""
            INSERT INTO cue_items
              (job_id, visit_type, due_date, priority, status,
               created_by_user_id, created_at, updated_at)
            VALUES (?, ?, ?, ?, 'Pending', ?, ?, ?)
        """, (row["id"], visit_type, today, priority,
              admin_user_id, now_str, now_str))

    cur.execute("""
        UPDATE jobs SET assigned_user_id = sub.agent_id, updated_at = ?
        FROM (
            SELECT s.job_id, s.assigned_to_user_id AS agent_id
            FROM schedules s
            INNER JOIN (
                SELECT job_id, MIN(scheduled_for) AS min_sf
                FROM schedules
                WHERE status NOT IN ('Cancelled', 'Completed')
                GROUP BY job_id
            ) ms ON ms.job_id = s.job_id AND ms.min_sf = s.scheduled_for
            WHERE s.status NOT IN ('Cancelled', 'Completed')
              AND s.assigned_to_user_id IS NOT NULL
        ) sub
        WHERE jobs.id = sub.job_id
          AND (jobs.assigned_user_id IS NULL OR jobs.assigned_user_id != sub.agent_id)
    """, (now_str,))

@app.get("/dashboard")
@login_required
def dashboard():
    user_id = session.get("user_id")
    role    = session.get("role")
    conn = db()
    cur  = conn.cursor()

    def jcount(where="", params=()):
        q = f"SELECT COUNT(*) AS c FROM jobs{(' WHERE ' + where) if where else ''}"
        cur.execute(q, params)
        return cur.fetchone()["c"]

    STATUS_LIST = [
        ("New",                       "new"),
        ("Active",                    "active"),
        ("Active - Phone work only",  "phone"),
        ("Suspended",                 "suspended"),
        ("Awaiting info from client", "awaiting"),
        ("Completed",                 "completed"),
        ("Invoiced",                  "invoiced"),
        ("Cancelled",                 "cancelled"),
    ]

    def jrows(status):
        agent_subq = """
            COALESCE(
                (SELECT u2.full_name FROM schedules sx
                 JOIN users u2 ON u2.id = sx.assigned_to_user_id
                 WHERE sx.job_id = j.id AND sx.status NOT IN ('Cancelled', 'Completed')
                 ORDER BY sx.scheduled_for ASC LIMIT 1),
                u.full_name
            ) AS assigned_name"""
        sched_subq = """
            (SELECT sx2.scheduled_for FROM schedules sx2
             WHERE sx2.job_id = j.id
               AND date(sx2.scheduled_for) >= date('now','localtime')
             ORDER BY sx2.scheduled_for ASC LIMIT 1) AS next_scheduled"""
        base_sel = f"""
                SELECT j.id, j.display_ref, j.status,
                       COALESCE(CASE WHEN cu.last_name IS NOT NULL THEN cu.last_name || COALESCE(' ' || cu.first_name, '') ELSE NULL END,
                                cu.company, 'No customer') AS customer_name,
                       COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
                       {agent_subq},
                       {sched_subq}
                FROM jobs j
                LEFT JOIN customers cu ON cu.id = j.customer_id
                LEFT JOIN users u ON u.id = j.assigned_user_id"""
        if role == "agent":
            sql = base_sel + """
                WHERE j.status = ? AND (j.assigned_user_id = ? OR EXISTS (
                    SELECT 1 FROM schedules s WHERE s.job_id = j.id
                    AND s.assigned_to_user_id = ? AND s.status NOT IN ('Cancelled')
                ))
                ORDER BY j.updated_at DESC"""
            cur.execute(sql, (status, user_id, user_id))
        else:
            sql = base_sel + """
                WHERE j.status = ?
                ORDER BY j.updated_at DESC"""
            cur.execute(sql, (status,))
        return cur.fetchall()

    _excl_arch = f"status NOT IN {ARCHIVED_STATUSES!r}"

    if role == "agent":
        base = """(assigned_user_id = ? OR EXISTS (
            SELECT 1 FROM schedules s WHERE s.job_id = jobs.id
            AND s.assigned_to_user_id = ? AND s.status NOT IN ('Cancelled')
        ))"""
        uu   = (user_id, user_id)
        jobs_all       = jcount(base + f" AND {_excl_arch}", uu)
        jobs_new       = jcount(base + " AND status = 'New'",                      (*uu,))
        jobs_active    = jcount(base + " AND status = 'Active'",                   (*uu,))
        jobs_phone     = jcount(base + " AND status = 'Active - Phone work only'", (*uu,))
        jobs_suspended = jcount(base + " AND status = 'Suspended'",                (*uu,))
        jobs_awaiting  = jcount(base + " AND status = 'Awaiting info from client'",(*uu,))
        jobs_completed = jcount(base + " AND status = 'Completed'",                (*uu,))
        jobs_invoiced  = jcount(base + " AND status = 'Invoiced'",                 (*uu,))
        jobs_archived  = jcount(base + " AND status = 'Archived - Invoiced'",      (*uu,))
    else:
        jobs_all       = jcount(_excl_arch)
        jobs_new       = jcount("status = 'New'")
        jobs_active    = jcount("status = 'Active'")
        jobs_phone     = jcount("status = 'Active - Phone work only'")
        jobs_suspended = jcount("status = 'Suspended'")
        jobs_awaiting  = jcount("status = 'Awaiting info from client'")
        jobs_completed = jcount("status = 'Completed'")
        jobs_invoiced  = jcount("status = 'Invoiced'")
        jobs_archived  = jcount("status = 'Archived - Invoiced'")

    jobs_unassigned = 0
    if role in ("admin", "both"):
        jobs_unassigned = jcount("assigned_user_id IS NULL AND status NOT IN ('Completed','Invoiced','Cancelled','Archived - Invoiced','Cold Stored')")

    agent_subq = """
        COALESCE(
            (SELECT u2.full_name FROM schedules sx
             JOIN users u2 ON u2.id = sx.assigned_to_user_id
             WHERE sx.job_id = j.id AND sx.status NOT IN ('Cancelled', 'Completed')
             ORDER BY sx.scheduled_for ASC LIMIT 1),
            u.full_name
        ) AS assigned_name"""
    sched_subq = """
        (SELECT sx2.scheduled_for FROM schedules sx2
         WHERE sx2.job_id = j.id
           AND date(sx2.scheduled_for) >= date('now','localtime')
         ORDER BY sx2.scheduled_for ASC LIMIT 1) AS next_scheduled"""
    base_sel = f"""
        SELECT j.id, j.display_ref, j.status,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name, 'No customer') AS customer_label,
               {agent_subq},
               {sched_subq},
               j.updated_at
        FROM jobs j
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN users u ON u.id = j.assigned_user_id"""

    if role == "agent":
        recent_sql = base_sel + f"""
            WHERE {_excl_arch} AND (j.assigned_user_id = ? OR EXISTS (
                SELECT 1 FROM schedules s WHERE s.job_id = j.id
                AND s.assigned_to_user_id = ? AND s.status NOT IN ('Cancelled')
            ))
            ORDER BY j.updated_at DESC LIMIT 15"""
        cur.execute(recent_sql, (user_id, user_id))
    else:
        recent_sql = base_sel + f"""
            WHERE {_excl_arch}
            ORDER BY j.updated_at DESC LIMIT 15"""
        cur.execute(recent_sql)
    recent_activity = cur.fetchall()

    sched_sel = f"""
        SELECT j.id, j.display_ref,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name, 'No customer') AS customer_label,
               {agent_subq},
               s_next.scheduled_for AS next_scheduled
        FROM jobs j
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN users u ON u.id = j.assigned_user_id
        JOIN (
            SELECT job_id, MIN(scheduled_for) AS scheduled_for
            FROM schedules
            WHERE date(scheduled_for) >= date('now','localtime')
            GROUP BY job_id
        ) s_next ON s_next.job_id = j.id
        WHERE {_excl_arch}"""
    if role == "agent":
        sched_sel += """
            AND (j.assigned_user_id = ? OR EXISTS (
                SELECT 1 FROM schedules s WHERE s.job_id = j.id
                AND s.assigned_to_user_id = ? AND s.status NOT IN ('Cancelled')
            ))"""
        sched_sel += " ORDER BY s_next.scheduled_for ASC LIMIT 6"
        cur.execute(sched_sel, (user_id, user_id))
    else:
        sched_sel += " ORDER BY s_next.scheduled_for ASC LIMIT 6"
        cur.execute(sched_sel)
    upcoming_schedules = cur.fetchall()

    _mel_now_d = datetime.now(_melbourne).date()
    completed_today = 0
    if role == "agent":
        cur.execute("""SELECT COUNT(*) AS c FROM jobs
            WHERE status = 'Completed' AND date(updated_at) = ?
            AND (assigned_user_id = ? OR EXISTS (
                SELECT 1 FROM schedules s WHERE s.job_id = jobs.id
                AND s.assigned_to_user_id = ? AND s.status NOT IN ('Cancelled')
            ))""", (_mel_now_d.isoformat(), user_id, user_id))
    else:
        cur.execute("SELECT COUNT(*) AS c FROM jobs WHERE status = 'Completed' AND date(updated_at) = ?",
                    (_mel_now_d.isoformat(),))
    completed_today = cur.fetchone()["c"]

    # Auto-flag overdue / today / tomorrow schedules into the job queue
    _admin_id = user_id if role in ("admin", "both") else None
    if _admin_id:
        auto_queue_schedule_alerts(cur, _admin_id)
        conn.commit()

    conn.close()
    from datetime import timedelta as _td
    _today = datetime.now().date()
    return render_template("index.html",
        jobs_all=jobs_all,
        jobs_new=jobs_new,       jobs_active=jobs_active,
        jobs_phone=jobs_phone,   jobs_suspended=jobs_suspended,
        jobs_awaiting=jobs_awaiting, jobs_completed=jobs_completed,
        jobs_invoiced=jobs_invoiced,
        jobs_archived=jobs_archived,
        recent_activity=recent_activity,
        upcoming_schedules=upcoming_schedules,
        completed_today=completed_today,
        jobs_unassigned=jobs_unassigned,
        today_iso=_today.isoformat(),
        tomorrow_iso=(_today + _td(days=1)).isoformat())


@app.get("/dashboard/jobs")
@login_required
def dashboard_jobs_api():
    category = request.args.get("category", "").strip()
    user_id = session.get("user_id")
    role    = session.get("role")
    conn = db()
    cur  = conn.cursor()

    _mel_now_d = datetime.now(_melbourne).date()
    STATUS_MAP = {
        "active":    "j.status IN ('Active','Active - Phone work only')",
        "suspended": "j.status = 'Suspended'",
        "awaiting":  "j.status = 'Awaiting info from client'",
        "completed": "j.status = 'Completed'",
        "completed_today": f"j.status = 'Completed' AND date(j.updated_at) = '{_mel_now_d.isoformat()}'",
    }
    where = STATUS_MAP.get(category)
    if not where:
        return jsonify(jobs=[], title="Unknown")

    TITLES = {"active": "Total Active", "suspended": "Needs Attention",
              "awaiting": "Awaiting Response", "completed": "Completed",
              "completed_today": "Completed Today"}

    agent_subq = """
        COALESCE(
            (SELECT u2.full_name FROM schedules sx
             JOIN users u2 ON u2.id = sx.assigned_to_user_id
             WHERE sx.job_id = j.id AND sx.status NOT IN ('Cancelled', 'Completed')
             ORDER BY sx.scheduled_for ASC LIMIT 1),
            u.full_name
        ) AS assigned_name"""
    sql = f"""
        SELECT j.id, j.display_ref, j.status,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name, 'No customer') AS customer_label,
               {agent_subq},
               j.updated_at,
               cu.address AS customer_address
        FROM jobs j
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN users u ON u.id = j.assigned_user_id
        WHERE {where}"""
    if role == "agent":
        sql += """ AND (j.assigned_user_id = ? OR EXISTS (
            SELECT 1 FROM schedules s WHERE s.job_id = j.id
            AND s.assigned_to_user_id = ? AND s.status NOT IN ('Cancelled')
        ))"""
        sql += " ORDER BY j.updated_at DESC LIMIT 25"
        cur.execute(sql, (user_id, user_id))
    else:
        sql += " ORDER BY j.updated_at DESC LIMIT 25"
        cur.execute(sql)
    rows = cur.fetchall()
    conn.close()
    jobs = []
    for r in rows:
        jobs.append({
            "id": r["id"],
            "display_ref": r["display_ref"],
            "status": r["status"],
            "customer_label": r["customer_label"],
            "assigned_name": r["assigned_name"] or "",
            "updated_at": r["updated_at"] or "",
            "customer_address": r["customer_address"] or "",
        })
    return jsonify(jobs=jobs, title=TITLES.get(category, category), count=len(jobs))


@app.get("/jobs")
@login_required
def jobs_list():
    lockout = _check_agent_lockout()
    if lockout:
        return lockout
    status       = request.args.get("status", "").strip()
    q            = request.args.get("q", "").strip()
    sort         = request.args.get("sort", "").strip()
    filter_client  = request.args.get("client_id", "").strip()
    filter_agent   = request.args.get("agent_id", "").strip()
    filter_btype   = request.args.get("booking_type_id", "").strip()
    filter_date_from = request.args.get("date_from", "").strip()
    filter_date_to   = request.args.get("date_to", "").strip()

    user_id = session.get("user_id")
    role = session.get("role")

    conn = db()
    cur = conn.cursor()

    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 25, type=int)
    if per_page not in (25, 50, 100):
        per_page = 25
    if page < 1:
        page = 1

    select_cols = """
    SELECT j.*,
           COALESCE(c.nickname, c.name) AS client_name,
           cu.last_name AS customer_last_name,
           COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
           (cu.first_name || ' ' || cu.last_name) AS customer_name,
           (SELECT ji.reg FROM job_items ji WHERE ji.job_id = j.id AND ji.item_type = 'vehicle' LIMIT 1) AS asset_reg,
           COALESCE(
               (SELECT u2.full_name FROM schedules s2
                JOIN users u2 ON u2.id = s2.assigned_to_user_id
                WHERE s2.job_id = j.id AND s2.status NOT IN ('Cancelled', 'Completed')
                ORDER BY s2.scheduled_for ASC LIMIT 1),
               u.full_name
           ) AS assigned_name,
           (SELECT s.scheduled_for FROM schedules s
            WHERE s.job_id = j.id AND s.status NOT IN ('Completed', 'Cancelled')
            ORDER BY s.scheduled_for ASC LIMIT 1) AS next_scheduled,
           (SELECT bt.name FROM schedules s
            JOIN booking_types bt ON bt.id = s.booking_type_id
            WHERE s.job_id = j.id AND s.status NOT IN ('Completed', 'Cancelled')
            ORDER BY s.scheduled_for ASC LIMIT 1) AS next_booking_type,
           (SELECT s.id FROM schedules s
            WHERE s.job_id = j.id AND s.status NOT IN ('Completed', 'Cancelled')
            ORDER BY s.scheduled_for ASC LIMIT 1) AS next_sched_id
    """

    from_where = """
    FROM jobs j
    LEFT JOIN clients c ON c.id = j.client_id
    LEFT JOIN customers cu ON cu.id = j.customer_id
    LEFT JOIN users u ON u.id = j.assigned_user_id
    WHERE 1=1
    """
    params = []

    include_archived = request.args.get("include_archived", "").strip()
    if status not in ARCHIVED_STATUSES and not include_archived:
        from_where += f" AND j.status NOT IN {ARCHIVED_STATUSES!r}"

    if role == "agent":
        from_where += """ AND (j.assigned_user_id = ? OR EXISTS (
            SELECT 1 FROM schedules s
            WHERE s.job_id = j.id AND s.assigned_to_user_id = ?
              AND s.status NOT IN ('Cancelled')
        ))"""
        params.extend([user_id, user_id])

    if status:
        if status == "Active":
            from_where += " AND j.status LIKE 'Active%'"
        else:
            from_where += " AND j.status = ?"
            params.append(status)

    if q:
        from_where += """
         AND (
           j.internal_job_number LIKE ? OR
           j.client_reference     LIKE ? OR
           j.display_ref          LIKE ? OR
           j.description          LIKE ? OR
           j.job_address          LIKE ? OR
           cu.first_name          LIKE ? OR
           cu.last_name           LIKE ? OR
           cu.company             LIKE ? OR
           c.name                 LIKE ? OR
           EXISTS (
             SELECT 1 FROM job_items ji
             WHERE ji.job_id = j.id
               AND (ji.reg LIKE ? OR ji.vin LIKE ? OR ji.description LIKE ?)
           )
         )"""
        like = f"%{q}%"
        params.extend([like] * 12)

    if filter_client:
        from_where += " AND j.client_id = ?"
        params.append(filter_client)

    filter_unassigned = request.args.get("unassigned", "").strip()
    if filter_unassigned == "1":
        from_where += " AND j.assigned_user_id IS NULL AND j.status NOT IN ('Completed','Invoiced','Cancelled','Archived - Invoiced','Cold Stored')"

    if filter_agent and role in ("admin", "both"):
        from_where += " AND (j.assigned_user_id = ? OR EXISTS (SELECT 1 FROM schedules sa WHERE sa.job_id = j.id AND sa.assigned_to_user_id = ? AND sa.status NOT IN ('Cancelled')))"
        params.extend([filter_agent, filter_agent])

    if filter_btype:
        from_where += " AND EXISTS (SELECT 1 FROM schedules sb JOIN booking_types btf ON btf.id = sb.booking_type_id WHERE sb.job_id = j.id AND sb.booking_type_id = ? AND sb.status NOT IN ('Completed', 'Cancelled'))"
        params.append(filter_btype)

    if filter_date_from:
        from_where += " AND date(j.updated_at) >= ?"
        params.append(filter_date_from)

    if filter_date_to:
        from_where += " AND date(j.updated_at) <= ?"
        params.append(filter_date_to)

    count_sql = "SELECT COUNT(*) " + from_where
    total_jobs = cur.execute(count_sql, params).fetchone()[0]
    total_pages = max(1, (total_jobs + per_page - 1) // per_page)
    if page > total_pages:
        page = total_pages

    _SORT_MAP = {
        "scheduled":  "CASE WHEN j.status='Invoiced' THEN 1 ELSE 0 END, CASE WHEN next_scheduled IS NULL THEN 1 ELSE 0 END, next_scheduled ASC, j.updated_at DESC",
        "agent":      "LOWER(COALESCE(u.full_name,'')) ASC, j.updated_at DESC",
        "active":     "CASE WHEN j.status LIKE 'Active%' THEN 0 ELSE 1 END ASC, j.updated_at DESC",
        "job_number": "CAST(j.internal_job_number AS INTEGER) DESC, j.updated_at DESC",
        "client_ref": "CAST(j.client_reference AS INTEGER) DESC, j.updated_at DESC",
    }
    if not sort:
        sort = "scheduled"
    order_clause = _SORT_MAP.get(sort, _SORT_MAP["scheduled"])
    offset = (page - 1) * per_page
    full_sql = select_cols + from_where + f" ORDER BY {order_clause} LIMIT ? OFFSET ?"
    cur.execute(full_sql, params + [per_page, offset])
    rows = cur.fetchall()

    clients_list = cur.execute(
        "SELECT id, COALESCE(nickname, name) AS name FROM clients ORDER BY name"
    ).fetchall()
    agents_list = cur.execute(
        "SELECT id, full_name FROM users WHERE role IN ('admin','agent','both') ORDER BY full_name"
    ).fetchall()
    btypes_list = cur.execute(
        "SELECT id, name FROM booking_types ORDER BY name"
    ).fetchall()

    conn.close()

    statuses = [
        "New",
        "Active",
        "Active - Phone work only",
        "Suspended",
        "Awaiting info from client",
        "Completed",
        "Invoiced",
        "Cancelled",
        "Archived - Invoiced",
    ]

    return render_template(
        "jobs.html",
        jobs=rows,
        statuses=statuses,
        status=status,
        q=q,
        sort=sort,
        clients_list=clients_list,
        agents_list=agents_list,
        btypes_list=btypes_list,
        filter_client=filter_client,
        filter_agent=filter_agent,
        filter_btype=filter_btype,
        filter_date_from=filter_date_from,
        filter_date_to=filter_date_to,
        page=page,
        per_page=per_page,
        total_jobs=total_jobs,
        total_pages=total_pages,
        filter_unassigned=filter_unassigned,
    )


@app.get("/api/jobs/search")
@login_required
def api_jobs_search():
    q = (request.args.get("q") or "").strip()
    if len(q) < 2:
        return jsonify([])
    like = f"%{q}%"
    conn = db()
    rows = conn.execute("""
        SELECT j.id, j.display_ref,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label
        FROM jobs j
        LEFT JOIN customers cu ON cu.id = j.customer_id
        WHERE j.display_ref LIKE ?
           OR j.client_reference LIKE ?
           OR j.client_job_number LIKE ?
           OR cu.last_name LIKE ?
           OR cu.company LIKE ?
        ORDER BY j.created_at DESC
        LIMIT 12
    """, (like, like, like, like, like)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.get("/jobs/search-reference")
@admin_required
def jobs_search_reference():
    q = request.args.get("q", "").strip()
    if len(q) < 2:
        return jsonify([])
    like = f"%{q}%"
    conn = db()
    rows = conn.execute("""
        SELECT j.id, j.display_ref, j.client_reference, j.client_job_number,
               j.status, j.job_type,
               (cu.first_name || ' ' || cu.last_name) AS customer_name,
               COALESCE(c.nickname, c.name) AS client_name
        FROM jobs j
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN clients c ON c.id = j.client_id
        WHERE j.client_reference LIKE ? OR j.client_job_number LIKE ?
        ORDER BY j.created_at DESC
        LIMIT 10
    """, (like, like)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.get("/jobs/<int:job_id>/clone-data")
@admin_required
def job_clone_data(job_id: int):
    conn = db()
    job = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({"error": "Not found"}), 404
    customer = None
    if job["customer_id"]:
        customer = conn.execute("SELECT * FROM customers WHERE id = ?", (job["customer_id"],)).fetchone()
    client = None
    if job["client_id"]:
        client = conn.execute("SELECT id, name FROM clients WHERE id = ?", (job["client_id"],)).fetchone()
    assets = conn.execute(
        "SELECT * FROM job_items WHERE job_id = ? ORDER BY id", (job_id,)
    ).fetchall()
    conn.close()

    def cents_to_str(c):
        return f"{c/100:.2f}" if c else ""

    return jsonify({
        "client_id":        job["client_id"],
        "client_name":      client["name"] if client else "",
        "customer_id":      job["customer_id"],
        "customer_name":    (customer["first_name"] + " " + customer["last_name"]) if customer else "",
        "customer_address": customer["address"] if customer else "",
        "client_reference": job["client_reference"] or "",
        "client_job_number": job["client_job_number"] or "",
        "job_type":         job["job_type"],
        "visit_type":       job["visit_type"],
        "status":           job["status"],
        "priority":         job["priority"],
        "job_address":      job["job_address"] or "",
        "description":      job["description"] or "",
        "lender_name":      job["lender_name"] or "",
        "account_number":   job["account_number"] or "",
        "regulation_type":  job["regulation_type"] or "",
        "arrears":          cents_to_str(job["arrears_cents"]),
        "costs":            cents_to_str(job["costs_cents"]),
        "mmp":              cents_to_str(job["mmp_cents"]),
        "job_due_date":     job["job_due_date"] or "",
        "display_ref":      job["display_ref"],
        "assets": [
            {
                "item_type":        a["item_type"],
                "description":      a["description"] or "",
                "reg":              a["reg"] or "",
                "vin":              a["vin"] or "",
                "make":             a["make"] or "",
                "model":            a["model"] or "",
                "year":             a["year"] or "",
                "property_address": a["property_address"] or "",
                "serial_number":    a["serial_number"] or "",
                "notes":            a["notes"] or "",
            }
            for a in assets
        ]
    })


@app.post("/jobs/new/autofill-upload")
@login_required
@admin_required
def job_new_autofill_upload():
    f = request.files.get("instruction_file")
    if not f or not f.filename:
        flash("No file selected.", "warning")
        return redirect(url_for("job_new"))

    ext = os.path.splitext(secure_filename(f.filename))[1].lower()
    if ext not in (".docx", ".pdf", ".doc"):
        flash("Auto-fill only supports Word (.doc, .docx) and PDF (.pdf) files.", "warning")
        return redirect(url_for("job_new"))

    try:
        original, mimetype, storage_key, path = _save_pending_upload(f)
    except Exception as e:
        flash(f"Could not save file: {e}", "danger")
        return redirect(url_for("job_new"))

    extracted_text = None
    try:
        if ext == ".docx":
            extracted_text = _extract_text_docx(path)
        elif ext == ".pdf":
            extracted_text = _extract_text_pdf(path)
        elif ext == ".doc":
            extracted_text = _extract_text_doc(path)
    except Exception as e:
        flash(f"Could not read document: {e}", "danger")
        try:
            os.remove(path)
        except Exception:
            pass
        return redirect(url_for("job_new"))

    extracted = _parse_instruction_text(extracted_text or "")

    now = now_ts()
    conn = db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO pending_uploads (original_filename, storage_key, content_type, uploaded_by_user_id, uploaded_at)
        VALUES (?, ?, ?, ?, ?)
    """, (original, storage_key, mimetype, session.get("user_id"), now))
    pending_id = cur.lastrowid

    cur.execute("""
        INSERT INTO document_extractions (pending_upload_id, status, provider_used, extracted_json, extracted_text, created_at)
        VALUES (?, 'success', 'rule_based', ?, ?, ?)
    """, (pending_id, json.dumps(extracted), extracted_text, now))
    extraction_id = cur.lastrowid

    conn.commit()
    conn.close()

    return redirect(url_for("job_new", autofill_id=extraction_id))


@app.get("/jobs/new")
@login_required
@admin_required
def job_new():
    conn = db()
    try:
        return _job_new_render(conn)
    except Exception:
        import traceback, sys
        traceback.print_exc(file=sys.stderr)
        traceback.print_exc()
        return render_template("error_500.html",
                               error_message="Unable to load the New Job form. Please contact support.",
                               path=request.path), 500
    finally:
        conn.close()


def _job_new_render(conn):
    cur = conn.cursor()
    cur.execute("SELECT id, name FROM clients ORDER BY name")
    clients = cur.fetchall()
    cur.execute("SELECT id, first_name, last_name, company, address FROM customers ORDER BY last_name, first_name")
    customers = cur.fetchall()
    cur.execute("SELECT id, full_name FROM users WHERE active = 1 ORDER BY full_name")
    users = cur.fetchall()
    cur.execute("SELECT * FROM system_settings WHERE id = 1")
    settings = cur.fetchone()
    if not settings:
        settings = {"job_prefix": datetime.now().strftime("%y%m"), "job_sequence": 0}

    new_client_id   = request.args.get("new_client_id",   type=int)
    new_customer_id = request.args.get("new_customer_id", type=int)
    new_user_id     = request.args.get("new_user_id",     type=int)
    autofill_id     = request.args.get("autofill_id",     type=int)

    autofill        = None
    autofill_notice = None
    autofill_filename = None
    if autofill_id:
        row = cur.execute("""
            SELECT de.extracted_json, pu.original_filename
            FROM document_extractions de
            JOIN pending_uploads pu ON pu.id = de.pending_upload_id
            WHERE de.id = ?
        """, (autofill_id,)).fetchone()
        if row:
            try:
                autofill = json.loads(row["extracted_json"])
                autofill_filename = row["original_filename"]
                autofill_notice = f"Auto-filled from \"{autofill_filename}\" — review and adjust fields before saving."
            except Exception:
                pass

    prefill_customer_address  = ""
    prefill_client_reference  = request.args.get("client_reference", "")
    prefill_lender_name       = request.args.get("lender_name", "")
    prefill_account_number    = request.args.get("account_number", "")
    prefill_deliver_to        = ""
    prefill_costs             = ""
    prefill_client_job_number = ""
    autofill_client_id        = None

    autofill_customer_id      = None
    autofill_customer_display = None
    autofill_customer_is_new  = False
    autofill_confidence       = {}
    autofill_filled_count     = 0

    if autofill:
        if not prefill_client_reference:
            prefill_client_reference = autofill.get("client_reference") or autofill.get("contract_number") or ""
        if not prefill_lender_name:
            prefill_lender_name = autofill.get("lender_name") or ""
        if not prefill_account_number:
            prefill_account_number = autofill.get("account_number") or ""
        if not prefill_customer_address:
            prefill_customer_address = autofill.get("job_address_full") or ""
        if not prefill_deliver_to:
            prefill_deliver_to = autofill.get("deliver_to") or ""
        if not prefill_costs:
            prefill_costs = autofill.get("costs") or ""
        if not prefill_client_job_number:
            prefill_client_job_number = autofill.get("our_ref") or ""

        autofill_confidence   = autofill.get("_confidence") or {}
        autofill_filled_count = autofill.get("_filled_count") or 0

        # ── Client lookup: try lender_name and from_name, multi-word matching ──
        lender_for_lookup = (autofill.get("lender_name") or autofill.get("from_name") or "").strip()
        if lender_for_lookup and not new_client_id:
            words = [w for w in lender_for_lookup.lower().split() if len(w) > 2]
            if words:
                # Try progressively shorter word sets for matching
                client_match = None
                for word in words[:4]:
                    candidate = cur.execute(
                        "SELECT id, name FROM clients WHERE LOWER(name) LIKE ? ORDER BY name LIMIT 1",
                        (f"%{word}%",)
                    ).fetchone()
                    if candidate:
                        match_lower = candidate["name"].lower()
                        if any(w in match_lower for w in words[:3]):
                            client_match = candidate
                            break
                if client_match:
                    autofill_client_id = client_match["id"]
                    autofill_confidence["lender_name"] = "matched"

        # ── Customer lookup: search by name or company ──────────────────────
        autofill_customer_is_new = False
        if not new_customer_id:
            cust_data = autofill.get("customer") or {}
            cust_full    = (cust_data.get("full_name") or "").strip()
            cust_company = (cust_data.get("company") or "").strip()
            search_name  = cust_full or cust_company
            autofill_customer_display = search_name or None

            if search_name:
                parts = search_name.lower().split()
                cust_match = None
                if len(parts) >= 2:
                    cust_match = cur.execute("""
                        SELECT id, first_name, last_name, company FROM customers
                        WHERE LOWER(last_name) = ? OR LOWER(company) LIKE ?
                        ORDER BY last_name LIMIT 1
                    """, (parts[-1], f"%{parts[-1]}%")).fetchone()
                if not cust_match and parts:
                    cust_match = cur.execute("""
                        SELECT id, first_name, last_name, company FROM customers
                        WHERE LOWER(last_name) LIKE ? OR LOWER(company) LIKE ? OR LOWER(first_name) LIKE ?
                        ORDER BY last_name LIMIT 1
                    """, (f"%{parts[0]}%", f"%{parts[0]}%", f"%{parts[0]}%")).fetchone()
                if cust_match:
                    autofill_customer_id = cust_match["id"]
                    disp_parts = [cust_match["first_name"] or "", cust_match["last_name"] or ""]
                    if cust_match["company"]:
                        disp_parts.append(f"({cust_match['company']})")
                    autofill_customer_display = " ".join(p for p in disp_parts if p)
                    autofill_confidence["customer_name"] = "matched"
                else:
                    try:
                        _af_now = now_ts()
                        _af_first = ""
                        _af_last  = ""
                        _af_comp  = ""
                        _COMPANY_KW = ("pty", "ltd", "limited", "inc", "llc", "trust", "trustee", "group", "corp")
                        if any(kw in search_name.lower() for kw in _COMPANY_KW):
                            _af_comp = search_name.strip()
                            if cust_full and cust_full != cust_company:
                                name_parts = cust_full.split()
                                _af_first = name_parts[0].title() if name_parts else ""
                                _af_last  = " ".join(name_parts[1:]).title() if len(name_parts) > 1 else ""
                        else:
                            name_parts = search_name.split()
                            _af_first = name_parts[0].title() if name_parts else ""
                            _af_last  = " ".join(name_parts[1:]).title() if len(name_parts) > 1 else ""
                        _af_addr = (autofill.get("job_address_full") or "").strip()
                        _af_dob  = (cust_data.get("dob") or "").strip()
                        cur.execute("""
                            INSERT INTO customers (first_name, last_name, company, address, dob, created_at, updated_at)
                            VALUES (?, ?, ?, ?, ?, ?, ?)
                        """, (_af_first, _af_last, _af_comp, _af_addr, _af_dob, _af_now, _af_now))
                        _af_cust_id = cur.lastrowid
                        _af_phone = (cust_data.get("mobile") or "").strip()
                        if _af_phone:
                            cur.execute("""
                                INSERT INTO contact_phone_numbers (entity_type, entity_id, label, phone_number, created_at)
                                VALUES ('customer', ?, 'Mobile', ?, ?)
                            """, (_af_cust_id, _af_phone, _af_now))
                        _af_email = (cust_data.get("email") or "").strip()
                        if _af_email:
                            cur.execute("""
                                INSERT INTO contact_emails (entity_type, entity_id, label, email, created_at)
                                VALUES ('customer', ?, 'Primary', ?, ?)
                            """, (_af_cust_id, _af_email, _af_now))
                        conn.commit()
                        autofill_customer_id = _af_cust_id
                        _af_disp = " ".join(p for p in [_af_first, _af_last] if p)
                        if _af_comp:
                            _af_disp = f"{_af_disp} ({_af_comp})" if _af_disp else _af_comp
                        autofill_customer_display = _af_disp
                        autofill_confidence["customer_name"] = "created"
                        autofill_customer_is_new = True
                    except Exception:
                        try:
                            conn.rollback()
                        except Exception:
                            pass
                        import logging as _aclog
                        _aclog.exception("Auto-create customer from autofill failed")

    if new_customer_id:
        cur.execute("SELECT address FROM customers WHERE id = ?", (new_customer_id,))
        row = cur.fetchone()
        if row:
            prefill_customer_address = row["address"] or ""

    try:
        cur.execute("SELECT id, name FROM job_types WHERE active = 1 ORDER BY name")
        job_types = cur.fetchall()
    except Exception:
        job_types = []
    try:
        cur.execute("SELECT DISTINCT lender_name FROM jobs WHERE lender_name IS NOT NULL ORDER BY lender_name")
        known_lenders = [r["lender_name"] for r in cur.fetchall()]
    except Exception:
        known_lenders = []
    try:
        cur.execute("SELECT * FROM booking_types WHERE active = 1 ORDER BY name")
        booking_types = cur.fetchall()
    except Exception:
        booking_types = []
    try:
        cur.execute("SELECT id, name FROM auction_yards WHERE active = 1 ORDER BY name")
        auction_yards = cur.fetchall()
    except Exception:
        auction_yards = []

    next_number = f"{settings['job_prefix']}{str(settings['job_sequence'] + 1).zfill(3)}"

    visit_types = ["New Visit", "Re-attend", "First Update", "Urgent Update", "Phone Follow-up", "Locate Only"]
    statuses = ["New", "Active", "Active - Phone work only", "Suspended", "Awaiting info from client", "Completed", "Invoiced", "Cancelled"]
    priorities = ["Low", "Normal", "High", "Urgent"]

    return render_template("job_new.html", clients=clients, customers=customers,
                           users=users, visit_types=visit_types, job_types=job_types,
                           statuses=statuses, priorities=priorities,
                           next_number=next_number,
                           new_client_id=new_client_id,
                           new_customer_id=new_customer_id,
                           new_user_id=new_user_id,
                           prefill_customer_address=prefill_customer_address,
                           prefill_client_reference=prefill_client_reference,
                           prefill_lender_name=prefill_lender_name,
                           prefill_account_number=prefill_account_number,
                           prefill_deliver_to=prefill_deliver_to,
                           prefill_costs=prefill_costs,
                           prefill_client_job_number=prefill_client_job_number,
                           autofill_client_id=autofill_client_id,
                           autofill_customer_id=autofill_customer_id,
                           autofill_customer_display=autofill_customer_display,
                           autofill_customer_is_new=autofill_customer_is_new,
                           autofill_confidence=autofill_confidence,
                           autofill_filled_count=autofill_filled_count,
                           known_lenders=known_lenders,
                           booking_types=booking_types,
                           auction_yards=auction_yards,
                           autofill=autofill,
                           autofill_id=autofill_id,
                           autofill_notice=autofill_notice,
                           autofill_filename=autofill_filename)


@app.post("/jobs/new")
@login_required
@admin_required
def job_create():
    internal_job_number = generate_internal_job_number()
    client_reference = request.form.get("client_reference", "").strip()
    client_job_number = request.form.get("client_job_number", "").strip() or None
    client_id = request.form.get("client_id") or None
    customer_id = request.form.get("customer_id") or None
    bill_to_client_id = request.form.get("bill_to_client_id") or None
    assigned_user_id = request.form.get("assigned_user_id") or None
    job_type = request.form.get("job_type", "Field Call").strip()
    visit_type = request.form.get("visit_type", "New Visit").strip()
    status = request.form.get("status", "New").strip()
    priority = request.form.get("priority", "Normal").strip()
    job_address = request.form.get("job_address", "").strip()
    description = request.form.get("description", "").strip()
    deliver_to  = request.form.get("deliver_to", "").strip() or None
    lender_name = request.form.get("lender_name", "").strip()
    account_number = request.form.get("account_number", "").strip() or client_reference or None
    regulation_type = request.form.get("regulation_type", "").strip()
    arrears_cents = money_to_cents(request.form.get("arrears"))
    costs_cents = money_to_cents(request.form.get("costs"))
    mmp_cents = money_to_cents(request.form.get("mmp"))
    job_due_date = request.form.get("job_due_date", "").strip() or None

    display_ref = internal_job_number
    if client_reference:
        display_ref = f"{internal_job_number} ({client_reference})"

    now = datetime.now().isoformat(timespec="seconds")

    conn = db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO jobs (
            internal_job_number, client_reference, client_job_number, display_ref,
            client_id, customer_id, bill_to_client_id, assigned_user_id,
            job_type, visit_type, status, priority,
            job_address, description, deliver_to,
            lender_name, account_number, regulation_type,
            arrears_cents, costs_cents, mmp_cents, job_due_date,
            created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        internal_job_number, client_reference or None, client_job_number, display_ref,
        client_id, customer_id, bill_to_client_id, assigned_user_id,
        job_type, visit_type, status, priority,
        job_address, description, deliver_to,
        lender_name or None, account_number or None, regulation_type or None,
        arrears_cents or None, costs_cents or None, mmp_cents or None, job_due_date,
        now, now
    ))
    job_id = cur.lastrowid

    if customer_id:
        cur.execute("""
            INSERT OR IGNORE INTO job_customers (job_id, customer_id, role, sort_order, created_at)
            VALUES (?, ?, 'Primary', 0, ?)
        """, (job_id, customer_id, now))

    cur.execute("""
        INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
        VALUES (?, ?, ?, ?, ?)
    """, (job_id, "System", f"Job created: {display_ref}. Status '{status}'. Visit type '{visit_type}'.", now, now))

    asset_types   = request.form.getlist("asset_type[]")
    descs         = request.form.getlist("asset_description[]")
    regos         = request.form.getlist("asset_rego[]")
    vins          = request.form.getlist("asset_vin[]")
    years         = request.form.getlist("asset_year[]")
    makes         = request.form.getlist("asset_make[]")
    models        = request.form.getlist("asset_model[]")
    colours       = request.form.getlist("asset_colour[]")
    engines       = request.form.getlist("asset_engine[]")
    addresses     = request.form.getlist("asset_address[]")
    serials       = request.form.getlist("asset_serial[]")
    asset_notes   = request.form.getlist("asset_notes[]")

    def _al(lst, i):
        return (lst[i] if i < len(lst) else "") or ""

    for i in range(len(asset_types)):
        a_type   = _al(asset_types, i).strip()
        a_desc   = _al(descs,       i).strip()
        a_rego   = _al(regos,       i).strip()
        a_vin    = _al(vins,        i).strip()
        a_year   = _al(years,       i).strip()
        a_make   = _al(makes,       i).strip()
        a_model  = _al(models,      i).strip()
        a_colour = _al(colours,     i).strip()
        a_engine = _al(engines,     i).strip()
        a_addr   = _al(addresses,   i).strip()
        a_ser    = _al(serials,     i).strip()
        a_note   = _al(asset_notes, i).strip()
        if not any([a_desc, a_rego, a_vin, a_addr, a_ser, a_note, a_make, a_model, a_year, a_engine, a_colour]):
            continue
        cur.execute("""
            INSERT INTO job_items
            (job_id, item_type, description, reg, vin, make, model, year, colour, engine_number, property_address, serial_number, notes, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (job_id, (a_type or "other").lower(),
              a_desc or None, a_rego or None, a_vin or None, a_make or None, a_model or None, a_year or None,
              a_colour or None, a_engine or None, a_addr or None, a_ser or None, a_note or None, now))

    sched_date = request.form.get("sched_date", "").strip()
    sched_time = request.form.get("sched_time", "").strip()
    sched_bt_id   = request.form.get("sched_booking_type_id", "").strip()
    sched_bt_name = request.form.get("sched_booking_type_name", "").strip()
    sched_user_id = request.form.get("sched_assigned_user_id", "").strip() or None
    sched_notes   = request.form.get("sched_notes", "").strip() or None

    if sched_date and (sched_bt_id or sched_bt_name):
        sched_dt = f"{sched_date}T{sched_time or '09:00'}:00"
        resolved_sched_bt = _resolve_booking_type(cur, sched_bt_id, sched_bt_name)
        if resolved_sched_bt:
            sched_assigned_int = int(sched_user_id) if sched_user_id else None
            cur.execute("""
                INSERT INTO schedules
                (job_id, booking_type_id, scheduled_for, status, notes, assigned_to_user_id, created_at, updated_at)
                VALUES (?, ?, ?, 'Scheduled', ?, ?, ?, ?)
            """, (job_id, resolved_sched_bt, sched_dt, sched_notes,
                  sched_assigned_int, now, now))
            _write_schedule_history(cur, cur.lastrowid, job_id, "created",
                                    new_scheduled_for=sched_dt, new_status="Scheduled",
                                    changed_by_user_id=session.get("user_id"))
            if sched_assigned_int:
                cur.execute("UPDATE jobs SET assigned_user_id = ?, updated_at = ? WHERE id = ?",
                            (sched_assigned_int, now, job_id))

    # Link autofill document to the new job
    autofill_id = request.form.get("autofill_id", "").strip()
    _had_autofill_doc = False
    if autofill_id and autofill_id.isdigit():
        pu_row = cur.execute("""
            SELECT pu.id, pu.original_filename, pu.storage_key, pu.content_type, pu.uploaded_by_user_id
            FROM document_extractions de
            JOIN pending_uploads pu ON pu.id = de.pending_upload_id
            WHERE de.id = ?
        """, (int(autofill_id),)).fetchone()
        if pu_row:
            import logging as _logging
            _af_log = _logging.getLogger(__name__)
            pending_path = os.path.join(PENDING_UPLOAD_DIR, pu_row["storage_key"])
            stored_name  = f"{job_id}_autofill_{pu_row['original_filename']}"
            note_copy_name = f"{job_id}_note_instruction_{pu_row['original_filename']}"
            try:
                cur.execute("""
                    INSERT INTO job_documents
                    (job_id, doc_type, title, original_filename, stored_filename, mime_type, uploaded_by_user_id, uploaded_at)
                    VALUES (?, 'Instruction', 'Auto-fill source document', ?, ?, ?, ?, ?)
                """, (job_id, pu_row["original_filename"], stored_name,
                      pu_row["content_type"], pu_row["uploaded_by_user_id"] or session.get("user_id"), now))
                if _uploads_container:
                    with open(pending_path, "rb") as fh:
                        _uploads_container.upload_blob(
                            name=stored_name, data=fh, overwrite=True,
                            content_settings=ContentSettings(content_type=pu_row["content_type"] or "application/octet-stream")
                        )
                    with open(pending_path, "rb") as fh:
                        _uploads_container.upload_blob(
                            name=note_copy_name, data=fh, overwrite=True,
                            content_settings=ContentSettings(content_type=pu_row["content_type"] or "application/octet-stream")
                        )
                else:
                    import shutil
                    shutil.copy2(pending_path, os.path.join(UPLOAD_FOLDER, stored_name))
                    shutil.copy2(pending_path, os.path.join(UPLOAD_FOLDER, note_copy_name))
                _autofill_uid = pu_row["uploaded_by_user_id"] or session.get("user_id")
                cur.execute("""
                    INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at)
                    VALUES (?, ?, ?, ?)
                """, (job_id, _autofill_uid,
                      f"Instruction document uploaded: {pu_row['original_filename']}", now))
                _af_note_id = cur.lastrowid
                cur.execute("""
                    INSERT INTO job_note_files (job_field_note_id, filename, filepath, uploaded_at)
                    VALUES (?, ?, ?, ?)
                """, (_af_note_id, note_copy_name, note_copy_name, now))
                _had_autofill_doc = True
            except Exception as _af_exc:
                _af_log.warning("Failed to save autofill document for job %s: %s", job_id, _af_exc)
            finally:
                try:
                    os.remove(pending_path)
                except Exception:
                    pass
            cur.execute("DELETE FROM pending_uploads WHERE id = ?", (pu_row["id"],))
            cur.execute("DELETE FROM document_extractions WHERE id = ?", (int(autofill_id),))

    conn.commit()
    conn.close()

    if job_address:
        _geocode_job_async(job_id, job_address)

    flash("Job created.", "success")
    if request.form.get("add_another"):
        params = {}
        if lender_name:    params["lender_name"]    = lender_name
        if account_number: params["account_number"] = account_number
        return redirect(url_for("job_new", **params))
    if _had_autofill_doc:
        return redirect(url_for("job_detail", job_id=job_id) + "?add_note=1#tab-notes")
    _redir_params = "?focus=lender"
    _final_assigned = False
    try:
        _chk = db()
        _fa = _chk.execute("SELECT assigned_user_id FROM jobs WHERE id = ?", (job_id,)).fetchone()
        _final_assigned = bool(_fa and _fa["assigned_user_id"])
        _chk.close()
    except Exception:
        pass
    if job_address and not _final_assigned:
        _redir_params += "&agent_suggest=1"
    return redirect(url_for("job_detail", job_id=job_id) + _redir_params)


@app.get("/jobs/<int:job_id>")
@login_required
def job_detail(job_id: int):
    lockout = _check_agent_lockout()
    if lockout:
        return lockout
    conn = db()
    cur = conn.cursor()

    cur.execute("""
    SELECT j.*,
           c.name AS client_name, c.nickname AS client_nickname, c.phone AS client_phone, c.email AS client_email, c.address AS client_address,
           (cu.first_name || ' ' || cu.last_name) AS customer_name, cu.last_name AS customer_last_name, cu.company AS customer_company, cu.email AS customer_email, cu.dob AS customer_dob, cu.address AS customer_address,
           u.full_name AS assigned_name, u.email AS assigned_email,
           btc.name AS bill_to_client_name
    FROM jobs j
    LEFT JOIN clients c ON c.id = j.client_id
    LEFT JOIN customers cu ON cu.id = j.customer_id
    LEFT JOIN users u ON u.id = j.assigned_user_id
    LEFT JOIN clients btc ON btc.id = j.bill_to_client_id
    WHERE j.id = ?
    """, (job_id,))
    job = cur.fetchone()
    if not job:
        conn.close()
        return ("Not found", 404)

    job = dict(job)
    total_dollars, include_mmp = calc_total_due_now(
        (job.get("arrears_cents") or 0) / 100,
        (job.get("costs_cents") or 0) / 100,
        (job.get("mmp_cents") or 0) / 100,
        job.get("job_due_date"),
        (job.get("costs2_cents") or 0) / 100,
    )
    _adv_due = advance_due_date_display(job.get("job_due_date"), job.get("payment_frequency"))
    job["due_date_display"] = format_ddmmyyyy(_adv_due)
    job["total_due_now_cents"] = int(round(total_dollars * 100))
    job["mmp_included_in_total"] = include_mmp

    # Payments received for this job
    cur.execute(
        "SELECT jp.*, u.full_name AS recorded_by_name FROM job_payments jp "
        "LEFT JOIN users u ON u.id = jp.recorded_by_user_id "
        "WHERE jp.job_id = ? ORDER BY jp.payment_date DESC, jp.id DESC",
        (job_id,)
    )
    job_payments = cur.fetchall()
    payments_total_cents = sum(p["amount_cents"] for p in job_payments)
    job["net_due_cents"] = max(0, job["total_due_now_cents"] - payments_total_cents)

    role = session.get("role")
    user_id = session.get("user_id")
    if role == "agent" and job["assigned_user_id"] != user_id:
        sched_check = conn.execute(
            """SELECT 1 FROM schedules WHERE job_id = ? AND assigned_to_user_id = ?
               AND status NOT IN ('Cancelled') LIMIT 1""",
            (job_id, user_id)
        ).fetchone()
        if not sched_check:
            conn.close()
            flash("You do not have access to that job.", "danger")
            return redirect(url_for("jobs_list"))

    cur.execute("""
    SELECT * FROM interactions
    WHERE job_id = ?
    ORDER BY occurred_at DESC, id DESC
    """, (job_id,))
    interactions = cur.fetchall()

    cur.execute("""
        SELECT phone_number, label FROM contact_phone_numbers
        WHERE entity_type = 'customer' AND entity_id = ? ORDER BY id LIMIT 1
    """, (job.get("customer_id"),))
    cphone_row = cur.fetchone()
    job["customer_phone"] = cphone_row["phone_number"] if cphone_row else None
    job["customer_phone_label"] = cphone_row["label"] if cphone_row else None

    cur.execute("SELECT * FROM job_items WHERE job_id = ? ORDER BY id", (job_id,))
    job_items = cur.fetchall()

    cur.execute("SELECT id, full_name FROM users WHERE active = 1 ORDER BY full_name")
    users = cur.fetchall()

    cur.execute("SELECT id, name, nickname FROM clients ORDER BY name")
    all_clients = cur.fetchall()

    cur.execute("SELECT id, first_name, last_name, company, address FROM customers ORDER BY last_name, first_name")
    all_customers_for_edit = cur.fetchall()

    cur.execute("""
        SELECT fn.*, u.full_name author_name
        FROM job_field_notes fn
        LEFT JOIN users u ON u.id = fn.created_by_user_id
        WHERE fn.job_id = ?
        ORDER BY fn.created_at DESC
    """, (job_id,))
    raw_notes = cur.fetchall()

    field_notes = []
    for note in raw_notes:
        cur.execute("SELECT * FROM job_note_files WHERE job_field_note_id = ?", (note["id"],))
        files = cur.fetchall()
        field_notes.append({"note": note, "files": files})

    cur.execute("SELECT * FROM booking_types WHERE active = 1 ORDER BY name")
    booking_types = cur.fetchall()

    cur.execute("""
        SELECT s.*, bt.name booking_type_name, u.full_name assigned_to_name
        FROM schedules s
        JOIN booking_types bt ON bt.id = s.booking_type_id
        LEFT JOIN users u ON u.id = s.assigned_to_user_id
        WHERE s.job_id = ?
        ORDER BY s.scheduled_for ASC
    """, (job_id,))
    schedules = cur.fetchall()
    next_schedule = next((s for s in schedules if s["status"] not in ("Cancelled", "Completed")), None)

    cur.execute("""
        SELECT d.*, u.full_name AS uploaded_by
        FROM job_documents d
        LEFT JOIN users u ON u.id = d.uploaded_by_user_id
        WHERE d.job_id = ?
        ORDER BY d.id DESC
    """, (job_id,))
    documents = cur.fetchall()

    cur.execute("""
        SELECT jc.id AS jc_id, jc.role, jc.sort_order,
               cu.id AS customer_id, cu.first_name, cu.last_name, cu.company,
               cu.email, cu.address,
               (SELECT cpn.phone_number FROM contact_phone_numbers cpn
                WHERE cpn.entity_type='customer' AND cpn.entity_id = cu.id
                ORDER BY CASE WHEN cpn.label='Mobile' THEN 0 ELSE 1 END LIMIT 1) AS primary_phone
        FROM job_customers jc
        JOIN customers cu ON cu.id = jc.customer_id
        WHERE jc.job_id = ?
        ORDER BY jc.sort_order, jc.id
    """, (job_id,))
    job_linked_customers = cur.fetchall()

    linked_ids = [r["customer_id"] for r in job_linked_customers]
    if linked_ids:
        placeholders = ",".join("?" * len(linked_ids))
        cur.execute(f"SELECT id, first_name, last_name, company FROM customers WHERE id NOT IN ({placeholders}) ORDER BY last_name, first_name", linked_ids)
    else:
        cur.execute("SELECT id, first_name, last_name, company FROM customers ORDER BY last_name, first_name")
    all_customers = cur.fetchall()

    conn2 = db()
    tow_operators = conn2.execute("SELECT * FROM tow_operators WHERE active=1 ORDER BY company_name").fetchall()
    auction_yards  = conn2.execute("SELECT * FROM auction_yards WHERE active=1 ORDER BY name").fetchall()
    form_templates = conn2.execute("SELECT * FROM form_templates WHERE active=1 ORDER BY name").fetchall()
    conn2.close()

    statuses = ["New", "Active", "Active - Phone work only", "Suspended", "Awaiting info from client", "Completed", "Invoiced", "Cancelled"]
    visit_types = ["New Visit", "Re-attend", "First Update", "Urgent Update", "Phone Follow-up", "Locate Only"]
    priorities = ["Low", "Normal", "High", "Urgent"]
    item_types = ["vehicle", "property", "equipment", "other"]
    doc_types = ["Instructions", "PPSR", "Contract", "Invoice", "Authority", "Form", "Other"]
    customer_roles = ["Primary", "Director", "Guarantor", "Borrower", "Spouse", "Other"]

    conn3 = db()
    job_types_rows = conn3.execute("SELECT name FROM job_types WHERE active=1 ORDER BY name").fetchall()
    conn3.close()
    job_types = [r["name"] for r in job_types_rows]

    conn4 = db()
    _lpr_sightings_ensure_table(conn4)
    job_lpr_sightings = conn4.execute("""
        SELECT s.*, u.full_name AS agent_name
        FROM lpr_sightings s
        LEFT JOIN users u ON u.id = s.user_id
        WHERE s.matched_job_id = ?
        ORDER BY s.created_at DESC
        LIMIT 30
    """, (job_id,)).fetchall()

    job_patrol_intel = None
    if job_lpr_sightings:
        reg = job_lpr_sightings[0]["registration_normalised"]
        if reg:
            try:
                _patrol_intelligence_ensure(conn4)
                pi_row = conn4.execute(
                    "SELECT * FROM lpr_patrol_intelligence WHERE registration_normalised=?",
                    (reg,)
                ).fetchone()
                if pi_row:
                    zone = None
                    if pi_row["likely_zone"]:
                        try:
                            zone = json.loads(pi_row["likely_zone"])
                        except Exception:
                            pass
                    factors = []
                    if pi_row["explanation"]:
                        try:
                            factors = json.loads(pi_row["explanation"])
                        except Exception:
                            pass
                    job_patrol_intel = {
                        "reg":         pi_row["registration_normalised"],
                        "repeat":      pi_row["repeat_count_30d"],
                        "agents":      pi_row["distinct_agent_count"],
                        "zone":        zone,
                        "day_bucket":  pi_row["likely_day_bucket"] or "unknown",
                        "time_window": pi_row["likely_time_window"] or "mixed",
                        "confidence":  pi_row["confidence_score"],
                        "priority":    pi_row["recommended_patrol_priority"] or "low",
                        "action":      pi_row["recommended_action"] or "",
                        "factors":     factors,
                        "watchlist":   bool(pi_row["watchlist_hit"]),
                        "result_type": pi_row["result_type"] or "no_match",
                        "computed_at": pi_row["last_computed_at"],
                    }
            except Exception:
                pass
    conn4.close()

    conn5 = db()
    _rl_rows = conn5.execute(
        "SELECT item_id, status FROM repo_lock_records WHERE job_id=?", (job_id,)
    ).fetchall()
    conn5.close()
    repo_lock_map = {r["item_id"]: (r["status"] or "Draft") for r in _rl_rows}

    return render_template("job_detail.html", job=job, interactions=interactions,
                           job_items=job_items, item_types=item_types,
                           statuses=statuses, visit_types=visit_types, priorities=priorities,
                           job_types=job_types, users=users,
                           field_notes=field_notes, documents=documents,
                           doc_types=doc_types, booking_types=booking_types,
                           schedules=schedules, next_schedule=next_schedule,
                           job_linked_customers=job_linked_customers,
                           all_customers=all_customers,
                           all_clients=all_clients,
                           all_customers_for_edit=all_customers_for_edit,
                           customer_roles=customer_roles,
                           tow_operators=tow_operators,
                           auction_yards=auction_yards,
                           form_templates=form_templates,
                           job_lpr_sightings=job_lpr_sightings,
                           job_patrol_intel=job_patrol_intel,
                           repo_lock_map=repo_lock_map,
                           job_payments=job_payments,
                           payments_total_cents=payments_total_cents)


@app.post("/jobs/<int:job_id>/customers/add")
@login_required
@admin_required
def job_customer_add(job_id: int):
    customer_id = request.form.get("customer_id") or None
    role = request.form.get("role", "Primary").strip()
    if not customer_id:
        flash("Please select a customer.", "warning")
        return redirect(url_for("job_detail", job_id=job_id))
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT MAX(sort_order) FROM job_customers WHERE job_id = ?", (job_id,))
    row = cur.fetchone()
    next_order = (row[0] or 0) + 1
    try:
        cur.execute("""
            INSERT INTO job_customers (job_id, customer_id, role, sort_order, created_at)
            VALUES (?, ?, ?, ?, ?)
        """, (job_id, int(customer_id), role, next_order, now_ts()))
        conn.commit()
        flash("Customer added to job.", "success")
    except Exception:
        flash("That customer is already linked to this job.", "warning")
    conn.close()
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/customers/<int:jc_id>/remove")
@login_required
@admin_required
def job_customer_remove(job_id: int, jc_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM job_customers WHERE job_id = ?", (job_id,))
    count = cur.fetchone()[0]
    if count <= 1:
        conn.close()
        flash("Cannot remove the only customer on this job.", "warning")
        return redirect(url_for("job_detail", job_id=job_id))
    cur.execute("DELETE FROM job_customers WHERE id = ? AND job_id = ?", (jc_id, job_id))
    conn.commit()
    conn.close()
    flash("Customer removed from job.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


def _resolve_booking_type(cur, bt_id: str, bt_name: str):
    """Return a valid booking_type id — looking up by id, then by name, creating if new."""
    if bt_id:
        try:
            row = cur.execute("SELECT id FROM booking_types WHERE id = ?", (int(bt_id),)).fetchone()
            if row:
                return row["id"]
        except Exception:
            pass
    if bt_name:
        bt_name = " ".join(bt_name.split()).strip()
        if not bt_name:
            return None
        row = cur.execute("SELECT id FROM booking_types WHERE LOWER(TRIM(name)) = LOWER(?)", (bt_name,)).fetchone()
        if row:
            return row["id"]
        cur.execute("INSERT INTO booking_types (name, active) VALUES (?, 1)", (bt_name,))
        return cur.lastrowid
    return None


@app.post("/jobs/<int:job_id>/schedule")
@login_required
@admin_required
def add_schedule(job_id: int):
    date_str    = request.form.get("schedule_date", "").strip()
    time_str    = request.form.get("schedule_time", "").strip()
    bt_id       = request.form.get("booking_type_id", "").strip()
    bt_name     = request.form.get("booking_type_name", "").strip()
    notes       = request.form.get("notes", "").strip() or None
    caller_id   = session.get("user_id")
    caller_role = session.get("role", "")
    is_ajax     = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    if caller_role in ("admin", "both"):
        assigned_to = request.form.get("assigned_to_user_id", "").strip() or None
        if assigned_to:
            assigned_to = int(assigned_to)
    else:
        assigned_to = caller_id

    if not date_str or not time_str or (not bt_id and not bt_name):
        if is_ajax:
            return jsonify({"ok": False, "error": "Date, time and booking type are required."})
        flash("Date, time and booking type are required.", "danger")
        return redirect(url_for("job_detail", job_id=job_id))

    try:
        dt_str = parse_interaction_datetime(date_str, time_str)
    except Exception:
        if is_ajax:
            return jsonify({"ok": False, "error": "Invalid date or time format."})
        flash("Invalid date or time format.", "danger")
        return redirect(url_for("job_detail", job_id=job_id))

    conn = db()
    cur = conn.cursor()
    resolved_bt_id = _resolve_booking_type(cur, bt_id, bt_name)
    if not resolved_bt_id:
        conn.close()
        if is_ajax:
            return jsonify({"ok": False, "error": "Invalid booking type."})
        flash("Invalid booking type.", "danger")
        return redirect(url_for("job_detail", job_id=job_id))

    bt_row = cur.execute("SELECT name FROM booking_types WHERE id = ?", (resolved_bt_id,)).fetchone()
    bt_status = bt_row["name"] if bt_row else "Active"
    ts = now_ts()
    cur.execute("""
        INSERT INTO schedules (job_id, booking_type_id, scheduled_for, status, notes,
                               assigned_to_user_id, created_by_user_id, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (job_id, resolved_bt_id, dt_str, bt_status, notes, assigned_to, caller_id, ts))
    new_sched_id = cur.lastrowid
    _write_schedule_history(cur, new_sched_id, job_id, "created",
                            new_scheduled_for=dt_str, new_status=bt_status,
                            changed_by_user_id=caller_id)
    _auto_complete_schedule_cues(cur, job_id, ts)
    _sync_visit_type_from_booking(cur, job_id, bt_status, ts)
    if assigned_to:
        cur.execute("UPDATE jobs SET assigned_user_id = ?, updated_at = ? WHERE id = ?",
                    (assigned_to, ts, job_id))
    conn.commit()
    conn.close()

    if is_ajax:
        return jsonify({"ok": True, "bt_id": resolved_bt_id, "bt_name": bt_row["name"] if bt_row else bt_name})
    flash("Booking added.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/schedule/ajax")
@login_required
@admin_required
def add_schedule_ajax(job_id: int):
    date_str  = request.form.get("schedule_date", "").strip()
    time_str  = request.form.get("schedule_time", "").strip()
    bt_name   = request.form.get("booking_label", "New Booking").strip() or "New Booking"
    notes     = request.form.get("booking_details", "").strip() or None
    caller_id = session.get("user_id")
    unassigned = request.form.get("unassigned") == "1"
    user_ids  = request.form.getlist("assigned_user_ids")
    now       = now_ts()

    if not date_str or not time_str:
        return jsonify({"ok": False, "error": "Date and time are required."})

    try:
        dt_str = parse_interaction_datetime(date_str, time_str)
    except Exception:
        return jsonify({"ok": False, "error": "Invalid date or time."})

    conn = db()
    cur = conn.cursor()

    bt_name = " ".join(bt_name.split()).strip()
    cur.execute("SELECT id FROM booking_types WHERE LOWER(TRIM(name)) = LOWER(?)", (bt_name,))
    bt_row = cur.fetchone()
    if bt_row:
        bt_id = bt_row["id"]
    else:
        cur.execute("INSERT INTO booking_types (name) VALUES (?)", (bt_name,))
        bt_id = cur.lastrowid

    created = []
    if unassigned or not user_ids:
        cur.execute("""INSERT INTO schedules
            (job_id, booking_type_id, scheduled_for, status, notes, assigned_to_user_id, created_by_user_id, created_at)
            VALUES (?, ?, ?, ?, ?, NULL, ?, ?)""",
            (job_id, bt_id, dt_str, bt_name, notes, caller_id, now))
        new_sid = cur.lastrowid
        created.append({"id": new_sid, "assigned_to": None})
        _write_schedule_history(cur, new_sid, job_id, "created",
                                new_scheduled_for=dt_str, new_status=bt_name,
                                changed_by_user_id=caller_id)
    else:
        for uid in user_ids:
            try:
                uid_int = int(uid)
            except ValueError:
                continue
            cur.execute("""INSERT INTO schedules
                (job_id, booking_type_id, scheduled_for, status, notes, assigned_to_user_id, created_by_user_id, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (job_id, bt_id, dt_str, bt_name, notes, uid_int, caller_id, now))
            new_sid = cur.lastrowid
            created.append({"id": new_sid, "assigned_to": uid_int})
            _write_schedule_history(cur, new_sid, job_id, "created",
                                    new_scheduled_for=dt_str, new_status=bt_name,
                                    changed_by_user_id=caller_id)

    cur.execute("""INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
                   VALUES (?, 'Schedule', ?, ?, ?)""",
                (job_id, f"Booking '{bt_name}' added for {dt_str[:10]}.", now, now))
    _auto_complete_schedule_cues(cur, job_id, now)
    _sync_visit_type_from_booking(cur, job_id, bt_name, now)
    last_assigned = next((c["assigned_to"] for c in reversed(created) if c["assigned_to"]), None)
    if last_assigned:
        cur.execute("UPDATE jobs SET assigned_user_id = ?, updated_at = ? WHERE id = ?",
                    (last_assigned, now, job_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "count": len(created)})


@app.post("/jobs/<int:job_id>/clone")
@login_required
@admin_required
def job_clone(job_id: int):
    conn = db()
    src = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not src:
        conn.close()
        flash("Job not found.", "danger")
        return redirect(url_for("jobs_list"))

    conn2 = db()
    cur2 = conn2.cursor()
    cur2.execute("SELECT * FROM system_settings WHERE id = 1")
    settings = cur2.fetchone()
    prefix   = settings["job_prefix"] if settings else "0000"
    seq      = (settings["job_sequence"] or 0) + 1
    internal = f"{prefix}{seq:03d}"
    cur2.execute("UPDATE system_settings SET job_sequence = ? WHERE id = 1", (seq,))

    now = now_ts()
    caller_id = session.get("user_id")
    cur2.execute("""INSERT INTO jobs
        (internal_job_number, display_ref, client_id, customer_id,
         client_reference,
         job_type, visit_type, status, priority,
         job_address, description, assigned_user_id,
         created_at, updated_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (internal, internal,
         src["client_id"], src["customer_id"],
         src["client_reference"],
         src["job_type"], src["visit_type"], "New", src["priority"],
         src["job_address"], src["description"], src["assigned_user_id"],
         now, now))
    new_id = cur2.lastrowid

    src_items = conn.execute("SELECT * FROM job_items WHERE job_id = ?", (job_id,)).fetchall()
    item_cols = [r[1] for r in conn.execute("PRAGMA table_info(job_items)").fetchall()
                 if r[1] not in ("id", "job_id")]
    for item in src_items:
        cols_present = [c for c in item_cols if c in dict(item)]
        if not cols_present:
            continue
        ph = ",".join("?" * (len(cols_present) + 1))
        col_str = "job_id," + ",".join(cols_present)
        vals = [new_id] + [item[c] for c in cols_present]
        cur2.execute(f"INSERT INTO job_items ({col_str}) VALUES ({ph})", vals)

    cur2.execute("""INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
                   VALUES (?, 'Create', ?, ?, ?)""",
                (new_id, f"Job cloned from {src['internal_job_number']}.", now, now))

    conn.close()
    conn2.commit()
    conn2.close()
    flash(f"Job cloned as {internal}.", "success")
    return redirect(url_for("job_detail", job_id=new_id))


@app.post("/jobs/<int:job_id>/activate")
@login_required
@admin_required
def job_activate(job_id):
    new_status  = request.form.get("new_status", "").strip()
    date_str    = request.form.get("schedule_date", "").strip()
    time_str    = request.form.get("schedule_time", "").strip()
    bt_id       = request.form.get("booking_type_id", "").strip()
    notes       = request.form.get("notes", "").strip() or None
    assigned_to = request.form.get("assigned_to_user_id", "").strip() or None
    caller_id   = session.get("user_id")
    now         = datetime.now().isoformat(timespec="seconds")

    allowed = ["Active", "Active - Phone work only", "Suspended", "Awaiting info from client"]
    if new_status not in allowed:
        flash("Invalid status.", "danger")
        return redirect(url_for("job_detail", job_id=job_id))

    conn = db()
    cur  = conn.cursor()

    cur.execute("UPDATE jobs SET status = ?, updated_at = ? WHERE id = ?", (new_status, now, job_id))
    cur.execute("""INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
                   VALUES (?, ?, ?, ?, ?)""",
                (job_id, "Status Update", f"Status changed to '{new_status}'.", now, now))

    bt_name = request.form.get("booking_type_name", "").strip()
    if date_str and time_str and (bt_id or bt_name):
        try:
            dt_str = parse_interaction_datetime(date_str, time_str)
            resolved_bt_id = _resolve_booking_type(cur, bt_id, bt_name)
            if resolved_bt_id:
                bt_row = cur.execute("SELECT name FROM booking_types WHERE id = ?", (resolved_bt_id,)).fetchone()
                bt_status = bt_row["name"] if bt_row else new_status
                assigned_int = int(assigned_to) if assigned_to else None
                cur.execute("""INSERT INTO schedules
                               (job_id, booking_type_id, scheduled_for, status, notes,
                                assigned_to_user_id, created_by_user_id, created_at)
                               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                            (job_id, resolved_bt_id, dt_str, bt_status, notes,
                             assigned_int, caller_id, now_ts()))
                _write_schedule_history(cur, cur.lastrowid, job_id, "created",
                                        new_scheduled_for=dt_str, new_status=bt_status,
                                        changed_by_user_id=caller_id)
                _auto_complete_schedule_cues(cur, job_id, now)
                if assigned_int:
                    cur.execute("UPDATE jobs SET assigned_user_id = ?, updated_at = ? WHERE id = ?",
                                (assigned_int, now_ts(), job_id))
        except Exception:
            flash("Schedule date/time invalid — status updated but no schedule created.", "warning")

    conn.commit()
    conn.close()
    flash("Job updated.", "success")
    return redirect(url_for("job_detail", job_id=job_id))

@app.get("/schedule")
@login_required
def schedule_index():
    caller_role = session.get("role", "")
    is_admin = caller_role in ("admin", "both")
    agents = []
    conn = db()
    cur = conn.cursor()
    if is_admin:
        cur.execute("SELECT id, full_name FROM users WHERE active = 1 AND role IN ('agent','both','admin') ORDER BY full_name")
        agents = [{"id": a["id"], "name": a["full_name"]} for a in cur.fetchall()]
    cur.execute("SELECT id, name FROM booking_types WHERE active = 1 ORDER BY name")
    booking_types = [{"id": bt["id"], "name": bt["name"]} for bt in cur.fetchall()]
    conn.close()
    return render_template("schedule/index.html",
                           is_admin=is_admin,
                           agents=agents,
                           booking_types=booking_types,
                           user_id=session.get("user_id"),
                           default_view="week" if is_admin else "day")


_BOOKING_TYPE_COLORS = {
    "New Visit": "#2563eb",
    "Re-attend": "#7c3aed",
    "Urgent New Visit": "#dc2626",
    "Update Required": "#d97706",
    "Urgent Update Required": "#b91c1c",
}

@app.get("/schedule/api/events")
@login_required
def schedule_api_events():
    caller_id   = session.get("user_id")
    caller_role = session.get("role", "")
    is_admin    = caller_role in ("admin", "both")

    start_str     = request.args.get("start", "")
    end_str       = request.args.get("end", "")
    agent_id_raw  = request.args.get("agent_id", "all")
    status_filter = request.args.get("status_filter", "all")

    if not start_str or not end_str:
        return jsonify({"error": "start and end query params required"}), 400

    conn = db()
    cur  = conn.cursor()

    where_clauses = ["s.scheduled_for >= ?", "s.scheduled_for <= ?"]
    params = [start_str, end_str]

    if status_filter == "upcoming":
        where_clauses.append("s.status NOT IN ('Completed', 'Cancelled')")
    elif status_filter == "past":
        where_clauses.append("s.status IN ('Completed', 'Cancelled')")

    if is_admin:
        if agent_id_raw == "unassigned":
            where_clauses.append("s.assigned_to_user_id IS NULL")
        elif agent_id_raw and agent_id_raw != "all":
            try:
                aid = int(agent_id_raw)
                where_clauses.append("s.assigned_to_user_id = ?")
                params.append(aid)
            except ValueError:
                pass
    else:
        where_clauses.append("s.assigned_to_user_id = ?")
        params.append(caller_id)

    where_sql = " AND ".join(where_clauses)

    cur.execute(f"""
        SELECT s.id, s.job_id, s.scheduled_for, s.status, s.notes,
               s.assigned_to_user_id, s.booking_type_id,
               bt.name AS booking_type_name,
               j.display_ref, j.job_address, j.job_type,
               j.lender_name,
               cl.name AS client_name,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_name,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''),
                        TRIM(COALESCE(cu.first_name,'') || ' ' || COALESCE(cu.last_name,''))) AS customer_label,
               u.full_name AS assigned_to_name
        FROM schedules s
        JOIN booking_types bt ON bt.id = s.booking_type_id
        JOIN jobs j ON j.id = s.job_id
        LEFT JOIN users u ON u.id = s.assigned_to_user_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN clients cl ON cl.id = j.client_id
        WHERE {where_sql} AND j.status NOT IN {ARCHIVED_STATUSES!r}
        ORDER BY s.scheduled_for ASC
    """, params)

    rows = cur.fetchall()

    events = []
    for r in rows:
        bt_name = r["booking_type_name"] or ""
        events.append({
            "id": r["id"],
            "job_id": r["job_id"],
            "display_ref": r["display_ref"] or f"Job #{r['job_id']}",
            "customer_name": r["customer_name"] or "",
            "customer_label": r["customer_label"] or "",
            "booking_type": bt_name,
            "booking_type_id": r["booking_type_id"],
            "scheduled_for": r["scheduled_for"],
            "assigned_to_name": r["assigned_to_name"] or "",
            "assigned_to_user_id": r["assigned_to_user_id"],
            "job_address": r["job_address"] or "",
            "job_type": r["job_type"] or "",
            "client_name": r["client_name"] or "",
            "lender_name": r["lender_name"] or "",
            "status": r["status"] or "Booked",
            "notes": r["notes"] or "",
            "booking_type_color": _BOOKING_TYPE_COLORS.get(bt_name, "#6b7280"),
        })

    agents = []
    if is_admin:
        cur.execute("SELECT id, full_name FROM users WHERE active = 1 AND role IN ('agent','both') ORDER BY full_name")
        agents = [{"id": a["id"], "name": a["full_name"]} for a in cur.fetchall()]

    conn.close()
    return jsonify({"events": events, "agents": agents})


@app.post("/schedule/api/<int:sched_id>/reschedule")
@login_required
def schedule_api_reschedule(sched_id):
    new_dt = (request.form.get("new_datetime") or "").strip()
    change_method = (request.form.get("change_method") or "").strip()
    if not new_dt:
        return jsonify({"ok": False, "error": "New date/time is required."}), 400
    try:
        datetime.strptime(new_dt[:16], "%Y-%m-%dT%H:%M")
    except (ValueError, IndexError):
        return jsonify({"ok": False, "error": "Invalid date/time format."}), 400
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM schedules WHERE id = ?", (sched_id,))
    sched = cur.fetchone()
    if not sched:
        conn.close()
        return jsonify({"ok": False, "error": "Booking not found."}), 404
    caller_id = session.get("user_id")
    caller_role = session.get("role", "")
    is_admin = caller_role in ("admin", "both")
    if not is_admin and sched["assigned_to_user_id"] != caller_id:
        conn.close()
        return jsonify({"ok": False, "error": "Not authorised."}), 403
    old_dt = sched["scheduled_for"]
    cur.execute("UPDATE schedules SET scheduled_for = ? WHERE id = ?", (new_dt, sched_id))
    method_note = " (drag & drop)" if change_method == "drag_drop" else ""
    _write_schedule_history(cur, sched_id, sched["job_id"], "rescheduled",
                            old_dt, new_dt, sched["status"], sched["status"],
                            caller_id, f"Rescheduled from {old_dt[:16]} to {new_dt[:16]}{method_note}")
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/schedule/api/<int:sched_id>/complete")
@login_required
def schedule_api_complete(sched_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM schedules WHERE id = ?", (sched_id,))
    sched = cur.fetchone()
    if not sched:
        conn.close()
        return jsonify({"ok": False, "error": "Booking not found."}), 404
    caller_id = session.get("user_id")
    caller_role = session.get("role", "")
    is_admin = caller_role in ("admin", "both")
    if not is_admin and sched["assigned_to_user_id"] != caller_id:
        conn.close()
        return jsonify({"ok": False, "error": "Not authorised."}), 403
    old_status = sched["status"]
    cur.execute("UPDATE schedules SET status = 'Completed' WHERE id = ?", (sched_id,))
    _write_schedule_history(cur, sched_id, sched["job_id"], "completed",
                            sched["scheduled_for"], sched["scheduled_for"],
                            old_status, "Completed", caller_id, "Marked complete from schedule")
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/schedule/api/<int:sched_id>/cancel")
@login_required
def schedule_api_cancel(sched_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM schedules WHERE id = ?", (sched_id,))
    sched = cur.fetchone()
    if not sched:
        conn.close()
        return jsonify({"ok": False, "error": "Booking not found."}), 404
    caller_id = session.get("user_id")
    caller_role = session.get("role", "")
    is_admin = caller_role in ("admin", "both")
    if not is_admin and sched["assigned_to_user_id"] != caller_id:
        conn.close()
        return jsonify({"ok": False, "error": "Not authorised."}), 403
    old_status = sched["status"]
    cur.execute("UPDATE schedules SET status = 'Cancelled' WHERE id = ?", (sched_id,))
    _write_schedule_history(cur, sched_id, sched["job_id"], "cancelled",
                            sched["scheduled_for"], sched["scheduled_for"],
                            old_status, "Cancelled", caller_id, "Cancelled from schedule")
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.get("/schedule/api/<int:sched_id>/history")
@login_required
def schedule_api_history(sched_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM schedules WHERE id = ?", (sched_id,))
    sched = cur.fetchone()
    if not sched:
        conn.close()
        return jsonify({"ok": False, "error": "Booking not found."}), 404
    caller_id = session.get("user_id")
    caller_role = session.get("role", "")
    is_admin = caller_role in ("admin", "both")
    if not is_admin and sched["assigned_to_user_id"] != caller_id:
        conn.close()
        return jsonify({"ok": False, "error": "Not authorised."}), 403
    cur.execute("""
        SELECT sh.*, u.full_name AS changed_by_name
        FROM schedule_history sh
        LEFT JOIN users u ON u.id = sh.changed_by_user_id
        WHERE sh.schedule_id = ?
        ORDER BY sh.created_at DESC
    """, (sched_id,))
    rows = cur.fetchall()
    history = []
    for r in rows:
        history.append({
            "action": r["action"],
            "old_scheduled_for": r["old_scheduled_for"] or "",
            "new_scheduled_for": r["new_scheduled_for"] or "",
            "old_status": r["old_status"] or "",
            "new_status": r["new_status"] or "",
            "changed_by": r["changed_by_name"] or "System",
            "created_at": r["created_at"] or "",
            "notes": r["notes"] or "",
        })
    conn.close()
    return jsonify({"ok": True, "history": history})


@app.post("/schedule/api/<int:sched_id>/update")
@login_required
def schedule_api_update(sched_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM schedules WHERE id = ?", (sched_id,))
    sched = cur.fetchone()
    if not sched:
        conn.close()
        return jsonify({"ok": False, "error": "Booking not found."}), 404
    caller_id = session.get("user_id")
    caller_role = session.get("role", "")
    is_admin = caller_role in ("admin", "both")
    if not is_admin and sched["assigned_to_user_id"] != caller_id:
        conn.close()
        return jsonify({"ok": False, "error": "Not authorised."}), 403

    new_dt = (request.form.get("scheduled_for") or "").strip()
    new_agent = request.form.get("assigned_to_user_id", "").strip()
    new_bt = request.form.get("booking_type_id", "").strip()
    new_notes = request.form.get("notes", "").strip()

    changes = []
    old_dt = sched["scheduled_for"] or ""

    if new_dt:
        try:
            datetime.strptime(new_dt[:16], "%Y-%m-%dT%H:%M")
        except (ValueError, IndexError):
            conn.close()
            return jsonify({"ok": False, "error": "Invalid date/time format."}), 400
        if new_dt[:16] != old_dt[:16]:
            cur.execute("UPDATE schedules SET scheduled_for = ? WHERE id = ?", (new_dt, sched_id))
            changes.append(f"Date changed from {old_dt[:16]} to {new_dt[:16]}")

    if is_admin and "assigned_to_user_id" in request.form:
        old_agent = sched["assigned_to_user_id"]
        new_agent_id = int(new_agent) if new_agent.isdigit() else None
        if new_agent_id != old_agent:
            cur.execute("UPDATE schedules SET assigned_to_user_id = ? WHERE id = ?", (new_agent_id, sched_id))
            changes.append("Agent reassigned")
            if new_agent_id:
                cur.execute("UPDATE jobs SET assigned_user_id = ?, updated_at = ? WHERE id = ?",
                            (new_agent_id, now_ts(), sched["job_id"]))

    if new_bt and new_bt.isdigit():
        old_bt = sched["booking_type_id"]
        if int(new_bt) != old_bt:
            bt_check = cur.execute("SELECT id, name FROM booking_types WHERE id = ? AND active = 1", (int(new_bt),)).fetchone()
            if not bt_check:
                conn.close()
                return jsonify({"ok": False, "error": "Invalid booking type."}), 400
            cur.execute("UPDATE schedules SET booking_type_id = ? WHERE id = ?", (int(new_bt), sched_id))
            changes.append("Booking type changed")
            _sync_visit_type_from_booking(cur, sched["job_id"], bt_check["name"], now_ts())

    if new_notes != (sched["notes"] or ""):
        cur.execute("UPDATE schedules SET notes = ? WHERE id = ?", (new_notes, sched_id))
        changes.append("Notes updated")

    if changes:
        action = "rescheduled" if any("Date" in c for c in changes) else "updated"
        _write_schedule_history(cur, sched_id, sched["job_id"], action,
                                old_dt, new_dt or old_dt, sched["status"], sched["status"],
                                caller_id, "; ".join(changes))

    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/booking-type")
@login_required
@admin_required
def add_booking_type():
    name = request.form.get("new_booking_type", "").strip()
    if name:
        conn = db()
        cur = conn.cursor()
        try:
            cur.execute("INSERT INTO booking_types (name) VALUES (?)", (name,))
            conn.commit()
            flash(f"Booking type '{name}' added.", "success")
        except Exception:
            flash("That booking type already exists.", "warning")
        finally:
            conn.close()
    referrer = request.referrer or url_for("jobs")
    return redirect(referrer)




@app.post("/booking-type/ajax")
@login_required
@admin_required
def add_booking_type_ajax():
    name = " ".join((request.form.get("name", "") or "").split()).strip()
    if not name:
        return jsonify({"ok": False, "error": "Name is required."})
    conn = db()
    cur = conn.cursor()
    existing = cur.execute("SELECT id, name FROM booking_types WHERE LOWER(TRIM(name)) = LOWER(?)", (name,)).fetchone()
    if existing:
        conn.close()
        return jsonify({"ok": True, "id": existing["id"], "name": existing["name"], "existing": True})
    try:
        cur.execute("INSERT INTO booking_types (name, active) VALUES (?, 1)", (name,))
        new_id = cur.lastrowid
        conn.commit()
    except Exception:
        conn.close()
        return jsonify({"ok": False, "error": "That booking type already exists."})
    conn.close()
    return jsonify({"ok": True, "id": new_id, "name": name})



@app.post("/booking-type/delete")
@login_required
@admin_required
def booking_type_delete():
    ids = request.form.getlist("bt_ids")
    if not ids:
        flash("No booking types selected.", "warning")
        return redirect(url_for("admin_settings"))
    conn = db()
    cur = conn.cursor()
    deleted, blocked = [], []
    for bt_id in ids:
        try:
            bt_id = int(bt_id)
        except ValueError:
            continue
        cur.execute("SELECT name FROM booking_types WHERE id = ?", (bt_id,))
        row = cur.fetchone()
        if not row:
            continue
        name = row["name"]
        cur.execute("SELECT COUNT(*) cnt FROM schedules WHERE booking_type_id = ?", (bt_id,))
        count = cur.fetchone()["cnt"]
        if count > 0:
            blocked.append(f"'{name}' ({count} schedule{'s' if count != 1 else ''})")
        else:
            cur.execute("DELETE FROM booking_types WHERE id = ?", (bt_id,))
            deleted.append(name)
    conn.commit()
    conn.close()
    if deleted:
        flash(f"Deleted: {', '.join(deleted)}.", "success")
    if blocked:
        flash(f"Cannot delete — booking type(s) with existing schedules: {', '.join(blocked)}.", "warning")
    return redirect(url_for("admin_settings"))

@app.post("/booking-type/<int:bt_id>/edit")
@login_required
@admin_required
def edit_booking_type(bt_id):
    name = request.form.get("name", "").strip()
    if not name:
        return jsonify({"ok": False, "error": "Name is required."})
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("UPDATE booking_types SET name = ? WHERE id = ?", (name, bt_id))
        conn.commit()
    except Exception:
        conn.close()
        return jsonify({"ok": False, "error": "That booking type name already exists."})
    conn.close()
    return jsonify({"ok": True, "name": name})


@app.post("/booking-type/<int:bt_id>/delete")
@login_required
@admin_required
def delete_booking_type_single(bt_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) cnt FROM schedules WHERE booking_type_id = ?", (bt_id,))
    count = cur.fetchone()["cnt"]
    if count > 0:
        conn.close()
        return jsonify({"ok": False, "error": f"Cannot delete — used by {count} schedule(s)."})
    cur.execute("DELETE FROM booking_types WHERE id = ?", (bt_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/settings/change-password")
@login_required
def settings_change_password():
    user_id = session.get("user_id")
    current = request.form.get("current_password", "").strip()
    new_pw  = request.form.get("new_password", "").strip()
    confirm = request.form.get("confirm_password", "").strip()
    if not current or not new_pw or not confirm:
        return jsonify({"ok": False, "error": "All fields are required."})
    if new_pw != confirm:
        return jsonify({"ok": False, "error": "New passwords do not match."})
    if len(new_pw) < 6:
        return jsonify({"ok": False, "error": "Password must be at least 6 characters."})
    if new_pw == current:
        return jsonify({"ok": False, "error": "New password must differ from current password."})
    conn = db()
    user = conn.execute("SELECT * FROM users WHERE id = ?", (user_id,)).fetchone()
    if not user or not check_password_hash(user["password"], current):
        conn.close()
        return jsonify({"ok": False, "error": "Current password is incorrect."})
    conn.execute("UPDATE users SET password = ? WHERE id = ?",
                 (generate_password_hash(new_pw), user_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/admin/settings/change-password")
@login_required
@admin_required
def admin_change_password():
    user_id = session.get("user_id")
    current = request.form.get("current_password", "").strip()
    new_pw  = request.form.get("new_password", "").strip()
    confirm = request.form.get("confirm_password", "").strip()
    if not current or not new_pw or not confirm:
        return jsonify({"ok": False, "error": "All fields are required."})
    if new_pw != confirm:
        return jsonify({"ok": False, "error": "New passwords do not match."})
    if len(new_pw) < 6:
        return jsonify({"ok": False, "error": "Password must be at least 6 characters."})
    conn = db()
    user = conn.execute("SELECT * FROM users WHERE id = ?", (user_id,)).fetchone()
    if not user or not check_password_hash(user["password"], current):
        conn.close()
        return jsonify({"ok": False, "error": "Current password is incorrect."})
    conn.execute("UPDATE users SET password = ? WHERE id = ?",
                 (generate_password_hash(new_pw), user_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/job-type")
@login_required
@admin_required
def add_job_type():
    name = request.form.get("new_job_type", "").strip()
    if name:
        conn = db()
        cur = conn.cursor()
        try:
            cur.execute("INSERT INTO job_types (name) VALUES (?)", (name,))
            conn.commit()
            flash(f"Job type '{name}' added.", "success")
        except Exception:
            flash("That job type already exists.", "warning")
        finally:
            conn.close()
    referrer = request.referrer or url_for("job_new")
    return redirect(referrer)


@app.post("/jobs/<int:job_id>/schedule/<int:sched_id>/status")
@login_required
@admin_required
def update_schedule_status(job_id: int, sched_id: int):
    new_status = request.form.get("status", "").strip()
    allowed = {"Completed", "Cancelled"}
    if new_status not in allowed:
        flash("Invalid status.", "danger")
        return redirect(url_for("job_detail", job_id=job_id))
    conn = db()
    cur = conn.cursor()
    sched = cur.execute("SELECT * FROM schedules WHERE id = ? AND job_id = ?", (sched_id, job_id)).fetchone()
    old_status = sched["status"] if sched else None
    old_scheduled = sched["scheduled_for"] if sched else None
    cur.execute("UPDATE schedules SET status = ? WHERE id = ? AND job_id = ?",
                (new_status, sched_id, job_id))
    action = "completed" if new_status == "Completed" else "cancelled"
    _write_schedule_history(cur, sched_id, job_id, action,
                            old_scheduled_for=old_scheduled, new_scheduled_for=old_scheduled,
                            old_status=old_status, new_status=new_status,
                            changed_by_user_id=session.get("user_id"))
    conn.commit()
    conn.close()
    flash("Booking updated.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/delete")
@login_required
@admin_required
def delete_job(job_id: int):
    conn = db()
    cur = conn.cursor()

    # Clean up note files (blobs + local)
    cur.execute("SELECT id FROM job_field_notes WHERE job_id = ?", (job_id,))
    note_ids = [r["id"] for r in cur.fetchall()]
    for nid in note_ids:
        cur.execute("SELECT filename, filepath FROM job_note_files WHERE job_field_note_id = ?", (nid,))
        for f in cur.fetchall():
            delete_blob_safely(f["filename"])
            try: os.remove(f["filepath"])
            except OSError: pass
        cur.execute("DELETE FROM job_note_files WHERE job_field_note_id = ?", (nid,))
    cur.execute("DELETE FROM job_field_notes WHERE job_id = ?", (job_id,))

    # Clean up uploaded documents (blobs)
    cur.execute("SELECT stored_filename FROM job_documents WHERE job_id = ?", (job_id,))
    for f in cur.fetchall():
        delete_blob_safely(f["stored_filename"])
    cur.execute("DELETE FROM job_documents WHERE job_id = ?", (job_id,))

    # Nullify matched_job_id in LPR tables (sightings are standalone records)
    for lpr_tbl in ("lpr_sightings", "lpr_patrol_intel"):
        try:
            cur.execute(f"UPDATE {lpr_tbl} SET matched_job_id=NULL WHERE matched_job_id=?", (job_id,))
        except Exception:
            pass

    # Delete all child rows
    for tbl in (
        "schedule_history",
        "job_items", "job_assets", "interactions", "cue_items", "schedules",
        "job_customers", "job_updates", "job_payments",
        "repo_lock_records", "repo_lock_queue",
    ):
        try:
            cur.execute(f"DELETE FROM {tbl} WHERE job_id = ?", (job_id,))
        except Exception:
            pass

    cur.execute("DELETE FROM jobs WHERE id = ?", (job_id,))
    conn.commit()
    conn.close()

    audit("job", job_id, "delete", "Job deleted", {})
    flash("Job deleted.", "success")
    return redirect(url_for("jobs_list"))


@app.post("/jobs/<int:job_id>/status")
@login_required
@admin_required
def job_status_update(job_id: int):
    status = request.form.get("status", "").strip()
    allowed = ["New", "Active", "Active - Phone work only", "Suspended",
               "Awaiting info from client", "Completed", "Invoiced", "Cancelled"]
    if status not in allowed:
        flash("Invalid status.", "danger")
        return redirect(url_for("jobs_list"))
    now = datetime.now().isoformat(timespec="seconds")
    conn = db()
    cur = conn.cursor()
    cur.execute("UPDATE jobs SET status = ?, updated_at = ? WHERE id = ?",
                (status, now, job_id))
    cur.execute("""
        INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
        VALUES (?, ?, ?, ?, ?)
    """, (job_id, "Status Update", f"Status changed to '{status}'.", now, now))
    if status in ("Completed", "Invoiced", "Cancelled"):
        cur.execute("UPDATE jobs SET assigned_user_id = NULL, updated_at = ? WHERE id = ?",
                    (now, job_id))
        pending_scheds = cur.execute(
            "SELECT id, scheduled_for, status FROM schedules WHERE job_id = ? AND status NOT IN ('Completed', 'Cancelled')",
            (job_id,)).fetchall()
        cur.execute("""
            UPDATE schedules SET status = 'Cancelled'
            WHERE job_id = ? AND status NOT IN ('Completed', 'Cancelled')
        """, (job_id,))
        for ps in pending_scheds:
            _write_schedule_history(cur, ps["id"], job_id, "cancelled",
                                    old_scheduled_for=ps["scheduled_for"],
                                    new_scheduled_for=ps["scheduled_for"],
                                    old_status=ps["status"], new_status="Cancelled",
                                    changed_by_user_id=session.get("user_id"),
                                    notes=f"Auto-cancelled — job marked '{status}'.")
        cur.execute("""
            INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
            VALUES (?, ?, ?, ?, ?)
        """, (job_id, "System",
                f"Agent unassigned and pending schedules cancelled — job marked '{status}'.",
                now, now))
        cur.execute("""
            UPDATE cue_items SET status='Completed', completed_at=?, updated_at=?
            WHERE job_id=? AND status IN ('Pending','In Progress')
        """, (now, now, job_id))
    conn.commit()
    conn.close()
    return redirect(url_for("job_detail", job_id=job_id))


@app.get("/api/agent-recommend/<int:job_id>")
@login_required
@admin_required
def agent_recommend(job_id: int):
    import re as _re
    conn = db()
    cur = conn.cursor()
    job = cur.execute("SELECT id, job_address, assigned_user_id FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job or not job["job_address"]:
        conn.close()
        return jsonify({"agents": [], "postcode": None})
    addr = job["job_address"]
    m = _re.search(r'\b(\d{4})\b', addr)
    if not m:
        conn.close()
        return jsonify({"agents": [], "postcode": None})
    postcode = m.group(1)
    rows = cur.execute("""
        SELECT u.id, u.full_name, COUNT(DISTINCT j2.id) AS job_count,
               GROUP_CONCAT(DISTINCT j2.display_ref) AS job_refs
        FROM jobs j2
        JOIN schedules s ON s.job_id = j2.id AND s.status NOT IN ('Cancelled', 'Completed')
        JOIN users u ON u.id = s.assigned_to_user_id
        WHERE j2.job_address LIKE ?
          AND j2.status NOT IN ('Completed', 'Invoiced', 'Cancelled', 'Archived - Invoiced', 'Cold Stored')
          AND j2.id != ?
          AND u.active = 1
        GROUP BY u.id
        ORDER BY job_count DESC
    """, (f"%{postcode}%", job_id)).fetchall()
    conn.close()
    agents = [{"id": r["id"], "name": r["full_name"], "job_count": r["job_count"],
               "job_refs": r["job_refs"]} for r in rows]
    return jsonify({"agents": agents, "postcode": postcode})


@app.post("/jobs/<int:job_id>/assign")
@login_required
@admin_required
def job_assign_agent(job_id: int):
    assigned_to = request.form.get("assigned_to_user_id", "").strip()
    if not assigned_to or not assigned_to.isdigit():
        return jsonify({"ok": False, "error": "No valid agent specified."}), 400
    agent_id = int(assigned_to)
    conn = db()
    cur = conn.cursor()
    job = cur.execute("SELECT id FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({"ok": False, "error": "Job not found."}), 404
    agent = cur.execute("SELECT id FROM users WHERE id = ? AND active = 1", (agent_id,)).fetchone()
    if not agent:
        conn.close()
        return jsonify({"ok": False, "error": "Agent not found or inactive."}), 404
    now = datetime.now().isoformat(timespec="seconds")
    cur.execute("UPDATE jobs SET assigned_user_id = ?, updated_at = ? WHERE id = ?",
                (agent_id, now, job_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/jobs/<int:job_id>/archive")
@login_required
@admin_required
def job_archive(job_id: int):
    conn = db()
    cur = conn.cursor()
    job = cur.execute("SELECT id, status, display_ref FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        conn.close()
        flash("Job not found.", "danger")
        return redirect(url_for("jobs_list"))
    archivable = ("Completed", "Invoiced", "Cancelled")
    if job["status"] in ARCHIVED_STATUSES:
        conn.close()
        flash("Job is already archived.", "warning")
        return redirect(url_for("job_detail", job_id=job_id))
    if job["status"] not in archivable:
        conn.close()
        flash(f"Only Completed, Invoiced or Cancelled jobs can be archived. Current status: {job['status']}.", "danger")
        return redirect(url_for("job_detail", job_id=job_id))
    uid = session.get("user_id")
    ts = now_ts()
    old_status = job["status"]
    cur.execute("""
        UPDATE jobs SET status = 'Archived - Invoiced', lifecycle_status = 'archived',
               archived_at = ?, archived_by_user_id = ?, updated_at = ?
        WHERE id = ?
    """, (ts, uid, ts, job_id))
    cur.execute("""
        UPDATE cue_items SET status='Completed', completed_at=?, updated_at=?
        WHERE job_id=? AND status IN ('Pending','In Progress')
    """, (ts, ts, job_id))
    cur.execute("""
        INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
        VALUES (?, ?, ?, ?, ?)
    """, (job_id, "Lifecycle", f"Job archived from '{old_status}'.", ts, ts))
    _log_lifecycle(cur, job_id, "archive", old_status, "Archived - Invoiced", uid)
    conn.commit()
    conn.close()
    audit("job", job_id, "archive", f"Job {job['display_ref']} archived.", {"from_status": old_status})
    flash(f"Job {job['display_ref']} has been archived.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/restore")
@login_required
@admin_required
def job_restore(job_id: int):
    restore_to = request.form.get("restore_to", "Invoiced").strip()
    if restore_to not in ("Invoiced", "Completed", "Active", "New"):
        restore_to = "Invoiced"
    conn = db()
    cur = conn.cursor()
    if restore_to in ("Active", "New"):
        settings = cur.execute("SELECT allow_restore_to_active FROM system_settings WHERE id = 1").fetchone()
        if settings and not settings["allow_restore_to_active"]:
            conn.close()
            flash("Restoring to active status is disabled by policy.", "danger")
            return redirect(url_for("job_detail", job_id=job_id))
    job = cur.execute("SELECT id, status, display_ref FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        conn.close()
        flash("Job not found.", "danger")
        return redirect(url_for("jobs_list"))
    if job["status"] not in ARCHIVED_STATUSES:
        conn.close()
        flash("Job is not archived.", "warning")
        return redirect(url_for("job_detail", job_id=job_id))
    uid = session.get("user_id")
    ts = now_ts()
    old_status = job["status"]
    cur.execute("""
        UPDATE jobs SET status = ?, lifecycle_status = 'active',
               archived_at = NULL, archived_by_user_id = NULL, updated_at = ?
        WHERE id = ?
    """, (restore_to, ts, job_id))
    cur.execute("""
        INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
        VALUES (?, ?, ?, ?, ?)
    """, (job_id, "Lifecycle", f"Job restored from archive to '{restore_to}'.", ts, ts))
    _log_lifecycle(cur, job_id, "restore", old_status, restore_to, uid)
    conn.commit()
    conn.close()
    audit("job", job_id, "restore", f"Job {job['display_ref']} restored to {restore_to}.", {"from_status": old_status, "to_status": restore_to})
    flash(f"Job {job['display_ref']} has been restored to {restore_to}.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/admin/archive/bulk")
@login_required
@admin_required
def archive_bulk():
    job_ids_raw = request.form.getlist("job_ids")
    if not job_ids_raw:
        data = request.get_json(silent=True)
        if data and "job_ids" in data:
            job_ids_raw = data["job_ids"]
    job_ids = []
    for jid in job_ids_raw:
        try:
            job_ids.append(int(jid))
        except (ValueError, TypeError):
            pass
    if not job_ids:
        if request.is_json:
            return jsonify({"ok": False, "error": "No valid job IDs provided."}), 400
        flash("No jobs selected.", "warning")
        return redirect(url_for("admin_archive"))
    uid = session.get("user_id")
    ts = now_ts()
    batch_id = f"bulk_{ts}_{uid}"
    conn = db()
    cur = conn.cursor()
    archived_count = 0
    for jid in job_ids:
        job = cur.execute("SELECT id, status, display_ref FROM jobs WHERE id = ?", (jid,)).fetchone()
        if not job or job["status"] in ARCHIVED_STATUSES or job["status"] not in ("Completed", "Invoiced"):
            continue
        old_status = job["status"]
        cur.execute("""
            UPDATE jobs SET status = 'Archived - Invoiced', lifecycle_status = 'archived',
                   archived_at = ?, archived_by_user_id = ?, updated_at = ?
            WHERE id = ?
        """, (ts, uid, ts, jid))
        cur.execute("""
            UPDATE cue_items SET status='Completed', completed_at=?, updated_at=?
            WHERE job_id=? AND status IN ('Pending','In Progress')
        """, (ts, ts, jid))
        cur.execute("""
            INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
            VALUES (?, ?, ?, ?, ?)
        """, (jid, "Lifecycle", f"Job bulk-archived from '{old_status}'.", ts, ts))
        _log_lifecycle(cur, jid, "bulk_archive", old_status, "Archived - Invoiced", uid, batch_id=batch_id)
        archived_count += 1
    conn.commit()
    conn.close()
    audit("system", None, "bulk_archive", f"Bulk archived {archived_count} jobs.", {"batch_id": batch_id, "count": archived_count})
    if request.is_json:
        return jsonify({"ok": True, "archived": archived_count, "batch_id": batch_id})
    flash(f"{archived_count} job(s) archived successfully.", "success")
    return redirect(url_for("admin_archive"))


@app.post("/admin/archive/bulk-all")
@login_required
@admin_required
def archive_bulk_all():
    if not request.is_json:
        return jsonify({"ok": False, "error": "Invalid request"}), 400
    data = request.get_json(silent=True) or {}
    days = 90
    try:
        days = int(data.get("days", 90))
    except (ValueError, TypeError):
        pass
    filter_client = data.get("client_id", "").strip()

    where = "WHERE j.status IN ('Invoiced', 'Completed')"
    params = []
    if days > 0:
        where += f"""
            AND (
                date(COALESCE(j.updated_at, j.created_at)) <= date('now', '-{days} days')
                OR j.geoop_source_description IS NOT NULL
            )"""
    if filter_client:
        where += " AND j.client_id = ?"
        params.append(filter_client)

    uid = session.get("user_id")
    ts = now_ts()
    batch_id = f"bulkall_{ts}_{uid}"
    conn = db()
    cur = conn.cursor()
    job_rows = cur.execute(f"SELECT id, status FROM jobs j {where}", params).fetchall()
    archived_count = 0
    for job in job_rows:
        jid = job["id"]
        old_status = job["status"]
        cur.execute("""
            UPDATE jobs SET status = 'Archived - Invoiced', lifecycle_status = 'archived',
                   archived_at = ?, archived_by_user_id = ?, updated_at = ?
            WHERE id = ?
        """, (ts, uid, ts, jid))
        cur.execute("""
            UPDATE cue_items SET status='Completed', completed_at=?, updated_at=?
            WHERE job_id=? AND status IN ('Pending','In Progress')
        """, (ts, ts, jid))
        cur.execute("""
            INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
            VALUES (?, ?, ?, ?, ?)
        """, (jid, "Lifecycle", f"Job bulk-archived from '{old_status}'.", ts, ts))
        _log_lifecycle(cur, jid, "bulk_archive", old_status, "Archived - Invoiced", uid, batch_id=batch_id)
        archived_count += 1
        if archived_count % 500 == 0:
            conn.commit()
    conn.commit()
    conn.close()
    audit("system", None, "bulk_archive_all", f"Bulk archived ALL {archived_count} eligible jobs.", {"batch_id": batch_id, "count": archived_count})
    return jsonify({"ok": True, "archived": archived_count, "batch_id": batch_id})


@app.get("/admin/archive")
@login_required
@admin_required
def admin_archive():
    conn = db()
    cur = conn.cursor()
    q = request.args.get("q", "").strip()
    filter_client = request.args.get("client_id", "").strip()
    filter_status = request.args.get("status", "").strip()
    date_from = request.args.get("date_from", "").strip()
    date_to = request.args.get("date_to", "").strip()
    mode = request.args.get("mode", "search").strip()

    clients_list = cur.execute(
        "SELECT id, COALESCE(nickname, name) AS name FROM clients ORDER BY name"
    ).fetchall()

    if mode == "eligible":
        where = "WHERE j.status IN ('Invoiced', 'Completed')"
        params = []
        days = request.args.get("days", "90").strip()
        try:
            days_int = int(days)
        except ValueError:
            days_int = 90
        if days_int > 0:
            where += f"""
                AND (
                    date(COALESCE(j.updated_at, j.created_at)) <= date('now', '-{days_int} days')
                    OR j.geoop_source_description IS NOT NULL
                )"""
        if filter_client:
            where += " AND j.client_id = ?"
            params.append(filter_client)

        page = request.args.get("page", "1").strip()
        try:
            page_int = max(1, int(page))
        except ValueError:
            page_int = 1
        per_page = 500

        total_eligible = cur.execute(f"""
            SELECT COUNT(*) FROM jobs j
            LEFT JOIN clients c ON c.id = j.client_id
            LEFT JOIN customers cu ON cu.id = j.customer_id
            {where}
        """, params).fetchone()[0]

        total_pages = max(1, (total_eligible + per_page - 1) // per_page)
        if page_int > total_pages:
            page_int = total_pages
        offset = (page_int - 1) * per_page

        cur.execute(f"""
            SELECT j.id, j.display_ref, j.status, j.client_id,
                   COALESCE(c.nickname, c.name) AS client_name,
                   COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
                   j.job_address, j.updated_at,
                   (SELECT ji.reg FROM job_items ji WHERE ji.job_id = j.id AND ji.item_type = 'vehicle' LIMIT 1) AS asset_reg,
                   (SELECT COUNT(*) FROM job_field_notes n WHERE n.job_id = j.id) AS note_count,
                   (SELECT COUNT(*) FROM job_documents d WHERE d.job_id = j.id) AS doc_count
            FROM jobs j
            LEFT JOIN clients c ON c.id = j.client_id
            LEFT JOIN customers cu ON cu.id = j.customer_id
            {where}
            ORDER BY j.updated_at ASC
            LIMIT ? OFFSET ?
        """, params + [per_page, offset])
        results = cur.fetchall()
        conn.close()
        return render_template("archive.html", results=results, mode="eligible",
                               q=q, filter_client=filter_client, filter_status=filter_status,
                               date_from=date_from, date_to=date_to, days=days_int,
                               clients_list=clients_list,
                               total_eligible=total_eligible, page=page_int,
                               total_pages=total_pages, per_page=per_page)

    where = "WHERE j.status IN ('Archived - Invoiced', 'Cold Stored')"
    params = []

    if filter_status and filter_status in ARCHIVED_STATUSES:
        where = f"WHERE j.status = ?"
        params.append(filter_status)
    if q:
        where += """
         AND (
           j.internal_job_number LIKE ? OR j.client_reference LIKE ? OR
           j.display_ref LIKE ? OR j.job_address LIKE ? OR
           cu.first_name LIKE ? OR cu.last_name LIKE ? OR cu.company LIKE ? OR
           c.name LIKE ? OR
           EXISTS (SELECT 1 FROM job_items ji WHERE ji.job_id = j.id AND (ji.reg LIKE ? OR ji.vin LIKE ?))
         )"""
        like = f"%{q}%"
        params.extend([like] * 10)
    if filter_client:
        where += " AND j.client_id = ?"
        params.append(filter_client)
    if date_from:
        where += " AND date(j.archived_at) >= ?"
        params.append(date_from)
    if date_to:
        where += " AND date(j.archived_at) <= ?"
        params.append(date_to)

    cur.execute(f"""
        SELECT j.id, j.display_ref, j.status, j.client_id, j.archived_at,
               COALESCE(c.nickname, c.name) AS client_name,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
               (cu.first_name || ' ' || cu.last_name) AS customer_name,
               j.job_address, j.updated_at, j.lifecycle_status,
               (SELECT ji.reg FROM job_items ji WHERE ji.job_id = j.id AND ji.item_type = 'vehicle' LIMIT 1) AS asset_reg,
               (SELECT ji.vin FROM job_items ji WHERE ji.job_id = j.id AND ji.item_type = 'vehicle' LIMIT 1) AS asset_vin,
               (SELECT COUNT(*) FROM job_field_notes n WHERE n.job_id = j.id) AS note_count,
               (SELECT COUNT(*) FROM job_documents d WHERE d.job_id = j.id) AS doc_count,
               u.full_name AS archived_by_name
        FROM jobs j
        LEFT JOIN clients c ON c.id = j.client_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN users u ON u.id = j.archived_by_user_id
        {where}
        ORDER BY j.archived_at DESC
        LIMIT 500
    """, params)
    results = cur.fetchall()

    archive_stats = cur.execute("""
        SELECT
            SUM(CASE WHEN status = 'Archived - Invoiced' THEN 1 ELSE 0 END) AS archived_count,
            SUM(CASE WHEN status = 'Cold Stored' THEN 1 ELSE 0 END) AS cold_stored_count
        FROM jobs WHERE status IN ('Archived - Invoiced', 'Cold Stored')
    """).fetchone()

    conn.close()
    return render_template("archive.html", results=results, mode="search",
                           q=q, filter_client=filter_client, filter_status=filter_status,
                           date_from=date_from, date_to=date_to, days=90,
                           clients_list=clients_list,
                           archive_stats=archive_stats)


@app.get("/admin/archive/lifecycle-log/<int:job_id>")
@login_required
@admin_required
def archive_lifecycle_log(job_id: int):
    conn = db()
    rows = conn.execute("""
        SELECT l.*, u.full_name AS performed_by_name
        FROM job_lifecycle_log l
        LEFT JOIN users u ON u.id = l.performed_by_user_id
        WHERE l.job_id = ?
        ORDER BY l.performed_at DESC
    """, (job_id,)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.post("/jobs/<int:job_id>/update")
@login_required
@admin_required
def job_update(job_id: int):
    status = request.form.get("status", "").strip()
    visit_type = request.form.get("visit_type", "").strip()
    assigned_user_id = request.form.get("assigned_user_id") or None
    now = datetime.now().isoformat(timespec="seconds")

    conn = db()
    cur = conn.cursor()
    cur.execute("UPDATE jobs SET status = ?, visit_type = ?, assigned_user_id = ?, updated_at = ? WHERE id = ?",
                (status, visit_type, assigned_user_id, now, job_id))

    cur.execute("""
        INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
        VALUES (?, ?, ?, ?, ?)
    """, (job_id, "Status/Visit Update", f"Status set to '{status}'. Visit type set to '{visit_type}'.", now, now))

    if status in ("Completed", "Invoiced", "Cancelled"):
        cur.execute("UPDATE jobs SET assigned_user_id = NULL, updated_at = ? WHERE id = ?",
                    (now, job_id))
        pending_scheds = cur.execute(
            "SELECT id, scheduled_for, status FROM schedules WHERE job_id = ? AND status NOT IN ('Completed', 'Cancelled')",
            (job_id,)).fetchall()
        cur.execute("""
            UPDATE schedules SET status = 'Cancelled'
            WHERE job_id = ? AND status NOT IN ('Completed', 'Cancelled')
        """, (job_id,))
        for ps in pending_scheds:
            _write_schedule_history(cur, ps["id"], job_id, "cancelled",
                                    old_scheduled_for=ps["scheduled_for"],
                                    new_scheduled_for=ps["scheduled_for"],
                                    old_status=ps["status"], new_status="Cancelled",
                                    changed_by_user_id=session.get("user_id"),
                                    notes=f"Auto-cancelled — job marked '{status}'.")
        cur.execute("""
            INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
            VALUES (?, ?, ?, ?, ?)
        """, (job_id, "System",
                f"Agent unassigned and pending schedules cancelled — job marked '{status}'.",
                now, now))
        cur.execute("""
            UPDATE cue_items SET status='Completed', completed_at=?, updated_at=?
            WHERE job_id=? AND status IN ('Pending','In Progress')
        """, (now, now, job_id))

    conn.commit()
    conn.close()
    flash("Job updated.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/edit")
@login_required
@admin_required
def job_edit(job_id: int):
    client_reference  = request.form.get("client_reference", "").strip() or None
    client_job_number = request.form.get("client_job_number", "").strip() or None
    client_id         = request.form.get("client_id") or None
    _cust_raw         = request.form.get("customer_id", "").strip()
    customer_id       = int(_cust_raw) if _cust_raw.isdigit() else None
    job_type          = request.form.get("job_type", "").strip() or "Recovery"
    visit_type        = request.form.get("visit_type", "").strip() or "Initial"
    status            = request.form.get("status", "").strip() or "Active"
    priority          = request.form.get("priority", "").strip() or "Normal"
    job_address       = request.form.get("job_address", "").strip() or None
    description       = request.form.get("description", "").strip() or None
    assigned_user_id  = request.form.get("assigned_user_id") or None
    now = now_ts()

    conn = db()
    cur  = conn.cursor()
    cur.execute("SELECT internal_job_number FROM jobs WHERE id = ?", (job_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        flash("Job not found.", "danger")
        return redirect(url_for("jobs_list"))

    internal = row["internal_job_number"]
    display_ref = internal
    if client_reference:
        display_ref = f"{internal} ({client_reference})"

    cur.execute("""
        UPDATE jobs SET
            client_reference=?, client_job_number=?, display_ref=?,
            client_id=?, customer_id=?,
            job_type=?, visit_type=?, status=?, priority=?,
            job_address=?, description=?, assigned_user_id=?,
            updated_at=?
        WHERE id=?
    """, (client_reference, client_job_number, display_ref,
          client_id, customer_id,
          job_type, visit_type, status, priority,
          job_address, description, assigned_user_id,
          now, job_id))

    if customer_id:
        cur.execute("SELECT id FROM job_customers WHERE job_id = ? AND role = 'Primary'", (job_id,))
        existing = cur.fetchone()
        if existing:
            cur.execute("UPDATE job_customers SET customer_id = ? WHERE id = ?",
                        (customer_id, existing["id"]))
        else:
            cur.execute("""INSERT INTO job_customers (job_id, customer_id, role, sort_order, created_at)
                           VALUES (?, ?, 'Primary', 0, ?)""", (job_id, customer_id, now))

    cur.execute("""
        INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
        VALUES (?, 'Edit', ?, ?, ?)
    """, (job_id, f"Job details updated. Status: {status}, Type: {job_type}, Visit: {visit_type}.", now, now))

    conn.commit()
    conn.close()

    if job_address:
        _geocode_job_async(job_id, job_address)

    flash("Job updated.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/link-client")
@login_required
@admin_required
def job_link_client(job_id: int):
    conn = db()
    cur = conn.cursor()
    job = cur.execute("SELECT id, internal_job_number, display_ref FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({"ok": False, "error": "Job not found."}), 404

    client_id = request.form.get("client_id", "").strip()
    new_client_name = " ".join(request.form.get("new_client_name", "").split()).strip()
    now = now_ts()
    action_label = ""

    try:
        if client_id and client_id.isdigit():
            client = cur.execute("SELECT id, name, nickname, email, phone FROM clients WHERE id=?", (int(client_id),)).fetchone()
            if not client:
                conn.close()
                return jsonify({"ok": False, "error": "Selected client not found."}), 404
            cur.execute("UPDATE jobs SET client_id=?, updated_at=? WHERE id=?", (int(client_id), now, job_id))
            action_label = f"Client linked: {client['name']}"
            result_client = {"id": client["id"], "name": client["name"], "nickname": client["nickname"] or "", "email": client["email"] or "", "phone": client["phone"] or ""}
        elif new_client_name:
            existing = cur.execute("SELECT id, name, nickname, email, phone FROM clients WHERE LOWER(TRIM(name))=LOWER(?)", (new_client_name,)).fetchone()
            if existing:
                cur.execute("UPDATE jobs SET client_id=?, updated_at=? WHERE id=?", (existing["id"], now, job_id))
                action_label = f"Client linked (existing match): {existing['name']}"
                result_client = {"id": existing["id"], "name": existing["name"], "nickname": existing["nickname"] or "", "email": existing["email"] or "", "phone": existing["phone"] or ""}
            else:
                cur.execute("INSERT INTO clients (name, created_at) VALUES (?, ?)", (new_client_name, now))
                new_id = cur.lastrowid
                cur.execute("UPDATE jobs SET client_id=?, updated_at=? WHERE id=?", (new_id, now, job_id))
                action_label = f"Client created and linked: {new_client_name}"
                result_client = {"id": new_id, "name": new_client_name, "nickname": "", "email": "", "phone": ""}
        else:
            conn.close()
            return jsonify({"ok": False, "error": "Please select an existing client or enter a name to create one."})

        cur.execute("""INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at)
                       VALUES (?, 'Client Link', ?, ?, ?)""", (job_id, action_label, now, now))
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        import traceback
        traceback.print_exc()
        return jsonify({"ok": False, "error": "An unexpected error occurred. Please try again."}), 500

    conn.close()
    return jsonify({"ok": True, "message": action_label, "client": result_client})


@app.post("/jobs/<int:job_id>/lender")
@login_required
@admin_required
def job_lender_update(job_id: int):
    lender_name     = request.form.get("lender_name", "").strip()
    account_number  = request.form.get("account_number", "").strip()
    regulation_type    = request.form.get("regulation_type", "").strip()
    arrears_cents      = money_to_cents(request.form.get("arrears", ""))
    costs_cents        = money_to_cents(request.form.get("costs", ""))
    costs2_cents       = money_to_cents(request.form.get("costs2", ""))
    mmp_cents          = money_to_cents(request.form.get("mmp", ""))
    job_due_date       = request.form.get("job_due_date", "").strip() or None
    payment_frequency  = request.form.get("payment_frequency", "").strip() or None
    conn = db()
    cur  = conn.cursor()
    cur.execute("""
        UPDATE jobs SET lender_name=?, account_number=?, regulation_type=?,
        arrears_cents=?, costs_cents=?, costs2_cents=?, mmp_cents=?, job_due_date=?,
        payment_frequency=?, updated_at=? WHERE id=?
    """, (lender_name or None, account_number or None, regulation_type or None,
          arrears_cents or None, costs_cents or None, costs2_cents or None,
          mmp_cents or None, job_due_date, payment_frequency, now_ts(), job_id))
    conn.commit()
    conn.close()
    flash("Lender details updated.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/description")
@login_required
@admin_required
def job_description_update(job_id: int):
    desc = request.form.get("description", "").strip()
    conn = db()
    conn.execute("UPDATE jobs SET description=?, updated_at=? WHERE id=?",
                 (desc or None, now_ts(), job_id))
    conn.commit()
    conn.close()
    flash("Description updated.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/payments")
@login_required
def job_payment_add(job_id: int):
    uid  = session.get("user_id")
    role = session.get("role", "")
    pmt_date   = request.form.get("payment_date", "").strip()
    pmt_amount = request.form.get("payment_amount", "").strip()
    pmt_note   = request.form.get("payment_note", "").strip()

    amount_cents = money_to_cents(pmt_amount)
    if not amount_cents or amount_cents <= 0:
        flash("Payment amount is required.", "danger")
        return redirect(url_for("job_detail", job_id=job_id))
    if not pmt_date:
        pmt_date = datetime.now(_melbourne).strftime("%Y-%m-%d")
    else:
        try:
            pmt_date = datetime.strptime(pmt_date, "%d/%m/%Y").strftime("%Y-%m-%d")
        except Exception:
            try:
                pmt_date = datetime.strptime(pmt_date, "%Y-%m-%d").strftime("%Y-%m-%d")
            except Exception:
                pmt_date = datetime.now(_melbourne).strftime("%Y-%m-%d")

    ts = now_ts()
    conn = db()
    cur  = conn.cursor()

    cur.execute(
        "INSERT INTO job_payments (job_id, payment_date, amount_cents, note, recorded_by_user_id, created_at) VALUES (?,?,?,?,?,?)",
        (job_id, pmt_date, amount_cents, pmt_note or None, uid, ts)
    )

    pmt_display_date = datetime.strptime(pmt_date, "%Y-%m-%d").strftime("%d/%m/%Y")
    note_body = f"Payment received: ${amount_cents/100:.2f} on {pmt_display_date}"
    if pmt_note:
        note_body += f" — {pmt_note}"
    cur.execute(
        "INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at) VALUES (?,?,?,?)",
        (job_id, uid, note_body, ts)
    )

    conn.commit()
    conn.close()
    flash("Payment recorded.", "success")
    return redirect(url_for("job_detail", job_id=job_id, _anchor="lender-payments"))


@app.post("/jobs/<int:job_id>/payments/<int:payment_id>/delete")
@login_required
@admin_required
def job_payment_delete(job_id: int, payment_id: int):
    conn = db()
    conn.execute("DELETE FROM job_payments WHERE id=? AND job_id=?", (payment_id, job_id))
    conn.commit()
    conn.close()
    flash("Payment removed.", "success")
    return redirect(url_for("job_detail", job_id=job_id, _anchor="lender-payments"))


@app.post("/jobs/<int:job_id>/interactions/new")
@login_required
@admin_required
def interaction_add(job_id: int):
    event_type = request.form.get("event_type", "Note").strip()
    narrative = request.form.get("narrative", "").strip()
    interaction_date = request.form.get("interaction_date", "").strip()
    interaction_time = request.form.get("interaction_time", "").strip()

    if not narrative:
        flash("Narrative text is required.", "danger")
        return redirect(url_for("job_detail", job_id=job_id))

    occurred_at = parse_interaction_datetime(interaction_date, interaction_time)

    photo_path = None
    photo = request.files.get("attendance_photo")
    if photo and photo.filename:
        ext = photo.filename.rsplit(".", 1)[-1].lower() if "." in photo.filename else ""
        if ext in {"png", "jpg", "jpeg", "webp", "heic"}:
            stored_name = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{secure_filename(photo.filename)}"
            blob_name   = f"interactions/{job_id}/{stored_name}"
            upload_to_blob(photo, blob_name)
            photo_path = blob_name

    now = datetime.now().isoformat(timespec="seconds")

    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT status FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    current_status = job_row["status"] if job_row else ""

    cur.execute("""
        INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at, photo_path)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (job_id, event_type, narrative, occurred_at, now, photo_path))

    advanced = maybe_auto_advance_status(
        cur, job_id, current_status, event_type, session.get("role", "")
    )
    cur.execute("UPDATE jobs SET updated_at = ? WHERE id = ?", (now, job_id))
    conn.commit()
    conn.close()

    if advanced:
        flash("Interaction added. Job status automatically set to Active.", "success")
    else:
        flash("Interaction added.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


# -------- Clients --------
@app.get("/clients")
@login_required
@admin_required
def clients_list():
    q = request.args.get("q", "").strip()
    conn = db()
    cur = conn.cursor()
    if q:
        like = f"%{q}%"
        cur.execute("""
            SELECT DISTINCT c.* FROM clients c
            LEFT JOIN contact_phone_numbers cp ON cp.entity_type='client' AND cp.entity_id=c.id
            LEFT JOIN contact_emails ce ON ce.entity_type='client' AND ce.entity_id=c.id
            WHERE c.name LIKE ? COLLATE NOCASE
               OR c.nickname LIKE ? COLLATE NOCASE
               OR c.phone LIKE ? COLLATE NOCASE
               OR c.email LIKE ? COLLATE NOCASE
               OR c.address LIKE ? COLLATE NOCASE
               OR c.notes LIKE ? COLLATE NOCASE
               OR cp.phone_number LIKE ? COLLATE NOCASE
               OR ce.email LIKE ? COLLATE NOCASE
            ORDER BY c.name
        """, (like, like, like, like, like, like, like, like))
    else:
        cur.execute("SELECT * FROM clients ORDER BY name")
    rows = cur.fetchall()
    conn.close()
    if request.headers.get("X-Requested-With") == "search":
        results = []
        for c in rows:
            results.append({"id": c["id"], "name": c["name"],
                            "email": c["email"] or "", "address": c["address"] or ""})
        return jsonify(results)
    return render_template("clients.html", clients=rows, search_q=q)



@app.post("/clients/delete")
@login_required
@admin_required
def clients_delete():
    ids = request.form.getlist("client_ids")
    if not ids:
        flash("No clients selected.", "warning")
        return redirect(url_for("clients_list"))

    conn = db()
    cur = conn.cursor()
    deleted, blocked = [], []
    for cid in ids:
        try:
            cid = int(cid)
        except ValueError:
            continue
        cur.execute("SELECT name FROM clients WHERE id = ?", (cid,))
        row = cur.fetchone()
        if not row:
            continue
        name = row["name"]
        cur.execute(
            "SELECT COUNT(*) cnt FROM jobs WHERE client_id = ? OR bill_to_client_id = ?",
            (cid, cid)
        )
        job_count = cur.fetchone()["cnt"]
        if job_count > 0:
            blocked.append(f"'{name}' ({job_count} job{'s' if job_count != 1 else ''})")
        else:
            cur.execute("DELETE FROM clients WHERE id = ?", (cid,))
            deleted.append(name)
    conn.commit()
    conn.close()

    if deleted:
        flash(f"Deleted: {', '.join(deleted)}.", "success")
    if blocked:
        flash(
            f"Could not delete — client(s) with existing jobs: {', '.join(blocked)}.",
            "warning"
        )
    return redirect(url_for("clients_list"))

@app.get("/clients/new")
@login_required
@admin_required
def client_new():
    next_url = request.args.get("next", "")
    return render_template("client_new.html", next_url=next_url)


@app.post("/clients/new")
@login_required
@admin_required
def client_create():
    name = request.form.get("name", "").strip()
    phone = request.form.get("phone", "").strip()
    email = request.form.get("email", "").strip()
    address = request.form.get("address", "").strip()
    notes = request.form.get("notes", "").strip()
    if not name:
        flash("Client name is required.", "danger")
        return redirect(url_for("client_new"))

    now = datetime.now().isoformat(timespec="seconds")
    conn = db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO clients (name, phone, email, address, notes, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (name, phone, email, address, notes, now))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    flash("Client created.", "success")
    next_url = request.form.get("next_url", "")
    if next_url:
        return redirect(f"{next_url}?new_client_id={new_id}")
    return redirect(url_for("clients_list"))



@app.get("/clients/new-popup")
@login_required
@admin_required
def client_new_popup():
    return render_template("partials/client_popup.html")


@app.post("/clients/new-popup")
@login_required
@admin_required
def client_create_popup():
    name = request.form.get("name", "").strip()
    if not name:
        return jsonify({"ok": False, "error": "Client name is required."})
    phone = request.form.get("phone", "").strip()
    email = request.form.get("email", "").strip()
    now = datetime.now().isoformat(timespec="seconds")
    conn = db()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO clients (name, phone, email, created_at) VALUES (?, ?, ?, ?)",
        (name, phone or None, email or None, now)
    )
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "id": new_id, "label": name})


@app.get("/customers/new-popup")
@login_required
@admin_required
def customer_new_popup():
    return render_template("partials/customer_popup.html")


@app.post("/customers/new-popup")
@login_required
@admin_required
def customer_create_popup():
    first_name = request.form.get("first_name", "").strip()
    last_name = request.form.get("last_name", "").strip()
    if not first_name or not last_name:
        return jsonify({"ok": False, "error": "First and last name are required."})
    company = request.form.get("company", "").strip()
    address = request.form.get("address", "").strip()
    if not address:
        return jsonify({"ok": False, "error": "Address is required."})

    phone_labels  = request.form.getlist("phone_label[]")
    phone_numbers = request.form.getlist("phone_number[]")
    email_labels  = request.form.getlist("email_label[]")
    email_addrs   = request.form.getlist("email_address[]")

    phones = [(lbl.strip(), num.strip()) for lbl, num in zip(phone_labels, phone_numbers) if num.strip()]
    emails = [(lbl.strip(), addr.strip()) for lbl, addr in zip(email_labels, email_addrs) if addr.strip()]

    first_email = emails[0][1] if emails else None

    now = datetime.now().isoformat(timespec="seconds")
    conn = db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO customers (first_name, last_name, company, email, address, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (first_name, last_name, company or None, first_email, address or None, now, now))
    new_id = cur.lastrowid

    for lbl, num in phones:
        cur.execute("""
            INSERT INTO contact_phone_numbers (entity_type, entity_id, label, phone_number, created_at)
            VALUES ('customer', ?, ?, ?, ?)
        """, (new_id, lbl or "Mobile", num, now))

    for lbl, addr in emails:
        cur.execute("""
            INSERT INTO contact_emails (entity_type, entity_id, label, email, created_at)
            VALUES ('customer', ?, ?, ?, ?)
        """, (new_id, lbl or "Email", addr, now))

    conn.commit()
    conn.close()
    label = f"{first_name} {last_name}"
    if company:
        label += f" ({company})"
    return jsonify({"ok": True, "id": new_id, "label": label, "address": address})

@app.get("/clients/<int:client_id>")
@login_required
def client_detail(client_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM clients WHERE id = ?", (client_id,))
    client = cur.fetchone()
    if not client:
        conn.close()
        return ("Not found", 404)
    cur.execute("""
        SELECT j.*,
               COALESCE(
                   (SELECT u2.full_name FROM schedules s2
                    JOIN users u2 ON u2.id = s2.assigned_to_user_id
                    WHERE s2.job_id = j.id AND s2.status NOT IN ('Cancelled', 'Completed')
                    ORDER BY s2.scheduled_for ASC LIMIT 1),
                   u.full_name
               ) AS agent_name,
               cu.last_name  AS cust_last_name,
               cu.company    AS cust_company
        FROM jobs j
        LEFT JOIN users     u  ON u.id  = j.assigned_user_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        WHERE j.client_id = ?
        ORDER BY j.created_at DESC
    """, (client_id,))
    jobs = cur.fetchall()
    cur.execute("""
        SELECT * FROM contact_phone_numbers
        WHERE entity_type = 'client' AND entity_id = ?
        ORDER BY id
    """, (client_id,))
    phones = cur.fetchall()
    conn.close()
    return render_template("client_detail.html", client=client, jobs=jobs, phones=phones)


@app.get("/clients/<int:client_id>/edit")
@login_required
@admin_required
def client_edit(client_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM clients WHERE id = ?", (client_id,))
    client = cur.fetchone()
    if not client:
        conn.close()
        return ("Not found", 404)
    cur.execute("""
        SELECT * FROM contact_phone_numbers
        WHERE entity_type = 'client' AND entity_id = ?
        ORDER BY id
    """, (client_id,))
    phones = cur.fetchall()
    conn.close()
    return render_template("client_edit.html", client=client, phones=phones)


@app.post("/clients/<int:client_id>/edit")
@login_required
@admin_required
def client_edit_post(client_id: int):
    name     = request.form.get("name", "").strip()
    nickname = request.form.get("nickname", "").strip() or None
    email    = request.form.get("email", "").strip()
    address  = request.form.get("address", "").strip()
    notes    = request.form.get("notes", "").strip()

    is_xhr = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    if not name:
        if is_xhr:
            return jsonify({"ok": False, "error": "Client name is required."})
        flash("Client name is required.", "danger")
        return redirect(url_for("client_edit", client_id=client_id))

    ts = now_ts()
    conn = db()
    cur = conn.cursor()

    cur.execute("""
        UPDATE clients
        SET name = ?, nickname = ?, email = ?, address = ?, notes = ?, updated_at = ?
        WHERE id = ?
    """, (name, nickname, email, address, notes, ts, client_id))

    cur.execute("""
        DELETE FROM contact_phone_numbers
        WHERE entity_type = 'client' AND entity_id = ?
    """, (client_id,))

    for label, field in [("Mobile", "phone_mobile"), ("Home", "phone_home"),
                         ("Work", "phone_work"), ("Other", "phone_other")]:
        number = request.form.get(field, "").strip()
        if number:
            cur.execute("""
                INSERT INTO contact_phone_numbers
                    (entity_type, entity_id, label, phone_number, created_at)
                VALUES ('client', ?, ?, ?, ?)
            """, (client_id, label, number, ts))

    conn.commit()
    conn.close()

    audit("client", client_id, "update", "Client details updated", {})

    if is_xhr:
        return jsonify({"ok": True, "name": name, "nickname": nickname, "email": email})

    flash("Client updated.", "success")
    return redirect(url_for("client_detail", client_id=client_id))


@app.get("/api/clients/<int:client_id>")
@login_required
@admin_required
def api_client_get(client_id: int):
    conn = db()
    client = conn.execute("SELECT * FROM clients WHERE id = ?", (client_id,)).fetchone()
    if not client:
        conn.close()
        return jsonify({"error": "Not found"}), 404
    phones = conn.execute(
        "SELECT label, phone_number FROM contact_phone_numbers WHERE entity_type='client' AND entity_id=? ORDER BY id",
        (client_id,)
    ).fetchall()
    conn.close()
    pm = {r["label"]: r["phone_number"] for r in phones}
    return jsonify({
        "id":       client["id"],
        "name":     client["name"] or "",
        "nickname": client["nickname"] or "",
        "email":    client["email"] or "",
        "address":  client["address"] or "",
        "notes":    client["notes"] or "",
        "phone_mobile": pm.get("Mobile", ""),
        "phone_home":   pm.get("Home", ""),
        "phone_work":   pm.get("Work", ""),
        "phone_other":  pm.get("Other", ""),
    })


# -------- Customers --------
@app.get("/customers")
@login_required
@admin_required
def customers_list():
    q = request.args.get("q", "").strip()
    conn = db()
    cur = conn.cursor()
    if q:
        like = f"%{q}%"
        cur.execute("""
            SELECT DISTINCT cu.* FROM customers cu
            LEFT JOIN contact_phone_numbers cp ON cp.entity_type='customer' AND cp.entity_id=cu.id
            LEFT JOIN contact_emails ce ON ce.entity_type='customer' AND ce.entity_id=cu.id
            LEFT JOIN jobs j ON j.customer_id=cu.id
            LEFT JOIN job_customers jc ON jc.customer_id=cu.id
            LEFT JOIN jobs j2 ON j2.id=jc.job_id
            WHERE cu.first_name LIKE ? COLLATE NOCASE
               OR cu.last_name LIKE ? COLLATE NOCASE
               OR (cu.first_name || ' ' || cu.last_name) LIKE ? COLLATE NOCASE
               OR (cu.last_name || ', ' || cu.first_name) LIKE ? COLLATE NOCASE
               OR cu.company LIKE ? COLLATE NOCASE
               OR cu.email LIKE ? COLLATE NOCASE
               OR cu.address LIKE ? COLLATE NOCASE
               OR cp.phone_number LIKE ? COLLATE NOCASE
               OR ce.email LIKE ? COLLATE NOCASE
               OR j.client_reference LIKE ? COLLATE NOCASE
               OR j.account_number LIKE ? COLLATE NOCASE
               OR j2.client_reference LIKE ? COLLATE NOCASE
               OR j2.account_number LIKE ? COLLATE NOCASE
            ORDER BY cu.last_name, cu.first_name
        """, (like,) * 13)
    else:
        cur.execute("SELECT * FROM customers ORDER BY last_name, first_name")
    rows = cur.fetchall()
    conn.close()
    if request.headers.get("X-Requested-With") == "search":
        results = []
        for c in rows:
            results.append({"id": c["id"], "first_name": c["first_name"],
                            "last_name": c["last_name"], "company": c["company"] or "",
                            "id_image_path": c["id_image_path"] or ""})
        return jsonify(results)
    return render_template("customers.html", customers=rows, search_q=q)


@app.post("/customers/delete")
@login_required
@admin_required
def customers_bulk_delete():
    customer_ids = request.form.getlist("customer_ids")
    if not customer_ids:
        flash("No customers selected.", "warning")
        return redirect(url_for("customers_list"))
    conn = db()
    cur = conn.cursor()
    deleted, skipped = 0, []
    for cid in customer_ids:
        try:
            cid = int(cid)
        except ValueError:
            continue
        cur.execute("SELECT COUNT(*) FROM jobs WHERE customer_id = ?", (cid,))
        if cur.fetchone()[0]:
            cur.execute("SELECT first_name, last_name FROM customers WHERE id = ?", (cid,))
            row = cur.fetchone()
            skipped.append(f"{row['first_name']} {row['last_name']}" if row else str(cid))
            continue
        cur.execute("DELETE FROM contact_phone_numbers WHERE entity_type='customer' AND entity_id=?", (cid,))
        cur.execute("DELETE FROM contact_emails WHERE entity_type='customer' AND entity_id=?", (cid,))
        cur.execute("DELETE FROM customers WHERE id=?", (cid,))
        deleted += 1
    conn.commit()
    conn.close()
    parts = []
    if deleted:
        parts.append(f"{deleted} customer(s) deleted.")
    if skipped:
        parts.append(f"{len(skipped)} skipped (has linked jobs): {', '.join(skipped)}.")
    flash(" ".join(parts) or "Nothing deleted.", "success" if deleted else "warning")
    return redirect(url_for("customers_list"))


@app.get("/customers/new")
@login_required
@admin_required
def customer_new():
    next_url = request.args.get("next", "")
    return render_template("customer_new.html", next_url=next_url)


@app.post("/customers/new")
@login_required
@admin_required
def customer_create():
    shared_company = request.form.get("shared_company", "").strip()
    first_names = request.form.getlist("first_name[]")
    last_names  = request.form.getlist("last_name[]")
    roles       = request.form.getlist("role[]")
    emails      = request.form.getlist("email[]")
    dobs        = request.form.getlist("dob[]")
    addresses   = request.form.getlist("address[]")
    notes_list  = request.form.getlist("notes[]")
    id_photos   = request.files.getlist("id_photo[]")

    if not first_names or not first_names[0].strip():
        flash("At least one person with a first and last name is required.", "danger")
        return redirect(url_for("customer_new"))

    ts = now_ts()
    conn = db()
    cur = conn.cursor()
    first_new_id = None
    created_count = 0

    for i in range(len(first_names)):
        fn = first_names[i].strip() if i < len(first_names) else ""
        ln = last_names[i].strip()  if i < len(last_names)  else ""
        if not fn or not ln:
            continue

        role    = roles[i].strip()    if i < len(roles)    else ""
        email   = emails[i].strip()   if i < len(emails)   else ""
        dob     = dobs[i].strip()     if i < len(dobs)     else ""
        address = addresses[i].strip() if i < len(addresses) else ""
        notes   = notes_list[i].strip() if i < len(notes_list) else ""

        id_image_filename = None
        id_image_path = None
        if i < len(id_photos):
            photo = id_photos[i]
            if photo and photo.filename:
                if not allowed_file(photo.filename):
                    flash(f"ID photo for {fn} {ln} must be PNG, JPG, JPEG, or WebP — skipped.", "warning")
                else:
                    safe_name = secure_filename(photo.filename)
                    safe_ts = ts.replace(":", "").replace("-", "").replace(" ", "")
                    stored_name = f"cust_{safe_ts}_{i}_{safe_name}"
                    upload_to_blob(photo, stored_name)
                    id_image_filename = safe_name
                    id_image_path = stored_name

        cur.execute("""
            INSERT INTO customers (first_name, last_name, company, role, email, dob, address, notes,
                                   id_image_filename, id_image_path, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (fn, ln, shared_company or None, role or None, email or None,
              dob or None, address or None, notes or None,
              id_image_filename, id_image_path, ts, ts))
        new_id = cur.lastrowid
        if first_new_id is None:
            first_new_id = new_id
        created_count += 1

    conn.commit()
    conn.close()

    if created_count == 0:
        flash("No valid customers to create.", "warning")
        return redirect(url_for("customer_new"))

    flash(f"{created_count} customer(s) created.", "success")
    next_url = request.form.get("next_url", "")
    if next_url:
        return redirect(f"{next_url}?new_customer_id={first_new_id}")
    return redirect(url_for("customers_list"))


@app.get("/customers/<int:customer_id>")
@login_required
def customer_detail(customer_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM customers WHERE id = ?", (customer_id,))
    customer = cur.fetchone()
    if not customer:
        conn.close()
        return ("Not found", 404)
    cur.execute("""
        SELECT j.*, c.name client_name
        FROM jobs j
        LEFT JOIN clients c ON c.id = j.client_id
        WHERE j.customer_id = ?
        ORDER BY j.created_at DESC
    """, (customer_id,))
    jobs = cur.fetchall()
    cur.execute("""
        SELECT * FROM contact_phone_numbers
        WHERE entity_type = 'customer' AND entity_id = ?
        ORDER BY id
    """, (customer_id,))
    phones = cur.fetchall()

    cur.execute("""
        SELECT * FROM contact_emails
        WHERE entity_type = 'customer' AND entity_id = ?
        ORDER BY id
    """, (customer_id,))
    emails = cur.fetchall()

    conn.close()
    return render_template("customer_detail.html", customer=customer, jobs=jobs, phones=phones, emails=emails)


@app.get("/customers/<int:customer_id>/edit")
@login_required
@admin_required
def customer_edit(customer_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM customers WHERE id = ?", (customer_id,))
    customer = cur.fetchone()
    if not customer:
        conn.close()
        return ("Not found", 404)
    cur.execute("""
        SELECT * FROM contact_phone_numbers
        WHERE entity_type = 'customer' AND entity_id = ?
        ORDER BY id
    """, (customer_id,))
    phones = cur.fetchall()
    cur.execute("""
        SELECT * FROM contact_emails
        WHERE entity_type = 'customer' AND entity_id = ?
        ORDER BY id
    """, (customer_id,))
    emails = cur.fetchall()
    conn.close()
    return render_template("customer_edit.html", customer=customer, phones=phones, emails=emails)


@app.post("/customers/<int:customer_id>/edit")
@login_required
@admin_required
def customer_edit_post(customer_id: int):
    first_name = request.form.get("first_name", "").strip()
    last_name = request.form.get("last_name", "").strip()
    company = request.form.get("company", "").strip()
    role = request.form.get("role", "").strip()
    email = request.form.get("email_personal", "").strip()
    dob = request.form.get("dob", "").strip()
    address = request.form.get("address", "").strip()
    notes = request.form.get("notes", "").strip()
    id_image = request.files.get("id_image")

    if not first_name or not last_name:
        flash("First and last name are required.", "danger")
        return redirect(url_for("customer_edit", customer_id=customer_id))

    ts = now_ts()
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT id_image_filename, id_image_path FROM customers WHERE id = ?", (customer_id,))
    existing = cur.fetchone()

    id_image_filename = existing["id_image_filename"] if existing else None
    id_image_path = existing["id_image_path"] if existing else None

    if id_image and id_image.filename:
        if not allowed_file(id_image.filename):
            conn.close()
            flash("Unsupported file type. Use PNG/JPG/PDF.", "danger")
            return redirect(url_for("customer_edit", customer_id=customer_id))
        filename = secure_filename(id_image.filename)
        safe_ts = ts.replace(":", "").replace("-", "").replace(" ", "")
        unique_name = f"customer_{customer_id}_id_{safe_ts}_{filename}"
        upload_to_blob(id_image, unique_name)
        id_image_filename = filename
        id_image_path = unique_name

    cur.execute("""
        UPDATE customers
        SET first_name = ?, last_name = ?, company = ?, role = ?, email = ?, dob = ?, address = ?, notes = ?,
            id_image_filename = ?, id_image_path = ?, updated_at = ?
        WHERE id = ?
    """, (first_name, last_name, company or None, role or None, email or None, dob or None, address or None, notes or None,
          id_image_filename, id_image_path, ts, customer_id))

    cur.execute("""
        DELETE FROM contact_phone_numbers
        WHERE entity_type = 'customer' AND entity_id = ?
    """, (customer_id,))

    for label, field in [("Mobile", "phone_mobile"), ("Home", "phone_home"),
                         ("Work", "phone_work"), ("Other", "phone_other")]:
        number = request.form.get(field, "").strip()
        if number:
            cur.execute("""
                INSERT INTO contact_phone_numbers
                    (entity_type, entity_id, label, phone_number, created_at)
                VALUES ('customer', ?, ?, ?, ?)
            """, (customer_id, label, number, ts))

    cur.execute("""
        DELETE FROM contact_emails
        WHERE entity_type = 'customer' AND entity_id = ?
    """, (customer_id,))

    for label, field in [("Personal", "email_personal"), ("Work", "email_work"),
                         ("Other", "email_other")]:
        em = request.form.get(field, "").strip()
        if em:
            cur.execute("""
                INSERT INTO contact_emails
                    (entity_type, entity_id, label, email, created_at)
                VALUES ('customer', ?, ?, ?, ?)
            """, (customer_id, label, em, ts))

    conn.commit()
    conn.close()

    audit("customer", customer_id, "update", "Customer details updated",
          {"id_image_updated": bool(id_image and id_image.filename)})

    flash("Customer updated.", "success")
    return redirect(url_for("customer_detail", customer_id=customer_id))


# -------- Customer Delete --------
@app.post("/customers/<int:customer_id>/delete")
@login_required
@admin_required
def customer_delete(customer_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM jobs WHERE customer_id = ?", (customer_id,))
    job_count = cur.fetchone()[0]
    if job_count:
        conn.close()
        flash(f"Cannot delete — this customer has {job_count} linked job(s).", "danger")
        return redirect(url_for("customer_detail", customer_id=customer_id))
    cur.execute("DELETE FROM contact_phone_numbers WHERE entity_type = 'customer' AND entity_id = ?", (customer_id,))
    cur.execute("DELETE FROM contact_emails WHERE entity_type = 'customer' AND entity_id = ?", (customer_id,))
    cur.execute("DELETE FROM customers WHERE id = ?", (customer_id,))
    conn.commit()
    conn.close()
    flash("Customer deleted.", "success")
    return redirect(url_for("customers_list"))


# -------- Job Items --------
@app.post("/jobs/<int:job_id>/items/new")
@login_required
@admin_required
def job_item_create(job_id: int):
    item_type        = request.form.get("item_type", "vehicle").strip()
    description      = request.form.get("description", "").strip()
    reg              = request.form.get("reg", "").strip()
    vin              = request.form.get("vin", "").strip()
    make             = request.form.get("make", "").strip()
    model            = request.form.get("model", "").strip()
    year             = request.form.get("year", "").strip()
    property_address = request.form.get("property_address", "").strip()
    lot_details      = request.form.get("lot_details", "").strip()
    serial_number    = request.form.get("serial_number", "").strip()
    identifier       = request.form.get("identifier", "").strip()
    notes            = request.form.get("notes", "").strip()
    item_lender      = request.form.get("item_lender_name", "").strip()
    item_account     = request.form.get("item_account_number", "").strip()
    item_regulation  = request.form.get("item_regulation_type", "").strip()
    item_deliver_to  = request.form.get("item_deliver_to", "").strip()

    conn = db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO job_items (
            job_id, item_type, description,
            reg, vin, make, model, year,
            property_address, lot_details,
            serial_number, identifier,
            notes, lender_name, account_number, regulation_type,
            deliver_to, created_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        job_id, item_type, description or None,
        reg or None, vin or None, make or None, model or None, year or None,
        property_address or None, lot_details or None,
        serial_number or None, identifier or None,
        notes or None, item_lender or None, item_account or None, item_regulation or None,
        item_deliver_to or None, now_ts()
    ))
    conn.commit()
    conn.close()

    flash("Item added.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/items/<int:item_id>/edit")
@login_required
@admin_required
def job_item_edit(job_id: int, item_id: int):
    item_type        = request.form.get("item_type", "vehicle").strip()
    description      = request.form.get("description", "").strip()
    reg              = request.form.get("reg", "").strip()
    vin              = request.form.get("vin", "").strip()
    serial_number    = request.form.get("serial_number", "").strip()
    identifier       = request.form.get("identifier", "").strip()
    property_address = request.form.get("property_address", "").strip()
    lot_details      = request.form.get("lot_details", "").strip()
    notes            = request.form.get("notes", "").strip()
    item_lender      = request.form.get("item_lender_name", "").strip()
    item_account     = request.form.get("item_account_number", "").strip()
    item_regulation  = request.form.get("item_regulation_type", "").strip()
    item_deliver_to  = request.form.get("item_deliver_to", "").strip()
    conn = db()
    cur  = conn.cursor()
    cur.execute("""
        UPDATE job_items SET
            item_type=?, description=?, reg=?, vin=?,
            serial_number=?, identifier=?, property_address=?, lot_details=?,
            notes=?, lender_name=?, account_number=?, regulation_type=?,
            deliver_to=?
        WHERE id=? AND job_id=?
    """, (
        item_type, description or None, reg or None, vin or None,
        serial_number or None, identifier or None,
        property_address or None, lot_details or None,
        notes or None, item_lender or None, item_account or None,
        item_regulation or None, item_deliver_to or None,
        item_id, job_id
    ))
    conn.commit()
    conn.close()
    flash("Security item updated.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


@app.post("/jobs/<int:job_id>/items/<int:item_id>/delete")
@login_required
@admin_required
def job_item_delete(job_id: int, item_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("DELETE FROM job_items WHERE id = ? AND job_id = ?", (item_id, job_id))
    conn.commit()
    conn.close()
    flash("Item removed.", "success")
    return redirect(url_for("job_detail", job_id=job_id))


# -------- Repo Lock --------

def _repo_lock_note(d):
    def _v(k):
        v = d.get(k) or ""
        return v if v else "—"

    lines = ["REPO LOCK SUBMITTED", "=" * 44]

    ref = d.get("swpi_ref") or ""
    if ref:
        lines.append(f"Reference:  {ref}")
    if d.get("client_name"):
        lines.append(f"Client:     {d['client_name']}")
    if d.get("finance_company"):
        lines.append(f"Finance:    {d['finance_company']}")

    time_part = ""
    if d.get("start_time"):
        time_part = f"  {d['start_time']}"
        if d.get("end_time"):
            time_part += f"–{d['end_time']}"
    lines.append(f"Date:       {_v('repo_date')}{time_part}")

    asset_parts = []
    if d.get("description"):  asset_parts.append(d["description"])
    if d.get("registration"): asset_parts.append(f"Reg: {d['registration']}")
    if d.get("vin"):          asset_parts.append(f"VIN: {d['vin']}")
    if d.get("engine_number"): asset_parts.append(f"Engine: {d['engine_number']}")
    lines.append(f"Asset:      {' | '.join(asset_parts) or '—'}")

    lines.append(f"Customer:   {_v('customer_name')}")
    if d.get("account_number"):
        lines[-1] += f"  (Acct: {d['account_number']})"
    lines.append(f"Repo Addr:  {_v('repo_address')}")

    lines.append("")
    lines.append("RECOVERY:")
    keys = _v("keys_obtained")
    keys_line = f"  Keys Obtained: {keys}"
    if keys == "Yes" and d.get("how_many_keys"):
        keys_line += f" ({d['how_many_keys']} key(s))"
    lines.append(keys_line)
    lines.append(f"  Vol. Surrender: {_v('vol_surrender')}")
    form13 = _v("form_13")
    form13_line = f"  Form 13: {form13}"
    if form13 == "Yes" and d.get("form_13_signed_by"):
        form13_line += f", Signed By: {d['form_13_signed_by']}"
    lines.append(form13_line)
    lines.append(f"  Repossessed From: {_v('repossessed_from')}")
    lines.append(f"  Lien Paid: {_v('lien_paid')}")
    lines.append(f"  Drivable: {_v('security_drivable')}")
    police = _v("police_notified")
    police_line = f"  Police Notified: {police}"
    if police == "Yes" and d.get("station_officer"):
        police_line += f" | {d['station_officer']}"
    lines.append(police_line)
    effects = _v("personal_effects_removed")
    effects_line = f"  Personal Effects Removed: {effects}"
    if effects == "Yes" and d.get("removed_by_who"):
        effects_line += f", By: {d['removed_by_who']}"
    lines.append(effects_line)
    if d.get("personal_effects_list"):
        lines.append(f"    Effects: {d['personal_effects_list']}")

    cond_parts = []
    for field, label in [("tyres","Tyres"), ("body","Body"), ("duco","Duco"),
                         ("interior","Interior"), ("fuel_level","Fuel")]:
        if d.get(field):
            cond_parts.append(f"{label}: {d[field]}")
    if cond_parts:
        lines.append("")
        lines.append(f"Condition:  {' | '.join(cond_parts)}")
    if d.get("any_damage") == "Yes" and d.get("damage_list"):
        lines.append(f"Damage:     {d['damage_list']}")

    lines.append("")
    lines.append("TOW & DELIVERY:")
    tow_name = d.get("tow_company_name") or ""
    tow_cost = d.get("tow_costs") or ""
    tow_line = f"  Tow:      {tow_name or '—'}"
    if tow_cost:
        tow_line += f" | Cost: {tow_cost}"
    lines.append(tow_line)
    deliver_parts = []
    if d.get("deliver_to"):       deliver_parts.append(d["deliver_to"])
    if d.get("delivery_address"): deliver_parts.append(d["delivery_address"])
    lines.append(f"  Deliver:  {' — '.join(deliver_parts) or '—'}")
    if d.get("expected_delivery_date"):
        lines.append(f"  Exp. Delivery: {d['expected_delivery_date']}")

    if d.get("customers_intention"):
        lines.append("")
        lines.append(f"Customer's Intention: {d['customers_intention']}")
    if d.get("other_info"):
        lines.append(f"Other Info: {d['other_info']}")

    lines.append("")
    lines.append(f"Submitted by: {_v('agent_name')}")
    return "\n".join(lines)


@app.get("/jobs/<int:job_id>/repo-lock/<int:item_id>")
@login_required
def repo_lock_get(job_id: int, item_id: int):
    conn = db()
    job  = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    item = conn.execute("SELECT * FROM job_items WHERE id=? AND job_id=?", (item_id, job_id)).fetchone()
    if not job or not item:
        conn.close()
        return jsonify({"error": "Not found"}), 404

    job  = dict(job)
    item = dict(item)

    rec  = conn.execute(
        "SELECT * FROM repo_lock_records WHERE job_id=? AND item_id=? ORDER BY id DESC LIMIT 1",
        (job_id, item_id)
    ).fetchone()

    client = conn.execute("SELECT * FROM clients WHERE id=?", (job["client_id"],)).fetchone() if job["client_id"] else None
    customer = conn.execute("SELECT * FROM customers WHERE id=?", (job["customer_id"],)).fetchone() if job["customer_id"] else None
    tow_ops = conn.execute("SELECT id, company_name, phone, mobile FROM tow_operators WHERE active=1 ORDER BY company_name").fetchall()
    auction_yards = conn.execute("SELECT id, name, address FROM auction_yards WHERE active=1 ORDER BY name").fetchall()
    user = conn.execute("SELECT full_name FROM users WHERE id=?", (session.get("user_id"),)).fetchone()
    conn.close()

    if customer:
        customer = dict(customer)

    from datetime import date as _date
    today = _date.today().strftime("%Y-%m-%d")

    prefill = {
        "client_name":      (dict(client)["name"] if client else ""),
        "client_reference": job.get("client_job_number") or "",
        "swpi_ref":         job.get("internal_job_number") or job.get("display_ref") or "",
        "finance_company":  job.get("lender_name") or item.get("lender_name") or "",
        "repo_date":        today,
        "account_number":   job.get("account_number") or item.get("account_number") or "",
        "description":      item.get("description") or "",
        "registration":     item.get("reg") or "",
        "vin":              item.get("vin") or "",
        "engine_number":    item.get("engine_number") or "",
        "deliver_to":       job.get("deliver_to") or "",
        "agent_name":       (user["full_name"] if user else ""),
        "agent_user_id":    session.get("user_id") or "",
    }
    if customer:
        name_parts = []
        if customer.get("first_name"): name_parts.append(customer["first_name"])
        if customer.get("last_name"):  name_parts.append(customer["last_name"])
        prefill["customer_name"] = " ".join(name_parts)
        prefill["repo_address"]  = customer.get("address") or ""

    existing = {}
    rec_status = None
    queue_id   = None
    if rec:
        existing   = dict(rec)
        rec_status = existing.get("status") or "Draft"

    conn2 = db()
    q_row = conn2.execute(
        "SELECT id FROM repo_lock_queue WHERE job_id=? AND item_id=? LIMIT 1",
        (job_id, item_id)
    ).fetchone()
    conn2.close()
    if q_row:
        queue_id = q_row["id"]

    return jsonify({
        "ok":        True,
        "has_rec":   rec is not None,
        "status":    rec_status,
        "queue_id":  queue_id,
        "prefill":   prefill,
        "existing":  existing,
        "tow_ops":   [{"id": t["id"], "name": t["company_name"], "phone": t["phone"] or t["mobile"] or ""} for t in tow_ops],
        "auction_yards": [{"id": y["id"], "name": y["name"], "address": y["address"] or ""} for y in auction_yards],
        "item_label": (item["reg"] or item["description"] or f"Item #{item_id}"),
    })


@app.post("/jobs/<int:job_id>/repo-lock/<int:item_id>/save")
@login_required
def repo_lock_save(job_id: int, item_id: int):
    conn = db()
    item = conn.execute("SELECT id FROM job_items WHERE id=? AND job_id=?", (item_id, job_id)).fetchone()
    if not item:
        conn.close()
        return jsonify({"error": "Not found"}), 404

    f = request.form
    fields = [
        "client_name", "client_reference", "swpi_ref", "finance_company",
        "repo_date", "start_time", "end_time",
        "customer_name", "account_number", "repo_address", "contact_number",
        "description", "registration", "rego_expiry", "registered", "insured", "insured_with",
        "vin", "engine_number", "speedometer",
        "person_present", "keys_obtained", "how_many_keys", "vol_surrender",
        "form_13", "form_13_signed_by", "repossessed_from", "lien_paid", "security_drivable",
        "police_notified", "station_officer",
        "personal_effects_removed", "removed_by_who", "personal_effects_list",
        "tyres", "body", "duco", "interior", "engine_condition", "transmission",
        "fuel_level", "any_damage", "damage_list",
        "tow_company_id", "tow_company_name", "tow_costs",
        "deliver_to", "delivery_address", "expected_delivery_date",
        "customers_intention", "other_info", "agent_name", "agent_user_id",
    ]

    values = {fld: (f.get(fld) or "").strip() or None for fld in fields}
    ts = now_ts()
    uid = session.get("user_id")

    existing = conn.execute(
        "SELECT id FROM repo_lock_records WHERE job_id=? AND item_id=? ORDER BY id DESC LIMIT 1",
        (job_id, item_id)
    ).fetchone()

    if existing:
        set_clause = ", ".join(f"{k}=?" for k in fields)
        set_clause += ", status='Draft', updated_by_user_id=?, updated_at=?"
        params = [values[k] for k in fields] + [uid, ts, existing["id"]]
        conn.execute(f"UPDATE repo_lock_records SET {set_clause} WHERE id=?", params)
        rec_id = existing["id"]
        is_new = False
    else:
        cols = ", ".join(fields) + ", job_id, item_id, status, created_by_user_id, created_at, updated_by_user_id, updated_at"
        placeholders = ", ".join("?" for _ in fields) + ", ?, ?, ?, ?, ?, ?, ?"
        params = [values[k] for k in fields] + [job_id, item_id, "Draft", uid, ts, uid, ts]
        cur = conn.execute(f"INSERT INTO repo_lock_records ({cols}) VALUES ({placeholders})", params)
        rec_id = cur.lastrowid
        is_new = True

    conn.commit()
    conn.close()
    return jsonify({"ok": True, "id": rec_id, "is_new": is_new, "status": "Draft"})


_RL_FIELDS = [
    "client_name", "client_reference", "swpi_ref", "finance_company",
    "repo_date", "start_time", "end_time",
    "customer_name", "account_number", "repo_address", "contact_number",
    "description", "registration", "rego_expiry", "registered", "insured", "insured_with",
    "vin", "engine_number", "speedometer",
    "person_present", "keys_obtained", "how_many_keys", "vol_surrender",
    "form_13", "form_13_signed_by", "repossessed_from", "lien_paid", "security_drivable",
    "police_notified", "station_officer",
    "personal_effects_removed", "removed_by_who", "personal_effects_list",
    "tyres", "body", "duco", "interior", "engine_condition", "transmission",
    "fuel_level", "any_damage", "damage_list",
    "tow_company_id", "tow_company_name", "tow_costs",
    "deliver_to", "delivery_address", "expected_delivery_date",
    "customers_intention", "other_info", "agent_name", "agent_user_id",
]


@app.post("/jobs/<int:job_id>/repo-lock/<int:item_id>/submit")
@login_required
def repo_lock_submit(job_id: int, item_id: int):
    conn = db()
    item = conn.execute("SELECT id FROM job_items WHERE id=? AND job_id=?", (item_id, job_id)).fetchone()
    if not item:
        conn.close()
        return jsonify({"error": "Not found"}), 404

    f = request.form
    values = {fld: (f.get(fld) or "").strip() or None for fld in _RL_FIELDS}
    ts  = now_ts()
    uid = session.get("user_id")

    errors = []
    if not values.get("repo_date"):
        errors.append("Repo Date is required.")
    if not values.get("agent_name"):
        errors.append("Agent name is required.")
    if not values.get("registration") and not values.get("description"):
        errors.append("Asset registration or description is required.")
    if errors:
        conn.close()
        return jsonify({"ok": False, "errors": errors}), 400

    existing = conn.execute(
        "SELECT id FROM repo_lock_records WHERE job_id=? AND item_id=? ORDER BY id DESC LIMIT 1",
        (job_id, item_id)
    ).fetchone()

    if existing:
        set_clause = ", ".join(f"{k}=?" for k in _RL_FIELDS)
        set_clause += ", status='Submitted', submitted_at=?, updated_by_user_id=?, updated_at=?"
        params = [values[k] for k in _RL_FIELDS] + [ts, uid, ts, existing["id"]]
        conn.execute(f"UPDATE repo_lock_records SET {set_clause} WHERE id=?", params)
        rec_id = existing["id"]
        is_new = False
    else:
        cols = ", ".join(_RL_FIELDS) + ", job_id, item_id, status, submitted_at, created_by_user_id, created_at, updated_by_user_id, updated_at"
        placeholders = ", ".join("?" for _ in _RL_FIELDS) + ", ?, ?, ?, ?, ?, ?, ?, ?"
        params = [values[k] for k in _RL_FIELDS] + [job_id, item_id, "Submitted", ts, uid, ts, uid, ts]
        cur = conn.execute(f"INSERT INTO repo_lock_records ({cols}) VALUES ({placeholders})", params)
        rec_id = cur.lastrowid
        is_new = True

    note_text = _repo_lock_note(values)
    conn.execute(
        "INSERT INTO interactions (job_id, event_type, narrative, occurred_at, created_at) VALUES (?,?,?,?,?)",
        (job_id, "Repo Lock Submitted", note_text, ts, ts)
    )
    conn.execute(
        "INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at) VALUES (?,?,?,?)",
        (job_id, uid, note_text, ts)
    )

    existing_q = conn.execute(
        "SELECT id, submission_count FROM repo_lock_queue WHERE job_id=? AND item_id=?",
        (job_id, item_id)
    ).fetchone()

    if existing_q:
        conn.execute(
            """UPDATE repo_lock_queue SET repo_lock_id=?, status='Pending',
               submission_count=?, submitted_at=?, submitted_by_user_id=?, updated_at=?
               WHERE id=?""",
            (rec_id, (existing_q["submission_count"] + 1), ts, uid, ts, existing_q["id"])
        )
        queue_id = existing_q["id"]
    else:
        qcur = conn.execute(
            """INSERT INTO repo_lock_queue
               (job_id, item_id, repo_lock_id, status, submission_count,
                submitted_at, submitted_by_user_id, created_at, updated_at)
               VALUES (?,?,?,'Pending',1,?,?,?,?)""",
            (job_id, item_id, rec_id, ts, uid, ts, ts)
        )
        queue_id = qcur.lastrowid

    conn.commit()
    conn.close()
    return jsonify({"ok": True, "id": rec_id, "is_new": is_new,
                    "queue_id": queue_id, "status": "Submitted"})


def _attach_pdf_to_job(conn, job_id: int, user_id, pdf_bytes: bytes,
                        original_filename: str, form_label: str, ts: str = None):
    """
    Save a generated PDF to job_documents and create an auto note.
    Does NOT call conn.commit() — caller is responsible.
    """
    ts = ts or now_ts()
    ts_safe = ts.replace(":", "").replace("-", "").replace(" ", "")
    stored_filename = f"job_{job_id}_{ts_safe}_{original_filename}"
    file_size = _save_bytes_to_storage(pdf_bytes, stored_filename, "application/pdf")

    cur = conn.cursor()
    cur.execute("""
        INSERT INTO job_documents
            (job_id, doc_type, title, original_filename, stored_filename,
             mime_type, file_size, uploaded_by_user_id, uploaded_at, notes)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (job_id, "Generated Form", form_label, original_filename, stored_filename,
          "application/pdf", file_size, user_id, ts, "Auto-generated by Forms module"))
    doc_id = cur.lastrowid

    note_text = f"{form_label} generated and attached to file."
    cur.execute("""
        INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at)
        VALUES (?, ?, ?, ?)
    """, (job_id, user_id, note_text, ts))
    note_id = cur.lastrowid

    cur.execute("""
        INSERT INTO job_note_files (job_field_note_id, filename, filepath, uploaded_at)
        VALUES (?, ?, ?, ?)
    """, (note_id, original_filename, stored_filename, ts))


def _rl_pdf_context(conn, rec, job_id):
    """Build merged data dict for PDF generation from all related records."""
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    job = dict(job) if job else {}

    item = {}
    if rec.get("item_id"):
        r = conn.execute("SELECT * FROM job_items WHERE id=?", (rec["item_id"],)).fetchone()
        if r:
            item = dict(r)

    client = {}
    if job.get("client_id"):
        r = conn.execute("SELECT * FROM clients WHERE id=?", (job["client_id"],)).fetchone()
        if r:
            client = dict(r)

    customer = {}
    if job.get("customer_id"):
        r = conn.execute("SELECT * FROM customers WHERE id=?", (job["customer_id"],)).fetchone()
        if r:
            customer = dict(r)

    tow_op = {}
    if rec.get("tow_company_id"):
        r = conn.execute("SELECT * FROM tow_operators WHERE id=?",
                         (rec["tow_company_id"],)).fetchone()
        if r:
            tow_op = dict(r)

    assigned_agent = ""
    if job.get("assigned_user_id"):
        r = conn.execute("SELECT full_name FROM users WHERE id=?",
                         (job["assigned_user_id"],)).fetchone()
        if r:
            assigned_agent = r["full_name"]

    cust_name = " ".join(filter(None, [
        customer.get("first_name"), customer.get("last_name")
    ])) or rec.get("customer_name") or ""

    agent_name = (rec.get("agent_name") or assigned_agent
                  or session.get("full_name") or "")

    d = {}
    d.update({k: (v if v else "") for k, v in rec.items()})
    d.update({
        "year":                 item.get("year") or "",
        "make":                 item.get("make") or "",
        "model":                item.get("model") or "",
        "colour":               item.get("colour") or rec.get("colour") or "",
        "client_name":          rec.get("client_name") or client.get("name") or "",
        "client_reference":     rec.get("client_reference") or job.get("client_job_number") or "",
        "swpi_ref":             rec.get("swpi_ref") or job.get("internal_job_number") or "",
        "finance_company":      rec.get("finance_company") or job.get("lender_name") or "",
        "customer_name":        cust_name,
        "account_number":       rec.get("account_number") or job.get("account_number") or "",
        "repo_address":         rec.get("repo_address") or customer.get("address") or job.get("job_address") or "",
        "deliver_to":           rec.get("deliver_to") or job.get("deliver_to") or "",
        "agent_name":           agent_name,
        "tow_company_name":     rec.get("tow_company_name") or tow_op.get("company_name") or "",
        "tow_company_name_db":  tow_op.get("company_name") or "",
        "tow_phone":            tow_op.get("phone") or tow_op.get("mobile") or "",
        "client_email":         client.get("email") or "",
        "item_make":            item.get("make") or "",
        "item_model":           item.get("model") or "",
        "item_year":            item.get("year") or "",
    })
    return d, agent_name, client, tow_op


@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/vir")
@login_required
def repo_lock_vir(job_id: int, rec_id: int):
    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)
    conn.close()
    return render_template("repo_lock_vir.html",
                           rec=rec, job_id=job_id,
                           agent_name=agent_name,
                           item_year=d.get("year",""), item_make=d.get("make",""),
                           item_model=d.get("model",""))


@app.post("/jobs/<int:job_id>/repo-lock/<int:rec_id>/vir")
@login_required
def repo_lock_vir_pdf(job_id: int, rec_id: int):
    from flask import send_file
    import pdf_gen as _pg

    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)

    agent_sig    = request.form.get("agent_sig", "").strip() or None
    customer_sig = request.form.get("customer_sig", "").strip() or None

    if not agent_sig:
        flash("Agent signature is required.", "error")
        conn.close()
        return redirect(url_for("repo_lock_vir", job_id=job_id, rec_id=rec_id))

    ts = now_ts()

    # Ensure optional columns exist
    for col in ("station_officer", "personal_effects_removed", "personal_effects_list",
                "colour", "make", "model", "year"):
        add_column_if_missing(conn, "repo_lock_records", col, "TEXT")

    def _f(name):
        return request.form.get(name, "").strip() or rec.get(name) or ""

    conn.execute("""UPDATE repo_lock_records
                    SET swpi_ref=?, finance_company=?, repo_date=?, account_number=?,
                        customer_name=?, repo_address=?,
                        registration=?, rego_expiry=?, vin=?, engine_number=?, speedometer=?,
                        colour=?, make=?, model=?, year=?,
                        person_present=?, keys_obtained=?, how_many_keys=?, vol_surrender=?,
                        form_13=?, security_drivable=?, police_notified=?, station_officer=?,
                        personal_effects_removed=?, personal_effects_list=?,
                        tyres=?, body=?, duco=?, interior=?, engine_condition=?,
                        transmission=?, fuel_level=?, any_damage=?, damage_list=?,
                        agent_name=?, updated_at=?
                    WHERE id=?""",
                 (_f("swpi_ref"), _f("finance_company"), _f("repo_date"), _f("account_number"),
                  _f("customer_name"), _f("repo_address"),
                  _f("registration"), _f("rego_expiry"), _f("vin"), _f("engine_number"), _f("speedometer"),
                  _f("colour"), _f("make"), _f("model"), _f("year"),
                  _f("person_present"), _f("keys_obtained"), _f("how_many_keys"), _f("vol_surrender"),
                  _f("form_13"), _f("security_drivable"), _f("police_notified"), _f("station_officer"),
                  _f("personal_effects_removed"), _f("personal_effects_list"),
                  _f("tyres"), _f("body"), _f("duco"), _f("interior"), _f("engine_condition"),
                  _f("transmission"), _f("fuel_level"), _f("any_damage"), _f("damage_list"),
                  _f("agent_name"), ts,
                  rec_id))
    conn.commit()

    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=?", (rec_id,)).fetchone()
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)

    # Signatures are session-only — pass directly from form, do not persist
    pdf_bytes = _pg.generate_vir_pdf(d, agent_sig=agent_sig, customer_sig=customer_sig)

    job_num   = (d.get("swpi_ref") or str(job_id)).replace("/", "-")
    date_str  = ts[:10].replace("-", "-") if ts else ""
    from datetime import datetime as _dt
    date_str  = _dt.now().strftime("%d-%m-%Y")
    form_label = "SWPI VIR"
    orig_filename = f"{job_num} - {form_label} - {date_str}.pdf"
    _attach_pdf_to_job(conn, job_id, session.get("user_id"), pdf_bytes,
                       orig_filename, form_label, ts)
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(pdf_bytes),
                     mimetype="application/pdf",
                     as_attachment=True,
                     download_name=orig_filename)


@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/transport-instructions")
@login_required
def repo_lock_transport(job_id: int, rec_id: int):
    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)
    conn.close()
    return render_template("repo_lock_transport.html",
                           rec=rec, job_id=job_id,
                           agent_name=agent_name,
                           item_make=d.get("make",""), item_model=d.get("model",""),
                           tow_company_name_db=d.get("tow_company_name_db",""),
                           tow_phone=d.get("tow_phone",""),
                           client_name=d.get("client_name",""),
                           client_email=d.get("client_email",""))


@app.post("/jobs/<int:job_id>/repo-lock/<int:rec_id>/transport-instructions")
@login_required
def repo_lock_transport_pdf(job_id: int, rec_id: int):
    from flask import send_file
    import pdf_gen as _pg

    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)

    agent_sig = request.form.get("agent_sig", "").strip() or None
    tow_sig   = request.form.get("tow_sig",   "").strip() or None

    if not agent_sig:
        flash("Agent signature is required.", "error")
        conn.close()
        return redirect(url_for("repo_lock_transport", job_id=job_id, rec_id=rec_id))

    for col in ("make", "model", "tow_phone"):
        add_column_if_missing(conn, "repo_lock_records", col, "TEXT")

    def _f(name):
        return request.form.get(name, "").strip() or rec.get(name) or ""

    ts = now_ts()

    conn.execute("""UPDATE repo_lock_records
                    SET swpi_ref=?, finance_company=?, repo_date=?,
                        customer_name=?, repo_address=?,
                        make=?, model=?, registration=?, vin=?,
                        tow_company_name=?, tow_phone=?, tow_costs=?,
                        deliver_to=?, delivery_address=?, updated_at=?
                    WHERE id=?""",
                 (_f("swpi_ref"), _f("finance_company"), _f("repo_date"),
                  _f("customer_name"), _f("repo_address"),
                  _f("make"), _f("model"), _f("registration"), _f("vin"),
                  _f("tow_company_name"), _f("tow_phone"), _f("tow_costs"),
                  _f("deliver_to"), _f("delivery_address"), ts, rec_id))
    conn.commit()

    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=?", (rec_id,)).fetchone()
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)

    # Signatures are session-only — use directly from form, do not persist
    pdf_bytes = _pg.generate_transport_pdf(d, agent_sig=agent_sig, tow_sig=tow_sig)

    job_num    = (d.get("swpi_ref") or str(job_id)).replace("/", "-")
    from datetime import datetime as _dt
    date_str   = _dt.now().strftime("%d-%m-%Y")
    form_label = "Transport Instructions"
    orig_filename = f"{job_num} - {form_label} - {date_str}.pdf"
    _attach_pdf_to_job(conn, job_id, session.get("user_id"), pdf_bytes,
                       orig_filename, form_label, ts)
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(pdf_bytes),
                     mimetype="application/pdf",
                     as_attachment=True,
                     download_name=orig_filename)


# ─────────────────────────── Wise VIR ──────────────────────────────────

@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/wise-vir")
@login_required
def repo_lock_wise_vir(job_id: int, rec_id: int):
    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)
    conn.close()
    return render_template("wise_vir.html",
                           rec=rec, job_id=job_id,
                           agent_name=agent_name,
                           item_year=d.get("year", ""), item_make=d.get("make", ""),
                           item_model=d.get("model", ""),
                           tow_company_name_db=d.get("tow_company_name_db", ""))


@app.post("/jobs/<int:job_id>/repo-lock/<int:rec_id>/wise-vir")
@login_required
def repo_lock_wise_vir_pdf(job_id: int, rec_id: int):
    from flask import send_file
    import pdf_gen as _pg

    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)

    agent_sig    = request.form.get("agent_sig", "").strip() or None
    customer_sig = request.form.get("customer_sig", "").strip() or None

    if not agent_sig:
        flash("Agent signature is required.", "error")
        conn.close()
        return redirect(url_for("repo_lock_wise_vir", job_id=job_id, rec_id=rec_id))

    for col in ("body_type", "bumpers", "glass", "accessories", "make", "model", "year", "colour"):
        add_column_if_missing(conn, "repo_lock_records", col, "TEXT")

    def _f(name):
        return request.form.get(name, "").strip() or rec.get(name) or ""

    ts = now_ts()
    conn.execute("""UPDATE repo_lock_records
                    SET customer_name=?,
                        year=?, make=?, model=?, body_type=?, colour=?,
                        registration=?, vin=?, engine_number=?,
                        body=?, duco=?, bumpers=?, glass=?, tyres=?,
                        security_drivable=?, engine_condition=?, interior=?,
                        speedometer=?, keys_obtained=?, accessories=?, damage_list=?,
                        tow_company_name=?, tow_costs=?, deliver_to=?, vol_surrender=?,
                        updated_at=?
                    WHERE id=?""",
                 (_f("customer_name"),
                  _f("year"), _f("make"), _f("model"), _f("body_type"), _f("colour"),
                  _f("registration"), _f("vin"), _f("engine_number"),
                  _f("body"), _f("duco"), _f("bumpers"), _f("glass"), _f("tyres"),
                  _f("security_drivable"), _f("engine_condition"), _f("interior"),
                  _f("speedometer"), _f("keys_obtained"), _f("accessories"), _f("damage_list"),
                  _f("tow_company_name"), _f("tow_costs"), _f("deliver_to"), _f("vol_surrender"),
                  ts, rec_id))
    conn.commit()

    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=?", (rec_id,)).fetchone()
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)

    # Signatures are session-only — pass directly from form, do not persist
    pdf_bytes = _pg.generate_wise_vir_pdf(d, agent_sig=agent_sig, customer_sig=customer_sig)

    job_num    = (d.get("swpi_ref") or str(job_id)).replace("/", "-")
    from datetime import datetime as _dt
    date_str   = _dt.now().strftime("%d-%m-%Y")
    form_label = "Wise VIR"
    orig_filename = f"{job_num} - {form_label} - {date_str}.pdf"
    _attach_pdf_to_job(conn, job_id, session.get("user_id"), pdf_bytes,
                       orig_filename, form_label, ts)
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=orig_filename)


# ─────────────────────────── Form 13 ──────────────────────────────────

@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/form-13")
@login_required
def repo_lock_form_13(job_id: int, rec_id: int):
    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)
    conn.close()
    return render_template("form_13.html",
                           rec=rec, job_id=job_id,
                           agent_name=agent_name,
                           item_year=d.get("year", ""), item_make=d.get("make", ""),
                           item_model=d.get("model", ""))


@app.post("/jobs/<int:job_id>/repo-lock/<int:rec_id>/form-13")
@login_required
def repo_lock_form_13_pdf(job_id: int, rec_id: int):
    from flask import send_file
    import pdf_gen as _pg

    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)

    occupant_sig     = request.form.get("occupant_sig", "").strip() or None
    agent_sig        = request.form.get("agent_sig", "").strip() or None
    occupant_refused = request.form.get("occupant_refused") == "1"

    if not agent_sig:
        flash("Agent signature is required.", "error")
        conn.close()
        return redirect(url_for("repo_lock_form_13", job_id=job_id, rec_id=rec_id))

    if not occupant_refused and not occupant_sig:
        flash("Occupier signature is required, or check 'Occupier refused to sign'.", "error")
        conn.close()
        return redirect(url_for("repo_lock_form_13", job_id=job_id, rec_id=rec_id))

    for col in ("make", "model", "year"):
        add_column_if_missing(conn, "repo_lock_records", col, "TEXT")

    def _f(name):
        return request.form.get(name, "").strip() or rec.get(name) or ""

    ts = now_ts()
    conn.execute("""UPDATE repo_lock_records
                    SET finance_company=?, customer_name=?, account_number=?,
                        repo_address=?, year=?, make=?, model=?, vin=?,
                        updated_at=?
                    WHERE id=?""",
                 (_f("finance_company"), _f("customer_name"), _f("account_number"),
                  _f("repo_address"), _f("year"), _f("make"), _f("model"), _f("vin"),
                  ts, rec_id))
    conn.commit()

    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=?", (rec_id,)).fetchone()
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)

    # Signatures are session-only — pass directly from form, do not persist
    pdf_bytes = _pg.generate_form_13_pdf(d, occupant_sig=occupant_sig, agent_sig=agent_sig)

    job_num    = (d.get("swpi_ref") or str(job_id)).replace("/", "-")
    from datetime import datetime as _dt
    date_str   = _dt.now().strftime("%d-%m-%Y")
    form_label = "Form 13"
    orig_filename = f"{job_num} - {form_label} - {date_str}.pdf"
    _attach_pdf_to_job(conn, job_id, session.get("user_id"), pdf_bytes,
                       orig_filename, form_label, ts)
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=orig_filename)


# ─────────────────────────── Voluntary Surrender ─────────────────────

@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/voluntary-surrender")
@login_required
def repo_lock_voluntary_surrender(job_id: int, rec_id: int):
    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)
    conn.close()
    return render_template("voluntary_surrender.html",
                           rec=rec, job_id=job_id,
                           agent_name=agent_name,
                           item_year=d.get("year", ""), item_make=d.get("make", ""),
                           item_model=d.get("model", ""))


@app.post("/jobs/<int:job_id>/repo-lock/<int:rec_id>/voluntary-surrender")
@login_required
def repo_lock_voluntary_surrender_pdf(job_id: int, rec_id: int):
    from flask import send_file
    import pdf_gen as _pg

    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)

    customer_sig = request.form.get("customer_sig", "").strip() or None
    agent_sig    = request.form.get("agent_sig", "").strip() or None

    if not agent_sig:
        flash("Agent signature is required.", "error")
        conn.close()
        return redirect(url_for("repo_lock_voluntary_surrender", job_id=job_id, rec_id=rec_id))

    if not customer_sig:
        flash("Customer signature is required for voluntary surrender.", "error")
        conn.close()
        return redirect(url_for("repo_lock_voluntary_surrender", job_id=job_id, rec_id=rec_id))

    for col in ("make", "model", "year"):
        add_column_if_missing(conn, "repo_lock_records", col, "TEXT")

    def _f(name):
        return request.form.get(name, "").strip() or rec.get(name) or ""

    ts = now_ts()
    conn.execute("""UPDATE repo_lock_records
                    SET finance_company=?, customer_name=?, account_number=?,
                        repo_address=?, year=?, make=?, model=?, vin=?,
                        deliver_to=?, updated_at=?
                    WHERE id=?""",
                 (_f("finance_company"), _f("customer_name"), _f("account_number"),
                  _f("repo_address"), _f("year"), _f("make"), _f("model"), _f("vin"),
                  _f("deliver_to"), ts, rec_id))
    conn.commit()

    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=?", (rec_id,)).fetchone()
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)

    # Signatures are session-only — pass directly from form, do not persist
    pdf_bytes = _pg.generate_voluntary_surrender_pdf(d, customer_sig=customer_sig,
                                                     agent_sig=agent_sig)

    job_num    = (d.get("swpi_ref") or str(job_id)).replace("/", "-")
    from datetime import datetime as _dt
    date_str   = _dt.now().strftime("%d-%m-%Y")
    form_label = "Voluntary Surrender"
    orig_filename = f"{job_num} - {form_label} - {date_str}.pdf"
    _attach_pdf_to_job(conn, job_id, session.get("user_id"), pdf_bytes,
                       orig_filename, form_label, ts)
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=orig_filename)


# ─────────────────────────── Complete Repo Pack ──────────────────────

@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/repo-pack")
@login_required
def repo_lock_repo_pack(job_id: int, rec_id: int):
    from flask import send_file
    import pdf_gen as _pg

    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)

    # Repo pack generates a reference copy of all forms (no live signatures —
    # sign the individual forms to produce the official signed PDFs)
    pdf_list = []

    # VIR always included
    pdf_list.append(_pg.generate_vir_pdf(d, agent_sig=None, customer_sig=None))

    # Transport instructions
    pdf_list.append(_pg.generate_transport_pdf(d, agent_sig=None, tow_sig=None))

    # Form 13 if the record indicates entry was required
    if d.get("form_13") == "YES":
        pdf_list.append(_pg.generate_form_13_pdf(d, occupant_sig=None, agent_sig=None))

    # Voluntary surrender if applicable
    if d.get("vol_surrender") == "YES":
        pdf_list.append(_pg.generate_voluntary_surrender_pdf(d, customer_sig=None, agent_sig=None))

    # Wise VIR if this is a Wise Group case
    if d.get("client_name", "").upper().startswith("WISE") or d.get("wise_case_number"):
        pdf_list.append(_pg.generate_wise_vir_pdf(d, agent_sig=None, customer_sig=None))

    merged = _pg.generate_repo_pack_pdf(pdf_list)

    ts         = now_ts()
    job_num    = (d.get("swpi_ref") or str(job_id)).replace("/", "-")
    from datetime import datetime as _dt
    date_str   = _dt.now().strftime("%d-%m-%Y")
    form_label = "Complete Repo Pack"
    orig_filename = f"{job_num} - {form_label} - {date_str}.pdf"
    _attach_pdf_to_job(conn, job_id, session.get("user_id"), merged,
                       orig_filename, form_label, ts)
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(merged), mimetype="application/pdf",
                     as_attachment=True, download_name=orig_filename)


# ─────────────────────────── Auction Letter ─────────────────────────

@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/auction-letter")
@login_required
def repo_lock_auction_letter(job_id: int, rec_id: int):
    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)
    conn.close()
    download_url = f"/jobs/{job_id}/repo-lock/{rec_id}/auction-letter/pdf"
    return render_template("letter_preview.html",
                           rec=rec, job_id=job_id,
                           agent_name=agent_name,
                           letter_title="Auction Manager Letter",
                           letter_type="auction",
                           addressee_label="Auction House",
                           addressee=rec.get("deliver_to") or d.get("delivery_address") or "Auction House",
                           deliver_to_fallback=d.get("delivery_address", ""),
                           download_url=download_url,
                           item_year=d.get("year", ""), item_make=d.get("make", ""),
                           item_model=d.get("model", ""))


@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/auction-letter/pdf")
@login_required
def repo_lock_auction_letter_pdf(job_id: int, rec_id: int):
    from flask import send_file
    import pdf_gen as _pg

    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)

    pdf_bytes = _pg.generate_auction_letter_pdf(d)

    ts = now_ts()
    job_num    = (d.get("swpi_ref") or str(job_id)).replace("/", "-")
    from datetime import datetime as _dt
    date_str   = _dt.now().strftime("%d-%m-%Y")
    form_label = "Auction Manager Letter"
    orig_filename = f"{job_num} - {form_label} - {date_str}.pdf"
    _attach_pdf_to_job(conn, job_id, session.get("user_id"), pdf_bytes,
                       orig_filename, form_label, ts)
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=orig_filename)


# ─────────────────────────── Tow Letter ─────────────────────────────

@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/tow-letter")
@login_required
def repo_lock_tow_letter(job_id: int, rec_id: int):
    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)
    conn.close()
    download_url = f"/jobs/{job_id}/repo-lock/{rec_id}/tow-letter/pdf"
    return render_template("letter_preview.html",
                           rec=rec, job_id=job_id,
                           agent_name=agent_name,
                           letter_title="Towing Contractor Letter",
                           letter_type="tow",
                           addressee_label="Tow Contractor",
                           addressee=rec.get("tow_company_name") or d.get("tow_company_name_db") or "Towing Contractor",
                           deliver_to_fallback=rec.get("deliver_to") or "",
                           download_url=download_url,
                           item_year=d.get("year", ""), item_make=d.get("make", ""),
                           item_model=d.get("model", ""))


@app.get("/jobs/<int:job_id>/repo-lock/<int:rec_id>/tow-letter/pdf")
@login_required
def repo_lock_tow_letter_pdf(job_id: int, rec_id: int):
    from flask import send_file
    import pdf_gen as _pg

    conn = db()
    rec_row = conn.execute("SELECT * FROM repo_lock_records WHERE id=? AND job_id=?",
                           (rec_id, job_id)).fetchone()
    if not rec_row:
        conn.close()
        return "Not found", 404
    rec = dict(rec_row)
    d, agent_name, client, tow_op = _rl_pdf_context(conn, rec, job_id)

    pdf_bytes = _pg.generate_tow_letter_pdf(d)

    ts = now_ts()
    job_num    = (d.get("swpi_ref") or str(job_id)).replace("/", "-")
    from datetime import datetime as _dt
    date_str   = _dt.now().strftime("%d-%m-%Y")
    form_label = "Tow Contractor Letter"
    orig_filename = f"{job_num} - {form_label} - {date_str}.pdf"
    _attach_pdf_to_job(conn, job_id, session.get("user_id"), pdf_bytes,
                       orig_filename, form_label, ts)
    conn.commit()
    conn.close()

    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=orig_filename)


# ──────────────────────────── Forms Module ────────────────────────────

_FORMS_CATALOGUE = [
    {
        "id":          "vir",
        "name":        "Vehicle Condition Report / Repossession Receipt",
        "description": "SWPI VIR — captures asset condition, recovery details, and signatures.",
        "short":       "VIR",
        "tags":        ["Repossession", "Asset Recovery", "SWPI"],
        "icon":        '<svg width="20" height="20" fill="none" viewBox="0 0 24 24" stroke="#2563eb" stroke-width="2"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg>',
        "available":   True,
        "generate_url": "/forms/generate?type=vir",
    },
    {
        "id":          "transport",
        "name":        "Transport Instructions / Tow Receipt",
        "description": "SWPI tow contractor dispatch notice and delivery instructions.",
        "short":       "Transport Instructions",
        "tags":        ["Transport", "Tow", "SWPI"],
        "icon":        '<svg width="20" height="20" fill="none" viewBox="0 0 24 24" stroke="#2563eb" stroke-width="2"><rect x="1" y="3" width="15" height="13" rx="1"/><path d="M16 8h4l3 5v3h-7V8z"/><circle cx="5.5" cy="18.5" r="2.5"/><circle cx="18.5" cy="18.5" r="2.5"/></svg>',
        "available":   True,
        "generate_url": "/forms/generate?type=transport",
    },
    {
        "id":          "voluntary_surrender",
        "name":        "Voluntary Surrender Form — Sec 78(1) NCC",
        "description": "Customer voluntarily surrenders mortgaged goods under the National Credit Code.",
        "short":       "Voluntary Surrender",
        "tags":        ["Repossession", "Surrender", "NCC"],
        "icon":        '<svg width="20" height="20" fill="none" viewBox="0 0 24 24" stroke="#2563eb" stroke-width="2"><polyline points="20 6 9 17 4 12"/></svg>',
        "available":   True,
        "generate_url": "/forms/generate?type=voluntary_surrender",
    },
    {
        "id":          "form_13",
        "name":        "Form 13 — Notice to Occupier of Premises",
        "description": "NCCP Act form required when entering private property to repossess goods under a consumer credit contract.",
        "short":       "Form 13",
        "tags":        ["Repossession", "NCCP", "Regulated"],
        "icon":        '<svg width="20" height="20" fill="none" viewBox="0 0 24 24" stroke="#2563eb" stroke-width="2"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/></svg>',
        "available":   True,
        "generate_url": "/forms/generate?type=form_13",
    },
    {
        "id":          "wise_vir",
        "name":        "Wise Group — Vehicle Inspection Report",
        "description": "Wise-branded VIR with condition checkboxes for Wise Group repossession jobs.",
        "short":       "Wise VIR",
        "tags":        ["Wise Group", "VIR", "Condition Report"],
        "icon":        '<svg width="20" height="20" fill="none" viewBox="0 0 24 24" stroke="#f59e0b" stroke-width="2"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg>',
        "available":   True,
        "generate_url": "/forms/generate?type=wise_vir",
    },
    {
        "id":          "auction_letter",
        "name":        "Auction Manager Letter",
        "description": "Auto-populated letter to the auction house confirming vehicle delivery.",
        "short":       "Auction Letter",
        "tags":        ["Letter", "Auction", "Delivery"],
        "icon":        '<svg width="20" height="20" fill="none" viewBox="0 0 24 24" stroke="#2563eb" stroke-width="2"><path d="M4 4h16c1.1 0 2 .9 2 2v12c0 1.1-.9 2-2 2H4c-1.1 0-2-.9-2-2V6c0-1.1.9-2 2-2z"/><polyline points="22,6 12,13 2,6"/></svg>',
        "available":   True,
        "generate_url": "/forms/generate?type=auction_letter",
    },
    {
        "id":          "tow_letter",
        "name":        "Towing Contractor Letter",
        "description": "Auto-populated towing instructions letter with collection and delivery details.",
        "short":       "Tow Letter",
        "tags":        ["Letter", "Tow", "Transport"],
        "icon":        '<svg width="20" height="20" fill="none" viewBox="0 0 24 24" stroke="#2563eb" stroke-width="2"><rect x="1" y="3" width="15" height="13" rx="1"/><path d="M16 8h4l3 5v3h-7V8z"/><circle cx="5.5" cy="18.5" r="2.5"/><circle cx="18.5" cy="18.5" r="2.5"/></svg>',
        "available":   True,
        "generate_url": "/forms/generate?type=tow_letter",
    },
    {
        "id":          "far",
        "name":        "Field Attendance Report",
        "description": "Document an agent field attendance and observations.",
        "short":       "FAR",
        "tags":        ["Field", "Attendance"],
        "icon":        '<svg width="20" height="20" fill="none" viewBox="0 0 24 24" stroke="#2563eb" stroke-width="2"><path d="M17 21v-2a4 4 0 00-4-4H5a4 4 0 00-4 4v2"/><circle cx="9" cy="7" r="4"/><path d="M23 21v-2a4 4 0 00-3-3.87"/><path d="M16 3.13a4 4 0 010 7.75"/></svg>',
        "available":   False,
        "generate_url": "#",
    },
]

_FORMS_META = {f["id"]: f for f in _FORMS_CATALOGUE}


def _forms_job_context(conn, job_id):
    """Return enriched job row dict with customer_name."""
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        return None
    j = dict(job)
    if j.get("customer_id"):
        c = conn.execute("SELECT first_name, last_name FROM customers WHERE id=?",
                         (j["customer_id"],)).fetchone()
        if c:
            j["customer_name"] = f"{c['first_name'] or ''} {c['last_name'] or ''}".strip()
    return j


@app.get("/forms")
@login_required
def forms_dashboard():
    job_id  = request.args.get("job_id", type=int)
    job     = None
    if job_id:
        conn = db()
        job = _forms_job_context(conn, job_id)
        conn.close()
    return render_template("forms.html",
                           forms=_FORMS_CATALOGUE,
                           job=job, job_id=job_id)


@app.get("/forms/generate")
@login_required
def forms_generate():
    form_type   = request.args.get("type", "").lower().strip()
    job_id      = request.args.get("job_id",  type=int)
    item_id     = request.args.get("item_id", type=int)

    meta = _FORMS_META.get(form_type)
    if not meta or not meta["available"]:
        return redirect(url_for("forms_dashboard"))

    form_title       = meta["name"]
    form_description = meta["description"]
    form_short       = meta["short"]

    conn = db()

    # ── If both job_id and item_id are specified, look for repo lock ──
    if job_id and item_id:
        job  = _forms_job_context(conn, job_id)
        item = conn.execute("SELECT * FROM job_items WHERE id=? AND job_id=?",
                            (item_id, job_id)).fetchone()
        item = dict(item) if item else None

        rl = conn.execute(
            """SELECT * FROM repo_lock_records
               WHERE job_id=? AND item_id=?
               ORDER BY CASE status WHEN 'Submitted' THEN 0 WHEN 'Reviewed' THEN 0 WHEN 'Processed' THEN 0 ELSE 1 END,
                        id DESC LIMIT 1""",
            (job_id, item_id)).fetchone()

        if rl:
            rec_id = rl["id"]
            _url_map = {
                "vir":                 f"/jobs/{job_id}/repo-lock/{rec_id}/vir",
                "transport":           f"/jobs/{job_id}/repo-lock/{rec_id}/transport-instructions",
                "wise_vir":            f"/jobs/{job_id}/repo-lock/{rec_id}/wise-vir",
                "form_13":             f"/jobs/{job_id}/repo-lock/{rec_id}/form-13",
                "voluntary_surrender": f"/jobs/{job_id}/repo-lock/{rec_id}/voluntary-surrender",
                "auction_letter":      f"/jobs/{job_id}/repo-lock/{rec_id}/auction-letter",
                "tow_letter":          f"/jobs/{job_id}/repo-lock/{rec_id}/tow-letter",
            }
            proceed_url = _url_map.get(form_type,
                                       f"/jobs/{job_id}/repo-lock/{rec_id}/vir")
            repo_pack_url = f"/jobs/{job_id}/repo-lock/{rec_id}/repo-pack"
            conn.close()
            return render_template("forms_selector.html",
                                   form_type=form_type, form_title=form_title,
                                   form_description=form_description, form_short=form_short,
                                   job=job, item=item,
                                   rl_rec=dict(rl), proceed_url=proceed_url,
                                   repo_pack_url=repo_pack_url,
                                   no_repo_lock=False, jobs=[], items=[],
                                   selected_job_id=job_id, selected_item_id=item_id,
                                   recent_rl=[])
        else:
            conn.close()
            return render_template("forms_selector.html",
                                   form_type=form_type, form_title=form_title,
                                   form_description=form_description, form_short=form_short,
                                   job=job, item=item,
                                   rl_rec=None, proceed_url=None,
                                   no_repo_lock=True, jobs=[], items=[],
                                   selected_job_id=job_id, selected_item_id=item_id,
                                   recent_rl=[])

    # ── Job selected, but no item yet — load items for that job ──
    items = []
    if job_id:
        rows = conn.execute("SELECT * FROM job_items WHERE job_id=? ORDER BY id",
                            (job_id,)).fetchall()
        items = [dict(r) for r in rows]
        # normalise: expose both .reg and .registration
        for it in items:
            it["registration"] = it.get("reg") or it.get("registration") or ""

    # ── Load all jobs for the selector ──
    job_rows = conn.execute(
        """SELECT j.id, j.internal_job_number, j.lender_name, j.client_job_number,
                  c.first_name, c.last_name, c.company
           FROM jobs j
           LEFT JOIN customers c ON c.id = j.customer_id
           ORDER BY j.id DESC LIMIT 200""").fetchall()
    jobs = []
    for r in job_rows:
        cname = f"{r['first_name'] or ''} {r['last_name'] or ''}".strip()
        clabel = (r["company"] or "").strip() or (r["last_name"] or "").strip() or None
        jobs.append({
            "id":                  r["id"],
            "internal_job_number": r["internal_job_number"],
            "lender_name":         r["lender_name"],
            "customer_name":       cname,
            "customer_label":      clabel,
        })

    # ── Recent submitted repo lock records for quick access ──
    rl_rows = conn.execute(
        """SELECT r.id, r.job_id, r.item_id, r.status, r.registration, r.description,
                  r.repo_date, r.agent_name, r.tow_company_name,
                  j.internal_job_number AS job_number,
                  i.make, i.model
           FROM repo_lock_records r
           LEFT JOIN jobs j ON j.id = r.job_id
           LEFT JOIN job_items i ON i.id = r.item_id
           WHERE r.status IN ('Submitted','Reviewed','Processed')
           ORDER BY r.id DESC LIMIT 10""").fetchall()
    recent_rl = [dict(r) for r in rl_rows]

    job = None
    if job_id:
        for j in jobs:
            if j["id"] == job_id:
                job = j
                break
    conn.close()

    return render_template("forms_selector.html",
                           form_type=form_type, form_title=form_title,
                           form_description=form_description, form_short=form_short,
                           job=job, item=None,
                           rl_rec=None, proceed_url=None,
                           no_repo_lock=False,
                           jobs=jobs, items=items,
                           selected_job_id=job_id, selected_item_id=item_id,
                           recent_rl=recent_rl)


# ── Mobile Forms ──
@app.get("/m/forms")
@login_required
def m_forms():
    job_id = request.args.get("job_id", type=int)
    job    = None
    if job_id:
        conn = db()
        job = _forms_job_context(conn, job_id)
        conn.close()
    # Reuse same catalogue; mobile template handles display
    forms_mobile = [
        {**f, "generate_url": f["generate_url"].replace("/forms/generate", "/forms/generate")}
        for f in _FORMS_CATALOGUE
    ]
    return render_template("mobile/forms.html",
                           forms=forms_mobile, job=job, job_id=job_id)


# ──────────────────────────────────────────────────────────────────────

# -------- Field Notes --------
@app.post("/jobs/<int:job_id>/notes/new")
@login_required
def add_job_note(job_id: int):
    note_text = request.form.get("note_text", "").strip()
    barcode   = request.form.get("barcode", "").strip()
    files = request.files.getlist("attachments")

    if barcode:
        note_text = f"[Barcode: {barcode}]\n{note_text}".strip()

    if not note_text and not any(f.filename for f in files):
        flash("A note or attachment is required.", "danger")
        return redirect(url_for("job_detail", job_id=job_id, _anchor="tab-notes"))

    # Admin can override staff and timestamp
    if session.get("role") in ("admin", "both"):
        staff_uid_raw = request.form.get("staff_user_id", "").strip()
        author_id = int(staff_uid_raw) if staff_uid_raw and staff_uid_raw.isdigit() else session.get("user_id")
        note_date  = request.form.get("note_date", "").strip()
        note_hour  = request.form.get("note_hour", "").strip()
        note_min   = request.form.get("note_minute", "").strip()
        if note_date and note_hour and note_min:
            try:
                ts = parse_interaction_datetime(note_date, f"{note_hour}:{note_min}")
            except Exception:
                ts = now_ts()
        else:
            ts = now_ts()
    else:
        author_id = session.get("user_id")
        ts = now_ts()

    conn = db()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at)
        VALUES (?, ?, ?, ?)
    """, (job_id, author_id, note_text, ts))

    note_id = cur.lastrowid

    import time as _time
    for file in files:
        if file and file.filename and allowed_file(file.filename):
            filename    = secure_filename(file.filename)
            unique_name = f"{job_id}_{note_id}_{int(_time.time())}_{filename}"
            upload_to_blob(file, unique_name)
            cur.execute("""
                INSERT INTO job_note_files (job_field_note_id, filename, filepath, uploaded_at)
                VALUES (?, ?, ?, ?)
            """, (note_id, unique_name, unique_name, ts))

    conn.commit()
    conn.close()

    audit("job_note", note_id, "create", "Field note added", {"job_id": job_id})

    if session.get("role") in ("agent", "both") and note_text:
        try:
            _today = datetime.now(_melbourne).date().isoformat()
            _ts    = now_ts()
            _conn  = db()
            _cur   = _conn.cursor()
            _cur.execute("""SELECT id FROM cue_items
                            WHERE job_id=? AND visit_type='Agent Note Review' AND status='Pending'""",
                         (job_id,))
            if not _cur.fetchone():
                _cur.execute("""
                    INSERT INTO cue_items
                      (job_id, visit_type, due_date, priority, status,
                       instructions, created_by_user_id, created_at, updated_at)
                    VALUES (?, 'Agent Note Review', ?, 'High', 'Pending', ?, ?, ?, ?)
                """, (job_id, _today, note_text[:200], session.get("user_id"), _ts, _ts))
                _conn.commit()
            _conn.close()
        except Exception:
            pass

    flash("Field note saved.", "success")
    if session.get("role") in ("admin", "both"):
        _sconn = db()
        has_active_schedule = _sconn.execute(
            "SELECT 1 FROM schedules WHERE job_id = ? AND status NOT IN ('Cancelled', 'Completed') LIMIT 1",
            (job_id,)
        ).fetchone()
        _sconn.close()
        if has_active_schedule:
            return redirect(url_for("job_detail", job_id=job_id, _anchor="tab-notes"))
        return redirect(url_for("job_detail", job_id=job_id) + "?schedule_prompt=1#tab-notes")
    return redirect(url_for("job_detail", job_id=job_id, _anchor="tab-notes"))


@app.post("/jobs/<int:job_id>/notes/<int:note_id>/delete")
@login_required
def delete_job_note(job_id: int, note_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM job_field_notes WHERE id = ? AND job_id = ?", (note_id, job_id))
    note = cur.fetchone()
    if not note:
        conn.close()
        flash("Note not found.", "danger")
        return redirect(url_for("job_detail", job_id=job_id, _anchor="tab-notes"))

    caller_id   = session.get("user_id")
    caller_role = session.get("role", "")
    if caller_role not in ("admin", "both") and note["created_by_user_id"] != caller_id:
        conn.close()
        flash("You can only delete your own notes.", "danger")
        return redirect(url_for("job_detail", job_id=job_id, _anchor="tab-notes"))

    cur.execute("SELECT filename, filepath FROM job_note_files WHERE job_field_note_id = ?", (note_id,))
    files = cur.fetchall()
    for f in files:
        delete_blob_safely(f["filename"])
        try: os.remove(f["filepath"])
        except OSError: pass
    cur.execute("DELETE FROM job_note_files WHERE job_field_note_id = ?", (note_id,))
    cur.execute("DELETE FROM job_field_notes WHERE id = ?", (note_id,))
    conn.commit()
    conn.close()

    audit("job_note", note_id, "delete", "Field note deleted", {"job_id": job_id})
    flash("Field note deleted.", "success")
    return redirect(url_for("job_detail", job_id=job_id, _anchor="tab-notes"))


@app.post("/jobs/<int:job_id>/notes/<int:note_id>/edit")
@login_required
def edit_job_note(job_id: int, note_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM job_field_notes WHERE id=? AND job_id=?", (note_id, job_id))
    note = cur.fetchone()
    if not note:
        conn.close()
        return jsonify({"ok": False, "error": "Note not found"}), 404

    caller_id   = session.get("user_id")
    caller_role = session.get("role", "")
    if caller_role not in ("admin", "both") and note["created_by_user_id"] != caller_id:
        conn.close()
        return jsonify({"ok": False, "error": "Permission denied"}), 403

    new_text = request.form.get("note_text", "").strip()
    if not new_text:
        conn.close()
        return jsonify({"ok": False, "error": "Note text cannot be empty"}), 400

    ts = now_ts()
    fields: dict = {"note_text": new_text, "updated_at": ts, "updated_by_user_id": caller_id}

    if caller_role in ("admin", "both"):
        note_date = request.form.get("note_date", "").strip()
        note_hour = request.form.get("note_hour", "").strip()
        note_min  = request.form.get("note_minute", "").strip()
        if note_date and note_hour and note_min:
            try:
                fields["created_at"] = parse_interaction_datetime(note_date, f"{note_hour}:{note_min}")
            except Exception:
                pass
        staff_uid_raw = request.form.get("staff_user_id", "").strip()
        if staff_uid_raw and staff_uid_raw.isdigit():
            fields["created_by_user_id"] = int(staff_uid_raw)

    set_clause = ", ".join(f"{k}=?" for k in fields)
    cur.execute(f"UPDATE job_field_notes SET {set_clause} WHERE id=?",
                list(fields.values()) + [note_id])

    import time as _time
    files = request.files.getlist("attachments")
    new_files = []
    for file in files:
        if file and file.filename and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            unique_name = f"{job_id}_{note_id}_{int(_time.time())}_{filename}"
            upload_to_blob(file, unique_name)
            cur.execute("""
                INSERT INTO job_note_files (job_field_note_id, filename, filepath, uploaded_at)
                VALUES (?, ?, ?, ?)
            """, (note_id, unique_name, unique_name, ts))
            new_files.append({"id": cur.lastrowid, "filename": unique_name, "original": filename, "uploaded_at": ts})

    conn.commit()

    cur.execute("SELECT created_at, updated_at FROM job_field_notes WHERE id=?", (note_id,))
    updated = cur.fetchone()

    cur.execute("SELECT id, filename, uploaded_at FROM job_note_files WHERE job_field_note_id=?", (note_id,))
    all_files = [dict(r) for r in cur.fetchall()]
    conn.close()

    audit("job_note", note_id, "edit", "Field note edited", {"job_id": job_id, "files_added": len(new_files)})
    return jsonify({
        "ok": True,
        "note_text": new_text,
        "created_at": updated["created_at"] if updated else note["created_at"],
        "updated_at": updated["updated_at"] if updated else ts,
        "files": all_files,
    })


@app.get("/jobs/<int:job_id>/notes/<int:note_id>/detail")
@login_required
def note_detail_api(job_id: int, note_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("""
        SELECT n.*, u.full_name AS author_name
        FROM job_field_notes n
        LEFT JOIN users u ON u.id = n.created_by_user_id
        WHERE n.id = ? AND n.job_id = ?
    """, (note_id, job_id))
    note = cur.fetchone()
    if not note:
        conn.close()
        return jsonify({"ok": False, "error": "Note not found"}), 404

    updated_by_name = None
    if note["updated_by_user_id"]:
        ub = cur.execute("SELECT full_name FROM users WHERE id=?", (note["updated_by_user_id"],)).fetchone()
        if ub:
            updated_by_name = ub["full_name"]

    cur.execute("SELECT id, filename, uploaded_at FROM job_note_files WHERE job_field_note_id=? ORDER BY id", (note_id,))
    files = [dict(r) for r in cur.fetchall()]
    conn.close()

    return jsonify({
        "ok": True,
        "note_id": note["id"],
        "note_text": note["note_text"] or "",
        "created_at": note["created_at"],
        "created_by_user_id": note["created_by_user_id"],
        "author_name": note["author_name"] or "Unknown",
        "updated_at": note["updated_at"],
        "updated_by_name": updated_by_name,
        "files": files,
    })


@app.post("/jobs/<int:job_id>/notes/<int:note_id>/attachments")
@login_required
def add_note_attachments(job_id: int, note_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM job_field_notes WHERE id=? AND job_id=?", (note_id, job_id))
    note = cur.fetchone()
    if not note:
        conn.close()
        return jsonify({"ok": False, "error": "Note not found"}), 404

    caller_id   = session.get("user_id")
    caller_role = session.get("role", "")
    if caller_role not in ("admin", "both") and note["created_by_user_id"] != caller_id:
        conn.close()
        return jsonify({"ok": False, "error": "Permission denied"}), 403

    files = request.files.getlist("attachments")
    if not files or not any(f.filename for f in files):
        conn.close()
        return jsonify({"ok": False, "error": "No files provided"}), 400

    import time as _time
    ts = now_ts()
    added = []
    rejected = []
    for file in files:
        if not file or not file.filename:
            continue
        if not allowed_file(file.filename):
            rejected.append(file.filename)
            continue
        filename = secure_filename(file.filename)
        unique_name = f"{job_id}_{note_id}_{int(_time.time())}_{filename}"
        upload_to_blob(file, unique_name)
        cur.execute("""
            INSERT INTO job_note_files (job_field_note_id, filename, filepath, uploaded_at)
            VALUES (?, ?, ?, ?)
        """, (note_id, unique_name, unique_name, ts))
        added.append({"id": cur.lastrowid, "filename": unique_name, "original": filename, "uploaded_at": ts})

    cur.execute("UPDATE job_field_notes SET updated_at=?, updated_by_user_id=? WHERE id=?",
                (ts, caller_id, note_id))
    conn.commit()

    cur.execute("SELECT id, filename, uploaded_at FROM job_note_files WHERE job_field_note_id=? ORDER BY id", (note_id,))
    all_files = [dict(r) for r in cur.fetchall()]
    conn.close()

    audit("job_note", note_id, "attach", f"{len(added)} file(s) added to note", {"job_id": job_id})
    return jsonify({"ok": True, "added": len(added), "rejected": rejected, "files": all_files})


@app.post("/jobs/<int:job_id>/notes/<int:note_id>/attachments/<int:file_id>/delete")
@login_required
def delete_note_attachment(job_id: int, note_id: int, file_id: int):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM job_field_notes WHERE id=? AND job_id=?", (note_id, job_id))
    note = cur.fetchone()
    if not note:
        conn.close()
        return jsonify({"ok": False, "error": "Note not found"}), 404

    caller_role = session.get("role", "")
    caller_id = session.get("user_id")
    if caller_role not in ("admin", "both") and note["created_by_user_id"] != caller_id:
        conn.close()
        return jsonify({"ok": False, "error": "Permission denied"}), 403

    cur.execute("SELECT * FROM job_note_files WHERE id=? AND job_field_note_id=?", (file_id, note_id))
    f = cur.fetchone()
    if not f:
        conn.close()
        return jsonify({"ok": False, "error": "File not found"}), 404

    delete_blob_safely(f["filename"])
    try:
        os.remove(os.path.join(UPLOAD_FOLDER, f["filename"]))
    except OSError:
        pass
    cur.execute("DELETE FROM job_note_files WHERE id=?", (file_id,))

    ts = now_ts()
    cur.execute("UPDATE job_field_notes SET updated_at=?, updated_by_user_id=? WHERE id=?",
                (ts, caller_id, note_id))
    conn.commit()
    conn.close()

    audit("job_note", note_id, "detach", f"Attachment removed: {f['filename']}", {"job_id": job_id, "file_id": file_id})
    return jsonify({"ok": True})


@app.get("/uploads/<path:filename>")
@login_required
def serve_upload(filename):
    import logging as _log
    if _uploads_container:
        try:
            blob_client = _uploads_container.get_blob_client(filename)
            download    = blob_client.download_blob()
            mime        = mimetypes.guess_type(filename)[0] or "application/octet-stream"
            return Response(download.readall(), mimetype=mime)
        except Exception as e:
            _log.warning("Blob fetch failed for %r: %s", filename, e)
    local_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    if os.path.exists(local_path):
        return send_from_directory(app.config["UPLOAD_FOLDER"], filename)
    conn = db()
    alt_row = conn.execute(
        "SELECT id, filename, filepath FROM job_note_files WHERE filename=? OR filepath=?",
        (filename, filename)
    ).fetchone()
    conn.close()
    if alt_row:
        alt_disk = alt_row["filepath"] or alt_row["filename"]
        if alt_disk and alt_disk != filename:
            alt_path = os.path.join(app.config["UPLOAD_FOLDER"], os.path.basename(alt_disk))
            if os.path.exists(alt_path):
                return send_from_directory(app.config["UPLOAD_FOLDER"], os.path.basename(alt_disk))
        _log.warning("Orphan note-file record for %r (id=%s) — file not in storage", filename, alt_row["id"])
        return render_template("error_500.html",
            error_message="This file is referenced in the database but the actual file data has not been imported into storage yet. "
                          "This typically happens with GeoOp-imported attachments that haven't been backfilled. "
                          "An admin can run the Attachment Backfill from the GeoOp Import page to resolve this.",
            path=f"/uploads/{filename}"), 404
    _log.warning("File not found locally or in blob: %r", filename)
    abort(404)


@app.get("/jobs/<int:job_id>/documents/<int:doc_id>/download")
@login_required
def download_job_document(job_id: int, doc_id: int):
    conn = db()
    cur = conn.cursor()

    if session.get("role") == "agent":
        user_id = session.get("user_id")
        access = conn.execute(
            """SELECT 1 FROM jobs WHERE id=? AND (
               assigned_user_id=? OR EXISTS (
                 SELECT 1 FROM schedules WHERE job_id=? AND assigned_to_user_id=?
                 AND status NOT IN ('Cancelled')
               )
             )""",
            (job_id, user_id, job_id, user_id),
        ).fetchone()
        if not access:
            conn.close()
            return ("Not found", 404)

    cur.execute(
        "SELECT original_filename, stored_filename, mime_type FROM job_documents WHERE id=? AND job_id=?",
        (doc_id, job_id),
    )
    doc = cur.fetchone()
    if not doc:
        conn.close()
        return ("Not found", 404)

    stored = doc["stored_filename"]
    mime = doc["mime_type"] or mimetypes.guess_type(stored)[0] or "application/octet-stream"

    if _uploads_container:
        try:
            data = _uploads_container.get_blob_client(stored).download_blob().readall()
            audit("job_document", doc_id, "download", f"Document downloaded: {doc['original_filename']}", {"job_id": job_id})
            conn.close()
            resp = make_response(data)
            resp.headers["Content-Type"] = mime
            resp.headers["Content-Disposition"] = f'attachment; filename="{doc["original_filename"]}"'
            return resp
        except Exception:
            pass

    local_path = os.path.join(app.config["UPLOAD_FOLDER"], stored)
    if os.path.exists(local_path):
        audit("job_document", doc_id, "download", f"Document downloaded: {doc['original_filename']}", {"job_id": job_id})
        conn.close()
        return send_from_directory(app.config["UPLOAD_FOLDER"], stored, as_attachment=True, download_name=doc["original_filename"])

    conn.close()
    abort(404)


@app.post("/jobs/<int:job_id>/documents/upload")
@login_required
@admin_required
def job_document_upload(job_id: int):
    doc_type = (request.form.get("doc_type") or "Other").strip()
    title = (request.form.get("title") or "").strip()
    notes = (request.form.get("notes") or "").strip()
    also_save_note = request.form.get("also_save_note") == "1"
    files = request.files.getlist("file")

    valid_files = [f for f in files if f and f.filename]
    if not valid_files:
        flash("Select at least one file to upload.", "danger")
        return redirect(url_for("job_detail", job_id=job_id, _anchor="tab-notes"))

    conn = db()
    cur = conn.cursor()
    uploaded = 0
    skipped = []
    ts = now_ts()
    for file in valid_files:
        if not allowed_file(file.filename):
            skipped.append(file.filename)
            continue
        original_filename = secure_filename(file.filename)
        ts_safe = ts.replace(":", "").replace("-", "").replace(" ", "")
        stored_filename = f"job_{job_id}_{ts_safe}_{original_filename}"
        mime_type = mimetypes.guess_type(original_filename)[0] or "application/octet-stream"
        file_size = upload_to_blob(file, stored_filename)
        cur.execute("""
            INSERT INTO job_documents
                (job_id, doc_type, title, original_filename, stored_filename,
                 mime_type, file_size, uploaded_by_user_id, uploaded_at, notes)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (job_id, doc_type, title or None, original_filename, stored_filename,
              mime_type, file_size, session.get("user_id"), ts, notes or None))
        doc_id = cur.lastrowid
        audit("job_document", doc_id, "create", "Job document uploaded",
              {"job_id": job_id, "doc_type": doc_type, "filename": original_filename})

        # Also save a reference in Notes & Docs so it appears in the notes tab
        if also_save_note:
            note_parts = [f"Document uploaded: {doc_type}"]
            if title:
                note_parts.append(f"— {title}")
            note_parts.append(f"({original_filename})")
            if notes:
                note_parts.append(f"\n{notes}")
            note_text = " ".join(note_parts)
            cur.execute("""
                INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at)
                VALUES (?, ?, ?, ?)
            """, (job_id, session.get("user_id"), note_text, ts))
            note_id = cur.lastrowid
            cur.execute("""
                INSERT INTO job_note_files (job_field_note_id, filename, filepath, uploaded_at)
                VALUES (?, ?, ?, ?)
            """, (note_id, original_filename, stored_filename, ts))

        uploaded += 1
    conn.commit()
    conn.close()

    if uploaded:
        msg = f"{uploaded} document{'s' if uploaded != 1 else ''} uploaded."
        if skipped:
            msg += f" {len(skipped)} skipped (unsupported type)."
        flash(msg, "success")
    else:
        flash("No supported files were uploaded.", "danger")
    return redirect(url_for("job_detail", job_id=job_id, _anchor="tab-notes"))


# -------- Update Builder (AI Assist) --------

def _ensure_draft_cue(conn, job_id, agent_id, draft_id, job_ref):
    ts = now_ts()
    melb = pytz.timezone("Australia/Melbourne")
    now_local = datetime.now(melb)
    five_pm_today = now_local.replace(hour=17, minute=0, second=0, microsecond=0)
    two_hours_later = now_local + _td(hours=2)
    due_time = min(five_pm_today, two_hours_later)
    due_date_str = due_time.strftime("%Y-%m-%d")
    cue_link = f"/jobs/{job_id}/update-builder"
    instructions = (
        f"An AI-generated attendance update was started but not applied to job {job_ref}. "
        "Please review and complete the update."
    )
    existing = conn.execute(
        """SELECT id FROM cue_items
           WHERE job_id=? AND assigned_user_id=?
             AND visit_type='Complete Attendance Update'
             AND status IN ('Pending','In Progress')""",
        (job_id, agent_id)
    ).fetchone()
    if not existing:
        conn.execute(
            """INSERT INTO cue_items
               (job_id, visit_type, due_date, priority, status, assigned_user_id,
                instructions, cue_link, created_by_user_id, created_at, updated_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
            (job_id, "Complete Attendance Update", due_date_str, "High", "Pending",
             agent_id, instructions, cue_link, agent_id, ts, ts)
        )
        conn.commit()


def _get_ai_client():
    conn = db()
    settings = conn.execute("SELECT * FROM system_settings WHERE id=1").fetchone()
    conn.close()
    use_own = settings and settings["ai_use_own_key"]
    own_key = settings and settings["openai_api_key"]
    if use_own and own_key:
        from openai import OpenAI
        return OpenAI(api_key=own_key), "own"
    from openai import OpenAI
    replit_base = os.getenv("AI_INTEGRATIONS_OPENAI_BASE_URL", "")
    replit_key  = os.getenv("AI_INTEGRATIONS_OPENAI_API_KEY", "")
    # Replit's AI integration only works inside Replit's dev environment.
    # If base_url points to localhost (i.e. running on Azure/external), fall back to own key.
    if replit_base and "localhost" not in replit_base and replit_key and not replit_key.startswith("_DUMMY"):
        return OpenAI(api_key=replit_key, base_url=replit_base), "replit"
    if not own_key:
        raise RuntimeError(
            "No OpenAI key configured. Go to Settings → AI to add your own OpenAI API key."
        )
    return OpenAI(api_key=own_key), "own"


def _calc_eta_date(from_date, poc):
    import datetime as _dt
    days_needed = 8 if poc >= 3 else 2
    d = from_date
    added = 0
    while added < days_needed:
        d += _dt.timedelta(days=1)
        if d.weekday() != 6:
            added += 1
    return d


def _fmt_attend_datetime(attend_date_str, attend_time_str):
    try:
        from datetime import datetime as _dt2
        dt = _dt2.strptime(f"{attend_date_str} {attend_time_str}", "%Y-%m-%d %H:%M")
        day   = dt.strftime("%d")
        month = dt.strftime("%m")
        year  = dt.strftime("%y")
        hour  = str(int(dt.strftime("%I")))
        minute = dt.strftime("%M")
        ampm  = dt.strftime("%p").lower()
        return f"{day}/{month}/{year} at {hour}:{minute}{ampm}"
    except Exception:
        return f"{attend_date_str} {attend_time_str}"


def _build_swpi_prompt(inputs, job_ctx):
    attend_date = inputs.get("attend_date", "")
    attend_time = inputs.get("attend_time", "")
    is_first    = inputs.get("is_first_attendance", False)
    prop_desc   = inputs.get("property_description", "")
    sec_sighted = inputs.get("security_sighted", False)
    sec_mm      = inputs.get("security_make_model", "")
    sec_reg     = inputs.get("security_reg", "")
    sec_loc     = inputs.get("security_location", "")
    cc          = inputs.get("calling_card", False)
    neighbour   = inputs.get("neighbour_outcome", "")
    call_made   = inputs.get("call_made", False)
    call_out    = inputs.get("call_outcome", "")
    voicemail   = inputs.get("voicemail_left", False)
    sms_sent    = inputs.get("sms_sent", False)
    phone_used  = (inputs.get("phone_number_used") or "").replace(" ", "")
    poc         = inputs.get("points_of_contact", 0)
    eta_date    = inputs.get("eta_next_date", "")
    agent_notes = (inputs.get("agent_notes") or "").strip()
    address     = job_ctx.get("job_address", "")
    is_regional = job_ctx.get("is_regional", False)

    formatted_dt = _fmt_attend_datetime(attend_date, attend_time)

    if is_first:
        if prop_desc:
            opening = f"{formatted_dt} Our agent attended at {address}, finding a {prop_desc}."
        else:
            opening = f"{formatted_dt} Our agent attended at {address}."
    else:
        opening = f"{formatted_dt} Our agent re-attended at {address}."

    if sec_sighted:
        sec_parts = []
        if sec_mm:
            sec_parts.append(sec_mm)
        if sec_reg:
            sec_parts.append(f"registration {sec_reg}")
        loc_phrase = f", located at {sec_loc}" if sec_loc else ""
        sec_sentence = f"The security was sighted at the premises{loc_phrase}."
        if sec_parts:
            sec_sentence = f"The security was sighted at the premises{loc_phrase} ({', '.join(sec_parts)})."
    else:
        sec_sentence = "The security was not sighted at or in the immediate vicinity."

    if cc:
        if is_first:
            cc_sentence = "A calling card was left in a sealed envelope addressed to the customer, marked 'Private and Confidential,' and wedged in the door, requesting urgent contact."
        else:
            cc_sentence = "A further calling card was left in a sealed envelope addressed to the customer, marked 'Private and Confidential,' and wedged in the door, requesting urgent contact."
    else:
        cc_sentence = ""

    if call_made:
        if call_out == "answered":
            call_sentence = f"While on site, our agent telephoned {phone_used}, the call was answered."
        elif call_out == "diverted to voicemail":
            vm_part = " where a message was left requesting urgent contact" if voicemail else ""
            call_sentence = f"While on site, our agent telephoned {phone_used}, the call diverted to voicemail{vm_part}."
        elif call_out == "no answer":
            call_sentence = f"While on site, our agent telephoned {phone_used}, however the call went unanswered."
        elif call_out == "disconnected":
            call_sentence = f"While on site, our agent telephoned {phone_used}, however the call was not connected."
        else:
            call_sentence = f"While on site, our agent telephoned {phone_used}. {call_out}."
    else:
        call_sentence = ""

    sms_sentence = f"Our agent also forwarded an SMS to {phone_used}." if sms_sent and phone_used else (
        "Our agent also forwarded an SMS to the customer." if sms_sent else "")

    poc_sentence = f"This attendance constitutes {poc} point{'s' if poc != 1 else ''} of contact."
    eta_sentence = f"ETA next attendance: {eta_date}."

    fixed_parts = []
    fixed_parts.append(sec_sentence)
    if cc_sentence:
        fixed_parts.append(cc_sentence)
    if call_sentence:
        fixed_parts.append(call_sentence)
    if sms_sentence:
        fixed_parts.append(sms_sentence)
    if neighbour:
        fixed_parts.append(neighbour)
    photos_count = inputs.get("photos_count", 0)
    if photos_count:
        fixed_parts.append(f"{photos_count} photo(s) were taken on-site and attached to this update for reference.")

    fixed_block = "\n".join(f"- {s}" for s in fixed_parts)

    has_notes = bool(agent_notes)

    if has_notes:
        prompt = f"""You are a compliance writer for SWPI, an Australian asset recovery and repossession company.
Your task is to write a single paragraph attendance update in SWPI's house style.

The agent has provided typed attendance notes describing what occurred during this visit. Your job is to rewrite those notes into a clean, professional narrative in third person, preserving every fact the agent recorded. Do not omit, invent, or alter any detail from the agent's notes.

MANDATORY RULES:
1. The narrative MUST begin with EXACTLY this sentence (copy it word for word, do not alter it):
   {opening}
2. Continue in third person throughout. Use "our agent" (lowercase after first use) for subsequent mentions.
3. British/Australian spelling. No acronyms.
4. Single block of continuous prose — no line breaks, no bullet points, no headings, no labels.
5. The FIXED SENTENCES below must appear in the output VERBATIM — do not reword, summarise, or paraphrase them. Weave them into the narrative in order.
6. The AGENT'S ATTENDANCE NOTES below are the primary source of narrative content. Rewrite them into polished professional prose while preserving every fact, person, interaction, confirmation, and observation recorded by the agent. Do not add any events, persons, or outcomes not present in the notes.
7. End the narrative with EXACTLY these two closing sentences (copy word for word, on the same line as the rest):
   {poc_sentence} {eta_sentence}
8. Do NOT invent, assume, or fabricate any events, interactions, persons, or observations not present in either the fixed sentences or the agent's notes.
9. The points of contact total is {poc}. Do not recalculate or change this number.
10. Do NOT add any labels, headings, or category prefixes (e.g. "Security:", "Note:", "Actions:") anywhere in the output.
11. Do NOT add a subject line, title, or reference number before the narrative.
12. Do NOT mention neighbours, occupant interactions, customer conversations, third-party responses, or any other event unless it is explicitly described in the agent's notes or fixed sentences below.

FIXED SENTENCES (include each one verbatim, woven into the narrative):
{fixed_block}

AGENT'S ATTENDANCE NOTES (rewrite into professional prose, preserving all facts):
{agent_notes}

MANDATORY CLOSING (copy exactly):
{poc_sentence} {eta_sentence}

Write the complete narrative now. Start with the mandatory opening, incorporate the fixed sentences verbatim, rewrite the agent's notes into professional third-person prose preserving every fact, and end with the mandatory closing."""
    else:
        neighbour_rule = "12. Do NOT mention any neighbour interaction, interview, or conversation. No neighbour content was provided." if not neighbour else "12. Include the neighbour content exactly as provided in the fixed sentences. Do not expand or invent additional neighbour details."

        prompt = f"""You are a compliance writer for SWPI, an Australian asset recovery and repossession company.
Your task is to write a single paragraph attendance update in SWPI's house style.

MANDATORY RULES:
1. The narrative MUST begin with EXACTLY this sentence (copy it word for word, do not alter it):
   {opening}
2. Continue in third person throughout. Use "our agent" (lowercase after first use) for subsequent mentions.
3. British/Australian spelling. No acronyms.
4. Single block of continuous prose — no line breaks, no bullet points, no headings, no labels.
5. Copy each fixed sentence below VERBATIM into the output. Do not reword, summarise, elaborate, or paraphrase any fixed sentence. Do NOT prefix them with labels such as "Security:" or "Note:".
6. End the narrative with EXACTLY these two closing sentences (copy word for word, on the same line as the rest):
   {poc_sentence} {eta_sentence}
7. Do NOT invent, assume, or fabricate any events, interactions, or observations not listed below. Only use information explicitly provided.
8. The points of contact total is {poc}. Do not recalculate or change this number.
9. Do NOT add any labels, headings, or category prefixes (e.g. "Security:", "Note:", "Actions:") anywhere in the output.
10. Do NOT add a subject line, title, or reference number before the narrative.
{neighbour_rule}
13. STRICT DETERMINISM: The output must contain ONLY the mandatory opening sentence, the fixed sentences listed below, and the mandatory closing sentences — concatenated in order with only whitespace and necessary punctuation between them. Do NOT add any additional observations, actions, descriptions, events, or narrative filler. Do NOT describe the property, surroundings, weather, vehicle presence, occupancy signs, or any other detail unless it appears verbatim in the fixed sentences.
14. FIELD-LEVEL PROHIBITION: If a calling card is not mentioned in the fixed sentences, do NOT mention a calling card. If a phone call is not mentioned, do NOT mention a phone call. If SMS is not mentioned, do NOT mention SMS. If neighbour interaction is not mentioned, do NOT mention neighbours. Only narrate what is explicitly provided below — nothing more.

FIXED SENTENCES (include each one verbatim, in this order, after the opening):
{fixed_block}

MANDATORY CLOSING (copy exactly):
{poc_sentence} {eta_sentence}

Write the complete narrative now. Output ONLY the mandatory opening, followed by each fixed sentence verbatim, followed by the mandatory closing. No other content."""

    return prompt


@app.get("/jobs/<int:job_id>/update-builder")
@login_required
def update_builder(job_id: int):
    conn = db()
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        abort(404)
    role = session.get("role", "")
    uid  = session.get("user_id")
    if role not in ("admin", "both"):
        if job["assigned_user_id"] != uid:
            conn.close()
            flash("Access denied.", "danger")
            return redirect(url_for("job_detail", job_id=job_id))

    customer = None
    customer_mobile = ""
    if job["customer_id"]:
        customer = conn.execute("SELECT * FROM customers WHERE id=?", (job["customer_id"],)).fetchone()
        phone_row = conn.execute(
            "SELECT phone_number FROM contact_phone_numbers WHERE entity_type='customer' AND entity_id=? AND label='Mobile' LIMIT 1",
            (job["customer_id"],)
        ).fetchone()
        customer_mobile = phone_row["phone_number"] if phone_row else ""

    client = None
    if job["client_id"]:
        client = conn.execute("SELECT name FROM clients WHERE id=?", (job["client_id"],)).fetchone()

    draft = conn.execute(
        "SELECT * FROM job_updates WHERE job_id=? AND created_by_user_id=? AND status='draft' ORDER BY id DESC LIMIT 1",
        (job_id, uid)
    ).fetchone()

    if not draft:
        ts = now_ts()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO job_updates (job_id, created_by_user_id, status, customer_mobile, created_at, updated_at)
            VALUES (?, ?, 'draft', ?, ?, ?)
        """, (job_id, uid, customer_mobile, ts, ts))
        draft_id = cur.lastrowid
        conn.commit()
        draft = conn.execute("SELECT * FROM job_updates WHERE id=?", (draft_id,)).fetchone()

    draft_photos = conn.execute(
        "SELECT id, tag FROM job_update_photos WHERE job_update_id=? ORDER BY id", (draft["id"],)
    ).fetchall()
    draft_photos_list = [{"id": p["id"], "url": f"/jobs/{job_id}/update-builder/photo/{p['id']}", "tag": p["tag"]} for p in draft_photos]
    conn.close()

    from datetime import datetime as _dt2
    mel_now = _dt2.now(_melbourne)
    now_date = mel_now.strftime("%Y-%m-%d")
    now_time = mel_now.strftime("%H:%M")
    return render_template("update_builder.html",
                           job=job, customer=customer, client=client,
                           customer_mobile=customer_mobile, draft=draft,
                           draft_photos=draft_photos_list,
                           now_date=now_date, now_time=now_time)


@app.post("/jobs/<int:job_id>/update-builder/generate")
@login_required
def update_builder_generate(job_id: int):
    conn = db()
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({"error": "Job not found"}), 404

    uid  = session.get("user_id")
    role = session.get("role", "")
    if role not in ("admin", "both") and job["assigned_user_id"] != uid:
        conn.close()
        return jsonify({"error": "Access denied"}), 403

    data = request.get_json(force=True)

    is_first    = bool(data.get("is_first_attendance"))
    cc          = bool(data.get("calling_card"))
    call_made   = bool(data.get("call_made"))
    voicemail   = bool(data.get("voicemail_left")) and call_made
    sms_sent    = bool(data.get("sms_sent"))
    phone_used  = (data.get("phone_number_used") or "").strip()
    confirmed_skip = bool(job["confirmed_skip"] if "confirmed_skip" in job.keys() else False)

    poc = 0
    if cc:
        poc += 1
    if call_made:
        poc += 1
    if sms_sent:
        poc += 1

    attend_date_str = data.get("attend_date", "")
    try:
        from datetime import datetime as _dt2
        attend_date_obj = _dt2.strptime(attend_date_str, "%Y-%m-%d").date()
    except Exception:
        attend_date_obj = _dt2.now(_melbourne).date()

    if not confirmed_skip:
        eta_obj = _calc_eta_date(attend_date_obj, poc)
        eta_str = eta_obj.strftime("%d/%m/%y")
    else:
        eta_str = "TBC"

    customer_name = ""
    if job["customer_id"]:
        c = conn.execute("SELECT first_name, last_name FROM customers WHERE id=?", (job["customer_id"],)).fetchone()
        if c:
            customer_name = f"{c['first_name']} {c['last_name']}"

    client_name = ""
    if job["client_id"]:
        cl = conn.execute("SELECT name FROM clients WHERE id=?", (job["client_id"],)).fetchone()
        if cl:
            client_name = cl["name"]

    job_ctx = {
        "job_ref":       job["display_ref"] or job["internal_job_number"],
        "customer_name": customer_name,
        "client_name":   client_name,
        "job_address":   job["job_address"] or "",
        "confirmed_skip": confirmed_skip,
        "is_regional":   bool(job["is_regional"] if "is_regional" in job.keys() else False),
    }
    draft_id_for_photos = data.get("draft_id")
    update_photos_count = 0
    if draft_id_for_photos:
        pc_row = conn.execute(
            "SELECT COUNT(*) c FROM job_update_photos WHERE job_update_id=?", (draft_id_for_photos,)
        ).fetchone()
        update_photos_count = pc_row["c"] if pc_row else 0

    inputs_for_prompt = dict(data)
    inputs_for_prompt["points_of_contact"] = poc
    inputs_for_prompt["eta_next_date"] = eta_str
    inputs_for_prompt["voicemail_left"] = voicemail
    inputs_for_prompt["phone_number_used"] = phone_used
    inputs_for_prompt["photos_count"] = update_photos_count
    inputs_for_prompt["agent_notes"] = (data.get("agent_notes") or "").strip()

    prompt = _build_swpi_prompt(inputs_for_prompt, job_ctx)

    try:
        ai_client, key_source = _get_ai_client()
        response = ai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.3,
        )
        narrative = response.choices[0].message.content.strip()
        tokens_used = response.usage.total_tokens if response.usage else None
        model_used = "gpt-4o-mini"
    except Exception as e:
        conn.close()
        return jsonify({"error": f"AI generation failed: {str(e)}"}), 500

    draft_id = data.get("draft_id")
    ts = now_ts()
    structured_json = json.dumps({k: v for k, v in data.items() if k != "draft_id"})
    if draft_id:
        conn.execute("""
            UPDATE job_updates SET
                attend_date=?, attend_time=?, is_first_attendance=?,
                property_description=?, security_sighted=?, security_make_model=?,
                security_reg=?, security_location=?, calling_card=?,
                neighbour_outcome=?, call_made=?, call_outcome=?,
                voicemail_left=?, sms_sent=?, customer_mobile=?,
                points_of_contact=?, eta_next_date=?,
                generated_narrative=?, ai_model_used=?, ai_tokens_used=?,
                structured_inputs_json=?, agent_notes=?, updated_at=?
            WHERE id=? AND created_by_user_id=?
        """, (
            data.get("attend_date"), data.get("attend_time"),
            1 if is_first else 0,
            data.get("property_description", ""),
            1 if data.get("security_sighted") else 0,
            data.get("security_make_model", ""), data.get("security_reg", ""), data.get("security_location", ""),
            1 if cc else 0,
            data.get("neighbour_outcome", ""),
            1 if call_made else 0, data.get("call_outcome", ""),
            1 if voicemail else 0, 1 if sms_sent else 0,
            phone_used,
            poc, eta_str,
            narrative, model_used, tokens_used,
            structured_json, (data.get("agent_notes") or "").strip(), ts,
            draft_id, uid
        ))
        conn.commit()

    conn.execute("""
        INSERT INTO ai_usage_log (user_id, job_id, feature, model, tokens_used, key_source, created_at)
        VALUES (?, ?, 'update_builder', ?, ?, ?, ?)
    """, (uid, job_id, model_used, tokens_used, key_source, ts))
    conn.commit()
    conn.close()

    audit("job_update_draft", draft_id or 0, "generate",
          f"AI narrative generated for job {job_id}",
          {"poc": poc, "eta": eta_str, "tokens": tokens_used})

    return jsonify({
        "narrative": narrative,
        "poc": poc,
        "eta": eta_str,
        "tokens": tokens_used,
    })


@app.post("/jobs/<int:job_id>/update-builder/save")
@login_required
def update_builder_save(job_id: int):
    conn = db()
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({"error": "Job not found"}), 404

    uid  = session.get("user_id")
    role = session.get("role", "")
    if role not in ("admin", "both") and job["assigned_user_id"] != uid:
        conn.close()
        return jsonify({"error": "Access denied"}), 403

    data = request.get_json(force=True)
    draft_id     = data.get("draft_id")
    final_text   = (data.get("final_narrative") or "").strip()
    gen_text     = (data.get("generated_narrative") or "").strip()
    was_edited   = 1 if final_text != gen_text else 0

    if not final_text:
        conn.close()
        return jsonify({"error": "No narrative to save"}), 400

    ts = now_ts()
    if draft_id:
        conn.execute("""
            UPDATE job_updates SET
                final_narrative=?, narrative_edited=?, status='complete',
                updated_at=?
            WHERE id=? AND created_by_user_id=?
        """, (final_text, was_edited, ts, draft_id, uid))
        conn.commit()

    photos_count = 0
    if draft_id:
        photos = conn.execute(
            "SELECT id, filename, filepath, tag FROM job_update_photos WHERE job_update_id=?", (draft_id,)
        ).fetchall()
        photos_count = len(photos)

    note_type = "text"
    if photos_count > 0:
        note_type = "photo_text" if final_text else "photo"

    cur = conn.cursor()
    cur.execute(
        "INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at, note_type) VALUES (?,?,?,?,?)",
        (job_id, uid, final_text, ts, note_type)
    )
    note_id = cur.lastrowid

    if photos_count > 0:
        for p in photos:
            cur.execute(
                "INSERT INTO job_note_files (job_field_note_id, filename, filepath, uploaded_at) VALUES (?,?,?,?)",
                (note_id, p["filename"], p["filepath"], ts)
            )
    conn.commit()

    for cue_type in ("Update Required", "Complete Attendance Update"):
        cue_today = conn.execute(
            "SELECT id FROM cue_items WHERE job_id=? AND assigned_user_id=? AND visit_type=? AND status IN ('Pending','In Progress')",
            (job_id, uid, cue_type)
        ).fetchone()
        if cue_today:
            conn.execute("UPDATE cue_items SET status='Completed', completed_at=?, updated_at=? WHERE id=?",
                         (ts, ts, cue_today["id"]))
            conn.commit()

    try:
        existing_review = conn.execute(
            "SELECT id FROM cue_items WHERE job_id=? AND visit_type='Agent Note Review' AND status='Pending'",
            (job_id,)
        ).fetchone()
        if not existing_review:
            from datetime import datetime as _dt2
            _today = _dt2.now(_melbourne).date().isoformat()
            conn.execute("""
                INSERT INTO cue_items
                  (job_id, visit_type, due_date, priority, status,
                   instructions, created_by_user_id, created_at, updated_at)
                VALUES (?, 'Agent Note Review', ?, 'High', 'Pending', ?, ?, ?, ?)
            """, (job_id, _today, final_text[:200], uid, ts, ts))
            conn.commit()
    except Exception:
        pass

    conn.close()
    audit("job_update", draft_id or 0, "save",
          f"Field update saved for job {job_id}",
          {"edited": was_edited, "photos": photos_count})
    return jsonify({"ok": True, "redirect": url_for("job_detail", job_id=job_id, _anchor="tab-notes")})


@app.get("/jobs/<int:job_id>/update-builder/draft-check")
@login_required
def update_builder_draft_check(job_id: int):
    uid = session.get("user_id")
    conn = db()
    draft = conn.execute(
        "SELECT id FROM job_updates WHERE job_id=? AND created_by_user_id=? AND status='draft' LIMIT 1",
        (job_id, uid)
    ).fetchone()
    conn.close()
    return jsonify({"has_draft": bool(draft), "draft_id": draft["id"] if draft else None})


@app.post("/jobs/<int:job_id>/update-builder/autosave")
@login_required
def update_builder_autosave(job_id: int):
    conn = db()
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({"ok": False}), 404

    uid  = session.get("user_id")
    role = session.get("role", "")
    if role not in ("admin", "both") and job["assigned_user_id"] != uid:
        conn.close()
        return jsonify({"ok": False}), 403

    data = request.get_json(force=True)
    draft_id = data.get("draft_id")
    ts = now_ts()
    structured_json = json.dumps({k: v for k, v in data.items() if k != "draft_id"})

    if draft_id:
        conn.execute("""
            UPDATE job_updates SET
                attend_date=?, attend_time=?, is_first_attendance=?,
                property_description=?, security_sighted=?, security_make_model=?,
                security_reg=?, security_location=?, calling_card=?,
                neighbour_outcome=?, call_made=?, call_outcome=?,
                voicemail_left=?, sms_sent=?, customer_mobile=?,
                structured_inputs_json=?, agent_notes=?, updated_at=?
            WHERE id=? AND created_by_user_id=? AND status='draft'
        """, (
            data.get("attend_date"), data.get("attend_time"),
            1 if data.get("is_first_attendance") else 0,
            data.get("property_description", ""),
            1 if data.get("security_sighted") else 0,
            data.get("security_make_model", ""), data.get("security_reg", ""),
            data.get("security_location", ""),
            1 if data.get("calling_card") else 0,
            data.get("neighbour_outcome", ""),
            1 if data.get("call_made") else 0, data.get("call_outcome", ""),
            1 if data.get("voicemail_left") else 0,
            1 if data.get("sms_sent") else 0,
            data.get("phone_number_used", ""),
            structured_json, (data.get("agent_notes") or "").strip(), ts,
            draft_id, uid
        ))
        conn.commit()

    leaving = request.args.get("leaving") == "1"
    if leaving:
        job_ref = job["display_ref"] or job["internal_job_number"] or str(job_id)
        _ensure_draft_cue(conn, job_id, uid, draft_id, job_ref)
    conn.close()
    return jsonify({"ok": True})


_UPDATE_PHOTO_DIR = os.path.join("uploads", "update_photos")
os.makedirs(_UPDATE_PHOTO_DIR, exist_ok=True)

_UPDATE_PHOTO_ALLOWED = {"png", "jpg", "jpeg", "webp", "heic", "heif"}
_UPDATE_PHOTO_MAX_BYTES = 25 * 1024 * 1024


@app.post("/jobs/<int:job_id>/update-builder/upload-photo")
@login_required
def update_builder_upload_photo(job_id):
    conn = db()
    job = conn.execute("SELECT id, assigned_user_id FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({"ok": False, "error": "Job not found"}), 404
    uid = session.get("user_id")
    role = session.get("role", "")
    if role not in ("admin", "both") and job["assigned_user_id"] != uid:
        conn.close()
        return jsonify({"ok": False, "error": "Access denied"}), 403

    draft_id = request.form.get("draft_id", type=int)
    if not draft_id:
        conn.close()
        return jsonify({"ok": False, "error": "No draft ID"}), 400

    draft = conn.execute(
        "SELECT id FROM job_updates WHERE id=? AND job_id=? AND created_by_user_id=? AND status='draft'",
        (draft_id, job_id, uid)
    ).fetchone()
    if not draft:
        conn.close()
        return jsonify({"ok": False, "error": "Draft not found"}), 404

    photo = request.files.get("photo")
    if not photo or not photo.filename:
        conn.close()
        return jsonify({"ok": False, "error": "No photo provided"}), 400

    ext = photo.filename.rsplit(".", 1)[-1].lower() if "." in photo.filename else ""
    if ext not in _UPDATE_PHOTO_ALLOWED:
        conn.close()
        return jsonify({"ok": False, "error": f"File type .{ext} not allowed"}), 400

    photo.seek(0, 2)
    size = photo.tell()
    photo.seek(0)
    if size > _UPDATE_PHOTO_MAX_BYTES:
        conn.close()
        return jsonify({"ok": False, "error": "Photo exceeds 25 MB limit"}), 400
    if size == 0:
        conn.close()
        return jsonify({"ok": False, "error": "Empty file"}), 400

    existing_count = conn.execute(
        "SELECT COUNT(*) c FROM job_update_photos WHERE job_update_id=?", (draft_id,)
    ).fetchone()["c"]
    if existing_count >= 10:
        conn.close()
        return jsonify({"ok": False, "error": "Maximum 10 photos per update"}), 400

    tag = (request.form.get("tag") or "general").strip()[:50]
    safe_fn = secure_filename(photo.filename)
    stored_name = f"{uuid.uuid4().hex}_{safe_fn}"
    filepath = os.path.join(_UPDATE_PHOTO_DIR, stored_name)
    photo.save(filepath)

    ts = now_ts()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO job_update_photos (job_update_id, job_id, filename, filepath, tag, uploaded_at) VALUES (?,?,?,?,?,?)",
        (draft_id, job_id, stored_name, filepath, tag, ts)
    )
    photo_id = cur.lastrowid
    conn.execute(
        "UPDATE job_updates SET photos_count = (SELECT COUNT(*) FROM job_update_photos WHERE job_update_id=?), updated_at=? WHERE id=?",
        (draft_id, ts, draft_id)
    )
    conn.commit()
    conn.close()

    return jsonify({"ok": True, "photo_id": photo_id, "url": f"/jobs/{job_id}/update-builder/photo/{photo_id}"})


@app.delete("/jobs/<int:job_id>/update-builder/photo/<int:photo_id>")
@login_required
def update_builder_delete_photo(job_id, photo_id):
    uid = session.get("user_id")
    conn = db()
    row = conn.execute(
        """SELECT p.id, p.filepath, p.job_update_id FROM job_update_photos p
           JOIN job_updates u ON u.id = p.job_update_id
           WHERE p.id=? AND p.job_id=? AND u.created_by_user_id=? AND u.status='draft'""",
        (photo_id, job_id, uid)
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "Photo not found"}), 404

    try:
        if os.path.exists(row["filepath"]):
            os.remove(row["filepath"])
    except OSError:
        pass

    ts = now_ts()
    conn.execute("DELETE FROM job_update_photos WHERE id=?", (photo_id,))
    conn.execute(
        "UPDATE job_updates SET photos_count = (SELECT COUNT(*) FROM job_update_photos WHERE job_update_id=?), updated_at=? WHERE id=?",
        (row["job_update_id"], ts, row["job_update_id"])
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.get("/jobs/<int:job_id>/update-builder/photo/<int:photo_id>")
@login_required
def update_builder_serve_photo(job_id, photo_id):
    uid = session.get("user_id")
    role = session.get("role", "")
    conn = db()
    row = conn.execute(
        """SELECT p.filepath FROM job_update_photos p
           JOIN job_updates u ON u.id = p.job_update_id
           WHERE p.id=? AND p.job_id=?""",
        (photo_id, job_id)
    ).fetchone()
    if not row:
        conn.close()
        abort(404)
    if role not in ("admin", "both"):
        job = conn.execute("SELECT assigned_user_id FROM jobs WHERE id=?", (job_id,)).fetchone()
        has_schedule = conn.execute(
            "SELECT 1 FROM schedules WHERE job_id=? AND assigned_to_user_id=? AND status NOT IN ('Cancelled','Completed')",
            (job_id, uid)
        ).fetchone()
        if not job or (job["assigned_user_id"] != uid and not has_schedule):
            conn.close()
            abort(403)
    conn.close()
    if not os.path.exists(row["filepath"]):
        abort(404)
    return send_file(row["filepath"])


@app.get("/my/drafts")
@login_required
def my_drafts():
    uid = session.get("user_id")
    conn = db()
    drafts = conn.execute("""
        SELECT ju.id AS draft_id, ju.job_id, ju.attend_date, ju.attend_time,
               ju.created_at, ju.updated_at,
               j.display_ref, j.internal_job_number, j.client_reference, j.job_address,
               (cu.first_name || ' ' || cu.last_name) AS customer_name
        FROM job_updates ju
        JOIN jobs j ON j.id = ju.job_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        WHERE ju.created_by_user_id = ? AND ju.status = 'draft'
        ORDER BY ju.updated_at DESC
    """, (uid,)).fetchall()
    conn.close()
    return render_template("my_drafts.html", drafts=drafts)


# -------- Users (admin only) --------
@app.get("/users")
@admin_required
def users_list():
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM users ORDER BY full_name")
    rows = cur.fetchall()
    conn.close()
    return render_template("users.html", users=rows)


@app.get("/users/new")
@admin_required
def user_new():
    next_url = request.args.get("next", "").strip()
    return render_template("user_new.html", next_url=next_url)


@app.post("/users/new")
@admin_required
def user_create():
    full_name = request.form.get("full_name", "").strip()
    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "").strip()
    role = request.form.get("role", "agent").strip()
    next_url = request.form.get("next_url", "").strip()

    if not full_name or not email or not password:
        flash("Name, email and password are required.", "danger")
        dest = url_for("user_new")
        if next_url:
            dest += "?next=" + next_url
        return redirect(dest)

    hashed = generate_password_hash(password)
    conn = db()
    cur = conn.cursor()
    user_id = None
    try:
        cur.execute("""
            INSERT INTO users (full_name, email, password, role, active, created_at)
            VALUES (?, ?, ?, ?, 1, ?)
        """, (full_name, email, hashed, role, now_ts()))
        user_id = cur.lastrowid
        conn.commit()
        audit("user", user_id, "create", f"User created: {full_name} ({role})", {"email": email, "role": role})
        flash("User created.", "success")
    except sqlite3.IntegrityError:
        flash("Email already in use.", "danger")
    finally:
        conn.close()

    if next_url and user_id:
        return redirect(f"{next_url}?new_user_id={user_id}")
    return redirect(url_for("users_list"))



@app.get("/users/new-popup")
@admin_required
def user_new_popup():
    return render_template("partials/user_popup.html")


@app.post("/users/new-popup")
@admin_required
def user_create_popup():
    full_name = request.form.get("full_name", "").strip()
    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "").strip()
    role = request.form.get("role", "agent").strip()
    if not full_name or not email or not password:
        return jsonify({"ok": False, "error": "Name, email and password are all required."})
    hashed = generate_password_hash(password)
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("""
            INSERT INTO users (full_name, email, password, role, active, created_at)
            VALUES (?, ?, ?, ?, 1, ?)
        """, (full_name, email, hashed, role, now_ts()))
        new_id = cur.lastrowid
        conn.commit()
        audit("user", new_id, "create", f"User created via popup: {full_name} ({role})", {"email": email, "role": role})
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({"ok": False, "error": "That email address is already in use."})
    conn.close()
    return jsonify({"ok": True, "id": new_id, "label": full_name})

@app.get("/admin/users/new")
@admin_required
def admin_user_new():
    return redirect(url_for("user_new"))


@app.post("/admin/users/new")
@admin_required
def admin_user_create():
    return redirect(url_for("user_create"))


@app.get("/users/<int:user_id>/edit")
@admin_required
def user_edit(user_id: int):
    conn = db()
    user = conn.execute("SELECT * FROM users WHERE id = ?", (user_id,)).fetchone()
    conn.close()
    if not user:
        flash("User not found.", "danger")
        return redirect(url_for("users_list"))
    return render_template("user_edit.html", u=user)


@app.post("/users/<int:user_id>/edit")
@admin_required
def user_edit_save(user_id: int):
    full_name = request.form.get("full_name", "").strip()
    email = request.form.get("email", "").strip().lower()
    role = request.form.get("role", "agent").strip()
    active = 1 if request.form.get("active") else 0
    new_password = request.form.get("new_password", "").strip()

    if not full_name or not email:
        flash("Name and email are required.", "danger")
        return redirect(url_for("user_edit", user_id=user_id))

    conn = db()
    try:
        if new_password:
            conn.execute(
                "UPDATE users SET full_name=?, email=?, role=?, active=?, password=? WHERE id=?",
                (full_name, email, role, active, generate_password_hash(new_password), user_id)
            )
        else:
            conn.execute(
                "UPDATE users SET full_name=?, email=?, role=?, active=? WHERE id=?",
                (full_name, email, role, active, user_id)
            )
        conn.commit()
        audit("user", user_id, "edit", f"User updated: {full_name} ({role})", {"email": email, "active": active})
        flash("User updated.", "success")
    except sqlite3.IntegrityError:
        flash("That email address is already in use.", "danger")
    finally:
        conn.close()
    return redirect(url_for("users_list"))


@app.post("/users/<int:user_id>/delete")
@admin_required
def user_delete(user_id: int):
    if user_id == session.get("user_id"):
        flash("You cannot delete your own account.", "danger")
        return redirect(url_for("users_list"))
    conn = db()
    user = conn.execute("SELECT full_name FROM users WHERE id = ?", (user_id,)).fetchone()
    if user:
        conn.execute("DELETE FROM users WHERE id = ?", (user_id,))
        conn.commit()
        audit("user", user_id, "delete", f"User deleted: {user['full_name']}", {})
        flash("User deleted.", "success")
    conn.close()
    return redirect(url_for("users_list"))


# -------- CSV Import --------
@app.get("/import/jobs")
@admin_required
def import_jobs_form():
    return redirect(url_for("admin_settings") + "#import-data")


@app.post("/import/jobs")
@admin_required
def import_jobs():
    file = request.files.get("file")
    if not file:
        flash("No file uploaded.", "danger")
        return redirect(url_for("admin_settings") + "#import-data")

    conn = db()
    cur = conn.cursor()

    # ── Pre-load lookup tables ────────────────────────────────────────────
    all_clients = conn.execute("SELECT id, name, nickname FROM clients").fetchall()
    all_users   = conn.execute("SELECT id, full_name FROM users WHERE active = 1").fetchall()

    def _client_initials(name):
        skip = {"and", "&", "the", "of", "for", "pty", "ltd", "inc", "co"}
        return "".join(w[0].upper() for w in name.split() if w.lower() not in skip and w.isalpha())

    client_by_initials  = {}
    client_by_name      = {}
    client_by_nickname  = {}
    for c in all_clients:
        client_by_initials[_client_initials(c["name"])] = c["id"]
        client_by_name[c["name"].lower()] = c["id"]
        if c["nickname"]:
            client_by_nickname[c["nickname"].strip().upper()] = c["id"]

    def _resolve_client(code):
        if not code:
            return None
        code = code.strip().upper()
        if code in client_by_nickname:
            return client_by_nickname[code]
        if code in client_by_initials:
            return client_by_initials[code]
        for n, cid in client_by_name.items():
            if code.lower() in n:
                return cid
        ts = datetime.now().isoformat(timespec="seconds")
        cur.execute(
            "INSERT INTO clients (name, created_at, updated_at) VALUES (?, ?, ?)",
            (code, ts, ts)
        )
        new_id = cur.lastrowid
        client_by_initials[code] = new_id
        client_by_name[code.lower()] = new_id
        client_by_nickname[code] = new_id
        client_created_names.append(code)
        return new_id

    def _resolve_user(staff_name):
        if not staff_name:
            return None
        parts = staff_name.strip().lower().split()
        for u in all_users:
            uname = u["full_name"].lower()
            if any(p in uname for p in parts if len(p) > 2):
                return u["id"]
        return None

    def _extract_job_type(raw):
        raw_lower = (raw or "").lower()
        if "repo/collect" in raw_lower:
            return "Repo/Collect"
        if "collect only" in raw_lower:
            return "Collect Only"
        if "field call" in raw_lower:
            return "Field Call"
        if "process serve" in raw_lower:
            return "Process Serve"
        if "repo only" in raw_lower or "upgraded" in raw_lower:
            return "Repo/Collect"
        return "Field Call"

    def _parse_date(raw):
        if not raw:
            return None
        raw = raw.strip()
        for fmt in ("%d-%b-%y", "%d/%m/%Y", "%Y-%m-%d", "%d-%b-%Y"):
            try:
                return datetime.strptime(raw, fmt).date().isoformat()
            except ValueError:
                pass
        return None

    def _parse_scheduled_for(date_raw, time_raw):
        d = _parse_date(date_raw)
        if not d:
            return None
        t = (time_raw or "").strip()
        if t:
            try:
                parsed_t = datetime.strptime(t, "%H:%M").strftime("%H:%M:%S")
            except ValueError:
                parsed_t = "09:00:00"
        else:
            parsed_t = "09:00:00"
        return f"{d} {parsed_t}"

    STATUS_MAP = {
        "awaiting advice from client": "Awaiting info from client",
        "awaiting info from client":   "Awaiting info from client",
        "active":     "Active",
        "new":        "New",
        "completed":  "Completed",
        "invoiced":   "Invoiced",
        "suspended":  "Suspended",
    }

    now_ts = datetime.now().isoformat(timespec="seconds")
    reader = csv.DictReader(file.stream.read().decode("utf-8-sig").splitlines())

    imported = 0
    skipped  = 0
    cust_created  = 0
    sched_created = 0
    client_created_names = []
    seen_job_numbers = set()

    new_visit_type_id = conn.execute(
        "SELECT id FROM booking_types WHERE name = 'New Visit' LIMIT 1"
    ).fetchone()
    new_visit_type_id = new_visit_type_id["id"] if new_visit_type_id else 1

    for row in reader:
        internal_job_number = (row.get("InternalJobNumber") or "").strip()
        if not internal_job_number:
            skipped += 1
            continue

        if internal_job_number in seen_job_numbers:
            skipped += 1
            continue
        seen_job_numbers.add(internal_job_number)

        client_reference = (row.get("ClientReference") or "").strip() or None
        display_ref = internal_job_number
        if client_reference:
            display_ref = f"{internal_job_number} ({client_reference})"

        addr_parts = [
            (row.get("Job Address 1") or "").strip(),
            (row.get("Job Address 2") or "").strip(),
            (row.get("Job Address City") or "").strip(),
            (row.get("Job Address State") or "").strip(),
            (row.get("Job Address Postcode") or "").strip(),
        ]
        job_address = ", ".join(p for p in addr_parts if p) or None

        raw_status = (row.get("Status") or "New").strip()
        status = STATUS_MAP.get(raw_status.lower(), raw_status)

        job_due_date   = _parse_date(row.get("Job Start Date"))
        scheduled_for  = _parse_scheduled_for(row.get("Job Start Date"), row.get("Job Start Time"))
        job_type       = _extract_job_type(row.get("JobType"))
        client_id      = _resolve_client(row.get("Bill Client Code"))
        assigned_uid = _resolve_user(row.get("Staff"))
        description  = (row.get("Job Description") or "").strip() or None
        priority     = (row.get("Priority") or "Normal").strip()

        # ── Resolve customer BEFORE inserting job so we can set customer_id ──
        cust_email   = (row.get("Customer Email") or "").strip() or None
        cust_company = (row.get("Customer Company") or "").strip() or None
        cust_first   = (row.get("Customer First Name") or "").strip() or None
        cust_last    = (row.get("Customer Last Name") or "").strip() or None
        cust_mobile  = (row.get("Customer Mobile") or "").strip() or None

        customer_id = None
        if cust_first or cust_last or cust_email:
            if cust_email:
                row2 = cur.execute("SELECT id FROM customers WHERE email = ?", (cust_email,)).fetchone()
                if row2:
                    customer_id = row2["id"]
            if not customer_id and cust_first and cust_last:
                row2 = cur.execute(
                    "SELECT id FROM customers WHERE LOWER(first_name)=? AND LOWER(last_name)=?",
                    (cust_first.lower(), cust_last.lower())
                ).fetchone()
                if row2:
                    customer_id = row2["id"]
            if not customer_id:
                cur.execute("""
                    INSERT INTO customers (first_name, last_name, company, email, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (cust_first or "", cust_last or "", cust_company, cust_email, now_ts, now_ts))
                customer_id = cur.lastrowid
                cust_created += 1
                if cust_mobile:
                    cur.execute("""
                        INSERT INTO contact_phone_numbers (entity_type, entity_id, label, phone_number, created_at)
                        VALUES ('customer', ?, 'Mobile', ?, ?)
                    """, (customer_id, cust_mobile, now_ts))

        cur.execute("""
            INSERT OR IGNORE INTO jobs (
                internal_job_number, client_reference, display_ref,
                job_type, visit_type, status, priority,
                job_address, description,
                client_id, customer_id, assigned_user_id,
                job_due_date,
                created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            internal_job_number, client_reference, display_ref,
            job_type, "New Visit", status, priority,
            job_address, description,
            client_id, customer_id, assigned_uid,
            job_due_date,
            now_ts, now_ts
        ))

        if not cur.rowcount:
            existing = cur.execute(
                "SELECT id, customer_id, client_id FROM jobs WHERE internal_job_number = ?",
                (internal_job_number,)
            ).fetchone()
            if not existing:
                skipped += 1
                continue
            job_id = existing["id"]
            patch = {}
            if customer_id and not existing["customer_id"]:
                patch["customer_id"] = customer_id
            if client_id and not existing["client_id"]:
                patch["client_id"] = client_id
            if patch:
                sets = ", ".join(f"{k} = ?" for k in patch)
                cur.execute(f"UPDATE jobs SET {sets} WHERE id = ?", (*patch.values(), job_id))
            skipped += 1
        else:
            job_id = cur.lastrowid
            imported += 1

        if customer_id:
            cur.execute("""
                INSERT OR IGNORE INTO job_customers (job_id, customer_id, role, sort_order, created_at)
                VALUES (?, ?, 'Debtor', 1, ?)
            """, (job_id, customer_id, now_ts))

        if scheduled_for:
            exists = cur.execute(
                "SELECT 1 FROM schedules WHERE job_id = ? AND scheduled_for = ?",
                (job_id, scheduled_for)
            ).fetchone()
            if not exists:
                cur.execute("""
                    INSERT INTO schedules
                        (job_id, booking_type_id, scheduled_for, status, assigned_to_user_id, created_at)
                    VALUES (?, ?, ?, 'Pending', ?, ?)
                """, (job_id, new_visit_type_id, scheduled_for, assigned_uid, now_ts))
                _write_schedule_history(cur, cur.lastrowid, job_id, "created",
                                        new_scheduled_for=scheduled_for, new_status="Pending",
                                        changed_by_user_id=session.get("user_id"),
                                        notes="Created via CSV import.")
                sched_created += 1

    conn.commit()
    conn.close()

    msg = (
        f"Import complete: {imported} jobs imported, {skipped} skipped (duplicates/blank). "
        f"{cust_created} new customer(s) created. {sched_created} schedule(s) created."
    )
    if client_created_names:
        names = ", ".join(client_created_names)
        msg += (
            f" {len(client_created_names)} new client(s) created from unrecognised codes: {names}. "
            f"Go to Clients to update their full names."
        )
    flash(msg, "success")
    return redirect(url_for("admin_settings") + "#import-data")


# -------- GeoOp Staged Import --------
import geoop_import as _geoop

@app.get("/admin/geoop-import")
@admin_required
def geoop_import_page():
    if not _geoop_is_unlocked():
        return render_template("geoop_login.html")
    conn = db()
    _geoop.ensure_staging_tables(conn)
    stats = {
        "staged_jobs": conn.execute("SELECT COUNT(*) c FROM geoop_staging_jobs").fetchone()["c"],
        "staged_notes": conn.execute("SELECT COUNT(*) c FROM geoop_staging_notes").fetchone()["c"],
        "staged_files": conn.execute("SELECT COUNT(*) c FROM geoop_staging_files").fetchone()["c"],
        "manifest_records": conn.execute("SELECT COUNT(*) c FROM geoop_staging_files WHERE source_type IN ('job_csv','note_csv')").fetchone()["c"],
        "physical_files": conn.execute("SELECT COUNT(*) c FROM geoop_staging_files WHERE found_on_disk=1").fetchone()["c"],
        "imported_jobs": conn.execute("SELECT COUNT(*) c FROM geoop_staging_jobs WHERE import_status='imported'").fetchone()["c"],
        "imported_notes": conn.execute("SELECT COUNT(*) c FROM geoop_staging_notes WHERE import_status='imported'").fetchone()["c"],
        "azure_attachments": conn.execute("SELECT COUNT(*) c FROM geoop_staging_files WHERE source_type='azure_blob' AND import_status='imported'").fetchone()["c"],
    }
    runs = conn.execute("SELECT * FROM geoop_import_runs ORDER BY id DESC LIMIT 10").fetchall()
    conn.close()
    return render_template("geoop_import.html", stats=stats, runs=runs)


@app.post("/admin/geoop-import/stage")
@geoop_required
def geoop_import_stage():
    jobs_file = request.files.get("jobs_csv")
    notes_file = request.files.get("notes_csv")

    conn = db()
    run_id = None
    try:
        _geoop.ensure_staging_tables(conn)
        ts = _geoop._now()
        uid = session.get("user_id")

        cur = conn.cursor()
        cur.execute("""
            INSERT INTO geoop_import_runs (run_type, status, started_at, run_by_user_id)
            VALUES ('stage', 'running', ?, ?)
        """, (ts, uid))
        run_id = cur.lastrowid
        conn.commit()

        total_jobs = 0
        total_notes = 0
        notes_result = None

        if jobs_file and jobs_file.filename:
            jobs_path = os.path.join(_geoop.GEOOP_IMPORT_DIR, f"jobs_{run_id}.csv")
            jobs_file.save(jobs_path)
            r = _geoop.stage_jobs_csv(jobs_path, conn)
            total_jobs = r["inserted"]

        if notes_file and notes_file.filename:
            notes_path = os.path.join(_geoop.GEOOP_IMPORT_DIR, f"notes_{run_id}.csv")
            notes_file.save(notes_path)
            notes_result = _geoop.stage_notes_csv(notes_path, conn)
            total_notes = notes_result["inserted"]

        _geoop.build_file_manifest_from_csv(conn)

        diag = _geoop.generate_diagnostics(conn)

        conn.execute("""
            UPDATE geoop_import_runs SET status='completed', total_jobs=?, total_notes=?,
            diagnostics_json=?, completed_at=? WHERE id=?
        """, (total_jobs, total_notes, json.dumps(diag), _geoop._now(), run_id))
        conn.commit()

        notes_breakdown = notes_result.get("breakdown", {}) if notes_result else {}

        msg = f"Staging complete: {total_jobs} jobs and {total_notes} notes staged."
        if notes_breakdown:
            parts = []
            for k, v in notes_breakdown.items():
                if v > 0:
                    parts.append(f"{k}: {v}")
            if parts:
                msg += f" Skips: {', '.join(parts)}."
            try:
                rej_count = conn.execute("SELECT COUNT(*) c FROM geoop_notes_rejects").fetchone()["c"]
                if rej_count:
                    msg += f" {rej_count} rejected rows logged for review."
            except Exception:
                pass

        flash(msg, "success")
    except Exception as e:
        try:
            if run_id:
                conn.execute(
                    "UPDATE geoop_import_runs SET status='failed', completed_at=? WHERE id=? AND status='running'",
                    (_geoop._now(), run_id)
                )
                conn.commit()
        except Exception:
            pass
        flash(f"Staging failed: {e}", "danger")
    finally:
        conn.close()
    return redirect(url_for("geoop_import_page"))


@app.get("/admin/geoop-import/diagnostics")
@geoop_required
def geoop_import_diagnostics():
    conn = db()
    _geoop.ensure_staging_tables(conn)
    diag = _geoop.generate_diagnostics(conn)

    sample = conn.execute("""
        SELECT geoop_job_id, reference_no, job_title, status_label,
               parsed_client_name, parsed_account_number, parsed_regulation_type,
               parsed_amount_type, parsed_amount_cents, parsed_costs_cents,
               parsed_reg, parsed_vin, parsed_security_make, parsed_security_model,
               parsed_security_colour, parsed_security_year, parsed_deliver_to
        FROM geoop_staging_jobs ORDER BY id LIMIT 20
    """).fetchall()

    unparsed = conn.execute("""
        SELECT geoop_job_id, reference_no, job_title, raw_description
        FROM geoop_staging_jobs
        WHERE parsed_client_name = '' AND parsed_reg = '' AND parsed_vin = ''
        LIMIT 10
    """).fetchall()

    conn.close()
    return jsonify({"diagnostics": diag, "sample_parsed": [dict(r) for r in sample],
                     "unparsed_samples": [dict(r) for r in unparsed]})


@app.post("/admin/geoop-import/execute")
@geoop_required
def geoop_import_execute():
    mode = request.form.get("mode", "insert_only")
    if mode not in ("insert_only", "update"):
        mode = "insert_only"
    import_notes = request.form.get("import_notes") == "1"

    conn = db()
    run_id = None
    try:
        _geoop.ensure_staging_tables(conn)
        ts = _geoop._now()
        uid = session.get("user_id")

        cur = conn.cursor()
        cur.execute("""
            INSERT INTO geoop_import_runs (run_type, status, started_at, run_by_user_id)
            VALUES ('import', 'running', ?, ?)
        """, (ts, uid))
        run_id = cur.lastrowid
        conn.commit()

        job_result = _geoop.import_staged_jobs(mode=mode, conn=conn)

        note_result = {"imported": 0, "skipped": 0, "unmatched": 0, "errors": 0}
        if import_notes:
            note_result = _geoop.import_staged_notes(conn)

        conn.execute("""
            UPDATE geoop_import_runs SET status='completed',
            jobs_imported=?, notes_imported=?,
            jobs_skipped=?, notes_skipped=?,
            errors=?, completed_at=?
            WHERE id=?
        """, (
            job_result["imported"], note_result["imported"],
            job_result["skipped"], note_result["skipped"],
            job_result["errors"] + note_result["errors"],
            _geoop._now(), run_id
        ))
        conn.commit()

        msg = (
            f"Import complete: {job_result['imported']} jobs imported, "
            f"{job_result['skipped']} skipped, {job_result['errors']} errors."
        )
        if import_notes:
            msg += (
                f" Notes: {note_result['imported']} imported, "
                f"{note_result['unmatched']} unmatched, {note_result['errors']} errors."
            )
        flash(msg, "success")
    except Exception as e:
        try:
            if run_id:
                conn.execute(
                    "UPDATE geoop_import_runs SET status='failed', completed_at=? WHERE id=? AND status='running'",
                    (_geoop._now(), run_id)
                )
                conn.commit()
        except Exception:
            pass
        flash(f"Import failed: {e}", "danger")
    finally:
        conn.close()
    return redirect(url_for("geoop_import_page"))


@app.get("/admin/geoop-import/rejects")
@geoop_required
def geoop_import_rejects():
    conn = db()
    try:
        _geoop._ensure_rejects_table(conn)
        summary = conn.execute("""
            SELECT reject_reason, COUNT(*) c FROM geoop_notes_rejects GROUP BY reject_reason ORDER BY c DESC
        """).fetchall()

        samples = {}
        for row in summary:
            reason = row["reject_reason"]
            rows = conn.execute("""
                SELECT csv_row_number, geoop_job_id, geoop_note_id, reject_reason,
                       raw_note_description, raw_fields_json
                FROM geoop_notes_rejects WHERE reject_reason = ? LIMIT 20
            """, (reason,)).fetchall()
            samples[reason] = [dict(r) for r in rows]

        total = conn.execute("SELECT COUNT(*) c FROM geoop_notes_rejects").fetchone()["c"]
        staged = conn.execute("SELECT COUNT(*) c FROM geoop_staging_notes").fetchone()["c"]
    except Exception:
        return jsonify({"error": "Rejects table not found. Run staging first.", "total_rejects": 0, "total_staged": 0, "summary": [], "samples": {}})
    finally:
        conn.close()

    return jsonify({
        "total_rejects": total,
        "total_staged": staged,
        "summary": [dict(r) for r in summary],
        "samples": samples,
    })


@app.post("/admin/geoop-import/scan-attachments")
@geoop_required
def geoop_import_scan_attachments():
    attach_path = (request.form.get("attachments_path") or "").strip()
    if not attach_path or not os.path.isdir(attach_path):
        flash(f"Invalid path: directory does not exist.", "danger")
        return redirect(url_for("geoop_import_page"))

    conn = db()
    try:
        _geoop.ensure_staging_tables(conn)
        result = _geoop.scan_attachment_dirs([attach_path], conn)
        flash(f"Attachment scan complete: {result['found']} files indexed, {result['skipped']} duplicates skipped.", "success")
    except Exception as e:
        flash(f"Attachment scan failed: {e}", "danger")
    finally:
        conn.close()
    return redirect(url_for("geoop_import_page"))


def _run_azure_scan_background(sas_url, run_id):
    import logging
    try:
        def _upload_to_storage_bg(data, blob_name, content_type):
            _save_bytes_to_storage(data, blob_name, content_type)

        result = _geoop.scan_azure_blob_attachments(
            sas_url, upload_fn=_upload_to_storage_bg, run_id=run_id
        )

        conn = db()
        try:
            final_status = "completed" if result.get("status") == "completed" else "failed"
            conn.execute("""
                UPDATE geoop_import_runs SET status=?,
                notes_imported=?, errors=?, completed_at=?,
                diagnostics_json=?
                WHERE id=?
            """, (
                final_status,
                result.get("attachments_linked", 0), result.get("errors", 0),
                _geoop._now(), json.dumps(result), run_id
            ))
            conn.commit()
        finally:
            conn.close()
    except Exception as e:
        logging.error("Azure Blob scan background job failed: %s", e, exc_info=True)
        try:
            fconn = db()
            try:
                fconn.execute(
                    "UPDATE geoop_import_runs SET status='failed', completed_at=? WHERE id=?",
                    (_geoop._now(), run_id)
                )
                fconn.commit()
            finally:
                fconn.close()
        except Exception:
            pass
        _geoop._persist_scan_progress(run_id, {
            "status": "failed",
            "error_message": str(e)[:500],
        })


@app.post("/admin/geoop-import/scan-azure")
@geoop_required
def geoop_import_scan_azure():
    sas_url = (request.form.get("azure_sas_url") or "").strip()
    if not sas_url:
        flash("Please provide a valid Azure Blob Storage container SAS URL.", "danger")
        return redirect(url_for("geoop_import_page"))

    from urllib.parse import urlparse
    parsed = urlparse(sas_url)
    if parsed.scheme != "https" or not parsed.hostname or not parsed.hostname.endswith(".blob.core.windows.net"):
        flash("Invalid SAS URL. Must be an HTTPS URL to *.blob.core.windows.net.", "danger")
        return redirect(url_for("geoop_import_page"))
    if not parsed.query or "sig=" not in parsed.query:
        flash("SAS URL appears to be missing authentication parameters (sig=).", "danger")
        return redirect(url_for("geoop_import_page"))

    conn = db()
    run_id = None
    try:
        _geoop.ensure_staging_tables(conn)
        ts = _geoop._now()
        uid = session.get("user_id")

        cur = conn.cursor()
        cur.execute("""
            INSERT INTO geoop_import_runs (run_type, status, started_at, run_by_user_id)
            VALUES ('azure_blob_scan', 'running', ?, ?)
        """, (ts, uid))
        run_id = cur.lastrowid
        conn.commit()
    except Exception as e:
        import logging
        logging.error("Failed to create Azure scan run record: %s", e, exc_info=True)
        flash("Failed to start Azure Blob scan. Check server logs.", "danger")
        return redirect(url_for("geoop_import_page"))
    finally:
        conn.close()

    import threading
    t = threading.Thread(
        target=_run_azure_scan_background,
        args=(sas_url, run_id),
        daemon=True,
    )
    t.start()

    flash(f"Azure Blob scan started (run #{run_id}). Progress updates will appear below.", "info")
    return redirect(url_for("geoop_import_page"))


@app.get("/admin/geoop-import/scan-azure/progress/<int:run_id>")
@geoop_required
def geoop_import_scan_azure_progress(run_id):
    progress = _geoop.get_azure_scan_progress(run_id)
    if progress:
        return jsonify(progress)
    return jsonify({"status": "not_found"}), 404


@app.get("/admin/geoop-import/unmatched-report/<int:run_id>")
@geoop_required
def geoop_unmatched_report(run_id):
    try:
        report = _geoop.get_unmatched_report(run_id)
    except Exception as e:
        return jsonify({"error": str(e)[:300], "total": 0, "by_reason": {}, "entries": []}), 500
    return jsonify(report)


@app.get("/admin/geoop-import/unmatched-report/<int:run_id>/csv")
@geoop_required
def geoop_unmatched_csv(run_id):
    import csv as csv_mod
    try:
        report = _geoop.get_unmatched_report(run_id)
    except Exception as e:
        return jsonify({"error": str(e)[:300]}), 500
    output = io.StringIO()
    writer = csv_mod.writer(output)
    writer.writerow(["zip_name", "entry_path", "filename", "geoop_job_id", "geoop_note_id", "reason"])
    for e in report["entries"]:
        writer.writerow([e["zip_name"], e["entry_path"], e["filename"],
                         e["geoop_job_id"], e["geoop_note_id"], e["reason"]])
    output.seek(0)
    return app.response_class(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment; filename=unmatched_attachments_run_{run_id}.csv"}
    )


@app.post("/admin/geoop-import/backfill-descriptions")
@geoop_required
def geoop_import_backfill():
    conn = db()
    run_id = None
    resume_checkpoint = None
    try:
        _geoop.ensure_staging_tables(conn)

        recovered = _geoop.recover_orphaned_backfill_runs()
        if recovered:
            for r in recovered:
                flash(f"Recovered orphaned backfill run #{r['run_id']} (was at batch {r['batch_number']}, staging ID {r['last_staging_id']}).", "info")

        existing = conn.execute(
            "SELECT id FROM geoop_import_runs WHERE run_type='description_backfill' AND status='running'"
        ).fetchone()
        if existing:
            flash(f"A backfill is already running (run #{existing['id']}). Please wait for it to finish.", "warning")
            return redirect(url_for("geoop_import_page"))

        resume_checkpoint = _geoop.get_last_backfill_checkpoint()

        ts = _geoop._now()
        uid = session.get("user_id")
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO geoop_import_runs (run_type, status, started_at, run_by_user_id)
            VALUES ('description_backfill', 'running', ?, ?)
        """, (ts, uid))
        run_id = cur.lastrowid
        conn.commit()
    except Exception as e:
        flash("Failed to start backfill. Check server logs.", "danger")
        return redirect(url_for("geoop_import_page"))
    finally:
        conn.close()

    import threading
    t = threading.Thread(
        target=_geoop.backfill_geoop_descriptions,
        args=(run_id,),
        kwargs={"resume_checkpoint": resume_checkpoint},
        daemon=True,
    )
    t.start()

    if resume_checkpoint:
        flash(f"Backfill resumed from run #{resume_checkpoint['run_id']} checkpoint (staging ID {resume_checkpoint['last_staging_id']}, {resume_checkpoint['jobs_processed']} already processed). Run #{run_id}.", "info")
    else:
        flash(f"Description backfill started (run #{run_id}). Progress updates will appear below.", "info")
    return redirect(url_for("geoop_import_page"))


@app.get("/admin/geoop-import/backfill-progress/<int:run_id>")
@geoop_required
def geoop_backfill_progress(run_id):
    progress = _geoop.get_backfill_progress(run_id)
    if progress:
        return jsonify(progress)
    return jsonify({"status": "not_found"}), 404


@app.get("/admin/geoop-import/scan-samples/<int:run_id>")
@geoop_required
def geoop_scan_samples(run_id):
    try:
        data = _geoop.get_scan_samples(run_id)
    except Exception as e:
        return jsonify({"error": str(e)[:300]}), 500
    return jsonify(data)


@app.get("/admin/geoop-import/backfill-samples/<int:run_id>")
@geoop_required
def geoop_backfill_samples(run_id):
    try:
        data = _geoop.get_backfill_samples(run_id)
    except Exception as e:
        return jsonify({"error": str(e)[:300]}), 500
    if not data:
        return jsonify({"error": "Run not found"}), 404
    return jsonify(data)


@app.post("/admin/geoop-import/backfill-clients")
@geoop_required
def geoop_import_backfill_clients():
    conn = db()
    try:
        existing = conn.execute(
            "SELECT id FROM geoop_import_runs WHERE run_type='client_backfill' AND status='running'"
        ).fetchone()
        if existing:
            flash(f"A client backfill is already running (run #{existing['id']}). Please wait.", "warning")
            return redirect(url_for("geoop_import_page"))

        cur = conn.cursor()
        cur.execute("""
            INSERT INTO geoop_import_runs (run_type, status, started_at, run_by_user_id)
            VALUES ('client_backfill', 'running', ?, ?)
        """, (now_ts(), session.get("user_id", 1)))
        run_id = cur.lastrowid
        conn.commit()
    except Exception as e:
        flash(f"Failed to start client backfill: {e}", "danger")
        return redirect(url_for("geoop_import_page"))
    finally:
        conn.close()

    import threading
    t = threading.Thread(
        target=_geoop.backfill_client_links,
        kwargs={"run_id": run_id},
        daemon=True
    )
    t.start()

    flash(f"Client backfill started (run #{run_id}). Progress updates will appear below.", "info")
    return redirect(url_for("geoop_import_page"))


@app.get("/admin/geoop-import/client-backfill-progress/<int:run_id>")
@geoop_required
def geoop_client_backfill_progress(run_id):
    progress = _geoop.get_backfill_progress(run_id)
    if not progress:
        return jsonify({"status": "not_found", "error": "Run not found"}), 404
    return jsonify(progress)


@app.get("/admin/geoop-import/client-gap-report")
@geoop_required
def geoop_client_gap_report():
    fmt = request.args.get("format", "json")
    report = _geoop.get_client_gap_report()

    if fmt == "csv":
        import io, csv as csv_mod
        output = io.StringIO()
        if report["rows"]:
            writer = csv_mod.DictWriter(output, fieldnames=report["rows"][0].keys())
            writer.writeheader()
            writer.writerows(report["rows"])
        resp = app.make_response(output.getvalue())
        resp.headers["Content-Type"] = "text/csv"
        resp.headers["Content-Disposition"] = "attachment; filename=client_gap_report.csv"
        return resp

    return jsonify(report)


@app.get("/admin/geoop-import/attachment-audit")
@geoop_required
def geoop_attachment_audit():
    action = request.args.get("action")
    if action == "start":
        started = _geoop.start_attachment_audit_background()
        if started:
            return jsonify({"status": "running", "message": "Audit started in background"})
        return jsonify({"status": "running", "message": "Audit already running"})

    cached = _geoop.get_attachment_audit_cached()
    if cached["result"]:
        resp = dict(cached["result"])
        resp["_audit_status"] = cached["status"]
        resp["_generated_at"] = cached["generated_at"]
        return jsonify(resp)

    if cached["status"] == "running":
        return jsonify({"_audit_status": "running", "_generated_at": None})

    if cached["status"] == "error":
        return jsonify({"_audit_status": "error", "_error": cached.get("error"), "_generated_at": None})

    return jsonify({"_audit_status": "idle", "_generated_at": None})


@app.post("/admin/geoop-import/backfill-attachments")
@geoop_required
def geoop_import_backfill_attachments():
    conn = db()
    try:
        existing = conn.execute(
            "SELECT id FROM geoop_import_runs WHERE run_type='attachment_backfill' AND status='running'"
        ).fetchone()
        if existing:
            flash(f"An attachment backfill is already running (run #{existing['id']}). Please wait.", "warning")
            return redirect(url_for("geoop_import_page"))

        cur = conn.cursor()
        cur.execute("""
            INSERT INTO geoop_import_runs (run_type, status, started_at, run_by_user_id)
            VALUES ('attachment_backfill', 'running', ?, ?)
        """, (now_ts(), session.get("user_id", 1)))
        run_id = cur.lastrowid
        conn.commit()
    except Exception as e:
        flash(f"Failed to start attachment backfill: {e}", "danger")
        return redirect(url_for("geoop_import_page"))
    finally:
        conn.close()

    import threading
    t = threading.Thread(
        target=_geoop.backfill_attachment_links,
        kwargs={"run_id": run_id},
        daemon=True
    )
    t.start()

    flash(f"Attachment backfill started (run #{run_id}). Progress updates will appear below.", "info")
    return redirect(url_for("geoop_import_page"))


_link_staged_lock = threading.Lock()

@app.post("/admin/geoop-import/link-staged-attachments")
@geoop_required
def geoop_link_staged_attachments():
    if not _link_staged_lock.acquire(blocking=False):
        flash("Link Staged Attachments is already running. Please wait.", "warning")
        return redirect(url_for("geoop_import_page"))

    def _run():
        try:
            conn = _geoop._db()
            try:
                linked = _geoop._link_staged_attachments(conn)
                print(f"[link-staged-attachments] Linked {linked} staged attachment(s)")
            except Exception as e:
                print(f"[link-staged-attachments] Error: {e}")
            finally:
                conn.close()
        finally:
            _link_staged_lock.release()

    t = threading.Thread(target=_run, daemon=True)
    t.start()
    flash("Link Staged Attachments started in background. Run the audit again in a moment to check progress.", "info")
    return redirect(url_for("geoop_import_page"))


@app.post("/admin/geoop-import/repair-dates")
@geoop_required
def geoop_repair_dates():
    target = request.args.get("target", "both")
    started_notes = False
    started_jobs = False
    if target in ("notes", "both"):
        started_notes = _geoop.repair_note_dates()
    if target in ("jobs", "both"):
        started_jobs = _geoop.repair_job_dates()
    if started_notes or started_jobs:
        parts = []
        if started_notes:
            parts.append("note dates")
        if started_jobs:
            parts.append("job dates")
        flash(f"Repair started for {' and '.join(parts)}. This runs in the background.", "info")
    else:
        flash("Repair is already running. Please wait for it to finish.", "warning")
    return redirect(url_for("geoop_import_page"))


@app.post("/admin/geoop-import/repair-phones")
@geoop_required
def geoop_repair_phones():
    started = _geoop.repair_phone_numbers()
    if started:
        flash("Phone number repair started. Restoring leading zeroes and international prefixes.", "info")
    else:
        flash("Phone repair is already running. Please wait.", "warning")
    return redirect(url_for("geoop_import_page"))


@app.post("/admin/geoop-import/backfill-job-ids")
@geoop_required
def geoop_backfill_job_ids():
    import logging as _log
    try:
        result = _geoop.backfill_geoop_job_ids()
        flash(f"GeoOp Job ID backfill complete: {result['updated']} updated, {result['skipped']} skipped.", "success")
    except Exception as e:
        _log.getLogger(__name__).error("Backfill geoop_job_ids error: %s", e)
        flash(f"Backfill error: {str(e)[:200]}", "danger")
    return redirect(url_for("geoop_import_page"))


@app.post("/admin/geoop-import/repair-registrations")
@geoop_required
def geoop_repair_registrations():
    started = _geoop.repair_registrations()
    if started:
        flash("Registration repair started. Re-parsing REG fields from original descriptions.", "info")
    else:
        flash("Registration repair is already running. Please wait.", "warning")
    return redirect(url_for("geoop_import_page"))


@app.post("/admin/geoop-import/repair-due-dates")
@geoop_required
def geoop_repair_due_dates():
    started = _geoop.repair_due_dates()
    if started:
        flash("Due date repair started. Re-parsing NMPD/Due dates from original descriptions.", "info")
    else:
        flash("Due date repair is already running. Please wait.", "warning")
    return redirect(url_for("geoop_import_page"))


@app.get("/admin/geoop-import/repair-progress")
@geoop_required
def geoop_repair_progress():
    return jsonify({
        "notes": _geoop.get_repair_dates_progress(),
        "jobs": _geoop.get_repair_job_dates_progress(),
        "phones": _geoop.get_repair_phones_progress(),
        "registrations": _geoop.get_repair_reg_progress(),
        "due_dates": _geoop.get_repair_due_dates_progress(),
    })


@app.get("/admin/geoop-import/repair-diagnostic")
@geoop_required
def geoop_repair_diagnostic():
    conn = _geoop._db()
    try:
        result = {}

        result["staging_total"] = conn.execute(
            "SELECT COUNT(*) FROM geoop_staging_jobs"
        ).fetchone()[0]
        result["staging_by_status"] = {
            r[0] or "NULL": r[1] for r in conn.execute(
                "SELECT import_status, COUNT(*) FROM geoop_staging_jobs GROUP BY import_status"
            ).fetchall()
        }
        result["staging_with_raw_desc"] = conn.execute(
            "SELECT COUNT(*) FROM geoop_staging_jobs WHERE raw_description IS NOT NULL AND raw_description != ''"
        ).fetchone()[0]
        result["staging_with_axion_job_id"] = conn.execute(
            "SELECT COUNT(*) FROM geoop_staging_jobs WHERE axion_job_id IS NOT NULL"
        ).fetchone()[0]
        result["staging_with_parsed_reg"] = conn.execute(
            "SELECT COUNT(*) FROM geoop_staging_jobs WHERE parsed_reg IS NOT NULL AND parsed_reg != ''"
        ).fetchone()[0]
        result["staging_with_parsed_nmpd"] = conn.execute(
            "SELECT COUNT(*) FROM geoop_staging_jobs WHERE parsed_nmpd_date IS NOT NULL AND parsed_nmpd_date != ''"
        ).fetchone()[0]

        result["jobs_total"] = conn.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
        result["jobs_with_geoop_source_desc"] = conn.execute(
            "SELECT COUNT(*) FROM jobs WHERE geoop_source_description IS NOT NULL AND geoop_source_description != ''"
        ).fetchone()[0]
        result["jobs_with_description"] = conn.execute(
            "SELECT COUNT(*) FROM jobs WHERE description IS NOT NULL AND description != ''"
        ).fetchone()[0]
        result["jobs_with_due_date"] = conn.execute(
            "SELECT COUNT(*) FROM jobs WHERE job_due_date IS NOT NULL AND job_due_date != ''"
        ).fetchone()[0]
        result["jobs_desc_contains_NPD"] = conn.execute(
            "SELECT COUNT(*) FROM jobs WHERE COALESCE(geoop_source_description, description) LIKE '%NPD%'"
        ).fetchone()[0]
        result["jobs_desc_contains_REG"] = conn.execute(
            "SELECT COUNT(*) FROM jobs WHERE COALESCE(geoop_source_description, description) LIKE '%REG%'"
        ).fetchone()[0]

        result["jobs_linked_to_staging"] = conn.execute(
            "SELECT COUNT(*) FROM jobs j JOIN geoop_staging_jobs sj ON sj.axion_job_id = j.id"
        ).fetchone()[0]

        result["job_items_vehicle"] = conn.execute(
            "SELECT COUNT(*) FROM job_items WHERE item_type = 'vehicle'"
        ).fetchone()[0]

        sample_npd = conn.execute("""
            SELECT j.id, j.job_due_date, j.mmp_cents,
                   substr(COALESCE(j.geoop_source_description, j.description), 1, 300) AS desc_preview
            FROM jobs j
            WHERE COALESCE(j.geoop_source_description, j.description) LIKE '%NPD%'
            LIMIT 5
        """).fetchall()
        result["sample_npd_jobs"] = [dict(r) for r in sample_npd]

        sample_raw_dates = conn.execute("""
            SELECT id, job_due_date FROM jobs
            WHERE job_due_date GLOB '[0-9][0-9]/[0-9][0-9]/[0-9][0-9]'
               OR job_due_date GLOB '[0-9]/[0-9]/[0-9][0-9]'
               OR job_due_date GLOB '[0-9][0-9]/[0-9][0-9]/[0-9][0-9][0-9][0-9]'
            LIMIT 10
        """).fetchall()
        result["jobs_with_raw_date_format"] = [dict(r) for r in sample_raw_dates]

        import re as _re
        npd_detail = conn.execute("""
            SELECT j.id, j.job_due_date, j.mmp_cents,
                   j.geoop_source_description AS desc_text
            FROM jobs j
            WHERE j.geoop_source_description LIKE '%NPD%'
              AND j.geoop_source_description IS NOT NULL
              AND j.geoop_source_description != ''
            LIMIT 10
        """).fetchall()
        repair_trace = []
        for r in npd_detail:
            desc = r["desc_text"]
            old_date = r["job_due_date"] or ""
            parsed = _geoop.parse_description(desc)
            new_date = parsed.get("parsed_nmpd_date", "")
            new_amount = parsed.get("parsed_nmpd_amount_cents", 0) or 0
            old_is_raw = bool(_re.match(r'^\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}$', old_date))
            old_normalised = _geoop._normalise_au_date(old_date) if old_is_raw else old_date
            would_update = bool(new_date) and new_date != old_normalised
            repair_trace.append({
                "job_id": r["id"],
                "old_date": old_date,
                "old_normalised": old_normalised,
                "new_date": new_date,
                "new_amount": new_amount,
                "old_is_raw": old_is_raw,
                "would_update": would_update,
                "skip_reason": (
                    "no_new_date" if not new_date else
                    "already_correct" if new_date == old_normalised else
                    "WILL_UPDATE"
                ),
                "desc_snippet": desc[:300],
            })
        result["repair_trace_npd_jobs"] = repair_trace

        return jsonify(result)
    finally:
        conn.close()


@app.post("/admin/geoop-import/recover-files")
@geoop_required
def geoop_recover_files():
    zip_dir = request.form.get("zip_dir", "").strip()
    if not zip_dir:
        zip_dir = os.path.join("uploads", "geoop_import")
    if not os.path.isdir(zip_dir):
        flash(f"Directory not found: {zip_dir}", "danger")
        return redirect(url_for("geoop_import_page"))

    zip_paths = sorted([
        os.path.join(zip_dir, f)
        for f in os.listdir(zip_dir)
        if f.lower().endswith(".zip") and os.path.isfile(os.path.join(zip_dir, f))
    ])
    if not zip_paths:
        flash(f"No ZIP files found in: {zip_dir}", "warning")
        return redirect(url_for("geoop_import_page"))

    started = _geoop.recover_files_from_zips(zip_paths)
    if started:
        flash(f"File recovery started from {len(zip_paths)} ZIP file(s) in {zip_dir}.", "info")
    else:
        flash("File recovery is already running. Please wait.", "warning")
    return redirect(url_for("geoop_import_page"))


@app.get("/admin/geoop-import/recover-files-progress")
@geoop_required
def geoop_recover_files_progress():
    return jsonify(_geoop.get_file_recovery_progress())


@app.get("/admin/geoop-import/repair-dates-progress")
@geoop_required
def geoop_repair_dates_progress():
    return jsonify({
        "notes": _geoop.get_repair_dates_progress(),
        "jobs": _geoop.get_repair_job_dates_progress(),
    })


@app.get("/admin/geoop-import/attachment-backfill-progress/<int:run_id>")
@geoop_required
def geoop_attachment_backfill_progress(run_id):
    progress = _geoop.get_backfill_progress(run_id)
    if not progress:
        return jsonify({"status": "not_found", "error": "Run not found"}), 404
    return jsonify(progress)


@app.post("/admin/geoop-import/reset")
@geoop_required
def geoop_import_reset():
    conn = db()
    _geoop.reset_staging(conn)
    conn.close()
    flash("All staging data cleared.", "info")
    return redirect(url_for("geoop_import_page"))


# -------- Admin dashboard --------
@app.get("/admin")
@admin_required
def admin_dashboard():
    today = datetime.now().date().isoformat()
    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT COUNT(*) c FROM cue_items WHERE due_date = ? AND status IN ('Pending','In Progress')", (today,))
    cues_today = cur.fetchone()["c"]

    cur.execute("SELECT COUNT(*) c FROM cue_items WHERE due_date < ? AND status IN ('Pending','In Progress')", (today,))
    cues_overdue = cur.fetchone()["c"]

    cur.execute("SELECT COUNT(*) c FROM cue_items WHERE assigned_user_id IS NULL AND status IN ('Pending','In Progress')")
    cues_unassigned = cur.fetchone()["c"]

    cur.execute("SELECT status, COUNT(*) c FROM jobs GROUP BY status ORDER BY c DESC")
    jobs_by_status = cur.fetchall()

    cur.execute("""
        SELECT a.*, u.full_name actor_name
        FROM audit_log a
        LEFT JOIN users u ON u.id = a.actor_user_id
        ORDER BY a.id DESC
        LIMIT 20
    """)
    recent = cur.fetchall()

    conn.close()
    return render_template("admin.html",
                           today=today,
                           cues_today=cues_today,
                           cues_overdue=cues_overdue,
                           cues_unassigned=cues_unassigned,
                           jobs_by_status=jobs_by_status,
                           recent=recent)


# -------- Admin Settings --------

@app.get("/admin/settings/popup")
@admin_required
def admin_settings_popup():
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM system_settings WHERE id = 1")
    settings = cur.fetchone()
    conn.close()
    return render_template("partials/settings_popup.html", settings=settings)


@app.post("/admin/settings/popup")
@admin_required
def admin_settings_popup_update():
    prefix   = request.form.get("job_prefix", "").strip()
    sequence = request.form.get("job_sequence", "0").strip()
    auto_enabled = 1 if request.form.get("auto_prefix_enabled") == "on" else 0

    if not prefix:
        return jsonify({"ok": False, "error": "Job prefix is required."})

    try:
        seq_int = int(sequence)
    except ValueError:
        return jsonify({"ok": False, "error": "Sequence must be a number."})

    conn = db()
    cur = conn.cursor()
    cur.execute("""
        UPDATE system_settings
        SET job_prefix = ?, job_sequence = ?, auto_prefix_enabled = ?, updated_at = ?
        WHERE id = 1
    """, (prefix, seq_int, auto_enabled, now_ts()))
    conn.commit()
    conn.close()

    audit("system", 1, "update", "Job numbering settings updated via popup",
          {"prefix": prefix, "sequence": sequence, "auto_enabled": auto_enabled})

    next_number = f"{prefix}{str(seq_int + 1).zfill(3)}"
    return jsonify({"ok": True, "next_number": next_number})

@app.get("/admin/settings")
@admin_required
def admin_settings():
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM system_settings WHERE id = 1")
    settings = cur.fetchone()
    cur.execute("SELECT * FROM booking_types ORDER BY name")
    booking_types = cur.fetchall()
    cur.execute("SELECT * FROM tow_operators ORDER BY company_name")
    tow_operators = cur.fetchall()
    cur.execute("SELECT * FROM auction_yards ORDER BY name")
    auction_yards = cur.fetchall()
    try:
        cur.execute("""
            SELECT al.*, u.full_name AS user_name, j.display_ref AS job_ref
            FROM ai_usage_log al
            LEFT JOIN users u ON u.id = al.user_id
            LEFT JOIN jobs j ON j.id = al.job_id
            ORDER BY al.created_at DESC LIMIT 50
        """)
        ai_usage = cur.fetchall()
    except Exception:
        ai_usage = []
    clients_list = cur.execute("SELECT id, name FROM clients ORDER BY name").fetchall()
    conn.close()
    return render_template("settings.html", settings=settings, booking_types=booking_types,
                           tow_operators=tow_operators, auction_yards=auction_yards,
                           ai_usage=ai_usage, clients_list=clients_list)


@app.get("/admin/api/duplicates")
@login_required
@admin_required
def admin_api_duplicates():
    conn = db()
    q = (request.args.get("q") or "").strip().lower()
    status_filter = (request.args.get("status") or "").strip()
    client_filter = (request.args.get("client") or "").strip()

    base_sel = """
        SELECT j.id, j.display_ref, j.internal_job_number, j.account_number,
               j.client_reference, j.client_job_number, j.lender_name, j.status,
               c.name AS client_name,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name,
                        cu.first_name || ' ' || cu.last_name) AS customer_name
        FROM jobs j
        LEFT JOIN clients c ON c.id = j.client_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
    """

    def _apply_filters(rows):
        out = []
        for _r in rows:
            r = dict(_r)
            if status_filter and r.get("status") != status_filter:
                continue
            if client_filter and str(r.get("client_name") or "").lower() != client_filter.lower():
                continue
            if q:
                haystack = " ".join(str(r.get(f) or "") for f in
                    ("display_ref", "internal_job_number", "account_number",
                     "client_reference", "client_job_number", "lender_name", "client_name", "customer_name")).lower()
                if q not in haystack:
                    continue
            out.append(r)
        return out

    # --- Group 1: same internal_job_number ---
    rows_by_jobnum = conn.execute(base_sel + """
        WHERE j.internal_job_number IS NOT NULL AND j.internal_job_number != ''
        ORDER BY LOWER(j.internal_job_number), j.id
    """).fetchall()

    jobnum_groups = {}
    for r in _apply_filters(rows_by_jobnum):
        key = (r["internal_job_number"] or "").strip().lower()
        if key:
            jobnum_groups.setdefault(key, []).append(r)
    dup_jobs_by_num = [
        {"key": v[0]["internal_job_number"].strip(), "match_type": "job_number", "jobs": v}
        for k, v in jobnum_groups.items() if len(v) > 1
    ]

    # --- Group 2: same non-null account_number ---
    rows_by_acct = conn.execute(base_sel + """
        WHERE j.account_number IS NOT NULL AND j.account_number != ''
        ORDER BY LOWER(j.account_number), j.id
    """).fetchall()

    acct_groups = {}
    for r in _apply_filters(rows_by_acct):
        key = (r["account_number"] or "").strip().lower()
        if key:
            acct_groups.setdefault(key, []).append(r)

    covered_ids = {j["id"] for grp in dup_jobs_by_num for j in grp["jobs"]}
    dup_jobs_by_acct = []
    for k, v in acct_groups.items():
        if len(v) > 1:
            uncovered = [j for j in v if j["id"] not in covered_ids]
            if len(uncovered) > 1:
                dup_jobs_by_acct.append(
                    {"key": v[0]["account_number"].strip(), "match_type": "account_number", "jobs": v}
                )

    # --- Group 3: same client_job_number ---
    rows_by_cjn = conn.execute(base_sel + """
        WHERE j.client_job_number IS NOT NULL AND j.client_job_number != ''
        ORDER BY LOWER(j.client_job_number), j.id
    """).fetchall()

    cjn_groups = {}
    for r in _apply_filters(rows_by_cjn):
        key = (r["client_job_number"] or "").strip().lower()
        if key:
            cjn_groups.setdefault(key, []).append(r)

    covered_ids_2 = covered_ids | {j["id"] for grp in dup_jobs_by_acct for j in grp["jobs"]}
    dup_jobs_by_cjn = []
    for k, v in cjn_groups.items():
        if len(v) > 1:
            uncovered = [j for j in v if j["id"] not in covered_ids_2]
            if len(uncovered) > 1:
                dup_jobs_by_cjn.append(
                    {"key": v[0]["client_job_number"].strip(), "match_type": "client_job_number", "jobs": v}
                )

    dup_jobs = dup_jobs_by_num + dup_jobs_by_acct + dup_jobs_by_cjn

    # --- Duplicate clients: same name ---
    dup_clients_raw = conn.execute("""
        SELECT id, name, email, phone
        FROM clients
        ORDER BY LOWER(name), id
    """).fetchall()
    name_groups = {}
    for r in dup_clients_raw:
        rname = (r["name"] or "").strip().lower()
        if rname:
            if q and q not in rname:
                continue
            name_groups.setdefault(rname, []).append(dict(r))
    dup_clients = [
        {"key": k, "clients": v}
        for k, v in name_groups.items() if len(v) > 1
    ]

    conn.close()
    return {"dup_jobs": dup_jobs, "dup_clients": dup_clients}


@app.post("/admin/api/duplicates/files")
@login_required
@admin_required
def admin_api_duplicates_files():
    files = request.files.getlist("files")
    if not files or all(f.filename == "" for f in files):
        return jsonify({"ok": False, "error": "No files selected."})

    conn = db()
    existing_jobs = {}
    for r in conn.execute("SELECT id, internal_job_number, display_ref, status FROM jobs WHERE internal_job_number IS NOT NULL AND internal_job_number != ''").fetchall():
        key = (r["internal_job_number"] or "").strip().lower()
        if key:
            existing_jobs.setdefault(key, []).append(dict(r))
    conn.close()

    file_records = []
    file_names = []
    parse_errors = []
    for f in files:
        if not f.filename:
            continue
        file_names.append(f.filename)
        try:
            content = f.stream.read().decode("utf-8-sig")
            reader = csv.DictReader(content.splitlines())
            for row_num, row in enumerate(reader, start=2):
                job_num = (row.get("InternalJobNumber") or "").strip()
                if not job_num:
                    continue
                client_ref = (row.get("ClientReference") or "").strip() or None
                customer_first = (row.get("Customer First Name") or "").strip()
                customer_last = (row.get("Customer Last Name") or "").strip()
                customer_name = f"{customer_first} {customer_last}".strip() or None
                file_records.append({
                    "file": f.filename,
                    "row": row_num,
                    "job_number": job_num,
                    "client_reference": client_ref,
                    "customer_name": customer_name,
                    "status": (row.get("Status") or "").strip() or None,
                    "job_type": (row.get("JobType") or "").strip() or None,
                })
        except Exception as e:
            parse_errors.append(f"{f.filename}: {str(e)}")

    dup_within_files = []
    job_num_groups = {}
    for rec in file_records:
        key = rec["job_number"].lower()
        job_num_groups.setdefault(key, []).append(rec)
    for key, recs in job_num_groups.items():
        if len(recs) > 1:
            dup_within_files.append({
                "key": recs[0]["job_number"],
                "match_type": "file_duplicate",
                "records": recs
            })

    dup_against_db = []
    for rec in file_records:
        key = rec["job_number"].lower()
        if key in existing_jobs:
            dup_against_db.append({
                "key": rec["job_number"],
                "match_type": "exists_in_database",
                "file_record": rec,
                "db_jobs": existing_jobs[key]
            })

    total_dups = len(dup_within_files) + len(dup_against_db)
    return jsonify({
        "ok": True,
        "file_names": file_names,
        "file_count": len(file_names),
        "total_records": len(file_records),
        "dup_within_files": dup_within_files,
        "dup_against_db": dup_against_db,
        "total_duplicates": total_dups,
        "parse_errors": parse_errors
    })


@app.post("/admin/api/duplicates/delete-job")
@login_required
@admin_required
def admin_api_delete_job_ajax():
    job_id = request.form.get("job_id")
    if not job_id:
        return jsonify({"ok": False, "error": "No job ID provided."})
    try:
        job_id = int(job_id)
    except ValueError:
        return jsonify({"ok": False, "error": "Invalid job ID."})

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT id, display_ref FROM jobs WHERE id = ?", (job_id,))
    job = cur.fetchone()
    if not job:
        conn.close()
        return jsonify({"ok": False, "error": "Job not found."})

    display_ref = job["display_ref"]

    cur.execute("SELECT id FROM job_field_notes WHERE job_id = ?", (job_id,))
    note_ids = [r["id"] for r in cur.fetchall()]
    for nid in note_ids:
        cur.execute("SELECT filename, filepath FROM job_note_files WHERE job_field_note_id = ?", (nid,))
        for f in cur.fetchall():
            delete_blob_safely(f["filename"])
            try: os.remove(f["filepath"])
            except OSError: pass
        cur.execute("DELETE FROM job_note_files WHERE job_field_note_id = ?", (nid,))
    cur.execute("DELETE FROM job_field_notes WHERE job_id = ?", (job_id,))

    cur.execute("SELECT stored_filename FROM job_documents WHERE job_id = ?", (job_id,))
    for f in cur.fetchall():
        delete_blob_safely(f["stored_filename"])
    cur.execute("DELETE FROM job_documents WHERE job_id = ?", (job_id,))

    for lpr_tbl in ("lpr_sightings", "lpr_patrol_intel"):
        try:
            cur.execute(f"UPDATE {lpr_tbl} SET matched_job_id=NULL WHERE matched_job_id=?", (job_id,))
        except Exception:
            pass

    for tbl in (
        "schedule_history",
        "job_items", "job_assets", "interactions", "cue_items", "schedules",
        "job_customers", "job_updates", "job_payments",
        "repo_lock_records", "repo_lock_queue",
    ):
        try:
            cur.execute(f"DELETE FROM {tbl} WHERE job_id = ?", (job_id,))
        except Exception:
            pass

    cur.execute("DELETE FROM jobs WHERE id = ?", (job_id,))
    conn.commit()
    conn.close()

    audit("job", job_id, "delete", f"Job {display_ref} deleted via Duplicate Finder", {})
    return jsonify({"ok": True, "job_id": job_id, "display_ref": display_ref})


@app.post("/admin/api/duplicates/delete-client")
@login_required
@admin_required
def admin_api_delete_client_ajax():
    client_id = request.form.get("client_id")
    if not client_id:
        return jsonify({"ok": False, "error": "No client ID provided."})
    try:
        client_id = int(client_id)
    except ValueError:
        return jsonify({"ok": False, "error": "Invalid client ID."})

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT name FROM clients WHERE id = ?", (client_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "Client not found."})

    name = row["name"]
    cur.execute(
        "SELECT COUNT(*) cnt FROM jobs WHERE client_id = ? OR bill_to_client_id = ?",
        (client_id, client_id)
    )
    job_count = cur.fetchone()["cnt"]
    if job_count > 0:
        conn.close()
        return jsonify({"ok": False, "error": f"Cannot delete '{name}' — {job_count} job{'s' if job_count != 1 else ''} linked."})

    cur.execute("DELETE FROM clients WHERE id = ?", (client_id,))
    conn.commit()
    conn.close()

    audit("client", client_id, "delete", f"Client '{name}' deleted via Duplicate Finder", {})
    return jsonify({"ok": True, "client_id": client_id, "name": name})


@app.post("/admin/settings")
@admin_required
def admin_settings_update():
    prefix = request.form.get("job_prefix", "").strip()
    sequence = request.form.get("job_sequence", "0").strip()
    auto_enabled = 1 if request.form.get("auto_prefix_enabled") == "on" else 0

    conn = db()
    cur = conn.cursor()
    cur.execute("""
        UPDATE system_settings
        SET job_prefix = ?, job_sequence = ?, auto_prefix_enabled = ?, updated_at = ?
        WHERE id = 1
    """, (prefix, int(sequence), auto_enabled, now_ts()))
    conn.commit()
    conn.close()

    audit("system", 1, "update",
          "Job numbering settings updated",
          {"prefix": prefix, "sequence": sequence, "auto_enabled": auto_enabled})

    flash("Settings saved.", "success")
    return redirect(url_for("admin_settings"))


@app.post("/admin/settings/ai")
@admin_required
def admin_settings_ai():
    use_own = 1 if request.form.get("ai_use_own_key") == "on" else 0
    own_key = request.form.get("openai_api_key", "").strip()
    conn = db()
    conn.execute("""
        UPDATE system_settings SET ai_use_own_key=?, openai_api_key=?, updated_at=? WHERE id=1
    """, (use_own, own_key if own_key else None, now_ts()))
    conn.commit()
    conn.close()
    audit("system", 1, "update", "AI settings updated", {"use_own_key": use_own})
    flash("AI settings saved.", "success")
    return redirect(url_for("admin_settings") + "#ai-settings")


@app.post("/admin/settings/archive")
@login_required
@admin_required
def admin_settings_archive():
    conn = db()
    conn.execute("""
        UPDATE system_settings SET
            archive_after_days = ?,
            cold_store_after_years = ?,
            archive_mode = ?,
            cold_storage_mode = ?,
            allow_restore_to_active = ?,
            allow_permanent_delete = ?,
            updated_at = ?
        WHERE id = 1
    """, (
        int(request.form.get("archive_after_days", 90)),
        int(request.form.get("cold_store_after_years", 3)),
        request.form.get("archive_mode", "manual"),
        request.form.get("cold_storage_mode", "manual"),
        int(request.form.get("allow_restore_to_active", 1)),
        int(request.form.get("allow_permanent_delete", 0)),
        now_ts()
    ))
    conn.commit()
    conn.close()
    audit("system", 1, "update", "Archive policy settings updated")
    flash("Archive policy settings saved.", "success")
    return redirect(url_for("admin_settings") + "#data-management")


# -------- Cues --------
def _queue_row_sql():
    return """
        SELECT ci.*,
               j.id AS job_id, j.internal_job_number, j.client_reference, j.display_ref,
               j.status AS job_status, j.job_address, j.assigned_user_id AS job_assigned_uid,
               c.name  AS client_name,
               (cu.first_name || ' ' || cu.last_name) AS customer_name,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
               (SELECT ji.reg FROM job_items ji
                WHERE ji.job_id = j.id AND ji.item_type = 'vehicle' LIMIT 1) AS asset_reg,
               ag.full_name AS agent_name,
               ag.email     AS agent_email,
               (SELECT ce.email FROM contact_emails ce
                WHERE ce.entity_type='client' AND ce.entity_id=j.client_id LIMIT 1) AS client_email
        FROM cue_items ci
        JOIN jobs j ON j.id = ci.job_id
        LEFT JOIN clients c  ON c.id  = j.client_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN users ag  ON ag.id = j.assigned_user_id
    """


@app.get("/queue")
@admin_required
def job_queue():
    import datetime as _dt
    mel_now = datetime.now(_melbourne)
    conn = db()
    cur = conn.cursor()
    try:
        admin_id = session.get("user_id")
        auto_queue_schedule_alerts(cur, admin_id)
        conn.commit()
    except Exception:
        pass

    overdue_types = ("Urgent: Schedule Overdue", "Schedule Due Today")
    tomorrow_type = "Schedule Due Tomorrow"
    note_type     = "Agent Note Review"

    _arch_excl = f"AND j.status NOT IN {ARCHIVED_STATUSES!r}"

    _note_excl = """AND ci.job_id NOT IN (
        SELECT job_id FROM cue_items
        WHERE visit_type = 'Agent Note Review' AND status = 'Pending'
    )"""

    cur.execute(_queue_row_sql() + f"""
        WHERE ci.visit_type IN (?,?) AND ci.status IN ('Pending','In Progress')
        {_arch_excl} {_note_excl}
        ORDER BY ci.priority DESC, ci.created_at DESC
    """, overdue_types)
    overdue = cur.fetchall()

    cur.execute(_queue_row_sql() + f"""
        WHERE ci.visit_type = ? AND ci.status IN ('Pending','In Progress')
        {_arch_excl} {_note_excl}
        ORDER BY ci.created_at DESC
    """, (tomorrow_type,))
    due_tomorrow = cur.fetchall()

    cur.execute(_queue_row_sql() + f"""
        WHERE ci.visit_type = ? AND ci.status = 'Pending'
        {_arch_excl}
        ORDER BY ci.updated_at DESC, ci.created_at DESC
    """, (note_type,))
    agent_notes = cur.fetchall()

    agents_list = cur.execute(
        "SELECT id, full_name, email FROM users WHERE role IN ('agent','both','admin') AND active=1 ORDER BY full_name"
    ).fetchall()
    clients_list = cur.execute(
        "SELECT id, name FROM clients ORDER BY name"
    ).fetchall()

    conn.close()
    return render_template("queue.html",
                           overdue=overdue,
                           due_tomorrow=due_tomorrow,
                           agent_notes=agent_notes,
                           agents_list=agents_list,
                           clients_list=clients_list,
                           now_melb=mel_now)


@app.post("/queue/new")
@admin_required
def cue_create():
    job_id = request.form.get("job_id", "").strip()
    visit_type = request.form.get("visit_type", "New Visit")
    due_date = request.form.get("due_date", "").strip()
    assigned_user_id = request.form.get("assigned_user_id") or None
    priority = request.form.get("priority", "Normal")
    instructions = request.form.get("instructions", "").strip()

    if not job_id or not due_date:
        flash("Job ID and due date are required.", "danger")
        return redirect(url_for("job_queue", date=due_date))

    ts = now_ts()
    conn = db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO cue_items (job_id, visit_type, due_date, priority, status, assigned_user_id, instructions, created_by_user_id, created_at, updated_at)
        VALUES (?, ?, ?, ?, 'Pending', ?, ?, ?, ?, ?)
    """, (job_id, visit_type, due_date, priority, assigned_user_id, instructions or None, session.get("user_id"), ts, ts))
    cue_id = cur.lastrowid
    conn.commit()
    conn.close()

    audit("cue", cue_id, "create",
          f"Cue created for job {job_id} due {due_date} ({visit_type}).",
          {"job_id": job_id, "due_date": due_date, "visit_type": visit_type, "assigned_user_id": assigned_user_id})

    flash("Item added to queue.", "success")
    return redirect(url_for("job_queue", date=due_date))


_ATTENDANCE_CUE_TYPES = {"New Visit", "Re-attend", "Urgent New Visit", "Re-Attend"}


@app.post("/cue/<int:cue_id>/complete")
@login_required
def cue_complete(cue_id: int):
    ts = now_ts()
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM cue_items WHERE id = ?", (cue_id,))
    cue = cur.fetchone()
    if not cue:
        conn.close()
        return ("Not found", 404)

    cur.execute("UPDATE cue_items SET status = 'Completed', completed_at = ?, updated_at = ? WHERE id = ?",
                (ts, ts, cue_id))
    conn.commit()
    job_id = cue["job_id"]
    visit_type = (cue["visit_type"] or "").strip()
    conn.close()

    audit("cue", cue_id, "status_change", f"Cue {cue_id} marked Completed.")

    if visit_type in _ATTENDANCE_CUE_TYPES:
        return redirect(url_for("update_builder", job_id=job_id))

    referrer = request.referrer or url_for("my_today")
    return redirect(referrer)


@app.post("/my/schedule/<int:sched_id>/attended")
@login_required
def my_schedule_attended(sched_id: int):
    conn = db()
    sched = conn.execute("SELECT * FROM schedules WHERE id=?", (sched_id,)).fetchone()
    if not sched:
        conn.close()
        abort(404)
    uid  = session.get("user_id")
    role = session.get("role", "")
    if role not in ("admin", "both") and sched["assigned_to_user_id"] != uid:
        conn.close()
        flash("Access denied.", "danger")
        return redirect(url_for("my_today"))
    ts = now_ts()
    old_status = sched["status"]
    old_scheduled = sched["scheduled_for"]
    conn.execute(
        "UPDATE schedules SET status='Completed' WHERE id=?",
        (sched_id,)
    )
    _write_schedule_history(conn.cursor(), sched_id, sched["job_id"], "completed",
                            old_scheduled_for=old_scheduled, new_scheduled_for=old_scheduled,
                            old_status=old_status, new_status="Completed",
                            changed_by_user_id=uid)
    conn.commit()
    conn.close()
    audit("schedule", sched_id, "status_change", "Schedule marked Attended via My Today.")
    return redirect(url_for("update_builder", job_id=sched["job_id"]))


@app.get("/queue/job-attachments/<int:job_id>")
@admin_required
def queue_job_attachments(job_id: int):
    conn = db()
    notes_rows = conn.execute("""
        SELECT fn.id, fn.note_text, fn.created_at,
               u.full_name AS staff_name
        FROM job_field_notes fn
        LEFT JOIN users u ON u.id = fn.created_by_user_id
        WHERE fn.job_id = ?
        ORDER BY fn.created_at DESC
    """, (job_id,)).fetchall()

    docs_rows = conn.execute("""
        SELECT d.id, d.title, d.original_filename, d.mime_type,
               d.uploaded_at, d.doc_type,
               u.full_name AS staff_name
        FROM job_documents d
        LEFT JOIN users u ON u.id = d.uploaded_by_user_id
        WHERE d.job_id = ?
        ORDER BY d.uploaded_at DESC
    """, (job_id,)).fetchall()

    forms_rows = conn.execute(
        "SELECT id, name, created_at FROM form_templates WHERE active=1 ORDER BY name"
    ).fetchall()

    job_row = conn.execute(
        "SELECT display_ref, description, job_address FROM jobs WHERE id=?", (job_id,)
    ).fetchone()

    sig_row = conn.execute("SELECT email_signature FROM system_settings WHERE id=1").fetchone()
    conn.close()

    def _fmt_dt(ts):
        try:
            return datetime.fromisoformat(ts).strftime("%d %b %Y %H:%M")
        except Exception:
            return ts or ""

    notes = [{
        "id": r["id"], "type": "note",
        "added": _fmt_dt(r["created_at"]),
        "description": (r["note_text"] or "")[:180],
        "staff": r["staff_name"] or ""
    } for r in notes_rows]

    docs = [{
        "id": r["id"], "type": "doc",
        "added": _fmt_dt(r["uploaded_at"]),
        "filename": r["original_filename"],
        "description": r["title"] or r["original_filename"],
        "mime": r["mime_type"] or "application/octet-stream",
        "doc_type": r["doc_type"] or "",
        "staff": r["staff_name"] or ""
    } for r in docs_rows]

    forms = [{
        "id": r["id"],
        "added": _fmt_dt(r["created_at"]),
        "description": r["name"],
        "status": "Available"
    } for r in forms_rows]

    smtp_from = os.environ.get("SMTP_FROM", os.environ.get("SMTP_USER", ""))
    return jsonify({
        "notes": notes, "docs": docs, "forms": forms,
        "job": dict(job_row) if job_row else {},
        "signature": (sig_row["email_signature"] or "") if sig_row else "",
        "smtp_from": smtp_from
    })


@app.post("/queue/<int:cue_id>/dismiss")
@admin_required
def queue_dismiss(cue_id: int):
    ts = now_ts()
    conn = db()
    conn.execute("UPDATE cue_items SET status='Completed', completed_at=?, updated_at=? WHERE id=?",
                 (ts, ts, cue_id))
    conn.commit()
    conn.close()
    audit("cue", cue_id, "dismiss", f"Queue item {cue_id} dismissed.")
    return jsonify({"ok": True})


@app.get("/queue/active-cue-ids")
@admin_required
def queue_active_cue_ids():
    conn = db()
    rows = conn.execute("""
        SELECT id FROM cue_items
        WHERE (
            (visit_type IN ('Urgent: Schedule Overdue', 'Schedule Due Today', 'Schedule Due Tomorrow')
             AND status IN ('Pending', 'In Progress'))
            OR
            (visit_type = 'Agent Note Review' AND status = 'Pending')
        )
    """).fetchall()
    conn.close()
    return jsonify({"ok": True, "ids": [r["id"] for r in rows]})


_VISIT_TYPES_FROM_BOOKING = {
    "new visit", "re-attend", "urgent new visit",
    "first update", "urgent update", "phone follow-up", "locate only",
    "update required", "urgent update required",
}

def _sync_visit_type_from_booking(cur, job_id, booking_type_name, ts=None):
    if ts is None:
        ts = now_ts()
    bt_lower = (booking_type_name or "").strip().lower()
    if bt_lower not in _VISIT_TYPES_FROM_BOOKING:
        return
    cur.execute("SELECT visit_type FROM jobs WHERE id = ?", (job_id,))
    row = cur.fetchone()
    if not row:
        return
    old_visit = row["visit_type"]
    if old_visit.strip().lower() == bt_lower:
        return
    cur.execute("UPDATE jobs SET visit_type = ?, updated_at = ? WHERE id = ?",
                (booking_type_name, ts, job_id))


def _auto_complete_schedule_cues(cur, job_id, ts=None):
    if ts is None:
        ts = now_ts()
    cur.execute("""
        UPDATE cue_items
        SET status = 'Completed', completed_at = ?, updated_at = ?
        WHERE job_id = ?
          AND visit_type IN ('Urgent: Schedule Overdue', 'Schedule Due Today', 'Schedule Due Tomorrow')
          AND status IN ('Pending', 'In Progress')
    """, (ts, ts, job_id))


@app.post("/queue/send-email")
@admin_required
def queue_send_email():
    import json as _json
    job_id           = request.form.get("job_id", "").strip()
    subject          = request.form.get("subject", "").strip()
    body             = request.form.get("body", "").strip()
    email_signature  = request.form.get("email_signature", "").strip()
    to_json          = request.form.get("to_recipients", "[]")
    cc_json          = request.form.get("cc_recipients", "[]")
    note_ids_json    = request.form.get("selected_note_ids", "[]")
    doc_ids_json     = request.form.get("selected_doc_ids", "[]")
    send_me_a_copy   = request.form.get("send_me_a_copy") == "1"

    if not job_id or not body:
        return jsonify({"ok": False, "error": "Job and message body are required."})

    try:
        to_list   = [e.strip() for e in _json.loads(to_json)  if e and "@" in e]
        cc_list   = [e.strip() for e in _json.loads(cc_json)  if e and "@" in e]
        note_ids  = [int(x) for x in _json.loads(note_ids_json)]
        doc_ids   = [int(x) for x in _json.loads(doc_ids_json)]
    except Exception:
        return jsonify({"ok": False, "error": "Invalid request data."})

    if send_me_a_copy:
        me = conn_user_email = db().execute(
            "SELECT email FROM users WHERE id=?", (session.get("user_id"),)
        ).fetchone()
        if me and me["email"] and me["email"] not in to_list and me["email"] not in cc_list:
            cc_list.append(me["email"])
        conn_user_email = None

    if not to_list:
        return jsonify({"ok": False, "error": "Please select or enter at least one recipient."})

    conn = db()
    job = conn.execute("""
        SELECT j.*, ag.email AS agent_email, ag.full_name AS agent_name
        FROM jobs j
        LEFT JOIN users ag ON ag.id = j.assigned_user_id
        WHERE j.id = ?
    """, (job_id,)).fetchone()

    if not job:
        conn.close()
        return jsonify({"ok": False, "error": "Job not found."})

    if not subject:
        subject = f"Job Update \u2014 {job['display_ref']}"

    full_body_txt = body
    if email_signature:
        full_body_txt += f"\n\n{email_signature}"

    if note_ids:
        note_rows = conn.execute(
            f"SELECT note_text, created_at FROM job_field_notes WHERE id IN ({','.join('?'*len(note_ids))}) AND job_id=?",
            note_ids + [int(job_id)]
        ).fetchall()
        if note_rows:
            full_body_txt += "\n\n--- Attached Notes ---\n"
            for nr in note_rows:
                note_body = nr['note_text']
                if note_body and note_body.startswith("[AI Update]\n"):
                    note_body = note_body[len("[AI Update]\n"):]
                full_body_txt += f"\n[{nr['created_at'][:16]}] {note_body}"

    body_html = f"""<div style="font-family:sans-serif;max-width:640px">
<p><strong>Job:</strong> {job['display_ref']}</p>
<p>{full_body_txt.replace(chr(10), '<br>')}</p>
<hr style="border:none;border-top:1px solid #e5e7eb">
<p style="color:#9ca3af;font-size:12px">Axion Field Operations Management</p>
</div>"""

    file_attachments = []
    if doc_ids:
        doc_rows = conn.execute(
            f"SELECT original_filename, stored_filename, mime_type FROM job_documents WHERE id IN ({','.join('?'*len(doc_ids))}) AND job_id=?",
            doc_ids + [int(job_id)]
        ).fetchall()
        for dr in doc_rows:
            local_path = os.path.join(UPLOAD_FOLDER, dr["stored_filename"])
            if os.path.exists(local_path):
                with open(local_path, "rb") as fh:
                    file_attachments.append((dr["original_filename"], fh.read(), dr["mime_type"] or "application/octet-stream"))

    conn.close()

    smtp_ok = bool(os.environ.get("SMTP_USER") and os.environ.get("SMTP_PASS"))
    smtp_skipped = False
    if smtp_ok:
        try:
            send_email(to_list, subject, full_body_txt, body_html,
                       cc_list=cc_list, attachments=file_attachments)
        except Exception as exc:
            return jsonify({"ok": False, "error": f"SMTP error: {exc}"})
    else:
        smtp_skipped = True

    mel_now = datetime.now(_melbourne)
    ts_str  = mel_now.strftime("%d/%m/%Y %H:%M")
    to_str  = ", ".join(to_list + (["CC: " + e for e in cc_list] if cc_list else []))
    doc_count  = len(file_attachments)
    note_count = len(note_ids)
    extras = []
    if note_count: extras.append(f"{note_count} note(s) appended")
    if doc_count:  extras.append(f"{doc_count} file(s) attached")
    extras_str = (" — " + "; ".join(extras)) if extras else ""

    if smtp_skipped:
        note_txt = f"Email queued (SMTP not configured) to {to_str}{extras_str} — {ts_str} — {session.get('user_name', 'Admin')}"
    else:
        note_txt = f"Email sent to {to_str}{extras_str} — {ts_str} — {session.get('user_name', 'Admin')}"

    conn2 = db()
    conn2.execute(
        "INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at) VALUES (?,?,?,?)",
        (job_id, session.get("user_id"), note_txt, now_ts())
    )
    conn2.commit()
    conn2.close()

    audit("job", int(job_id), "email_sent", note_txt)
    return jsonify({"ok": True, "sent_to": ", ".join(to_list), "smtp_skipped": smtp_skipped})


@app.post("/queue/email-agent-queue")
@admin_required
def queue_email_agent_queue():
    agent_id = request.form.get("agent_id", "").strip()
    filter_client = request.form.get("filter_client", "").strip()
    if not agent_id:
        return jsonify({"ok": False, "error": "Please select an agent."})

    conn = db()
    agent = conn.execute("SELECT id, full_name, email FROM users WHERE id=?", (agent_id,)).fetchone()
    if not agent or not agent["email"]:
        conn.close()
        return jsonify({"ok": False, "error": "Agent not found or has no email address."})

    cur = conn.cursor()
    overdue_types = ("Urgent: Schedule Overdue", "Schedule Due Today")
    tomorrow_type = "Schedule Due Tomorrow"
    note_type = "Agent Note Review"

    all_items = []
    for label, sql_where, params in [
        ("OVERDUE", "ci.visit_type IN (?,?) AND ci.status IN ('Pending','In Progress')", overdue_types),
        ("CURRENTLY DUE", "ci.visit_type = ? AND ci.status IN ('Pending','In Progress')", (tomorrow_type,)),
        ("AGENT NOTES", "ci.visit_type = ? AND ci.status = 'Pending'", (note_type,)),
    ]:
        cur.execute(_queue_row_sql() + " WHERE " + sql_where + " ORDER BY ci.priority DESC, ci.created_at DESC", params)
        rows = cur.fetchall()
        for r in rows:
            if str(r["job_assigned_uid"] or "") != str(agent_id):
                continue
            if filter_client and str(r["client_name"] or "") != filter_client:
                continue
            all_items.append((label, r))

    conn.close()

    if not all_items:
        return jsonify({"ok": False, "error": f"No queue items found for {agent['full_name']}."})

    mel_now = datetime.now(_melbourne)
    date_str = mel_now.strftime("%A %d %B %Y")

    from markupsafe import escape as _h
    rows_html = ""
    current_section = ""
    for section, item in all_items:
        if section != current_section:
            current_section = section
            rows_html += f'<tr><td colspan="6" style="background:#f3f4f6;font-weight:700;font-size:13px;padding:8px 10px;border-top:2px solid #d1d5db">{_h(section)}</td></tr>'
        ref = _h(item["display_ref"] or item["internal_job_number"] or "")
        client = _h(item["client_name"] or "—")
        borrower = _h(item["customer_label"] or item["customer_name"] or "—")
        address = _h(item["job_address"] or "—")
        status = _h(item["job_status"] or "—")
        action = _h(item["instructions"] or item["visit_type"] or "—")
        rows_html += f'''<tr>
          <td style="padding:6px 10px;border-bottom:1px solid #e5e7eb;font-weight:600">{ref}</td>
          <td style="padding:6px 10px;border-bottom:1px solid #e5e7eb">{client}</td>
          <td style="padding:6px 10px;border-bottom:1px solid #e5e7eb">{borrower}</td>
          <td style="padding:6px 10px;border-bottom:1px solid #e5e7eb;font-size:12px">{address}</td>
          <td style="padding:6px 10px;border-bottom:1px solid #e5e7eb">{status}</td>
          <td style="padding:6px 10px;border-bottom:1px solid #e5e7eb;font-size:12px">{action}</td>
        </tr>'''

    body_html = f"""<div style="font-family:sans-serif;max-width:800px">
<h2 style="margin:0 0 4px;font-size:18px">Your Queue — {_h(date_str)}</h2>
<p style="color:#6b7280;font-size:13px;margin:0 0 16px">Hi {_h(agent['full_name'])}, here is your current queue ({len(all_items)} item{'s' if len(all_items) != 1 else ''}):</p>
<table style="width:100%;border-collapse:collapse;font-size:13px;font-family:sans-serif">
<thead><tr style="background:#1e3a5f;color:#fff">
  <th style="padding:8px 10px;text-align:left">Job</th>
  <th style="padding:8px 10px;text-align:left">Client</th>
  <th style="padding:8px 10px;text-align:left">Borrower</th>
  <th style="padding:8px 10px;text-align:left">Address</th>
  <th style="padding:8px 10px;text-align:left">Status</th>
  <th style="padding:8px 10px;text-align:left">Action</th>
</tr></thead>
<tbody>{rows_html}</tbody>
</table>
<hr style="border:none;border-top:1px solid #e5e7eb;margin:20px 0 10px">
<p style="color:#9ca3af;font-size:12px">Axion Field Operations Management</p>
</div>"""

    body_txt = f"Your Queue — {date_str}\n\nHi {agent['full_name']}, here is your current queue ({len(all_items)} items):\n\n"
    for section, item in all_items:
        ref = item["display_ref"] or ""
        body_txt += f"[{section}] {ref} | {item['client_name'] or '—'} | {item['customer_label'] or item['customer_name'] or '—'} | {item['job_address'] or '—'} | {item['instructions'] or item['visit_type'] or '—'}\n"

    subject = f"Your AxionX Queue — {date_str} ({len(all_items)} items)"
    to_list = [agent["email"]]

    smtp_ok = bool(os.environ.get("SMTP_USER") and os.environ.get("SMTP_PASS"))
    smtp_skipped = False
    if smtp_ok:
        try:
            send_email(to_list, subject, body_txt, body_html)
        except Exception as exc:
            return jsonify({"ok": False, "error": f"SMTP error: {exc}"})
    else:
        smtp_skipped = True

    audit("queue", 0, "queue_emailed", f"Queue emailed to {agent['full_name']} ({agent['email']}) — {len(all_items)} items")
    return jsonify({"ok": True, "sent_to": agent["email"], "agent_name": agent["full_name"],
                    "item_count": len(all_items), "smtp_skipped": smtp_skipped})


# -------- Assignment board --------
@app.post("/cue/<int:cue_id>/assign")
@admin_required
def cue_assign(cue_id: int):
    assigned_user_id = request.form.get("assigned_user_id") or None
    ts = now_ts()

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT assigned_user_id, job_id FROM cue_items WHERE id = ?", (cue_id,))
    before = cur.fetchone()
    if not before:
        conn.close()
        return ("Not found", 404)

    cur.execute("UPDATE cue_items SET assigned_user_id = ?, updated_at = ? WHERE id = ?",
                (assigned_user_id, ts, cue_id))
    conn.commit()
    conn.close()

    audit("cue", cue_id, "assign",
          f"Cue {cue_id} assigned to user {assigned_user_id or 'Unassigned'}.",
          {"from": before["assigned_user_id"], "to": assigned_user_id, "job_id": before["job_id"]})

    return ("OK", 200)


# -------- Monthly report --------
@app.get("/reports/monthly")
@admin_required
def report_monthly():
    prefix = request.args.get("prefix") or datetime.now().strftime("%y%m")

    conn = db()
    cur = conn.cursor()

    cur.execute("""
        SELECT COUNT(*) c FROM jobs
        WHERE strftime('%y%m', created_at) = ?
    """, (prefix,))
    total_jobs = cur.fetchone()["c"]

    cur.execute("""
        SELECT status, COUNT(*) c FROM jobs
        WHERE strftime('%y%m', created_at) = ?
        GROUP BY status ORDER BY c DESC
    """, (prefix,))
    by_status = cur.fetchall()

    cur.execute("""
        SELECT u.full_name, COUNT(*) c
        FROM cue_items ci
        JOIN users u ON u.id = ci.assigned_user_id
        JOIN jobs j ON j.id = ci.job_id
        WHERE strftime('%y%m', j.created_at) = ? AND ci.status = 'Completed'
        GROUP BY u.full_name
        ORDER BY c DESC
    """, (prefix,))
    completed_by_agent = cur.fetchall()

    conn.close()
    return render_template("report_monthly.html",
                           prefix=prefix,
                           total_jobs=total_jobs,
                           by_status=by_status,
                           completed_by_agent=completed_by_agent)


# -------- Agent: My Today --------
@app.get("/my/today")
@login_required
def my_today():
    lockout = _check_agent_lockout()
    if lockout:
        return lockout
    today = datetime.now().date().isoformat()
    today_display = today[8:10] + "/" + today[5:7] + "/" + today[:4]
    user_id = session.get("user_id")

    conn = db()
    cur = conn.cursor()
    cur.execute(f"""
        SELECT ci.*, j.internal_job_number, j.client_reference, j.job_address,
               (cu.first_name || ' ' || cu.last_name) customer_name,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
               (SELECT ji.reg FROM job_items ji WHERE ji.job_id = j.id AND ji.item_type = 'vehicle' LIMIT 1) asset_reg
        FROM cue_items ci
        JOIN jobs j ON j.id = ci.job_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        WHERE ci.due_date = ? AND ci.assigned_user_id = ?
          AND ci.status IN ('Pending','In Progress')
          AND j.status NOT IN {ARCHIVED_STATUSES!r}
        ORDER BY ci.priority DESC, ci.id
    """, (today, user_id))
    cues = cur.fetchall()

    cur.execute(f"""
        SELECT s.id, s.job_id, s.scheduled_for, s.status, s.notes,
               bt.name AS booking_type_name,
               j.internal_job_number, j.client_reference, j.display_ref, j.job_address,
               (cu.first_name || ' ' || cu.last_name) AS customer_name,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
               (SELECT ji.reg FROM job_items ji WHERE ji.job_id = j.id AND ji.item_type='vehicle' LIMIT 1) AS asset_reg
        FROM schedules s
        JOIN booking_types bt ON bt.id = s.booking_type_id
        JOIN jobs j ON j.id = s.job_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        WHERE date(s.scheduled_for) = ? AND s.assigned_to_user_id = ?
          AND s.status NOT IN ('Cancelled', 'Completed')
          AND j.status NOT IN {ARCHIVED_STATUSES!r}
        ORDER BY s.scheduled_for
    """, (today, user_id))
    schedules = cur.fetchall()

    cur.execute("""
        SELECT ju.id AS draft_id, ju.job_id, ju.created_at,
               j.internal_job_number, j.client_reference, j.job_address,
               (cu.first_name || ' ' || cu.last_name) AS customer_name,
               COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label
        FROM job_updates ju
        JOIN jobs j ON j.id = ju.job_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        WHERE ju.created_by_user_id = ? AND ju.status = 'draft'
        ORDER BY ju.updated_at DESC
    """, (user_id,))
    update_drafts = cur.fetchall()
    conn.close()

    return render_template("my_today.html", cues=cues, schedules=schedules,
                           today=today, today_display=today_display,
                           update_drafts=update_drafts)


@app.get("/my/settings")
def my_settings():
    if not session.get("user_id"):
        return _login_redirect()
    return render_template("my_settings.html")


@app.post("/my/settings/password")
def my_settings_password():
    user_id = session.get("user_id")
    if not user_id:
        return _login_redirect()
    current = request.form.get("current_password", "").strip()
    new_pw = request.form.get("new_password", "").strip()
    confirm = request.form.get("confirm_password", "").strip()

    if not current or not new_pw or not confirm:
        flash("All fields are required.", "danger")
        return redirect(url_for("my_settings"))
    if new_pw != confirm:
        flash("New passwords do not match.", "danger")
        return redirect(url_for("my_settings"))
    if len(new_pw) < 6:
        flash("Password must be at least 6 characters.", "danger")
        return redirect(url_for("my_settings"))

    conn = db()
    user = conn.execute("SELECT * FROM users WHERE id = ?", (user_id,)).fetchone()
    if not user or not check_password_hash(user["password"], current):
        conn.close()
        flash("Current password is incorrect.", "danger")
        return redirect(url_for("my_settings"))

    conn.execute("UPDATE users SET password = ? WHERE id = ?",
                 (generate_password_hash(new_pw), user_id))
    conn.commit()
    conn.close()
    flash("Password updated successfully.", "success")
    return redirect(url_for("my_settings"))


@app.post("/jobs/<int:job_id>/note-update-emailed")
def note_update_emailed(job_id: int):
    user_id = session.get("user_id")
    user_name = session.get("user_name", "Unknown")
    role = session.get("role")
    if not user_id:
        return _login_redirect()

    conn = db()
    job = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        conn.close()
        flash("Job not found.", "danger")
        return redirect(url_for("jobs_list"))

    if role == "agent" and job["assigned_user_id"] != user_id:
        sched_check = conn.execute(
            """SELECT 1 FROM schedules WHERE job_id = ? AND assigned_to_user_id = ?
               AND status NOT IN ('Cancelled') LIMIT 1""",
            (job_id, user_id)
        ).fetchone()
        if not sched_check:
            conn.close()
            flash("You do not have access to that job.", "danger")
            return redirect(url_for("jobs_list"))

    now_melb = datetime.now(_melbourne)
    ts = now_melb.strftime("%d/%m/%Y %H:%M")
    note_text = f"Update emailed — {ts} — {user_name}"
    conn.execute(
        """INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at)
           VALUES (?, ?, ?, ?)""",
        (job_id, user_id, note_text, now_ts())
    )
    conn.commit()
    conn.close()

    flash("Update emailed note added.", "success")
    ref = request.referrer or url_for("jobs_list")
    return redirect(ref)


@app.get("/resources")
@login_required
def resources():
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM tow_operators WHERE active=1 ORDER BY company_name")
    tow_operators = cur.fetchall()
    cur.execute("SELECT * FROM auction_yards WHERE active=1 ORDER BY name")
    auction_yards = cur.fetchall()
    conn.close()
    resp = make_response(render_template("resources.html", tow_operators=tow_operators, auction_yards=auction_yards))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    return resp


@app.post("/resources/tow-operators/add")
@login_required
def tow_operator_add():
    company_name = request.form.get("company_name", "").strip()
    phone = request.form.get("phone", "").strip() or None
    address = request.form.get("address", "").strip() or None
    if not company_name:
        return jsonify({"ok": False, "error": "Company name is required."})
    conn = db()
    cur = conn.cursor()
    cur.execute("INSERT INTO tow_operators (company_name, phone, address, created_at) VALUES (?,?,?,?)",
                (company_name, phone, address, now_ts()))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "id": new_id, "company_name": company_name, "phone": phone or "", "address": address or ""})


@app.post("/resources/tow-operators/<int:op_id>/edit")
@login_required
def tow_operator_edit(op_id):
    company_name = request.form.get("company_name", "").strip()
    phone = request.form.get("phone", "").strip() or None
    address = request.form.get("address", "").strip() or None
    if not company_name:
        return jsonify({"ok": False, "error": "Company name is required."})
    conn = db()
    conn.execute("UPDATE tow_operators SET company_name=?, phone=?, address=? WHERE id=?",
                 (company_name, phone, address, op_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/resources/tow-operators/<int:op_id>/delete")
@login_required
def tow_operator_delete(op_id):
    conn = db()
    conn.execute("UPDATE tow_operators SET active=0 WHERE id=?", (op_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/resources/auction-yards/add")
@login_required
def auction_yard_add():
    name = request.form.get("name", "").strip()
    address = request.form.get("address", "").strip() or None
    if not name:
        return jsonify({"ok": False, "error": "Auction yard name is required."})
    conn = db()
    cur = conn.cursor()
    cur.execute("INSERT INTO auction_yards (name, address, created_at) VALUES (?,?,?)",
                (name, address, now_ts()))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "id": new_id, "name": name, "address": address or ""})


@app.post("/resources/auction-yards/<int:yard_id>/edit")
@login_required
def auction_yard_edit(yard_id):
    name = request.form.get("name", "").strip()
    address = request.form.get("address", "").strip() or None
    if not name:
        return jsonify({"ok": False, "error": "Name is required."})
    conn = db()
    conn.execute("UPDATE auction_yards SET name=?, address=? WHERE id=?", (name, address, yard_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/resources/auction-yards/<int:yard_id>/delete")
@login_required
def auction_yard_delete(yard_id):
    conn = db()
    conn.execute("UPDATE auction_yards SET active=0 WHERE id=?", (yard_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/form-templates/add")
@login_required
@admin_required
def form_template_add():
    name = request.form.get("name", "").strip()
    field_list = request.form.get("field_list", "").strip()
    if not name or not field_list:
        return jsonify({"ok": False, "error": "Name and at least one field are required."})
    caller_id = session.get("user_id")
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("INSERT INTO form_templates (name, field_list, created_by, created_at) VALUES (?,?,?,?)",
                    (name, field_list, caller_id, now_ts()))
        new_id = cur.lastrowid
        conn.commit()
    except Exception:
        conn.close()
        return jsonify({"ok": False, "error": "A template with that name already exists."})
    conn.close()
    return jsonify({"ok": True, "id": new_id, "name": name, "field_list": field_list})


@app.post("/form-templates/<int:tmpl_id>/delete")
@login_required
@admin_required
def form_template_delete(tmpl_id):
    conn = db()
    conn.execute("UPDATE form_templates SET active=0 WHERE id=?", (tmpl_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ──────────────────── Contacts hub ─────────────────────────────────

@app.get("/contacts")
@admin_required
def contacts_hub():
    conn = db()
    client_count   = conn.execute("SELECT COUNT(*) FROM clients").fetchone()[0]
    customer_count = conn.execute("SELECT COUNT(*) FROM customers").fetchone()[0]
    staff_count    = conn.execute("SELECT COUNT(*) FROM users WHERE active=1").fetchone()[0]
    tow_count      = conn.execute("SELECT COUNT(*) FROM tow_operators WHERE active=1").fetchone()[0]
    yard_count     = conn.execute("SELECT COUNT(*) FROM auction_yards WHERE active=1").fetchone()[0]
    conn.close()
    return render_template("contacts.html",
        client_count=client_count,
        customer_count=customer_count,
        staff_count=staff_count,
        tow_count=tow_count,
        yard_count=yard_count)


# ──────────────────── Geomap ────────────────────────────────────────

@app.get("/map")
@login_required
def geomap_page():
    is_admin = session.get("role") in ("admin", "both")
    agents = []
    if is_admin:
        conn = db()
        agents = conn.execute(
            "SELECT id, full_name FROM users WHERE role IN ('agent', 'both') AND active=1 ORDER BY full_name"
        ).fetchall()
        conn.close()
    return render_template("map.html", agents=agents, is_admin=is_admin)


@app.get("/api/map/data")
@login_required
def api_map_data():
    is_admin = session.get("role") in ("admin", "both")
    uid = session.get("user_id")
    conn = db()

    if is_admin:
        jobs = conn.execute("""
            SELECT j.id, j.display_ref, j.job_address, j.status, j.lat, j.lng,
                   (cu.first_name || ' ' || cu.last_name) AS customer_name,
                   COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
                   c.name AS client_name,
                   ag.full_name AS agent_name
            FROM jobs j
            LEFT JOIN customers cu ON cu.id = j.customer_id
            LEFT JOIN clients   c  ON c.id  = j.client_id
            LEFT JOIN users     ag ON ag.id = j.assigned_user_id
            WHERE j.status NOT IN ('Closed','Cancelled','Completed')
              AND j.job_address IS NOT NULL AND j.job_address != ''
            ORDER BY j.updated_at DESC
        """).fetchall()
    else:
        jobs = conn.execute("""
            SELECT j.id, j.display_ref, j.job_address, j.status, j.lat, j.lng,
                   (cu.first_name || ' ' || cu.last_name) AS customer_name,
                   COALESCE(NULLIF(TRIM(COALESCE(cu.company,'')), ''), cu.last_name) AS customer_label,
                   c.name AS client_name,
                   ag.full_name AS agent_name
            FROM jobs j
            LEFT JOIN customers cu ON cu.id = j.customer_id
            LEFT JOIN clients   c  ON c.id  = j.client_id
            LEFT JOIN users     ag ON ag.id = j.assigned_user_id
            WHERE j.status NOT IN ('Closed','Cancelled','Completed')
              AND j.job_address IS NOT NULL AND j.job_address != ''
              AND j.assigned_user_id = ?
            ORDER BY j.updated_at DESC
        """, (uid,)).fetchall()

    two_hours_ago = (datetime.now(_melbourne) - _td(hours=2)).isoformat()

    agents = []
    if is_admin:
        agents = conn.execute("""
            SELECT u.id, u.full_name, al.lat, al.lng, al.accuracy, al.updated_at
            FROM users u
            JOIN agent_locations al ON al.user_id = u.id
            WHERE u.role IN ('agent', 'both') AND u.active = 1
              AND al.updated_at >= ?
            ORDER BY u.full_name
        """, (two_hours_ago,)).fetchall()
    conn.close()

    def initials(name):
        parts = (name or "?").split()
        return (parts[0][0] + (parts[-1][0] if len(parts) > 1 else "")).upper()

    return jsonify({
        "jobs": [{
            "id": r["id"],
            "ref": r["display_ref"],
            "address": r["job_address"],
            "status": r["status"],
            "lat": r["lat"],
            "lng": r["lng"],
            "customer": r["customer_label"] or r["customer_name"] or "",
            "client": r["client_name"] or "",
            "agent": r["agent_name"] or ""
        } for r in jobs],
        "agents": [{
            "id": r["id"],
            "name": r["full_name"],
            "initials": initials(r["full_name"]),
            "lat": r["lat"],
            "lng": r["lng"],
            "accuracy": r["accuracy"],
            "updated_at": r["updated_at"]
        } for r in agents]
    })


@app.post("/api/agent/location")
@login_required
def api_agent_location():
    data = request.get_json(silent=True) or {}
    lat  = data.get("lat")
    lng  = data.get("lng")
    acc  = data.get("accuracy")
    if lat is None or lng is None:
        return jsonify({"ok": False, "error": "lat and lng required"}), 400
    uid = session.get("user_id")
    ts  = now_ts()
    conn = db()
    conn.execute("""
        INSERT INTO agent_locations (user_id, lat, lng, accuracy, updated_at)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(user_id) DO UPDATE SET
            lat=excluded.lat, lng=excluded.lng,
            accuracy=excluded.accuracy, updated_at=excluded.updated_at
    """, (uid, float(lat), float(lng), float(acc) if acc is not None else None, ts))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


def _geocode_address(address: str):
    """Geocode an address via Nominatim. Returns (lat, lng) or None.
    Rate-limit compliant: caller must sleep ≥1 s between calls."""
    import urllib.request
    import urllib.parse
    try:
        params = urllib.parse.urlencode({
            "q": address,
            "format": "json",
            "limit": 1,
            "countrycodes": "au",
        })
        url = f"https://nominatim.openstreetmap.org/search?{params}"
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "AxionX/1.0 field-ops (contact@swpirecoveries.com.au)"}
        )
        with urllib.request.urlopen(req, timeout=8) as resp:
            import json as _json
            data = _json.loads(resp.read())
        if data:
            return float(data[0]["lat"]), float(data[0]["lon"])
    except Exception:
        pass
    return None


def _geocode_job_async(job_id: int, address: str):
    """Fire-and-forget: geocode one job address and persist the result."""
    def _worker():
        result = _geocode_address(address)
        if result:
            lat, lng = result
            try:
                conn = db()
                conn.execute("UPDATE jobs SET lat=?, lng=? WHERE id=?", (lat, lng, job_id))
                conn.commit()
                conn.close()
            except Exception:
                pass
    threading.Thread(target=_worker, daemon=True).start()


@app.post("/api/jobs/<int:job_id>/geocode")
@login_required
def api_job_geocode(job_id: int):
    data = request.get_json(silent=True) or {}
    lat  = data.get("lat")
    lng  = data.get("lng")
    if lat is None or lng is None:
        return jsonify({"ok": False}), 400
    conn = db()
    conn.execute("UPDATE jobs SET lat=?, lng=? WHERE id=?",
                 (float(lat), float(lng), job_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/m/api/jobs/geocode-pending")
def m_geocode_pending():
    """Batch-geocode up to 5 jobs that have an address but no lat/lng.
    Returns {ok, updated:[{id,lat,lng}], remaining} so the client can
    keep calling until remaining==0."""
    if not session.get("user_id"):
        return jsonify({"ok": False}), 401
    import time as _time
    conn = db()
    pending = conn.execute(
        "SELECT id, job_address FROM jobs"
        " WHERE job_address IS NOT NULL AND job_address != ''"
        "   AND (lat IS NULL OR lng IS NULL)"
        " LIMIT 5"
    ).fetchall()
    conn.close()

    updated = []
    for job in pending:
        result = _geocode_address(job["job_address"])
        if result:
            lat, lng = result
            try:
                c2 = db()
                c2.execute("UPDATE jobs SET lat=?, lng=? WHERE id=?",
                           (lat, lng, job["id"]))
                c2.commit()
                c2.close()
                updated.append({"id": job["id"], "lat": lat, "lng": lng})
            except Exception:
                pass
        _time.sleep(1.05)  # Nominatim hard rate limit: 1 req/s

    remaining_conn = db()
    remaining = remaining_conn.execute(
        "SELECT COUNT(*) FROM jobs"
        " WHERE job_address IS NOT NULL AND job_address != ''"
        "   AND (lat IS NULL OR lng IS NULL)"
    ).fetchone()[0]
    remaining_conn.close()
    return jsonify({"ok": True, "updated": updated, "remaining": remaining})


# ─────────────────────────────────────────────────────────────────────────────
# MOBILE ROUTES  (/m)
# ─────────────────────────────────────────────────────────────────────────────

def mobile_login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            next_path = request.path
            return redirect(url_for("m_login") + f"?next={next_path}")
        return f(*args, **kwargs)
    return decorated


def agent_mobile_guard(f):
    """Blocks agents from mutating routes (job edits, schedule edits, status changes).
    Admins and 'both' roles pass through freely."""
    @wraps(f)
    def decorated(*args, **kwargs):
        role = session.get("role", "")
        if role not in ("admin", "both"):
            flash("Access restricted to administrators.", "warning")
            return redirect(url_for("m_today"))
        return f(*args, **kwargs)
    return decorated


@app.get("/m")
def m_root():
    return redirect(url_for("m_login"))


@app.get("/m/login")
def m_login():
    if session.get("user_id"):
        return redirect(url_for("m_today"))
    next_path = request.args.get("next", "")
    return render_template("mobile/login.html", error=None, prefill_email="", next=next_path)


@app.post("/m/login")
def m_login_post():
    email    = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "")
    next_path = request.args.get("next", "").strip()
    ip       = request.headers.get("X-Forwarded-For", request.remote_addr or "").split(",")[0].strip()
    ip_key   = f"ip:{ip}"

    conn = db()
    allowed, locked_until = throttle_check(conn, ip_key)
    if not allowed:
        conn.close()
        return render_template("mobile/login.html",
                               error=f"Too many failed attempts. Try again after {locked_until} UTC.",
                               prefill_email=email, next=next_path)

    user = conn.execute("SELECT * FROM users WHERE LOWER(email)=? AND active=1", (email,)).fetchone()

    if not user or not check_password_hash(user["password"], password):
        throttle_fail(conn, ip_key)
        conn.commit()
        conn.close()
        return render_template("mobile/login.html", error="Invalid email or password.",
                               prefill_email=email, next=next_path)

    throttle_success(conn, ip_key)
    conn.commit()
    conn.close()
    session.permanent = True
    session["user_id"]   = user["id"]
    session["user_name"] = user["full_name"]
    session["role"]      = user["role"]
    if next_path and next_path.startswith("/m/"):
        return redirect(next_path, code=302)
    return redirect(url_for("m_today"), code=302)


@app.get("/m/logout")
def m_logout():
    token = request.args.get("token", "").strip()
    if token:
        conn = db()
        conn.execute(
            "UPDATE mobile_auth_tokens SET revoked_at=? WHERE token=? AND revoked_at IS NULL",
            (now_ts(), token))
        conn.commit()
        conn.close()
    session.clear()
    return redirect(url_for("m_login"))


@app.post("/m/api/auth/create-token")
@mobile_login_required
def m_api_auth_create_token():
    uid = session.get("user_id")
    device_name = request.get_json(silent=True) or {}
    device_name = device_name.get("device_name", "iOS Device")
    token = secrets.token_urlsafe(48)
    ts = now_ts()
    expires = (datetime.now(_melbourne) + _td(days=90)).strftime("%Y-%m-%d %H:%M:%S")
    conn = db()
    conn.execute(
        "INSERT INTO mobile_auth_tokens (token, user_id, device_name, created_at, expires_at) VALUES (?,?,?,?,?)",
        (token, uid, device_name, ts, expires))
    conn.commit()
    conn.close()
    return jsonify({"token": token, "expires_at": expires})


@app.post("/m/api/auth/token-login")
def m_api_auth_token_login():
    data = request.get_json(silent=True) or {}
    token = data.get("token", "").strip()
    if not token:
        return jsonify({"success": False, "error": "Token required"}), 400

    conn = db()
    row = conn.execute(
        "SELECT t.*, u.full_name, u.role, u.active FROM mobile_auth_tokens t "
        "JOIN users u ON u.id = t.user_id "
        "WHERE t.token=? AND t.revoked_at IS NULL", (token,)).fetchone()

    if not row:
        conn.close()
        return jsonify({"success": False, "error": "Invalid or revoked token"}), 401

    now = datetime.now(_melbourne).strftime("%Y-%m-%d %H:%M:%S")
    if row["expires_at"] < now:
        conn.execute("UPDATE mobile_auth_tokens SET revoked_at=? WHERE token=?", (now_ts(), token))
        conn.commit()
        conn.close()
        return jsonify({"success": False, "error": "Token expired"}), 401

    if not row["active"]:
        conn.close()
        return jsonify({"success": False, "error": "Account disabled"}), 401

    conn.close()
    session.permanent = True
    session["user_id"] = row["user_id"]
    session["user_name"] = row["full_name"]
    session["role"] = row["role"]
    return jsonify({"success": True, "user_name": row["full_name"], "role": row["role"]})


@app.post("/m/api/auth/revoke-token")
def m_api_auth_revoke_token():
    data = request.get_json(silent=True) or {}
    token = data.get("token", "").strip()
    if not token:
        return jsonify({"success": False}), 400
    conn = db()
    conn.execute(
        "UPDATE mobile_auth_tokens SET revoked_at=? WHERE token=? AND revoked_at IS NULL",
        (now_ts(), token))
    conn.commit()
    conn.close()
    return jsonify({"success": True})


@app.get("/m/api/auth/biometric-status")
@mobile_login_required
def m_api_auth_biometric_status():
    uid = session.get("user_id")
    now = datetime.now(_melbourne).strftime("%Y-%m-%d %H:%M:%S")
    conn = db()
    row = conn.execute(
        "SELECT created_at, expires_at FROM mobile_auth_tokens "
        "WHERE user_id=? AND revoked_at IS NULL AND expires_at>? ORDER BY created_at DESC LIMIT 1",
        (uid, now)).fetchone()
    conn.close()
    if row:
        return jsonify({
            "has_active_token": True,
            "token_created_at": row["created_at"],
            "token_expires_at": row["expires_at"]
        })
    return jsonify({"has_active_token": False})


@app.post("/m/api/auth/revoke-all-tokens")
@mobile_login_required
def m_api_auth_revoke_all_tokens():
    uid = session.get("user_id")
    conn = db()
    conn.execute(
        "UPDATE mobile_auth_tokens SET revoked_at=? WHERE user_id=? AND revoked_at IS NULL",
        (now_ts(), uid))
    conn.commit()
    conn.close()
    return jsonify({"success": True})


@app.get("/m/schedule/today")
@mobile_login_required
def m_today():
    lockout = _check_agent_lockout()
    if lockout:
        return lockout
    uid   = session.get("user_id")
    today = datetime.now(_melbourne).date().isoformat()
    today_display = today[8:10] + "/" + today[5:7] + "/" + today[:4]

    conn = db()

    cues = conn.execute(f"""
        SELECT ci.*, j.internal_job_number, j.client_reference, j.display_ref,
               j.job_address, j.id AS jid,
               (cu.first_name || ' ' || cu.last_name) AS customer_name,
               (SELECT ji.reg FROM job_items ji WHERE ji.job_id=j.id AND ji.item_type='vehicle' LIMIT 1) AS asset_reg
        FROM cue_items ci
        JOIN jobs j ON j.id = ci.job_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        WHERE ci.due_date = ? AND ci.assigned_user_id = ?
          AND ci.status IN ('Pending','In Progress')
          AND j.status NOT IN {ARCHIVED_STATUSES!r}
        ORDER BY ci.priority DESC, ci.id
    """, (today, uid)).fetchall()

    schedules = conn.execute(f"""
        SELECT s.id, s.job_id, s.scheduled_for, s.status, s.notes,
               bt.name AS booking_type_name,
               j.internal_job_number, j.client_reference, j.display_ref, j.job_address,
               (cu.first_name || ' ' || cu.last_name) AS customer_name,
               (SELECT ji.reg FROM job_items ji WHERE ji.job_id=j.id AND ji.item_type='vehicle' LIMIT 1) AS asset_reg
        FROM schedules s
        JOIN booking_types bt ON bt.id = s.booking_type_id
        JOIN jobs j ON j.id = s.job_id
        LEFT JOIN customers cu ON cu.id = j.customer_id
        WHERE date(s.scheduled_for,'localtime') = ? AND s.assigned_to_user_id = ?
          AND s.status NOT IN ('Cancelled','Completed')
          AND j.status NOT IN {ARCHIVED_STATUSES!r}
        ORDER BY s.scheduled_for
    """, (today, uid)).fetchall()

    drafts_raw = conn.execute("""
        SELECT ju.id AS draft_id, ju.job_id, ju.created_at,
               j.display_ref, j.internal_job_number, j.job_address
        FROM job_updates ju
        JOIN jobs j ON j.id = ju.job_id
        WHERE ju.created_by_user_id = ? AND ju.status = 'draft'
        ORDER BY ju.updated_at DESC
    """, (uid,)).fetchall()
    conn.close()

    draft_job_ids = {d["job_id"] for d in drafts_raw}

    return render_template("m/today.html",
                           today=today, today_display=today_display,
                           cues=cues, schedules=schedules,
                           drafts=drafts_raw, draft_job_ids=draft_job_ids)


def _mobile_jobs_query(uid, role, params_in):
    """Shared jobs-list query engine for mobile. Returns (jobs, draft_job_ids, prefs_used).
    params_in: dict with optional keys sort, dir, scope, status_filter, show_completed, q
    This function is intentionally decoupled so it can be reused by future desktop dispatcher views.
    """
    is_admin = role in ("admin", "both")
    conn = db()

    # Load saved prefs as baseline
    prefs = conn.execute(
        "SELECT * FROM user_mobile_settings WHERE user_id=?", (uid,)
    ).fetchone()

    def pref(key, default):
        v = params_in.get(key)
        if v is not None:
            return v
        return (prefs[key] if prefs and key in prefs.keys() else None) or default

    def pref_col(param_key, col_key, default):
        v = params_in.get(param_key)
        if v is not None:
            return v
        try:
            cv = prefs[col_key] if prefs else None
        except (IndexError, KeyError):
            cv = None
        return cv or default

    sort           = pref_col("sort",       "list_sort",  "distance")
    direction      = pref_col("dir",        "list_dir",   "asc")
    scope          = pref_col("scope",     "job_scope",      "all" if is_admin else "mine")
    status_filter  = pref("status_filter", "")
    show_completed = pref_col("show_completed", "show_completed", "week")
    q              = params_in.get("q", "").strip()
    distance_unit  = (prefs["distance_unit"] if prefs else None) or "km"

    # Validate
    if sort not in ("visit_date", "status", "created", "distance"):
        sort = "visit_date"
    if direction not in ("asc", "desc"):
        direction = "asc"
    if show_completed not in ("day", "week", "month", "all", "none"):
        show_completed = "week"

    dir_sql = "ASC" if direction == "asc" else "DESC"

    # ── Ownership scope ──
    where_clauses = [f"j.status NOT IN {ARCHIVED_STATUSES!r}"]
    params = []

    has_search = bool(q)

    if is_admin and (scope != "mine" or has_search):
        pass
    else:
        where_clauses.append(
            "(j.assigned_user_id = ? OR EXISTS ("
            "  SELECT 1 FROM schedules s"
            "  WHERE s.job_id=j.id AND s.assigned_to_user_id=?"
            "  AND s.status NOT IN ('Cancelled','Completed')"
            "))"
        )
        params += [uid, uid]

    # ── Scheduled / Unscheduled scope (skip when searching) ──
    if not has_search:
        if scope == "scheduled":
            where_clauses.append(
                "EXISTS (SELECT 1 FROM schedules s WHERE s.job_id=j.id AND s.status NOT IN ('Cancelled','Completed'))"
            )
        elif scope == "unscheduled":
            where_clauses.append(
                "NOT EXISTS (SELECT 1 FROM schedules s WHERE s.job_id=j.id AND s.status NOT IN ('Cancelled','Completed'))"
            )

    # ── Status filter (skip completed restriction when searching) ──
    completed_statuses = ("Completed", "Invoiced")
    if has_search:
        pass
    elif status_filter and status_filter not in ("", "all"):
        where_clauses.append("j.status = ?")
        params.append(status_filter)
    else:
        if show_completed == "none":
            where_clauses.append(f"j.status NOT IN {completed_statuses!r}")
        elif show_completed == "day":
            where_clauses.append(
                f"(j.status NOT IN {completed_statuses!r} OR"
                " date(j.updated_at) >= date('now','localtime'))"
            )
        elif show_completed == "week":
            where_clauses.append(
                f"(j.status NOT IN {completed_statuses!r} OR"
                " date(j.updated_at) >= date('now','-6 days','localtime'))"
            )
        elif show_completed == "month":
            where_clauses.append(
                f"(j.status NOT IN {completed_statuses!r} OR"
                " date(j.updated_at) >= date('now','-29 days','localtime'))"
            )
        # show_completed == "all" → no extra filter

    # ── Search ──
    if q:
        where_clauses.append(
            "(j.display_ref LIKE ? OR j.internal_job_number LIKE ? OR"
            " j.client_reference LIKE ? OR j.job_address LIKE ? OR"
            " j.lender_name LIKE ? OR j.client_job_number LIKE ? OR"
            " j.account_number LIKE ? OR"
            " (cu.first_name || ' ' || cu.last_name) LIKE ? OR"
            " EXISTS (SELECT 1 FROM job_items ji WHERE ji.job_id=j.id"
            "   AND (ji.reg LIKE ? OR ji.vin LIKE ?)))"
        )
        like = f"%{q}%"
        params += [like]*10

    where_sql = " AND ".join(where_clauses) if where_clauses else "1=1"

    # ── Sort ──
    if sort == "status":
        order_sql = f"j.status {dir_sql}, next_scheduled ASC NULLS LAST"
    elif sort == "created":
        order_sql = f"j.created_at {dir_sql}"
    elif sort == "distance":
        # Client-side distance sort; server sorts by due date as secondary
        order_sql = f"next_scheduled ASC NULLS LAST, j.job_due_date ASC NULLS LAST"
    else:  # visit_date
        order_sql = f"next_scheduled {dir_sql} NULLS LAST, j.job_due_date {dir_sql} NULLS LAST, j.updated_at DESC"

    jobs = conn.execute(f"""
        SELECT j.*,
               (cu.first_name || ' ' || cu.last_name)  AS customer_name,
               (SELECT ji.reg  FROM job_items ji WHERE ji.job_id=j.id AND ji.item_type='vehicle' ORDER BY ji.id LIMIT 1) AS asset_reg,
               (SELECT ji.vin  FROM job_items ji WHERE ji.job_id=j.id AND ji.item_type='vehicle' ORDER BY ji.id LIMIT 1) AS asset_vin,
               (SELECT s.scheduled_for FROM schedules s WHERE s.job_id=j.id
                AND s.status NOT IN ('Cancelled','Completed') ORDER BY s.scheduled_for LIMIT 1) AS next_scheduled,
               (SELECT cpn.phone_number FROM contact_phone_numbers cpn
                WHERE cpn.entity_type='customer' AND cpn.entity_id=j.customer_id
                ORDER BY CASE WHEN cpn.label='Mobile' THEN 0 ELSE 1 END LIMIT 1) AS customer_phone,
               au.full_name AS assigned_agent_name
        FROM jobs j
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN users au ON au.id = j.assigned_user_id
        WHERE {where_sql}
        ORDER BY {order_sql}
        LIMIT 300
    """, params).fetchall()

    # Draft job IDs for this user
    draft_job_ids = {r["job_id"] for r in conn.execute(
        "SELECT DISTINCT job_id FROM job_updates WHERE created_by_user_id=? AND status='draft'",
        (uid,)
    ).fetchall()}

    conn.close()

    prefs_used = {
        "sort": sort, "dir": direction, "scope": scope,
        "status_filter": status_filter, "show_completed": show_completed,
        "distance_unit": distance_unit, "q": q,
    }
    return jobs, draft_job_ids, prefs_used


@app.get("/m/jobs")
@mobile_login_required
def m_jobs():
    lockout = _check_agent_lockout()
    if lockout:
        return lockout
    uid  = session.get("user_id")
    role = session.get("role", "")
    params_in = {
        k: request.args.get(k)
        for k in ("sort", "dir", "scope", "status_filter", "show_completed", "q")
        if request.args.get(k) is not None
    }
    jobs, draft_job_ids, prefs = _mobile_jobs_query(uid, role, params_in)
    # Prepare for client-side distance calc: pass lat/lng as JSON
    import json as _json
    jobs_geo = _json.dumps([
        {"id": j["id"], "lat": j["lat"], "lng": j["lng"]}
        for j in jobs if j["lat"] and j["lng"]
    ])
    is_admin = role in ("admin", "both")
    return render_template("mobile/jobs.html",
                           jobs=jobs, draft_job_ids=draft_job_ids, prefs=prefs,
                           jobs_geo=jobs_geo, is_admin=is_admin)


@app.get("/m/api/jobs/search")
@mobile_login_required
def m_api_jobs_search():
    uid  = session.get("user_id")
    role = session.get("role", "")
    q    = request.args.get("q", "").strip()
    if not q or len(q) < 2:
        return jsonify({"jobs": []})
    is_admin = role in ("admin", "both")
    conn = db()
    like = f"%{q}%"
    scope_sql = "1=1" if is_admin else (
        "(j.assigned_user_id = ? OR EXISTS ("
        "  SELECT 1 FROM schedules s"
        "  WHERE s.job_id=j.id AND s.assigned_to_user_id=?"
        "  AND s.status NOT IN ('Cancelled','Completed')"
        "))")
    params = [] if is_admin else [uid, uid]
    params += [like] * 10
    rows = conn.execute(f"""
        SELECT j.id, j.display_ref, j.internal_job_number, j.client_reference,
               j.account_number, j.status, j.job_address, j.lat, j.lng,
               j.lender_name,
               (cu.first_name || ' ' || cu.last_name) AS customer_name,
               (SELECT ji.reg FROM job_items ji WHERE ji.job_id=j.id AND ji.item_type='vehicle' ORDER BY ji.id LIMIT 1) AS asset_reg,
               (SELECT ji.vin FROM job_items ji WHERE ji.job_id=j.id AND ji.item_type='vehicle' ORDER BY ji.id LIMIT 1) AS asset_vin,
               (SELECT s.scheduled_for FROM schedules s WHERE s.job_id=j.id
                AND s.status NOT IN ('Cancelled','Completed') ORDER BY s.scheduled_for LIMIT 1) AS next_scheduled,
               (SELECT cpn.phone_number FROM contact_phone_numbers cpn
                WHERE cpn.entity_type='customer' AND cpn.entity_id=j.customer_id
                ORDER BY CASE WHEN cpn.label='Mobile' THEN 0 ELSE 1 END LIMIT 1) AS customer_phone,
               au.full_name AS assigned_agent_name
        FROM jobs j
        LEFT JOIN customers cu ON cu.id = j.customer_id
        LEFT JOIN users au ON au.id = j.assigned_user_id
        WHERE {scope_sql}
          AND (j.display_ref LIKE ? OR j.internal_job_number LIKE ? OR
               j.client_reference LIKE ? OR j.job_address LIKE ? OR
               j.lender_name LIKE ? OR j.client_job_number LIKE ? OR
               j.account_number LIKE ? OR
               (cu.first_name || ' ' || cu.last_name) LIKE ? OR
               EXISTS (SELECT 1 FROM job_items ji WHERE ji.job_id=j.id
                 AND (ji.reg LIKE ? OR ji.vin LIKE ?)))
        ORDER BY j.created_at DESC
        LIMIT 50
    """, params).fetchall()
    conn.close()
    results = []
    for r in rows:
        results.append({
            "id": r["id"],
            "display_ref": r["display_ref"] or r["internal_job_number"] or "",
            "internal_job_number": r["internal_job_number"] or "",
            "client_reference": r["client_reference"] or "",
            "account_number": r["account_number"] or "",
            "status": r["status"] or "",
            "job_address": r["job_address"] or "",
            "lat": r["lat"],
            "lng": r["lng"],
            "lender_name": r["lender_name"] or "",
            "customer_name": r["customer_name"] or "",
            "asset_reg": r["asset_reg"] or "",
            "asset_vin": r["asset_vin"] or "",
            "next_scheduled": r["next_scheduled"] or "",
            "customer_phone": r["customer_phone"] or "",
            "assigned_agent_name": r["assigned_agent_name"] or "" if is_admin else "",
        })
    return jsonify({"jobs": results})


@app.post("/m/jobs/prefs/save")
@mobile_login_required
def m_jobs_prefs_save():
    uid = session.get("user_id")
    sort           = request.form.get("sort", "visit_date")
    direction      = request.form.get("dir", "asc")
    scope          = request.form.get("scope", "mine")
    show_completed = request.form.get("show_completed", "week")
    dist_unit      = request.form.get("distance_unit", "km")
    ts = now_ts()
    conn = db()
    conn.execute("""
        INSERT INTO user_mobile_settings
            (user_id, list_sort, list_dir, job_scope, show_completed, distance_unit, updated_at)
        VALUES (?,?,?,?,?,?,?)
        ON CONFLICT(user_id) DO UPDATE SET
            list_sort=excluded.list_sort, list_dir=excluded.list_dir,
            job_scope=excluded.job_scope, show_completed=excluded.show_completed,
            distance_unit=excluded.distance_unit, updated_at=excluded.updated_at
    """, (uid, sort, direction, scope, show_completed, dist_unit, ts))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.get("/m/job/<int:job_id>")
@mobile_login_required
def m_job_detail(job_id):
    lockout = _check_agent_lockout()
    if lockout:
        return lockout
    uid  = session.get("user_id")
    role = session.get("role", "")
    conn = db()
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        abort(404)

    if role not in ("admin", "both"):
        has_access = job["assigned_user_id"] == uid or conn.execute(
            "SELECT 1 FROM schedules WHERE job_id=? AND assigned_to_user_id=? AND status NOT IN ('Cancelled','Completed')",
            (job_id, uid)
        ).fetchone()
        if not has_access:
            conn.close()
            flash("Access denied.", "danger")
            return redirect(url_for("m_today"))

    customer = None
    customer_mobile = ""
    if job["customer_id"]:
        customer = conn.execute("SELECT * FROM customers WHERE id=?", (job["customer_id"],)).fetchone()
        phone_row = conn.execute(
            "SELECT phone_number FROM contact_phone_numbers WHERE entity_type='customer' AND entity_id=? AND label='Mobile' LIMIT 1",
            (job["customer_id"],)
        ).fetchone()
        customer_mobile = phone_row["phone_number"] if phone_row else ""
        if not customer_mobile:
            phone_row2 = conn.execute(
                "SELECT phone_number FROM contact_phone_numbers WHERE entity_type='customer' AND entity_id=? LIMIT 1",
                (job["customer_id"],)
            ).fetchone()
            customer_mobile = phone_row2["phone_number"] if phone_row2 else ""

    client      = conn.execute("SELECT * FROM clients WHERE id=?", (job["client_id"],)).fetchone() if job["client_id"] else None
    _bill_to_id = job["bill_to_client_id"] if "bill_to_client_id" in job.keys() else None
    bill_client = conn.execute("SELECT * FROM clients WHERE id=?", (_bill_to_id,)).fetchone() if _bill_to_id else None
    assets      = conn.execute("SELECT * FROM job_items WHERE job_id=? ORDER BY id", (job_id,)).fetchall()
    notes       = conn.execute("""
        SELECT jfn.*, u.full_name AS agent_name
        FROM job_field_notes jfn
        LEFT JOIN users u ON u.id = jfn.created_by_user_id
        WHERE jfn.job_id=?
        ORDER BY jfn.created_at DESC
        LIMIT 30
    """, (job_id,)).fetchall()
    _note_ids = [n["id"] for n in notes]
    _note_files_map: dict = {}
    if _note_ids:
        placeholders = ",".join("?" * len(_note_ids))
        _file_rows = conn.execute(
            f"SELECT * FROM job_note_files WHERE job_field_note_id IN ({placeholders}) ORDER BY id",
            _note_ids
        ).fetchall()
        for fr in _file_rows:
            _note_files_map.setdefault(fr["job_field_note_id"], []).append(fr)

    has_draft = bool(conn.execute(
        "SELECT 1 FROM job_updates WHERE job_id=? AND created_by_user_id=? AND status='draft' LIMIT 1",
        (job_id, uid)
    ).fetchone())
    _rl_rows = conn.execute(
        "SELECT item_id, status FROM repo_lock_records WHERE job_id=?", (job_id,)
    ).fetchall()
    repo_lock_map = {r["item_id"]: (r["status"] or "Draft") for r in _rl_rows}
    assigned_agent_name = None
    if job["assigned_user_id"]:
        _ag = conn.execute("SELECT full_name FROM users WHERE id=?", (job["assigned_user_id"],)).fetchone()
        if _ag:
            assigned_agent_name = _ag["full_name"]
    conn.close()

    is_admin = role in ("admin", "both")
    return render_template("mobile/job_detail.html",
                           job=job, customer=customer, customer_mobile=customer_mobile,
                           client=client, bill_client=bill_client,
                           assets=assets, notes=notes, note_files_map=_note_files_map,
                           has_draft=has_draft, repo_lock_map=repo_lock_map,
                           assigned_agent_name=assigned_agent_name, is_admin=is_admin)


@app.get("/m/job/<int:job_id>/note/new")
@mobile_login_required
def m_update_builder(job_id):
    uid  = session.get("user_id")
    role = session.get("role", "")
    conn = db()
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        abort(404)

    if role not in ("admin", "both"):
        has_access = job["assigned_user_id"] == uid or conn.execute(
            "SELECT 1 FROM schedules WHERE job_id=? AND assigned_to_user_id=? AND status NOT IN ('Cancelled','Completed')",
            (job_id, uid)
        ).fetchone()
        if not has_access:
            conn.close()
            flash("Access denied.", "danger")
            return redirect(url_for("m_today"))

    customer = None
    customer_mobile = ""
    if job["customer_id"]:
        customer = conn.execute("SELECT * FROM customers WHERE id=?", (job["customer_id"],)).fetchone()
        phone_row = conn.execute(
            "SELECT phone_number FROM contact_phone_numbers WHERE entity_type='customer' AND entity_id=? AND label='Mobile' LIMIT 1",
            (job["customer_id"],)
        ).fetchone()
        customer_mobile = phone_row["phone_number"] if phone_row else ""

    draft = conn.execute(
        "SELECT * FROM job_updates WHERE job_id=? AND created_by_user_id=? AND status='draft' ORDER BY id DESC LIMIT 1",
        (job_id, uid)
    ).fetchone()

    if not draft:
        ts = now_ts()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO job_updates (job_id, created_by_user_id, status, customer_mobile, created_at, updated_at)
            VALUES (?, ?, 'draft', ?, ?, ?)
        """, (job_id, uid, customer_mobile, ts, ts))
        draft_id = cur.lastrowid
        conn.commit()
        draft = conn.execute("SELECT * FROM job_updates WHERE id=?", (draft_id,)).fetchone()

    first_asset = conn.execute(
        "SELECT * FROM job_items WHERE job_id=? AND item_type='vehicle' LIMIT 1", (job_id,)
    ).fetchone()
    asset_make_model = ""
    asset_reg = ""
    if first_asset:
        parts = [p for p in [first_asset["year"], first_asset["make"], first_asset["model"]] if p]
        asset_make_model = " ".join(parts)
        asset_reg = first_asset["reg"] or ""

    draft_photos = conn.execute(
        "SELECT id, tag FROM job_update_photos WHERE job_update_id=? ORDER BY id", (draft["id"],)
    ).fetchall()
    draft_photos_list = [{"id": p["id"], "url": f"/jobs/{job_id}/update-builder/photo/{p['id']}", "tag": p["tag"]} for p in draft_photos]
    conn.close()

    mel_now = datetime.now(_melbourne)
    return render_template("m/update_builder.html",
                           job=job, customer=customer, customer_mobile=customer_mobile,
                           draft=draft,
                           asset_make_model=asset_make_model,
                           asset_reg=asset_reg,
                           draft_photos=draft_photos_list,
                           now_date=mel_now.strftime("%Y-%m-%d"),
                           now_time=mel_now.strftime("%H:%M"))


# ─────────────────────────────────────────────────────────────────────────────
# Quick Field Note — text / audio / photo (inline, no page reload)
# ─────────────────────────────────────────────────────────────────────────────
_NOTE_AUDIO_DIR = os.path.join("uploads", "notes", "audio")
_NOTE_PHOTO_DIR = os.path.join("uploads", "notes", "photos")

@app.post("/m/job/<int:job_id>/quick-note")
@mobile_login_required
def m_quick_note_save(job_id):
    uid      = session.get("user_id")
    username = session.get("user_name", "Unknown")
    conn = db()
    job = conn.execute("SELECT id FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({"ok": False, "error": "Job not found"}), 404

    note_text  = (request.form.get("note_text") or "").strip() or None
    audio_file = request.files.get("audio_file")
    photo_file = request.files.get("photo_file")

    audio_filename = None
    if audio_file and audio_file.filename:
        ct  = (audio_file.content_type or "").lower()
        ext = ".webm" if "webm" in ct else ".m4a" if ("mp4" in ct or "m4a" in ct) else ".ogg" if "ogg" in ct else ".wav"
        audio_filename = f"{uuid.uuid4().hex}{ext}"
        os.makedirs(_NOTE_AUDIO_DIR, exist_ok=True)
        audio_file.save(os.path.join(_NOTE_AUDIO_DIR, audio_filename))

    has_text  = note_text is not None
    has_audio = bool(audio_filename)
    has_photo = bool(photo_file and photo_file.filename)

    if has_audio and has_text:       note_type = "audio_text"
    elif has_audio:                  note_type = "audio"
    elif has_photo and has_text:     note_type = "photo_text"
    elif has_photo:                  note_type = "photo"
    else:                            note_type = "text"

    if not has_text and not has_audio and not has_photo:
        conn.close()
        return jsonify({"ok": False, "error": "Nothing to save"}), 400

    ts  = now_ts()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO job_field_notes (job_id, created_by_user_id, note_text, created_at, note_type, audio_filename) VALUES (?,?,?,?,?,?)",
        (job_id, uid, note_text, ts, note_type, audio_filename)
    )
    note_id = cur.lastrowid

    photo_file_id = None
    if has_photo:
        safe_fn      = secure_filename(photo_file.filename or "photo.jpg")
        photo_name   = f"{uuid.uuid4().hex}_{safe_fn}"
        photo_path   = os.path.join(_NOTE_PHOTO_DIR, photo_name)
        os.makedirs(_NOTE_PHOTO_DIR, exist_ok=True)
        photo_file.save(photo_path)
        cur.execute(
            "INSERT INTO job_note_files (job_field_note_id, filename, filepath, uploaded_at) VALUES (?,?,?,?)",
            (note_id, photo_name, photo_path, ts)
        )
        photo_file_id = cur.lastrowid

    conn.commit()
    conn.close()

    return jsonify({
        "ok":           True,
        "note_id":      note_id,
        "note_type":    note_type,
        "created_at":   ts[:16].replace("T", " "),
        "agent_name":   username,
        "note_text":    note_text or "",
        "audio_url":    f"/m/note-audio/{note_id}" if audio_filename else None,
        "photo_url":    f"/m/note-photo/{photo_file_id}" if photo_file_id else None,
    })


@app.get("/m/note-audio/<int:note_id>")
@mobile_login_required
def m_note_audio(note_id):
    conn = db()
    note = conn.execute("SELECT audio_filename FROM job_field_notes WHERE id=?", (note_id,)).fetchone()
    conn.close()
    if not note or not note["audio_filename"]:
        abort(404)
    audio_path = os.path.join(_NOTE_AUDIO_DIR, note["audio_filename"])
    if not os.path.exists(audio_path):
        abort(404)
    ext      = os.path.splitext(note["audio_filename"])[1].lower().lstrip(".")
    mimetypes_map = {"webm": "audio/webm", "m4a": "audio/mp4", "ogg": "audio/ogg", "wav": "audio/wav"}
    mimetype = mimetypes_map.get(ext, "audio/mpeg")
    return send_file(audio_path, mimetype=mimetype)


@app.get("/m/note-photo/<int:file_id>")
@mobile_login_required
def m_note_photo(file_id):
    conn = db()
    row = conn.execute("SELECT filepath FROM job_note_files WHERE id=?", (file_id,)).fetchone()
    conn.close()
    if not row or not os.path.exists(row["filepath"]):
        abort(404)
    return send_file(row["filepath"])


@app.get("/m/tow-operators")
@mobile_login_required
def m_tow_operators_list():
    conn = db()
    rows = conn.execute(
        "SELECT * FROM tow_operators WHERE active=1 ORDER BY company_name"
    ).fetchall()
    conn.close()
    return render_template("mobile/tow_operators.html", items=rows)


@app.get("/m/tow-operators/new")
@mobile_login_required
def m_tow_operator_new():
    return render_template("mobile/tow_operator_form.html", item=None, form={})


@app.post("/m/tow-operators/new")
@mobile_login_required
def m_tow_operator_new_post():
    uid = session.get("user_id")
    f = request.form
    company_name = f.get("company_name", "").strip()
    if not company_name:
        flash("Company name is required.", "danger")
        return render_template("mobile/tow_operator_form.html", item=None, form=f)
    conn = db()
    cur = conn.execute("""
        INSERT INTO tow_operators
            (company_name, contact_name, mobile, phone, other_phone, email,
             address, suburb, state, postcode, notes, created_at, created_by_user_id)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        company_name,
        f.get("contact_name","").strip() or None,
        f.get("mobile","").strip() or None,
        f.get("phone","").strip() or None,
        f.get("other_phone","").strip() or None,
        f.get("email","").strip() or None,
        f.get("address","").strip() or None,
        f.get("suburb","").strip() or None,
        f.get("state","").strip() or None,
        f.get("postcode","").strip() or None,
        f.get("notes","").strip() or None,
        now_ts(), uid,
    ))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    try:
        _log_audit(uid, "create", "tow_operator", new_id)
    except Exception:
        pass
    flash(f"Tow operator \"{company_name}\" added.", "success")
    return redirect(url_for("m_tow_operators_list"))


@app.get("/m/auction-yards")
@mobile_login_required
def m_auction_yards_list():
    conn = db()
    rows = conn.execute(
        "SELECT * FROM auction_yards WHERE active=1 ORDER BY name"
    ).fetchall()
    conn.close()
    return render_template("mobile/auction_yards.html", items=rows)


@app.get("/m/auction-yards/new")
@mobile_login_required
def m_auction_yard_new():
    return render_template("mobile/auction_yard_form.html", item=None, form={})


@app.post("/m/auction-yards/new")
@mobile_login_required
def m_auction_yard_new_post():
    uid = session.get("user_id")
    f = request.form
    name = f.get("name", "").strip()
    if not name:
        flash("Yard name is required.", "danger")
        return render_template("mobile/auction_yard_form.html", item=None, form=f)
    conn = db()
    cur = conn.execute("""
        INSERT INTO auction_yards
            (name, contact_name, mobile, phone, other_phone, email,
             address, suburb, state, postcode, notes, created_at, created_by_user_id)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        name,
        f.get("contact_name","").strip() or None,
        f.get("mobile","").strip() or None,
        f.get("phone","").strip() or None,
        f.get("other_phone","").strip() or None,
        f.get("email","").strip() or None,
        f.get("address","").strip() or None,
        f.get("suburb","").strip() or None,
        f.get("state","").strip() or None,
        f.get("postcode","").strip() or None,
        f.get("notes","").strip() or None,
        now_ts(), uid,
    ))
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    try:
        _log_audit(uid, "create", "auction_yard", new_id)
    except Exception:
        pass
    flash(f"Auction yard \"{name}\" added.", "success")
    return redirect(url_for("m_auction_yards_list"))


def _log_audit(user_id, action, record_type, record_id):
    """Lightweight audit helper for mobile create actions."""
    conn = db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS mobile_audit_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER, action TEXT, record_type TEXT,
            record_id INTEGER, created_at TEXT
        )
    """)
    conn.execute(
        "INSERT INTO mobile_audit_log (user_id, action, record_type, record_id, created_at) VALUES (?,?,?,?,?)",
        (user_id, action, record_type, record_id, now_ts())
    )
    conn.commit()
    conn.close()


# ── LPR helpers ───────────────────────────────────────────────────────────────

_LPR_PRIVILEGED = {"admin", "both"}

def _lpr_ensure_table(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_audit_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            username TEXT NOT NULL,
            user_role TEXT NOT NULL,
            searched_registration TEXT NOT NULL,
            normalised_registration TEXT NOT NULL,
            result_type TEXT NOT NULL,
            matched_job_id INTEGER,
            matched_job_number TEXT,
            is_allocated_to_user INTEGER DEFAULT 0,
            search_method TEXT DEFAULT 'manual',
            created_at TEXT NOT NULL
        )
    """)
    add_column_if_missing(conn.cursor(), "lpr_audit_logs", "search_method", "TEXT DEFAULT 'manual'")


def _lpr_watchlist_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_watchlist (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            registration TEXT NOT NULL,
            registration_normalised TEXT NOT NULL,
            matched_job_id INTEGER,
            reason TEXT,
            priority TEXT DEFAULT 'normal',
            active INTEGER DEFAULT 1,
            created_by INTEGER,
            created_at TEXT NOT NULL,
            FOREIGN KEY(matched_job_id) REFERENCES jobs(id),
            FOREIGN KEY(created_by) REFERENCES users(id)
        )
    """)


def _lpr_watchlist_check(reg_norm: str) -> dict:
    conn = db()
    _lpr_watchlist_ensure(conn)
    row = conn.execute("""
        SELECT reason, priority FROM lpr_watchlist
        WHERE registration_normalised = ? AND active = 1
        ORDER BY CASE priority
            WHEN 'urgent' THEN 1 WHEN 'high' THEN 2
            WHEN 'normal' THEN 3 ELSE 4 END
        LIMIT 1
    """, (reg_norm,)).fetchone()
    conn.close()
    if row:
        return {
            "watchlist_hit":      True,
            "watchlist_reason":   row["reason"],
            "watchlist_priority": row["priority"],
        }
    return {"watchlist_hit": False, "watchlist_reason": None, "watchlist_priority": None}


def lookup_registration_for_lpr(uid: int, role: str, username: str, reg_input: str) -> dict:
    reg_norm = normalise_registration(reg_input)
    if not reg_norm:
        return {"result_type": "invalid", "message": "Enter a valid registration."}

    wl = _lpr_watchlist_check(reg_norm)

    conn = db()
    _lpr_ensure_table(conn)
    rows = conn.execute("""
        SELECT ji.id AS item_id, ji.reg, ji.make, ji.model, ji.year, ji.vin,
               j.id AS job_id, j.internal_job_number, j.display_ref,
               j.assigned_user_id, j.client_id, j.status,
               u.full_name AS assigned_agent_name
        FROM job_items ji
        JOIN jobs j ON j.id = ji.job_id
        LEFT JOIN users u ON u.id = j.assigned_user_id
        WHERE j.status NOT IN ('Completed','Invoiced','Cancelled')
          AND ji.reg IS NOT NULL AND ji.reg != ''
    """).fetchall()

    matches = [row for row in rows if normalise_registration(row["reg"]) == reg_norm]

    if not matches:
        conn.close()
        return {
            "result_type": "no_match",
            "searched_registration": reg_norm,
            "message": "No active registration found.",
            **wl,
        }

    if len(matches) > 1:
        conn.close()
        return {
            "result_type": "conflict",
            "searched_registration": reg_norm,
            "match_count": len(matches),
            "message": f"{len(matches)} active files share this registration. Contact the office for instructions.",
            **wl,
        }

    matched     = matches[0]
    job_id      = matched["job_id"]
    job_number  = matched["display_ref"] or matched["internal_job_number"]
    allocated   = matched["assigned_user_id"] == uid
    privileged  = role in _LPR_PRIVILEGED

    client_name = None
    if matched["client_id"]:
        c = conn.execute("SELECT name FROM clients WHERE id=?", (matched["client_id"],)).fetchone()
        client_name = c["name"] if c else None

    asset = {
        "registration": matched["reg"],
        "year":  matched["year"],
        "make":  matched["make"],
        "model": matched["model"],
        "vin":   matched["vin"],
    }

    conn.close()

    if allocated or privileged:
        return {
            "result_type": "allocated_match",
            "searched_registration": reg_norm,
            "matched_job_id": job_id,
            "matched_job_number": job_number,
            "is_allocated_to_user": bool(allocated),
            "asset": asset,
            "open_url": url_for("m_job_detail", job_id=job_id),
            **wl,
        }

    return {
        "result_type": "restricted_match",
        "searched_registration": reg_norm,
        "matched_job_id": job_id,
        "matched_job_number": job_number,
        "is_allocated_to_user": False,
        "asset": asset,
        "client_name": client_name,
        "assigned_agent_name": matched["assigned_agent_name"] or None,
        "notice": "Contact office for allocation before action.",
        **wl,
    }


def _log_lpr_search(uid: int, role: str, username: str, reg_input: str, result: dict,
                    search_method: str = "manual"):
    conn = db()
    _lpr_ensure_table(conn)
    conn.execute("""
        INSERT INTO lpr_audit_logs
            (user_id, username, user_role, searched_registration, normalised_registration,
             result_type, matched_job_id, matched_job_number, is_allocated_to_user,
             search_method, created_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,?)
    """, (
        uid,
        username,
        role,
        reg_input,
        normalise_registration(reg_input),
        result.get("result_type"),
        result.get("matched_job_id"),
        result.get("matched_job_number"),
        int(result.get("is_allocated_to_user", False)),
        search_method,
        now_ts(),
    ))
    conn.commit()
    conn.close()


def _default_prefs():
    return {
        "list_sort": "visit_date", "list_dir": "asc", "distance_unit": "km",
        "gps_foreground": 1, "gps_bg": 0, "gps_interval_mins": 5,
        "job_scope": "mine", "show_completed": "week",
        "mobile_default_view": "schedule", "show_status_on_visits": 1,
    }


@app.get("/m/settings")
@mobile_login_required
def m_settings():
    uid = session.get("user_id")
    conn = db()
    row = conn.execute("SELECT * FROM user_mobile_settings WHERE user_id=?", (uid,)).fetchone()
    conn.close()
    prefs = {**_default_prefs(), **(dict(row) if row else {})}
    return render_template("mobile/settings.html", prefs=prefs)


@app.post("/m/settings")
@mobile_login_required
def m_settings_post():
    uid = session.get("user_id")
    f = request.form
    ts = now_ts()
    conn = db()
    conn.execute("""
        INSERT INTO user_mobile_settings
            (user_id, list_sort, list_dir, distance_unit, gps_foreground, gps_bg, gps_interval_mins,
             job_scope, show_completed, mobile_default_view, show_status_on_visits, updated_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
        ON CONFLICT(user_id) DO UPDATE SET
            list_sort=excluded.list_sort, list_dir=excluded.list_dir,
            distance_unit=excluded.distance_unit, gps_foreground=excluded.gps_foreground,
            gps_bg=excluded.gps_bg, gps_interval_mins=excluded.gps_interval_mins,
            job_scope=excluded.job_scope, show_completed=excluded.show_completed,
            mobile_default_view=excluded.mobile_default_view,
            show_status_on_visits=excluded.show_status_on_visits,
            updated_at=excluded.updated_at
    """, (
        uid,
        f.get("list_sort", "visit_date"),
        f.get("list_dir", "asc"),
        f.get("distance_unit", "km"),
        1 if f.get("gps_foreground") else 0,
        1 if f.get("gps_bg") else 0,
        int(f.get("gps_interval_mins", 5)),
        f.get("job_scope", "mine"),
        f.get("show_completed", "week"),
        f.get("mobile_default_view", "schedule"),
        1 if f.get("show_status_on_visits") else 0,
        ts,
    ))
    conn.commit()
    conn.close()
    flash("Preferences saved.", "success")
    return redirect(url_for("m_settings"))


@app.post("/m/api/location/update")
@mobile_login_required
def m_api_location_update():
    uid  = session.get("user_id")
    data = request.get_json(silent=True) or {}
    lat  = data.get("lat")
    lng  = data.get("lng")
    acc  = data.get("accuracy")
    if not lat or not lng:
        return jsonify({"ok": False, "error": "lat/lng required"}), 400
    ts = now_ts()
    conn = db()
    conn.execute("""
        INSERT INTO agent_locations (user_id, lat, lng, accuracy, updated_at)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(user_id) DO UPDATE SET
            lat=excluded.lat, lng=excluded.lng,
            accuracy=excluded.accuracy, updated_at=excluded.updated_at
    """, (uid, float(lat), float(lng), float(acc) if acc else None, ts))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.post("/m/api/location/ping")
@mobile_login_required
def m_api_location_ping():
    """
    Richer location ping from AgentLocationService.
    Stores into agent_movement (full history) and upserts agent_locations (latest).
    No customer or finance data in payload — only position + operational context.
    """
    uid  = session.get("user_id")
    data = request.get_json(silent=True) or {}
    lat         = data.get("lat")
    lng         = data.get("lng")
    accuracy    = data.get("accuracy")
    captured_at = (data.get("captured_at") or "").strip() or now_ts()
    source      = (data.get("source") or "unknown").strip()
    battery     = (data.get("battery_state") or "unknown").strip()
    context     = (data.get("context") or "unknown").strip()

    if not lat or not lng:
        return jsonify({"ok": False, "error": "lat/lng required"}), 400

    received_at = now_ts()
    conn = db()
    _agent_movement_ensure(conn)
    conn.execute("""
        INSERT INTO agent_movement
            (user_id, latitude, longitude, captured_at, received_at,
             source, accuracy_m, battery_state, context)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (uid, float(lat), float(lng), captured_at, received_at,
          source, float(accuracy) if accuracy else None, battery, context))
    # Also keep agent_locations current for backward compat
    conn.execute("""
        INSERT INTO agent_locations (user_id, lat, lng, accuracy, updated_at)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(user_id) DO UPDATE SET
            lat=excluded.lat, lng=excluded.lng,
            accuracy=excluded.accuracy, updated_at=excluded.updated_at
    """, (uid, float(lat), float(lng),
          float(accuracy) if accuracy else None, received_at))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.get("/m/api/map/jobs")
@mobile_login_required
def m_api_map_jobs():
    """Reusable mobile job-map JSON endpoint with date-range filter.
    date_filter: today | tomorrow | week | month | all (default: today for agents, all for admins)
    """
    uid         = session.get("user_id")
    role        = session.get("role", "")
    is_admin    = role in ("admin", "both")
    date_filter = request.args.get("date_filter", "today" if not is_admin else "all")

    date_clauses = {
        "today":    "date(j.job_due_date) = date('now','localtime')",
        "tomorrow": "date(j.job_due_date) = date('now','+1 day','localtime')",
        "week":     "date(j.job_due_date) BETWEEN date('now','localtime') AND date('now','+6 days','localtime')",
        "month":    "date(j.job_due_date) BETWEEN date('now','localtime') AND date('now','+29 days','localtime')",
    }
    date_sql = f"AND {date_clauses[date_filter]}" if date_filter in date_clauses else ""

    base_where = f"j.status NOT IN ('Closed','Cancelled') {date_sql}"

    conn = db()
    extra_cols = """,
                   j.job_type, j.created_at,
                   (SELECT COUNT(*) FROM job_field_notes fn WHERE fn.job_id = j.id) AS note_count,
                   (SELECT 1 FROM job_updates ju
                    WHERE ju.job_id = j.id AND ju.status = 'draft'
                    LIMIT 1) AS has_draft"""
    if is_admin:
        rows = conn.execute(f"""
            SELECT j.id, j.display_ref, j.job_address, j.status, j.lat, j.lng,
                   j.job_due_date, j.lender_name, j.client_job_number,
                   (cu.first_name || ' ' || cu.last_name) AS customer_name,
                   c.name  AS client_name,
                   ag.full_name AS agent_name,
                   ag.id        AS agent_id
                   {extra_cols}
            FROM jobs j
            LEFT JOIN customers cu ON cu.id = j.customer_id
            LEFT JOIN clients   c  ON c.id  = j.client_id
            LEFT JOIN users     ag ON ag.id = j.assigned_user_id
            WHERE {base_where}
            ORDER BY j.job_due_date ASC
            LIMIT 500
        """).fetchall()
    else:
        rows = conn.execute(f"""
            SELECT j.id, j.display_ref, j.job_address, j.status, j.lat, j.lng,
                   j.job_due_date, j.lender_name, j.client_job_number,
                   (cu.first_name || ' ' || cu.last_name) AS customer_name,
                   c.name  AS client_name,
                   ag.full_name AS agent_name,
                   ag.id        AS agent_id
                   {extra_cols}
            FROM jobs j
            LEFT JOIN customers cu ON cu.id = j.customer_id
            LEFT JOIN clients   c  ON c.id  = j.client_id
            LEFT JOIN users     ag ON ag.id = j.assigned_user_id
            WHERE j.assigned_user_id = ? AND {base_where}
            ORDER BY j.job_due_date ASC
            LIMIT 200
        """, (uid,)).fetchall()

    conn.close()
    jobs_out = []
    for r in rows:
        jobs_out.append({
            "id":           r["id"],
            "ref":          r["display_ref"],
            "address":      r["job_address"] or "",
            "status":       r["status"] or "",
            "lat":          r["lat"],
            "lng":          r["lng"],
            "due_date":     r["job_due_date"] or "",
            "lender":       r["lender_name"] or "",
            "client_ref":   r["client_job_number"] or "",
            "customer":     r["customer_name"] or "",
            "client":       r["client_name"] or "",
            "agent":        r["agent_name"] or "",
            "agent_id":     r["agent_id"],
            "job_type":     r["job_type"] or "",
            "note_count":   r["note_count"] or 0,
            "created_at":   r["created_at"] or "",
            "has_draft":    bool(r["has_draft"]),
        })
    return jsonify({"jobs": jobs_out, "filter": date_filter})


@app.get("/m/map")
@mobile_login_required
def m_map():
    uid  = session.get("user_id")
    conn = db()
    prefs = conn.execute(
        "SELECT distance_unit, gps_foreground, gps_interval_mins FROM user_mobile_settings WHERE user_id=?",
        (uid,)
    ).fetchone()
    conn.close()
    distance_unit   = prefs["distance_unit"]    if prefs else "km"
    gps_foreground  = prefs["gps_foreground"]   if prefs else 1
    gps_interval    = prefs["gps_interval_mins"] if prefs and prefs["gps_interval_mins"] else 5
    role = session.get("role", "")
    is_admin = role in ("admin", "both")
    return render_template("mobile/map.html",
                           distance_unit=distance_unit,
                           gps_foreground=gps_foreground,
                           gps_interval_mins=gps_interval,
                           is_admin=is_admin,
                           current_user_id=uid)


# ─────────────────────────────────────────────────────────────────────────────
# LPR Admin Audit Screen
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/admin/lpr-audit")
@admin_required
def admin_lpr_audit():
    conn   = db()
    _lpr_ensure_table(conn)

    f_from   = request.args.get("from_date", "")
    f_to     = request.args.get("to_date", "")
    f_user   = request.args.get("user_id", "")
    f_reg    = (request.args.get("registration") or "").strip().upper()
    f_result = request.args.get("result_type", "")
    f_method = request.args.get("search_method", "")

    where, params = [], []
    if f_from:
        where.append("l.created_at >= ?"); params.append(f_from)
    if f_to:
        where.append("l.created_at <= ?"); params.append(f_to + "T23:59:59")
    if f_user:
        where.append("l.user_id = ?"); params.append(int(f_user))
    if f_reg:
        where.append("l.normalised_registration = ?"); params.append(normalise_registration(f_reg))
    if f_result:
        where.append("l.result_type = ?"); params.append(f_result)
    if f_method:
        where.append("l.search_method = ?"); params.append(f_method)

    sql = """
        SELECT l.*, u.full_name AS agent_name
        FROM lpr_audit_logs l
        LEFT JOIN users u ON u.id = l.user_id
        {}
        ORDER BY l.created_at DESC
        LIMIT 500
    """.format("WHERE " + " AND ".join(where) if where else "")

    rows  = conn.execute(sql, params).fetchall()
    users = conn.execute("SELECT id, full_name FROM users WHERE active=1 ORDER BY full_name").fetchall()
    conn.close()

    return render_template("lpr_audit.html",
                           rows=rows, users=users,
                           f_from=f_from, f_to=f_to,
                           f_user=f_user, f_reg=f_reg,
                           f_result=f_result, f_method=f_method)


# ─────────────────────────────────────────────────────────────────────────────
# Mobile LPR — Stage 1: manual plate lookup
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/m/lpr/capture", methods=["GET", "POST"])
@mobile_login_required
def m_lpr_capture():
    if request.method == "POST":
        file = request.files.get("image")
        if not file or not file.filename:
            return render_template("mobile/lpr_capture.html", error="No image uploaded. Please select or take a photo.")

        if file.content_length and file.content_length > 25 * 1024 * 1024:
            return render_template("mobile/lpr_capture.html", error="Image too large. Maximum file size is 25 MB.")

        safe_name = secure_filename(file.filename) or "plate.jpg"
        tmp_path = os.path.join(_LPR_TMP, f"{uuid.uuid4().hex}_{safe_name}")

        try:
            file.save(tmp_path)
        except Exception:
            return render_template("mobile/lpr_capture.html", error="Failed to save uploaded image. Please try again.")

        if os.path.getsize(tmp_path) == 0:
            try:
                os.remove(tmp_path)
            except OSError:
                pass
            return render_template("mobile/lpr_capture.html", error="The uploaded image is empty. Please take or select a photo and try again.")

        try:
            detected_plate = extract_plate_from_image(tmp_path)
        except Exception:
            detected_plate = ""

        try:
            os.remove(tmp_path)
        except OSError:
            pass

        return render_template("mobile/lpr_confirm.html",
                               detected_plate=detected_plate,
                               search_method="photo_ocr")

    return render_template("mobile/lpr_capture.html", error=None)


@app.route("/m/lpr", methods=["GET", "POST"])
@mobile_login_required
def m_lpr():
    uid      = session.get("user_id")
    role     = session.get("role", "")
    username = session.get("user_name", "")

    if request.method == "POST":
        reg_input     = (request.form.get("registration") or "").strip()
        search_method = (request.form.get("method") or "manual").strip()
        result        = lookup_registration_for_lpr(uid, role, username, reg_input)
        _log_lpr_search(uid, role, username, reg_input, result, search_method=search_method)
        if result.get("watchlist_hit") and result.get("watchlist_priority") in ("urgent", "high"):
            _notify_admins(
                "Watchlist Plate Detected",
                "An LPR scan matched a watchlist entry.",
                "watchlist_lookup",
                {"watchlist_priority": result.get("watchlist_priority")},
                exclude_uid=uid,
            )
        return render_template("mobile/lpr_result.html", result=result, search_method=search_method)

    plate_param = (request.args.get("plate") or "").strip()
    return render_template("mobile/lpr_search.html", plate_param=plate_param)


@app.get("/m/lpr/patrol-mode")
@mobile_login_required
def m_lpr_patrol_mode():
    conn = db()
    settings = conn.execute("SELECT lpr_patrol_mode_enabled FROM system_settings WHERE id=1").fetchone()
    conn.close()
    if settings and settings["lpr_patrol_mode_enabled"] == 0:
        from flask import abort
        abort(403)
    return render_template("mobile/lpr_patrol_mode.html")


@app.post("/m/api/lpr/patrol-scan")
@mobile_login_required
def m_api_lpr_patrol_scan():
    uid      = session.get("user_id")
    role     = session.get("role", "")
    username = session.get("user_name", "")

    # Path A — JSON with pre-extracted plate (iOS native bridge, Swift has already done OCR)
    if request.is_json:
        body  = request.get_json(force=True, silent=True) or {}
        plate = normalise_registration(body.get("plate", ""))
        if not plate or len(plate) < 3:
            return jsonify({"ok": True, "plate": None}), 200

    # Path B — multipart form with a camera frame (web browser fallback, Tesseract OCR)
    elif "frame" in request.files:
        frame    = request.files["frame"]
        tmp_path = os.path.join(_LPR_TMP, f"patrol_{uuid.uuid4().hex}.jpg")
        frame.save(tmp_path)
        try:
            plate = extract_plate_from_image(tmp_path)
        except Exception:
            plate = ""
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
        if not plate or len(plate) < 3:
            return jsonify({"ok": True, "plate": None}), 200

    else:
        return jsonify({"ok": False, "error": "No frame or plate provided"}), 400

    result = lookup_registration_for_lpr(uid, role, username, plate)
    _log_lpr_search(uid, role, username, plate, result, search_method="patrol_scan")

    if result.get("watchlist_hit") and result.get("watchlist_priority") in ("urgent", "high"):
        _notify_admins(
            "Watchlist Plate Detected (Patrol Scan)",
            "An LPR patrol scan matched a watchlist entry.",
            "watchlist_lookup",
            {"watchlist_priority": result.get("watchlist_priority")},
            exclude_uid=uid,
        )

    return jsonify({"ok": True, "plate": plate, "result": result}), 200


@app.post("/m/api/lpr/lookup")
@mobile_login_required
def m_api_lpr_lookup():
    uid      = session.get("user_id")
    role     = session.get("role", "")
    username = session.get("user_name", "")

    data          = request.get_json(silent=True) or {}
    reg_input     = (data.get("registration") or "").strip()
    search_method = (data.get("method") or "live_scan").strip()

    result = lookup_registration_for_lpr(uid, role, username, reg_input)
    _log_lpr_search(uid, role, username, reg_input, result, search_method=search_method)
    if result.get("watchlist_hit") and result.get("watchlist_priority") in ("urgent", "high"):
        _notify_admins(
            "Watchlist Plate Detected",
            "An LPR scan matched a watchlist entry.",
            "watchlist_lookup",
            {"watchlist_priority": result.get("watchlist_priority")},
            exclude_uid=uid,
        )
    return jsonify(result), 200


# ─────────────────────────────────────────────────────────────────────────────
# LPR Sightings — save + admin view + agent history
# ─────────────────────────────────────────────────────────────────────────────

def _lpr_sightings_ensure_table(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_sightings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            registration_raw TEXT NOT NULL,
            registration_normalised TEXT NOT NULL,
            search_method TEXT DEFAULT 'live_scan',
            result_type TEXT NOT NULL,
            matched_job_id INTEGER,
            matched_job_number TEXT,
            latitude REAL,
            longitude REAL,
            photo_path TEXT,
            notes TEXT,
            escalated_to_office INTEGER DEFAULT 0,
            watchlist_hit INTEGER DEFAULT 0,
            reviewed INTEGER DEFAULT 0,
            reviewed_by INTEGER,
            reviewed_at TEXT,
            office_note TEXT,
            follow_up_status TEXT,
            client_action_id TEXT,
            FOREIGN KEY(user_id) REFERENCES users(id)
        )
    """)
    cur = conn.cursor()
    add_column_if_missing(cur, "lpr_sightings", "watchlist_hit",      "INTEGER DEFAULT 0")
    add_column_if_missing(cur, "lpr_sightings", "reviewed",           "INTEGER DEFAULT 0")
    add_column_if_missing(cur, "lpr_sightings", "reviewed_by",        "INTEGER")
    add_column_if_missing(cur, "lpr_sightings", "reviewed_at",        "TEXT")
    add_column_if_missing(cur, "lpr_sightings", "office_note",        "TEXT")
    add_column_if_missing(cur, "lpr_sightings", "follow_up_status",   "TEXT")
    add_column_if_missing(cur, "lpr_sightings", "client_action_id",   "TEXT")


@app.post("/m/api/lpr/sighting")
@mobile_login_required
def m_api_lpr_sighting_save():
    uid      = session.get("user_id")
    username = session.get("user_name", "")

    data = request.get_json(silent=True) or {}

    reg_raw   = (data.get("registration_raw") or "").strip()
    reg_norm  = normalise_registration(reg_raw)
    if not reg_norm:
        return jsonify({"ok": False, "error": "Invalid registration"}), 400

    result_type   = (data.get("result_type") or "").strip()
    search_method = (data.get("search_method") or "live_scan").strip()

    # Only accept safe (non-customer) fields
    matched_job_id     = data.get("matched_job_id")    or None
    matched_job_number = data.get("matched_job_number") or None

    lat  = data.get("latitude")
    lng  = data.get("longitude")
    notes       = (data.get("notes") or "").strip() or None
    escalated   = 1 if data.get("escalated_to_office") else 0
    watchlist_h = 1 if data.get("watchlist_hit") else 0

    client_action_id = (data.get("client_action_id") or "").strip() or None

    conn = db()
    _lpr_sightings_ensure_table(conn)

    # Idempotency: if this client_action_id was already saved by this user, return it
    if client_action_id:
        existing = conn.execute(
            "SELECT id FROM lpr_sightings WHERE client_action_id=? AND user_id=?",
            (client_action_id, uid)
        ).fetchone()
        if existing:
            conn.close()
            return jsonify({"ok": True, "sighting_id": existing["id"], "duplicate": True}), 200

    cur = conn.execute("""
        INSERT INTO lpr_sightings
            (created_at, user_id, registration_raw, registration_normalised,
             search_method, result_type, matched_job_id, matched_job_number,
             latitude, longitude, notes, escalated_to_office, watchlist_hit,
             client_action_id)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (now_ts(), uid, reg_raw, reg_norm, search_method, result_type,
          matched_job_id, matched_job_number,
          float(lat) if lat is not None else None,
          float(lng) if lng is not None else None,
          notes, escalated, watchlist_h, client_action_id))
    sighting_id = cur.lastrowid
    # proximity check before close
    prox_hits = []
    if lat is not None and lng is not None and watchlist_h:
        prox_hits = _proximity_check(float(lat), float(lng), conn)
    conn.commit()
    conn.close()

    _log_audit(uid, "save", "lpr_sighting", sighting_id)

    if watchlist_h or escalated:
        _notify_admins(
            "LPR Sighting Requires Review",
            "A sighting has been flagged that requires office attention.",
            "sighting_alert",
            {"sighting_id": sighting_id,
             "watchlist_hit": bool(watchlist_h),
             "escalated": bool(escalated)},
            exclude_uid=uid,
        )
    for ph in prox_hits:
        rule = ph["rule"]
        _notify_admins(
            "Proximity Zone Alert",
            f"LPR sighting inside zone \"{rule['name']}\" ({ph['distance_m']}m from centre).",
            "proximity_alert",
            {"sighting_id": sighting_id,
             "rule_id": rule["id"],
             "rule_name": rule["name"]},
            exclude_uid=uid,
        )

    return jsonify({"ok": True, "sighting_id": sighting_id}), 200


@app.get("/m/api/lpr/sync")
@mobile_login_required
def m_api_lpr_sync():
    uid  = session.get("user_id")
    conn = db()
    _lpr_notifications_ensure(conn)
    _lpr_followups_ensure(conn)
    unread = conn.execute(
        "SELECT COUNT(*) AS n FROM lpr_notifications WHERE user_id=? AND read_at IS NULL",
        (uid,)
    ).fetchone()["n"]
    followups = conn.execute("""
        SELECT COUNT(*) AS n FROM lpr_followups
        WHERE assigned_user_id=? AND status='open'
    """, (uid,)).fetchone()["n"]
    conn.close()
    return jsonify({"unread_notification_count": unread,
                    "assigned_followup_count": followups,
                    "server_time": now_ts()}), 200


@app.get("/m/api/lpr/assigned-followups")
@mobile_login_required
def m_api_lpr_assigned_followups():
    uid  = session.get("user_id")
    conn = db()
    _lpr_followups_ensure(conn)
    rows = conn.execute("""
        SELECT f.id, f.action_type, f.priority, f.status, f.due_at,
               s.registration_normalised, s.latitude, s.longitude
        FROM lpr_followups f
        LEFT JOIN lpr_sightings s ON s.id = f.sighting_id
        WHERE f.assigned_user_id=?
          AND f.status NOT IN ('completed', 'cancelled')
        ORDER BY f.created_at DESC
        LIMIT 20
    """, (uid,)).fetchall()
    count = len(rows)
    conn.close()
    items = [{"id": r["id"], "action_type": r["action_type"],
              "priority": r["priority"], "status": r["status"],
              "due_at": r["due_at"],
              "registration": r["registration_normalised"],
              "latitude": r["latitude"], "longitude": r["longitude"]} for r in rows]
    return jsonify({"count": count, "items": items}), 200


@app.get("/m/api/lpr/followup/<int:followup_id>")
@mobile_login_required
def m_api_lpr_followup_detail(followup_id: int):
    """
    Return dispatch detail for one follow-up assigned to the current agent.
    Returns only operational fields — no customer name, address, arrears, or file data.
    """
    uid  = session.get("user_id")
    conn = db()
    _lpr_followups_ensure(conn)
    _lpr_sightings_ensure_table(conn)
    f = conn.execute("""
        SELECT f.id, f.action_type, f.priority, f.status,
               f.due_at, f.office_note, f.assigned_user_id,
               f.sighting_id,
               s.registration_normalised, s.result_type,
               s.latitude, s.longitude, s.created_at AS sighting_at
        FROM lpr_followups f
        LEFT JOIN lpr_sightings s ON s.id = f.sighting_id
        WHERE f.id = ? AND f.assigned_user_id = ?
    """, (followup_id, uid)).fetchone()
    conn.close()
    if not f:
        return jsonify({"error": "Not found"}), 404
    return jsonify({
        "id":          f["id"],
        "action_type": f["action_type"],
        "priority":    f["priority"],
        "status":      f["status"],
        "due_at":      f["due_at"] or "",
        "office_note": f["office_note"] or "",
        "sighting": {
            "id":           f["sighting_id"],
            "registration": f["registration_normalised"] or "",
            "result_type":  f["result_type"] or "",
            "latitude":     f["latitude"],
            "longitude":    f["longitude"],
            "sighting_at":  (f["sighting_at"] or "")[:16].replace("T", " "),
        },
    }), 200


@app.route("/m/api/lpr/followup/<int:followup_id>/status", methods=["PATCH", "POST"])
@mobile_login_required
def m_api_lpr_followup_status(followup_id: int):
    """
    Agent-side status transition for an assigned follow-up.
    Valid statuses: assigned → en_route → near_target → arrived → completed (or cancelled).
    Logs the transition timestamp in the appropriate column.
    """
    uid  = session.get("user_id")
    data = request.get_json(silent=True) or {}
    new_status = (data.get("status") or "").strip().lower()
    ts   = now_ts()
    conn = db()
    _lpr_followups_ensure(conn)
    row = conn.execute(
        "SELECT status, assigned_user_id FROM lpr_followups WHERE id=?",
        (followup_id,)
    ).fetchone()
    if not row or row["assigned_user_id"] != uid:
        conn.close()
        return jsonify({"ok": False, "error": "Not found or not assigned to you"}), 403
    current = row["status"] or "open"
    allowed = _FOLLOWUP_VALID_TRANSITIONS.get(current, set())
    if new_status not in allowed:
        conn.close()
        return jsonify({"ok": False,
                        "error": f"Cannot transition from '{current}' to '{new_status}'"}), 422
    ts_col_map = {
        "assigned":    "assigned_at",
        "en_route":    "en_route_at",
        "arrived":     "arrived_at",
        "completed":   "completed_at",
    }
    ts_col     = ts_col_map.get(new_status)
    set_parts  = ["status=?"]
    params     = [new_status]
    if ts_col:
        set_parts.append(f"{ts_col}=?")
        params.append(ts)
    params.append(followup_id)
    try:
        conn.execute(f"UPDATE lpr_followups SET {', '.join(set_parts)} WHERE id=?", params)
    except Exception:
        cur = conn.cursor()
        if ts_col:
            add_column_if_missing(cur, "lpr_followups", ts_col, "TEXT")
        conn.execute(f"UPDATE lpr_followups SET {', '.join(set_parts)} WHERE id=?", params)
    if new_status == "near_target":
        try:
            _dispatch_geofences_ensure(conn)
            conn.execute("""
                UPDATE dispatch_geofences
                SET triggered=1, triggered_at=?
                WHERE followup_id=? AND triggered=0
            """, (ts, followup_id))
        except Exception:
            pass
    conn.commit()
    conn.close()
    _log_audit(uid, f"followup_{new_status}", "lpr_followup", followup_id)
    return jsonify({"ok": True, "status": new_status}), 200


@app.post("/m/api/lpr/dispatch/sequence")
@mobile_login_required
def m_api_lpr_dispatch_sequence():
    """
    Return an optimised attendance sequence for a list of sighting IDs.
    Greedy nearest-neighbour from their centroid — no customer/finance data.
    """
    data = request.get_json(silent=True) or {}
    raw_ids = data.get("sighting_ids", [])
    ids = []
    for x in raw_ids:
        try:
            ids.append(int(x))
        except (TypeError, ValueError):
            pass
    if not ids:
        return jsonify({"ok": False, "error": "sighting_ids required"}), 400
    conn = db()
    _lpr_sightings_ensure_table(conn)
    _lpr_followups_ensure(conn)
    ordered = _dispatch_sequence(ids, conn)
    conn.close()
    clean = [{"id":       r["id"],
              "sequence": r["sequence"],
              "lat":      r["latitude"],
              "lng":      r["longitude"],
              "priority": r["priority"]} for r in ordered]
    return jsonify({"ok": True, "sequence": clean}), 200


@app.get("/m/lpr/history")
@mobile_login_required
def m_lpr_history():
    uid  = session.get("user_id")
    conn = db()
    _lpr_sightings_ensure_table(conn)
    rows = conn.execute("""
        SELECT s.*, u.full_name AS agent_name
        FROM lpr_sightings s
        LEFT JOIN users u ON u.id = s.user_id
        WHERE s.user_id = ?
        ORDER BY s.created_at DESC
        LIMIT 100
    """, (uid,)).fetchall()
    conn.close()
    return render_template("mobile/lpr_history.html", rows=rows)


@app.get("/admin/lpr-sightings")
@admin_required
def admin_lpr_sightings():
    conn = db()
    _lpr_sightings_ensure_table(conn)

    f_from      = request.args.get("from_date", "")
    f_to        = request.args.get("to_date", "")
    f_user      = request.args.get("user_id", "")
    f_escalated = request.args.get("escalated", "")
    f_result    = request.args.get("result_type", "")

    where, params = [], []
    if f_from:
        where.append("s.created_at >= ?"); params.append(f_from)
    if f_to:
        where.append("s.created_at <= ?"); params.append(f_to + "T23:59:59")
    if f_user:
        where.append("s.user_id = ?"); params.append(int(f_user))
    if f_escalated == "1":
        where.append("s.escalated_to_office = 1")
    if f_result:
        where.append("s.result_type = ?"); params.append(f_result)

    where_clause = "WHERE " + " AND ".join(where) if where else ""
    sql = """
        WITH plate_counts AS (
            SELECT registration_normalised, COUNT(*) AS cnt
            FROM lpr_sightings
            WHERE created_at >= datetime('now', '-30 days')
            GROUP BY registration_normalised
        )
        SELECT s.*, u.full_name AS agent_name,
               COALESCE(pc.cnt, 1) AS plate_count
        FROM lpr_sightings s
        LEFT JOIN users u ON u.id = s.user_id
        LEFT JOIN plate_counts pc ON pc.registration_normalised = s.registration_normalised
        {where}
        ORDER BY s.created_at DESC
        LIMIT 500
    """.format(where=where_clause)

    rows  = conn.execute(sql, params).fetchall()
    users = conn.execute("SELECT id, full_name FROM users WHERE active=1 ORDER BY full_name").fetchall()
    conn.close()

    return render_template("lpr_sightings.html",
                           rows=rows, users=users,
                           f_from=f_from, f_to=f_to,
                           f_user=f_user, f_escalated=f_escalated,
                           f_result=f_result)


@app.post("/admin/lpr-sightings/<int:sighting_id>/review")
@admin_required
def admin_lpr_sighting_review(sighting_id: int):
    reviewer_id = session.get("user_id")
    reviewer    = session.get("user_name", "")
    office_note       = (request.form.get("office_note") or "").strip() or None
    follow_up_status  = (request.form.get("follow_up_status") or "").strip() or None

    conn = db()
    _lpr_sightings_ensure_table(conn)
    conn.execute("""
        UPDATE lpr_sightings
        SET reviewed=1, reviewed_by=?, reviewed_at=?, office_note=?, follow_up_status=?
        WHERE id=?
    """, (reviewer_id, now_ts(), office_note, follow_up_status, sighting_id))
    conn.commit()
    sighting = conn.execute(
        "SELECT user_id FROM lpr_sightings WHERE id=?", (sighting_id,)
    ).fetchone()
    conn.close()
    _log_audit(reviewer_id, "review", "lpr_sighting", sighting_id)
    if sighting and sighting["user_id"] != reviewer_id:
        _notify_user(
            sighting["user_id"],
            "Sighting Reviewed by Office",
            "The office has reviewed your LPR sighting.",
            "sighting_reviewed",
            {"sighting_id": sighting_id, "follow_up_status": follow_up_status},
        )
    return redirect(request.referrer or url_for("admin_lpr_sightings"))


@app.get("/admin/lpr-sightings/map")
@admin_required
def admin_lpr_sightings_map():
    import json as _json
    conn = db()
    _lpr_sightings_ensure_table(conn)
    rows = conn.execute("""
        WITH plate_counts AS (
            SELECT registration_normalised, COUNT(*) AS cnt
            FROM lpr_sightings
            WHERE created_at >= datetime('now', '-30 days')
            GROUP BY registration_normalised
        )
        SELECT s.id, s.created_at, s.registration_normalised,
               s.result_type, s.search_method,
               s.latitude, s.longitude,
               s.escalated_to_office, s.watchlist_hit,
               s.matched_job_number, s.reviewed, s.follow_up_status,
               u.full_name AS agent_name,
               COALESCE(pc.cnt, 1) AS plate_count
        FROM lpr_sightings s
        LEFT JOIN users u ON u.id = s.user_id
        LEFT JOIN plate_counts pc ON pc.registration_normalised = s.registration_normalised
        WHERE s.latitude IS NOT NULL AND s.longitude IS NOT NULL
        ORDER BY s.created_at DESC
        LIMIT 1000
    """).fetchall()

    # Agent locations — prefer agent_movement (richer), fallback to agent_locations
    agent_rows = []
    try:
        _agent_movement_ensure(conn)
        agent_rows = conn.execute("""
            SELECT am.user_id, u.full_name,
                   am.latitude AS lat, am.longitude AS lng, am.received_at AS updated_at,
                   am.source, am.battery_state
            FROM agent_movement am
            JOIN users u ON u.id = am.user_id
            WHERE am.received_at >= datetime('now', '-8 hours')
              AND am.received_at = (
                  SELECT MAX(am2.received_at)
                  FROM agent_movement am2
                  WHERE am2.user_id = am.user_id
                    AND am2.received_at >= datetime('now', '-8 hours')
              )
            ORDER BY am.received_at DESC
        """).fetchall()
    except Exception:
        pass
    if not agent_rows:
        try:
            agent_rows = conn.execute("""
                SELECT al.user_id, u.full_name, al.lat, al.lng, al.updated_at,
                       NULL AS source, NULL AS battery_state
                FROM agent_locations al
                JOIN users u ON u.id = al.user_id
                WHERE al.lat IS NOT NULL AND al.lng IS NOT NULL
                  AND al.updated_at >= datetime('now', '-8 hours')
                ORDER BY al.updated_at DESC
            """).fetchall()
        except Exception:
            pass
    conn.close()

    features = []
    for r in rows:
        features.append({
            "type": "Feature",
            "geometry": {"type": "Point", "coordinates": [r["longitude"], r["latitude"]]},
            "properties": {
                "id":           r["id"],
                "reg":          r["registration_normalised"],
                "result_type":  r["result_type"],
                "method":       r["search_method"],
                "agent":        r["agent_name"] or "",
                "date":         r["created_at"][:10] if r["created_at"] else "",
                "time":         r["created_at"][11:16] if r["created_at"] else "",
                "job":          r["matched_job_number"] or "",
                "escalated":    bool(r["escalated_to_office"]),
                "watchlist":    bool(r["watchlist_hit"]),
                "reviewed":     bool(r["reviewed"]),
                "follow_up":    r["follow_up_status"] or "",
                "plate_count":  r["plate_count"] or 1,
            }
        })

    agents = [
        {
            "name":       r["full_name"],
            "lat":        r["lat"],
            "lng":        r["lng"],
            "updated_at": r["updated_at"],
        }
        for r in agent_rows
    ]

    geojson = _json.dumps({"type": "FeatureCollection", "features": features})
    agents_json = _json.dumps(agents)
    return render_template("lpr_sightings_map.html",
                           geojson=geojson, count=len(features),
                           agents_json=agents_json)


@app.get("/admin/lpr-watchlist")
@admin_required
def admin_lpr_watchlist():
    conn = db()
    _lpr_watchlist_ensure(conn)
    f_active = request.args.get("active", "1")
    where = "WHERE w.active=1" if f_active == "1" else ("WHERE w.active=0" if f_active == "0" else "")
    rows = conn.execute(f"""
        SELECT w.*, u.full_name AS creator_name,
               j.display_ref AS job_ref, j.internal_job_number AS job_num
        FROM lpr_watchlist w
        LEFT JOIN users u ON u.id = w.created_by
        LEFT JOIN jobs  j ON j.id = w.matched_job_id
        {where}
        ORDER BY w.created_at DESC
        LIMIT 500
    """).fetchall()
    jobs  = conn.execute("""
        SELECT id, COALESCE(display_ref, internal_job_number) AS ref
        FROM jobs WHERE status NOT IN ('Completed','Invoiced','Cancelled','Archived - Invoiced','Cold Stored')
        ORDER BY ref
    """).fetchall()
    conn.close()
    return render_template("lpr_watchlist.html", rows=rows, jobs=jobs, f_active=f_active)


@app.post("/admin/lpr-watchlist/add")
@admin_required
def admin_lpr_watchlist_add():
    uid = session.get("user_id")
    reg_raw = (request.form.get("registration") or "").strip()
    reg_norm = normalise_registration(reg_raw)
    if not reg_norm:
        flash("Invalid registration plate.", "danger")
        return redirect(url_for("admin_lpr_watchlist"))

    matched_job_id = request.form.get("matched_job_id") or None
    reason   = (request.form.get("reason") or "").strip() or None
    priority = (request.form.get("priority") or "normal").strip()

    conn = db()
    _lpr_watchlist_ensure(conn)
    conn.execute("""
        INSERT INTO lpr_watchlist
            (registration, registration_normalised, matched_job_id,
             reason, priority, active, created_by, created_at)
        VALUES (?,?,?,?,?,1,?,?)
    """, (reg_raw, reg_norm, matched_job_id, reason, priority, uid, now_ts()))
    conn.commit()
    conn.close()
    flash(f"Watchlist entry added for {reg_norm}.", "success")
    return redirect(url_for("admin_lpr_watchlist"))


@app.post("/admin/lpr-watchlist/<int:entry_id>/toggle")
@admin_required
def admin_lpr_watchlist_toggle(entry_id: int):
    conn = db()
    _lpr_watchlist_ensure(conn)
    current = conn.execute("SELECT active FROM lpr_watchlist WHERE id=?", (entry_id,)).fetchone()
    if current:
        conn.execute("UPDATE lpr_watchlist SET active=? WHERE id=?",
                     (0 if current["active"] else 1, entry_id))
        conn.commit()
    conn.close()
    return redirect(request.referrer or url_for("admin_lpr_watchlist"))


# ─────────────────────────────────────────────────────────────────────────────
# Stage 7: Push Notifications · Follow-up Dispatch · Proximity Rules
# ─────────────────────────────────────────────────────────────────────────────

import math as _math
import time as _time


# ── Table helpers ──────────────────────────────────────────────────────────────

def _lpr_device_tokens_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_device_tokens (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            platform TEXT DEFAULT 'ios',
            token TEXT NOT NULL,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            UNIQUE(user_id, token)
        )
    """)
    cur = conn.cursor()
    add_column_if_missing(cur, "lpr_device_tokens", "platform", "TEXT DEFAULT 'ios'")


def _lpr_notifications_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_notifications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            body TEXT NOT NULL,
            notification_type TEXT,
            data_json TEXT,
            read_at TEXT
        )
    """)


def _lpr_followups_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_followups (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT NOT NULL,
            created_by INTEGER NOT NULL,
            sighting_id INTEGER NOT NULL,
            matched_job_id INTEGER,
            assigned_user_id INTEGER,
            priority TEXT DEFAULT 'normal',
            action_type TEXT NOT NULL,
            status TEXT DEFAULT 'open',
            due_at TEXT,
            office_note TEXT
        )
    """)
    cur = conn.cursor()
    add_column_if_missing(cur, "lpr_followups", "status",       "TEXT DEFAULT 'open'")
    add_column_if_missing(cur, "lpr_followups", "assigned_at",  "TEXT")
    add_column_if_missing(cur, "lpr_followups", "en_route_at",  "TEXT")
    add_column_if_missing(cur, "lpr_followups", "arrived_at",   "TEXT")
    add_column_if_missing(cur, "lpr_followups", "completed_at", "TEXT")


def _agent_movement_ensure(conn):
    """Dedicated agent movement table — no customer/finance data, only position + context."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS agent_movement (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id       INTEGER NOT NULL,
            latitude      REAL NOT NULL,
            longitude     REAL NOT NULL,
            captured_at   TEXT NOT NULL,
            received_at   TEXT NOT NULL,
            source        TEXT,
            accuracy_m    REAL,
            battery_state TEXT,
            context       TEXT
        )
    """)
    conn.execute("""
        CREATE INDEX IF NOT EXISTS idx_agent_movement_user_time
        ON agent_movement (user_id, received_at DESC)
    """)
    conn.commit()


def _lpr_proximity_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_proximity_rules (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            latitude REAL NOT NULL,
            longitude REAL NOT NULL,
            radius_m REAL NOT NULL DEFAULT 500,
            active INTEGER DEFAULT 1,
            priority TEXT DEFAULT 'normal',
            created_by INTEGER,
            created_at TEXT NOT NULL
        )
    """)
    cur = conn.cursor()
    add_column_if_missing(cur, "lpr_proximity_rules", "priority", "TEXT DEFAULT 'normal'")


# ── Haversine distance + proximity check + dispatch intelligence ───────────────

def _haversine_m(lat1: float, lng1: float, lat2: float, lng2: float) -> float:
    R = 6_371_000.0
    phi1, phi2 = _math.radians(lat1), _math.radians(lat2)
    dphi = _math.radians(lat2 - lat1)
    dlam = _math.radians(lng2 - lng1)
    a = (_math.sin(dphi / 2) ** 2
         + _math.cos(phi1) * _math.cos(phi2) * _math.sin(dlam / 2) ** 2)
    return 2 * R * _math.asin(_math.sqrt(a))


def _proximity_check(lat: float, lng: float, conn) -> list:
    _lpr_proximity_ensure(conn)
    rules = conn.execute(
        "SELECT * FROM lpr_proximity_rules WHERE active=1"
    ).fetchall()
    triggered = []
    for rule in rules:
        dist = _haversine_m(lat, lng, rule["latitude"], rule["longitude"])
        if dist <= rule["radius_m"]:
            triggered.append({"rule": rule, "distance_m": round(dist)})
    return triggered


def _dist_label(metres: float) -> str:
    if metres < 1000:
        return f"{round(metres)} m"
    return f"{metres / 1000:.1f} km"


def _nearest_agents(lat: float, lng: float, conn, limit: int = 3,
                    max_hours: int = 8) -> list:
    """Return up to `limit` agents with a recent GPS ping, sorted by distance.
    Queries agent_movement (richer data) with fallback to agent_locations."""
    from datetime import timedelta as _td2
    cutoff = (datetime.utcnow() - _td2(hours=max_hours)).strftime("%Y-%m-%dT%H:%M:%S")

    rows = []
    try:
        rows = conn.execute("""
            SELECT am.user_id, u.full_name,
                   am.latitude AS lat, am.longitude AS lng,
                   am.received_at AS updated_at,
                   am.source, am.battery_state
            FROM agent_movement am
            JOIN users u ON u.id = am.user_id
            WHERE am.received_at >= ?
              AND am.received_at = (
                SELECT MAX(am2.received_at)
                FROM agent_movement am2
                WHERE am2.user_id = am.user_id
                  AND am2.received_at >= ?
              )
            ORDER BY am.received_at DESC
        """, (cutoff, cutoff)).fetchall()
    except Exception:
        pass

    if not rows:
        try:
            rows = conn.execute("""
                SELECT al.user_id, u.full_name, al.lat AS lat, al.lng AS lng,
                       al.updated_at, NULL AS source, NULL AS battery_state
                FROM agent_locations al
                JOIN users u ON u.id = al.user_id
                WHERE al.updated_at >= ? AND al.lat IS NOT NULL AND al.lng IS NOT NULL
                ORDER BY al.updated_at DESC
            """, (cutoff,)).fetchall()
        except Exception:
            return []

    agents = []
    seen_ids: set = set()
    for r in rows:
        uid = r["user_id"]
        if uid in seen_ids:
            continue
        seen_ids.add(uid)
        dist = _haversine_m(lat, lng, r["lat"], r["lng"])
        agents.append({
            "user_id":      uid,
            "name":         r["full_name"],
            "lat":          r["lat"],
            "lng":          r["lng"],
            "dist_m":       round(dist),
            "dist_label":   _dist_label(dist),
            "updated_at":   r["updated_at"],
            "source":       r["source"] or "",
            "battery":      r["battery_state"] or "",
        })
    agents.sort(key=lambda x: x["dist_m"])
    return agents[:limit]


def _lpr_repeat_info(reg_norm: str, lat, lng, conn,
                     exclude_id=None, days: int = 30,
                     radius_m: float = 1000.0) -> dict:
    """Repeat-sighting intelligence for a given normalised registration."""
    from datetime import timedelta as _td2
    cutoff = (datetime.utcnow() - _td2(days=days)).strftime("%Y-%m-%dT%H:%M:%S")
    q = """
        SELECT s.id, s.created_at, s.result_type, s.latitude, s.longitude,
               s.watchlist_hit, s.user_id, u.full_name AS agent_name
        FROM lpr_sightings s
        LEFT JOIN users u ON u.id = s.user_id
        WHERE s.registration_normalised=? AND s.created_at>=?
    """
    params = [reg_norm, cutoff]
    if exclude_id:
        q += " AND s.id != ?"
        params.append(exclude_id)
    q += " ORDER BY s.created_at DESC"
    rows = conn.execute(q, params).fetchall()

    total_count     = len(rows)
    agents_seen     = len(set(r["user_id"] for r in rows))
    watchlist_count = sum(1 for r in rows if r["watchlist_hit"])

    nearby = []
    if lat is not None and lng is not None:
        for r in rows:
            if r["latitude"] is not None and r["longitude"] is not None:
                d = _haversine_m(lat, lng, r["latitude"], r["longitude"])
                if d <= radius_m:
                    nearby.append({
                        "id":       r["id"],
                        "dist_m":   round(d),
                        "dist_label": _dist_label(d),
                        "date":     r["created_at"][:10] if r["created_at"] else "",
                        "agent":    r["agent_name"] or "",
                        "result_type": r["result_type"],
                    })
    nearby.sort(key=lambda x: x["dist_m"])

    return {
        "total_count":      total_count,
        "agents_count":     agents_seen,
        "nearby_count":     len(nearby),
        "nearby":           nearby[:5],
        "watchlist_count":  watchlist_count,
        "recent": [
            {
                "id":          r["id"],
                "date":        r["created_at"][:10] if r["created_at"] else "",
                "agent":       r["agent_name"] or "",
                "result_type": r["result_type"],
            }
            for r in rows[:5]
        ],
    }


_ACTION_LABELS_RESULT = {
    "allocated_match":  "Allocated",
    "restricted_match": "Restricted",
    "conflict":         "Conflict",
    "no_match":         "No Match",
}


def _lpr_dispatch_score(result_type: str, watchlist_h: int, escalated: int,
                        repeat_info: dict, nearest_agents: list,
                        prox_hits: list) -> dict:
    """
    Compute a priority band and recommended action from observable signals.
    Returns no customer/finance data — only operational intelligence.
    """
    score = 0
    if watchlist_h:            score += 40
    if prox_hits:              score += 20
    if result_type == "conflict":         score += 25
    elif result_type == "restricted_match": score += 15
    elif result_type == "allocated_match":  score += 10
    if escalated:              score += 15
    tc = repeat_info.get("total_count", 0)
    ac = repeat_info.get("agents_count", 0)
    nc = repeat_info.get("nearby_count", 0)
    if tc >= 5:   score += 20
    elif tc >= 2: score += 10
    if ac >= 2:   score += 10
    if nc >= 2:   score += 10

    if score >= 70:
        priority, color = "Urgent",  "danger"
    elif score >= 45:
        priority, color = "High",    "warning"
    elif score >= 20:
        priority, color = "Medium",  "info"
    else:
        priority, color = "Low",     "secondary"

    if watchlist_h and prox_hits:
        action = "Dispatch urgently — watchlist plate inside a proximity zone"
    elif watchlist_h and tc >= 2:
        action = "Dispatch — watchlist plate seen multiple times recently"
    elif watchlist_h:
        action = "Review and dispatch — watchlist hit"
    elif result_type == "conflict":
        action = "Investigate — multiple active files for this plate"
    elif result_type == "restricted_match" and nc >= 2:
        action = "Dispatch to area — plate repeatedly sighted nearby"
    elif result_type == "restricted_match":
        action = "Phone contact — restricted plate sighted"
    elif escalated and nearest_agents:
        action = f"Assign to {nearest_agents[0]['name']} — {nearest_agents[0]['dist_label']} from sighting"
    elif tc >= 3:
        action = "Investigate pattern — plate seen frequently in the area"
    else:
        action = "Monitor — no immediate dispatch required"

    return {"score": score, "priority": priority, "color": color, "action": action}


# ── Route / ETA / Dispatch intelligence ────────────────────────────────────────

def _eta_minutes(dist_m: float, source: str = "") -> float:
    """Rough road ETA in minutes from straight-line distance.
    Uses a 1.35× road-factor and assumes different speeds based on movement source."""
    speed_kph = 52.0 if source == "active_job" else 42.0
    speed_mps = speed_kph * 1000.0 / 3600.0
    return (dist_m * 1.35) / speed_mps / 60.0


def _eta_label(minutes: float) -> str:
    if minutes < 2:
        return "< 2 min"
    if minutes < 60:
        return f"~{round(minutes)} min"
    h = int(minutes // 60)
    m = int(minutes % 60)
    return f"~{h}h {m}m" if m else f"~{h}h"


def _agent_route_recommendation(lat: float, lng: float, conn,
                                 limit: int = 5) -> list:
    """Like _nearest_agents but sorted by estimated drive ETA, not straight-line distance.
    Returns source and battery alongside each agent — no customer data."""
    agents = _nearest_agents(lat, lng, conn, limit=limit * 2)
    for a in agents:
        eta = _eta_minutes(a["dist_m"], a.get("source", ""))
        a["eta_min"]   = round(eta, 1)
        a["eta_label"] = _eta_label(eta)
    agents.sort(key=lambda x: x["eta_min"])
    return agents[:limit]


def _diversion_score(agent_lat: float, agent_lng: float,
                     dest_lat: float, dest_lng: float,
                     sighting_lat: float, sighting_lng: float) -> dict:
    """Score the cost of diverting from an agent's current trajectory to a new sighting.
    Inputs and outputs contain no customer or finance data — only coordinates and labels."""
    direct_m  = _haversine_m(agent_lat, agent_lng, dest_lat, dest_lng)
    via_m     = (_haversine_m(agent_lat, agent_lng, sighting_lat, sighting_lng)
                 + _haversine_m(sighting_lat, sighting_lng, dest_lat, dest_lng))
    extra_m   = max(0.0, via_m - direct_m)
    extra_min = _eta_minutes(extra_m)
    worthwhile = extra_min <= 8.0

    if extra_min < 1.0:
        label = "On the way — no meaningful detour"
    elif extra_min <= 8.0:
        label = f"+{round(extra_min)} min detour — worth diverting"
    elif extra_min <= 20.0:
        label = f"+{round(extra_min)} min detour — consider a closer agent"
    else:
        label = f"+{round(extra_min)} min detour — not recommended"

    return {
        "extra_dist_m":  round(extra_m),
        "extra_eta_min": round(extra_min, 1),
        "worthwhile":    worthwhile,
        "label":         label,
    }


def _dispatch_sequence(sighting_ids: list, conn) -> list:
    """Greedy nearest-neighbour sequencing for attending multiple open sightings.
    No customer or finance data — only sighting IDs, coordinates, and priority."""
    if not sighting_ids:
        return []
    placeholders = ",".join("?" * len(sighting_ids))
    rows = conn.execute(f"""
        SELECT s.id, s.latitude, s.longitude,
               COALESCE(f.priority, 'normal') AS priority
        FROM lpr_sightings s
        LEFT JOIN lpr_followups f ON f.sighting_id = s.id
        WHERE s.id IN ({placeholders})
          AND s.latitude IS NOT NULL AND s.longitude IS NOT NULL
    """, sighting_ids).fetchall()
    if not rows:
        return []
    remaining = [dict(r) for r in rows]
    mean_lat  = sum(r["latitude"]  for r in remaining) / len(remaining)
    mean_lng  = sum(r["longitude"] for r in remaining) / len(remaining)
    ordered   = []
    cur_lat, cur_lng = mean_lat, mean_lng
    while remaining:
        closest = min(
            remaining,
            key=lambda r: _haversine_m(cur_lat, cur_lng, r["latitude"], r["longitude"])
        )
        remaining.remove(closest)
        ordered.append(closest)
        cur_lat, cur_lng = closest["latitude"], closest["longitude"]
    for i, item in enumerate(ordered):
        item["sequence"] = i + 1
    return ordered


def _dispatch_geofences_ensure(conn):
    """Temporary geofence records for high-priority dispatch assignments.
    Triggered when the agent's device enters the monitored region (near_target)."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS dispatch_geofences (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            followup_id  INTEGER NOT NULL,
            latitude     REAL NOT NULL,
            longitude    REAL NOT NULL,
            radius_m     REAL NOT NULL DEFAULT 150,
            created_at   TEXT NOT NULL,
            expires_at   TEXT,
            triggered    INTEGER DEFAULT 0,
            triggered_at TEXT
        )
    """)
    conn.commit()


_FOLLOWUP_VALID_TRANSITIONS: dict = {
    "open":        {"assigned", "en_route", "cancelled"},
    "assigned":    {"en_route", "cancelled"},
    "en_route":    {"near_target", "arrived", "completed", "cancelled"},
    "near_target": {"arrived", "completed", "cancelled"},
    "arrived":     {"completed", "cancelled"},
    "completed":   set(),
    "cancelled":   set(),
}


# ── APNs push delivery ─────────────────────────────────────────────────────────

def _apns_send(device_token: str, title: str, body: str, data: dict = None) -> bool:
    try:
        import jwt as _jwt
        import httpx as _httpx
    except ImportError:
        return False

    key_id      = os.environ.get("APNS_KEY_ID")
    team_id     = os.environ.get("APNS_TEAM_ID")
    bundle_id   = os.environ.get("APNS_BUNDLE_ID", "com.axionx.ios")
    private_key = os.environ.get("APNS_PRIVATE_KEY", "").replace("\\n", "\n")
    is_sandbox  = os.environ.get("APNS_SANDBOX", "1") == "1"

    if not all([key_id, team_id, private_key]):
        return False

    try:
        apns_token = _jwt.encode(
            {"iss": team_id, "iat": int(_time.time())},
            private_key,
            algorithm="ES256",
            headers={"kid": key_id},
        )
        host    = "api.sandbox.push.apple.com" if is_sandbox else "api.push.apple.com"
        payload = {"aps": {"alert": {"title": title, "body": body}, "sound": "default"}}
        if data:
            payload.update(data)
        with _httpx.Client(http2=True, timeout=10) as client:
            resp = client.post(
                f"https://{host}/3/device/{device_token}",
                json=payload,
                headers={
                    "authorization": f"bearer {apns_token}",
                    "apns-topic":    bundle_id,
                    "apns-push-type": "alert",
                    "apns-priority":  "10",
                },
            )
        return resp.status_code == 200
    except Exception:
        return False


# ── In-app notification helpers ────────────────────────────────────────────────

def _notify_user(user_id: int, title: str, body: str, notif_type: str,
                 data: dict = None):
    try:
        conn = db()
        _lpr_notifications_ensure(conn)
        _lpr_device_tokens_ensure(conn)
        conn.execute("""
            INSERT INTO lpr_notifications
                (created_at, user_id, title, body, notification_type, data_json)
            VALUES (?,?,?,?,?,?)
        """, (now_ts(), user_id, title, body, notif_type,
              json.dumps(data) if data else None))
        conn.commit()
        tokens = conn.execute(
            "SELECT token FROM lpr_device_tokens WHERE user_id=?", (user_id,)
        ).fetchall()
        conn.close()
        for tok in tokens:
            _apns_send(tok["token"], title, body, data)
    except Exception:
        pass


def _notify_admins(title: str, body: str, notif_type: str, data: dict = None,
                   exclude_uid: int = None):
    try:
        conn = db()
        q    = "SELECT id FROM users WHERE role IN ('admin','both') AND active=1"
        args = []
        if exclude_uid:
            q += " AND id != ?"
            args.append(exclude_uid)
        admins = conn.execute(q, args).fetchall()
        conn.close()
        for admin in admins:
            _notify_user(admin["id"], title, body, notif_type, data)
    except Exception:
        pass


# ── Device token registration ──────────────────────────────────────────────────

@app.post("/m/api/device/register")
@mobile_login_required
def m_api_device_register():
    uid      = session.get("user_id")
    data     = request.get_json(silent=True) or {}
    token    = (data.get("token") or "").strip()
    platform = (data.get("platform") or "ios").strip()
    if not token:
        return jsonify({"ok": False, "error": "Missing token"}), 400
    conn = db()
    _lpr_device_tokens_ensure(conn)
    conn.execute("""
        INSERT INTO lpr_device_tokens (user_id, platform, token, created_at, updated_at)
        VALUES (?,?,?,?,?)
        ON CONFLICT(user_id, token) DO UPDATE SET updated_at=excluded.updated_at
    """, (uid, platform, token, now_ts(), now_ts()))
    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


# ── Agent notification feed (mobile) ──────────────────────────────────────────

@app.get("/m/lpr/notifications")
@mobile_login_required
def m_lpr_notifications():
    uid  = session.get("user_id")
    conn = db()
    _lpr_notifications_ensure(conn)
    rows = conn.execute("""
        SELECT * FROM lpr_notifications
        WHERE user_id=?
        ORDER BY created_at DESC
        LIMIT 60
    """, (uid,)).fetchall()
    unread = conn.execute("""
        SELECT COUNT(*) AS n FROM lpr_notifications
        WHERE user_id=? AND read_at IS NULL
    """, (uid,)).fetchone()["n"]
    conn.execute(
        "UPDATE lpr_notifications SET read_at=? WHERE user_id=? AND read_at IS NULL",
        (now_ts(), uid)
    )
    conn.commit()
    conn.close()
    return render_template("mobile/lpr_notifications.html", rows=rows, unread=unread)


@app.post("/m/api/lpr/notifications/read")
@mobile_login_required
def m_api_lpr_notifications_read():
    uid  = session.get("user_id")
    conn = db()
    _lpr_notifications_ensure(conn)
    conn.execute(
        "UPDATE lpr_notifications SET read_at=? WHERE user_id=? AND read_at IS NULL",
        (now_ts(), uid)
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


@app.get("/m/api/lpr/notifications/unread-count")
@mobile_login_required
def m_api_lpr_notifications_count():
    uid  = session.get("user_id")
    conn = db()
    _lpr_notifications_ensure(conn)
    n = conn.execute(
        "SELECT COUNT(*) AS n FROM lpr_notifications WHERE user_id=? AND read_at IS NULL",
        (uid,)
    ).fetchone()["n"]
    conn.close()
    return jsonify({"count": n}), 200


# ── Proximity rules (admin) ────────────────────────────────────────────────────

@app.get("/admin/lpr-proximity")
@admin_required
def admin_lpr_proximity():
    conn     = db()
    _lpr_proximity_ensure(conn)
    f_active = request.args.get("active", "1")
    where    = (
        "WHERE p.active=1" if f_active == "1"
        else ("WHERE p.active=0" if f_active == "0" else "")
    )
    rows = conn.execute(f"""
        SELECT p.*, u.full_name AS creator_name
        FROM lpr_proximity_rules p
        LEFT JOIN users u ON u.id = p.created_by
        {where}
        ORDER BY p.created_at DESC
    """).fetchall()
    conn.close()
    return render_template("lpr_proximity.html", rows=rows, f_active=f_active)


@app.post("/admin/lpr-proximity/add")
@admin_required
def admin_lpr_proximity_add():
    uid  = session.get("user_id")
    name = (request.form.get("name") or "").strip()
    try:
        lat = float(request.form.get("latitude", ""))
        lng = float(request.form.get("longitude", ""))
        rad = float(request.form.get("radius_m", "500"))
    except (ValueError, TypeError):
        flash("Invalid coordinates or radius.", "danger")
        return redirect(url_for("admin_lpr_proximity"))
    priority = (request.form.get("priority") or "normal").strip()
    if not name:
        flash("Zone name is required.", "danger")
        return redirect(url_for("admin_lpr_proximity"))
    conn = db()
    _lpr_proximity_ensure(conn)
    conn.execute("""
        INSERT INTO lpr_proximity_rules
            (name, latitude, longitude, radius_m, active, priority, created_by, created_at)
        VALUES (?,?,?,?,1,?,?,?)
    """, (name, lat, lng, rad, priority, uid, now_ts()))
    conn.commit()
    conn.close()
    flash(f"Proximity zone \"{name}\" created.", "success")
    return redirect(url_for("admin_lpr_proximity"))


@app.post("/admin/lpr-proximity/<int:rule_id>/toggle")
@admin_required
def admin_lpr_proximity_toggle(rule_id: int):
    conn    = db()
    _lpr_proximity_ensure(conn)
    current = conn.execute(
        "SELECT active FROM lpr_proximity_rules WHERE id=?", (rule_id,)
    ).fetchone()
    if current:
        conn.execute(
            "UPDATE lpr_proximity_rules SET active=? WHERE id=?",
            (0 if current["active"] else 1, rule_id)
        )
        conn.commit()
    conn.close()
    return redirect(request.referrer or url_for("admin_lpr_proximity"))


# ── Follow-up dispatch (admin) ─────────────────────────────────────────────────

_FOLLOWUP_LABELS = {
    "investigate":   "Investigate Vehicle",
    "phone_contact": "Phone Contact Required",
    "re_attendance": "Re-Attendance Required",
    "dispatch":      "Dispatch to Location",
    "close":         "Close — No Action",
}


@app.post("/admin/lpr-sightings/<int:sighting_id>/followup")
@admin_required
def admin_lpr_sighting_followup(sighting_id: int):
    creator_id       = session.get("user_id")
    action_type      = (request.form.get("action_type") or "").strip()
    assigned_user_id = request.form.get("assigned_user_id") or None
    priority         = (request.form.get("priority") or "normal").strip()
    due_at           = (request.form.get("due_at") or "").strip() or None
    office_note      = (request.form.get("office_note") or "").strip() or None
    if not action_type:
        flash("Action type is required.", "danger")
        return redirect(request.referrer or url_for("admin_lpr_sightings"))

    conn = db()
    _lpr_sightings_ensure_table(conn)
    _lpr_followups_ensure(conn)
    sighting = conn.execute(
        "SELECT user_id, matched_job_id, latitude, longitude FROM lpr_sightings WHERE id=?",
        (sighting_id,)
    ).fetchone()
    if not sighting:
        conn.close()
        flash("Sighting not found.", "danger")
        return redirect(url_for("admin_lpr_sightings"))

    assigned_uid = int(assigned_user_id) if assigned_user_id else None
    status       = "assigned" if assigned_uid else "open"
    assigned_at  = now_ts() if assigned_uid else None
    ts           = now_ts()
    conn.execute("""
        INSERT INTO lpr_followups
            (created_at, created_by, sighting_id, matched_job_id,
             assigned_user_id, priority, action_type, status, due_at, office_note,
             assigned_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,?)
    """, (ts, creator_id, sighting_id, sighting["matched_job_id"],
          assigned_uid, priority, action_type, status, due_at, office_note, assigned_at))
    followup_id = conn.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]

    # Create a temporary dispatch geofence for urgent/high priority dispatches with GPS
    if (assigned_uid and priority in ("urgent", "high")
            and sighting["latitude"] and sighting["longitude"]):
        try:
            _dispatch_geofences_ensure(conn)
            from datetime import timedelta as _td3
            expires = (datetime.utcnow() + _td3(hours=12)).strftime("%Y-%m-%dT%H:%M:%S")
            conn.execute("""
                INSERT INTO dispatch_geofences
                    (followup_id, latitude, longitude, radius_m, created_at, expires_at)
                VALUES (?,?,?,150,?,?)
            """, (followup_id, sighting["latitude"], sighting["longitude"], ts, expires))
        except Exception:
            pass

    conn.commit()
    conn.close()
    _log_audit(creator_id, "followup_create", "lpr_sighting", sighting_id)

    if assigned_uid and assigned_uid != creator_id:
        label = _FOLLOWUP_LABELS.get(action_type, action_type.replace("_", " ").title())
        _notify_user(
            assigned_uid,
            "LPR Follow-up Assigned",
            f"Action required: {label}",
            "followup_assigned",
            {"followup_id": followup_id, "sighting_id": sighting_id,
             "action_type": action_type, "priority": priority},
        )

    flash("Follow-up created.", "success")
    return redirect(request.referrer or url_for("admin_lpr_sightings"))


@app.get("/admin/lpr-sightings/<int:sighting_id>/intelligence")
@admin_required
def admin_lpr_sighting_intelligence(sighting_id: int):
    import json as _json
    conn = db()
    _lpr_sightings_ensure_table(conn)
    s = conn.execute("""
        SELECT s.id, s.registration_normalised, s.result_type,
               s.watchlist_hit, s.escalated_to_office,
               s.latitude, s.longitude, s.search_method, s.created_at,
               u.full_name AS agent_name
        FROM lpr_sightings s
        LEFT JOIN users u ON u.id = s.user_id
        WHERE s.id=?
    """, (sighting_id,)).fetchone()
    if not s:
        conn.close()
        return jsonify({"error": "Not found"}), 404

    lat = s["latitude"]
    lng = s["longitude"]

    nearest   = _nearest_agents(lat, lng, conn) if (lat and lng) else []
    repeat    = _lpr_repeat_info(s["registration_normalised"], lat, lng,
                                  conn, exclude_id=sighting_id)
    prox_hits = []
    if lat and lng and s["watchlist_hit"]:
        prox_hits = _proximity_check(lat, lng, conn)

    score = _lpr_dispatch_score(
        s["result_type"], s["watchlist_hit"], s["escalated_to_office"],
        repeat, nearest, prox_hits
    )

    prox_clean = [{"zone": p["rule"]["zone_name"], "dist_label": _dist_label(p["distance_m"])}
                  for p in prox_hits]

    route_rec = _agent_route_recommendation(lat, lng, conn) if (lat and lng) else []

    conn.close()
    return jsonify({
        "sighting_id":          sighting_id,
        "nearest_agents":       nearest,
        "route_recommendation": route_rec,
        "repeat_info":          repeat,
        "proximity_hits":       prox_clean,
        "dispatch_score":       score,
    }), 200


@app.get("/admin/lpr-sightings/<int:sighting_id>/diversion")
@admin_required
def admin_lpr_sighting_diversion(sighting_id: int):
    """
    Compute the detour cost for diverting an en-route agent to a new sighting.
    Requires: agent_id (int), dest_sighting_id (int — their current destination).
    Returns no customer or finance data — only coordinates, ETA labels, and the score.
    """
    agent_id        = request.args.get("agent_id", type=int)
    dest_sighting_id = request.args.get("dest_sighting_id", type=int)

    if not agent_id or not dest_sighting_id:
        return jsonify({"error": "agent_id and dest_sighting_id required"}), 400

    conn = db()
    _lpr_sightings_ensure_table(conn)
    _agent_movement_ensure(conn)

    # Agent's last known position
    agent_pos = conn.execute("""
        SELECT latitude, longitude, source
        FROM agent_movement
        WHERE user_id = ?
        ORDER BY received_at DESC
        LIMIT 1
    """, (agent_id,)).fetchone()
    if not agent_pos:
        conn.close()
        return jsonify({"error": "No recent position for this agent"}), 404

    # Agent's current destination
    dest = conn.execute(
        "SELECT latitude, longitude FROM lpr_sightings WHERE id=?",
        (dest_sighting_id,)
    ).fetchone()

    # New sighting
    new_s = conn.execute(
        "SELECT latitude, longitude FROM lpr_sightings WHERE id=?",
        (sighting_id,)
    ).fetchone()
    conn.close()

    if not dest or not dest["latitude"]:
        return jsonify({"error": "Current destination sighting has no GPS"}), 404
    if not new_s or not new_s["latitude"]:
        return jsonify({"error": "New sighting has no GPS"}), 404

    result = _diversion_score(
        agent_lat=agent_pos["latitude"],  agent_lng=agent_pos["longitude"],
        dest_lat=dest["latitude"],        dest_lng=dest["longitude"],
        sighting_lat=new_s["latitude"],   sighting_lng=new_s["longitude"],
    )
    result["agent_id"]         = agent_id
    result["dest_sighting_id"] = dest_sighting_id
    result["new_sighting_id"]  = sighting_id
    return jsonify(result), 200


@app.get("/admin/lpr/agent-map")
@admin_required
def admin_lpr_agent_map():
    """
    Admin view: recent agent positions for dispatch support.
    Shows last-known position and a short trail (last 10 pings, 8 h window) per agent.
    No customer/finance data — only agent name, position, context, battery, and timestamp.
    """
    import json as _json
    conn = db()
    _agent_movement_ensure(conn)

    # Latest ping per agent (last 8 h)
    latest_rows = conn.execute("""
        SELECT am.user_id, u.full_name,
               am.latitude, am.longitude,
               am.received_at, am.source, am.battery_state, am.context
        FROM agent_movement am
        JOIN users u ON u.id = am.user_id
        WHERE am.received_at >= datetime('now', '-8 hours')
          AND am.received_at = (
              SELECT MAX(am2.received_at)
              FROM agent_movement am2
              WHERE am2.user_id = am.user_id
                AND am2.received_at >= datetime('now', '-8 hours')
          )
        ORDER BY am.received_at DESC
    """).fetchall()

    # Trail rows (last 10 pings per agent, 8 h window)
    trail_rows = conn.execute("""
        SELECT am.user_id, am.latitude, am.longitude, am.received_at,
               am.source, am.battery_state
        FROM agent_movement am
        WHERE am.received_at >= datetime('now', '-8 hours')
        ORDER BY am.user_id, am.received_at DESC
    """).fetchall()
    conn.close()

    # Build trails dict: user_id → last 10 pings
    trails: dict = {}
    for r in trail_rows:
        uid = r["user_id"]
        if uid not in trails:
            trails[uid] = []
        if len(trails[uid]) < 10:
            trails[uid].append({
                "lat":  r["latitude"],
                "lng":  r["longitude"],
                "at":   r["received_at"][:16].replace("T", " ") if r["received_at"] else "",
                "src":  r["source"] or "",
            })

    agents = []
    for r in latest_rows:
        uid  = r["user_id"]
        at   = r["received_at"]
        agents.append({
            "name":    r["full_name"],
            "lat":     r["latitude"],
            "lng":     r["longitude"],
            "at":      at[:16].replace("T", " ") if at else "",
            "source":  r["source"] or "",
            "battery": r["battery_state"] or "unknown",
            "context": r["context"] or "",
            "trail":   trails.get(uid, []),
        })

    agents_json = _json.dumps(agents)
    return render_template("lpr_agent_map.html",
                           agents_json=agents_json,
                           count=len(agents))


# ─────────────────────────────────────────────────────────────────────────────
# Stage 12 — Patrol Intelligence
# ─────────────────────────────────────────────────────────────────────────────

import math as _math_pi
from collections import defaultdict as _defaultdict

def _patrol_intelligence_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_patrol_intelligence (
            id                          INTEGER PRIMARY KEY AUTOINCREMENT,
            registration_normalised     TEXT    NOT NULL UNIQUE,
            matched_job_id              INTEGER,
            repeat_count_30d            INTEGER NOT NULL DEFAULT 0,
            distinct_agent_count        INTEGER NOT NULL DEFAULT 0,
            likely_zone                 TEXT,
            likely_day_bucket           TEXT,
            likely_time_window          TEXT,
            confidence_score            INTEGER NOT NULL DEFAULT 0,
            recommended_patrol_priority TEXT    DEFAULT 'low',
            recommended_action          TEXT,
            explanation                 TEXT,
            watchlist_hit               INTEGER DEFAULT 0,
            result_type                 TEXT,
            last_computed_at            TEXT    NOT NULL
        )
    """)
    conn.commit()


def _recompute_patrol_intelligence(conn, registration_filter=None):
    """
    Scan lpr_sightings for the last 30 days and build/update patrol
    intelligence records.  No customer or finance data is emitted —
    only plate, coordinates, time-pattern signals, and recommended action.
    """
    from datetime import timedelta as _td_pi
    cutoff = (datetime.utcnow() - _td_pi(days=30)).strftime("%Y-%m-%dT%H:%M:%S")

    q = """
        SELECT s.registration_normalised,
               s.latitude, s.longitude,
               s.created_at, s.result_type,
               s.user_id, s.watchlist_hit,
               s.matched_job_id
        FROM lpr_sightings s
        WHERE s.created_at >= ?
          AND s.registration_normalised IS NOT NULL
          AND s.registration_normalised != ''
    """
    params = [cutoff]
    if registration_filter:
        q += " AND s.registration_normalised = ?"
        params.append(registration_filter)
    q += " ORDER BY s.registration_normalised, s.created_at DESC"

    rows = conn.execute(q, params).fetchall()

    groups = _defaultdict(list)
    for r in rows:
        groups[r["registration_normalised"]].append(dict(r))

    now = now_ts()

    for reg, sightings in groups.items():
        if len(sightings) < 2:
            continue

        repeat_count     = len(sightings)
        agent_ids        = {s["user_id"] for s in sightings if s["user_id"]}
        distinct_agents  = len(agent_ids)
        watchlist_hit    = any(s["watchlist_hit"] for s in sightings)
        result_type      = sightings[0]["result_type"] or "no_match"

        matched_job_id = None
        for s in sightings:
            if s.get("matched_job_id"):
                matched_job_id = s["matched_job_id"]
                break

        gps = [s for s in sightings if s["latitude"] and s["longitude"]]
        likely_zone_json = None
        best_count       = 0

        if gps:
            best_center = None
            for s in gps:
                cnt = sum(
                    1 for t in gps
                    if _haversine_m(s["latitude"], s["longitude"],
                                    t["latitude"], t["longitude"]) <= 2000
                )
                if cnt > best_count:
                    best_count  = cnt
                    best_center = s
            zone_lat = best_center["latitude"]  if best_center else gps[0]["latitude"]
            zone_lng = best_center["longitude"] if best_center else gps[0]["longitude"]
            likely_zone_json = json.dumps({
                "lat":           zone_lat,
                "lng":           zone_lng,
                "cluster_count": best_count,
                "total_gps":     len(gps),
            })

        day_counts = {"weekday": 0, "weekend": 0}
        for s in sightings:
            try:
                dt = datetime.fromisoformat(s["created_at"][:19])
                key = "weekday" if dt.weekday() < 5 else "weekend"
                day_counts[key] += 1
            except Exception:
                pass

        total_days = day_counts["weekday"] + day_counts["weekend"]
        if total_days == 0:
            day_bucket = "unknown"
        elif day_counts["weekday"] / total_days >= 0.75:
            day_bucket = "weekday"
        elif day_counts["weekend"] / total_days >= 0.75:
            day_bucket = "weekend"
        else:
            day_bucket = "both"

        time_buckets = {"morning": 0, "afternoon": 0, "evening": 0, "night": 0}
        for s in sightings:
            try:
                dt = datetime.fromisoformat(s["created_at"][:19])
                h  = dt.hour
                if   6 <= h < 12: time_buckets["morning"]   += 1
                elif 12 <= h < 18: time_buckets["afternoon"] += 1
                elif 18 <= h < 22: time_buckets["evening"]   += 1
                else:              time_buckets["night"]      += 1
            except Exception:
                pass

        total_time = sum(time_buckets.values())
        dominant   = max(time_buckets, key=lambda k: time_buckets[k])
        if total_time > 0 and time_buckets[dominant] / total_time >= 0.55:
            time_window = dominant
        else:
            time_window = "mixed"

        conf    = 0
        factors = []

        if repeat_count >= 10:
            conf += 55; factors.append(f"Seen {repeat_count} times in 30 days")
        elif repeat_count >= 5:
            conf += 40; factors.append(f"Seen {repeat_count} times in 30 days")
        elif repeat_count >= 3:
            conf += 25; factors.append(f"Seen {repeat_count} times in 30 days")
        else:
            conf += 15; factors.append(f"Seen {repeat_count} times in 30 days")

        if watchlist_hit:
            conf += 20; factors.append("On the watchlist")

        if distinct_agents >= 2:
            conf += 10; factors.append(f"Confirmed by {distinct_agents} different agents")

        if gps and best_count and len(gps) > 0:
            zone_pct = best_count / len(gps)
            if zone_pct >= 0.6:
                conf += 15
                factors.append(f"Consistent area — {round(zone_pct * 100)}% of sightings cluster together")

        if time_window != "mixed" and total_time > 0:
            tw_pct = time_buckets[time_window] / total_time
            if tw_pct >= 0.55:
                conf += 15
                factors.append(f"Typically seen in the {time_window}")

        if result_type == "allocated_match":
            conf += 10; factors.append("Linked to an active job")
        elif result_type == "conflict":
            conf += 15; factors.append("Multiple file conflict")
        elif result_type == "restricted_match":
            conf += 5; factors.append("Restricted plate")

        conf = min(conf, 100)

        if conf >= 75 or (watchlist_hit and repeat_count >= 3):
            priority = "urgent"
        elif conf >= 50 or (watchlist_hit and repeat_count >= 2):
            priority = "high"
        elif conf >= 30:
            priority = "medium"
        else:
            priority = "low"

        day_label  = {"weekday": "weekdays", "weekend": "weekends",
                      "both": "all days", "unknown": "various days"}.get(day_bucket, "")
        time_label = {"morning":   "mornings (6 am – noon)",
                      "afternoon": "afternoons (noon – 6 pm)",
                      "evening":   "evenings (6 – 10 pm)",
                      "night":     "nights (10 pm – 6 am)",
                      "mixed":     "at varying times"}.get(time_window, "")

        if priority == "urgent" and watchlist_hit:
            action = (f"Patrol area — watchlist plate seen {repeat_count}\u00d7 recently, "
                      f"typically on {day_label} {time_label}").strip()
        elif priority == "urgent":
            action = (f"High-frequency plate — deploy patrol {day_label} {time_label} "
                      "in cluster area").strip()
        elif priority == "high" and watchlist_hit:
            action = f"Target patrol — watchlist plate seen {repeat_count}\u00d7 on {day_label}".strip()
        elif priority == "high":
            action = (f"Patrol opportunity — plate seen {repeat_count}\u00d7 "
                      f"on {day_label} {time_label}").strip()
        elif priority == "medium":
            action = (f"Monitor area — plate seen {repeat_count} times; "
                      f"patrol recommended on {day_label}").strip()
        else:
            action = "Low confidence — include in general patrol coverage"

        conn.execute("""
            INSERT INTO lpr_patrol_intelligence
                (registration_normalised, matched_job_id, repeat_count_30d,
                 distinct_agent_count, likely_zone, likely_day_bucket,
                 likely_time_window, confidence_score,
                 recommended_patrol_priority, recommended_action,
                 explanation, watchlist_hit, result_type, last_computed_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(registration_normalised) DO UPDATE SET
                matched_job_id              = excluded.matched_job_id,
                repeat_count_30d            = excluded.repeat_count_30d,
                distinct_agent_count        = excluded.distinct_agent_count,
                likely_zone                 = excluded.likely_zone,
                likely_day_bucket           = excluded.likely_day_bucket,
                likely_time_window          = excluded.likely_time_window,
                confidence_score            = excluded.confidence_score,
                recommended_patrol_priority = excluded.recommended_patrol_priority,
                recommended_action          = excluded.recommended_action,
                explanation                 = excluded.explanation,
                watchlist_hit               = excluded.watchlist_hit,
                result_type                 = excluded.result_type,
                last_computed_at            = excluded.last_computed_at
        """, (reg, matched_job_id, repeat_count, distinct_agents,
              likely_zone_json, day_bucket, time_window,
              conf, priority, action,
              json.dumps(factors), 1 if watchlist_hit else 0,
              result_type, now))
    conn.commit()


# ─────────────────────────────────────────────────────────────────────────────
# Stage 13 — ML-assisted prediction refinement
# ─────────────────────────────────────────────────────────────────────────────

def _lpr_prediction_scores_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_prediction_scores (
            id                      INTEGER PRIMARY KEY AUTOINCREMENT,
            registration_normalised TEXT    NOT NULL UNIQUE,
            matched_job_id          INTEGER,
            rule_confidence_score   INTEGER NOT NULL DEFAULT 0,
            ml_confidence_score     INTEGER,
            combined_score          INTEGER NOT NULL DEFAULT 0,
            prediction_window       TEXT    DEFAULT '72h',
            model_version           TEXT    DEFAULT 'unscored',
            last_scored_at          TEXT    NOT NULL
        )
    """)
    add_column_if_missing(conn, "lpr_prediction_scores", "blend_rule_weight", "REAL")
    add_column_if_missing(conn, "lpr_prediction_scores", "blend_ml_weight",   "REAL")
    add_column_if_missing(conn, "lpr_prediction_scores", "ranking_config_id", "INTEGER")
    add_column_if_missing(conn, "lpr_prediction_scores", "experiment_id",     "INTEGER")
    add_column_if_missing(conn, "lpr_prediction_scores", "experiment_arm",    "TEXT")
    conn.commit()


# ── Adaptive ranking config ─────────────────────────────────────────────────────

def _lpr_ranking_config_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_ranking_config (
            id                     INTEGER PRIMARY KEY AUTOINCREMENT,
            model_version          TEXT    NOT NULL DEFAULT '*',
            prediction_window      TEXT    NOT NULL DEFAULT '72h',
            confidence_band        TEXT    NOT NULL DEFAULT '*',
            priority_band          TEXT    NOT NULL DEFAULT '*',
            rule_weight            REAL    NOT NULL DEFAULT 0.40,
            ml_weight              REAL    NOT NULL DEFAULT 0.60,
            min_combined_threshold INTEGER NOT NULL DEFAULT 30,
            active                 INTEGER NOT NULL DEFAULT 1,
            source                 TEXT    NOT NULL DEFAULT 'default',
            reason                 TEXT,
            effective_from         TEXT    NOT NULL,
            created_by             INTEGER,
            created_at             TEXT    NOT NULL
        )
    """)
    n = conn.execute("SELECT COUNT(*) FROM lpr_ranking_config").fetchone()[0]
    if n == 0:
        now = now_ts()
        conn.execute("""
            INSERT INTO lpr_ranking_config
                (model_version, prediction_window, confidence_band, priority_band,
                 rule_weight, ml_weight, min_combined_threshold, active, source,
                 reason, effective_from, created_at)
            VALUES ('*','72h','*','*',0.40,0.60,30,1,'default',
                    'System default — 40/60 blend',?,?)
        """, (now, now))
    conn.commit()


def _get_active_ranking_config(conn, model_version, conf_band, pri_band):
    """
    Lookup active blend config with fallback hierarchy:
      1. Exact: model_version=x  conf_band=y  pri_band=z
      2. Model+conf: model_version=x  conf_band=y  pri_band='*'
      3. Model only: model_version=x  conf_band='*'  pri_band='*'
      4. Global default: model_version='*'  conf_band='*'  pri_band='*'
    Returns (rule_weight, ml_weight, min_combined_threshold, config_id).
    """
    candidates = [
        (model_version, conf_band, pri_band),
        (model_version, conf_band, "*"),
        (model_version, "*",       "*"),
        ("*",           "*",       "*"),
    ]
    for mv, cb, pb in candidates:
        row = conn.execute("""
            SELECT id, rule_weight, ml_weight, min_combined_threshold
            FROM lpr_ranking_config
            WHERE active=1
              AND model_version=? AND confidence_band=? AND priority_band=?
            ORDER BY effective_from DESC
            LIMIT 1
        """, (mv, cb, pb)).fetchone()
        if row:
            return (row["rule_weight"], row["ml_weight"],
                    row["min_combined_threshold"], row["id"])
    return (0.40, 0.60, 30, None)


def _score_to_band(score):
    if score is None:
        return "low"
    if score >= 75:
        return "urgent"
    if score >= 50:
        return "high"
    if score >= 30:
        return "medium"
    return "low"


def _run_recalibration(conn, min_sample=25):
    """
    Analyse lpr_prediction_outcomes grouped by model version + confidence band.
    Returns a list of recommendation dicts. Does NOT persist changes.
    """
    _lpr_ranking_config_ensure(conn)
    _lpr_prediction_outcomes_ensure(conn)

    POS = ("'confirmed_present','repeat_area_confirmed',"
           "'recovery_progressed','recovery_completed','followup_required'")

    rows = conn.execute(f"""
        SELECT
            COALESCE(model_version,'unscored') AS model_version,
            CASE
                WHEN combined_score >= 75 THEN 'urgent'
                WHEN combined_score >= 50 THEN 'high'
                WHEN combined_score >= 30 THEN 'medium'
                ELSE 'low'
            END AS conf_band,
            COUNT(*) AS total,
            SUM(CASE WHEN actual_outcome IN ({POS}) THEN 1 ELSE 0 END) AS positive_count,
            SUM(CASE WHEN actual_outcome = 'false_positive' THEN 1 ELSE 0 END) AS fp_count,
            SUM(CASE WHEN actual_outcome = 'no_locate'      THEN 1 ELSE 0 END) AS nl_count
        FROM lpr_prediction_outcomes
        WHERE combined_score IS NOT NULL
        GROUP BY model_version, conf_band
        ORDER BY model_version, conf_band
    """).fetchall()

    recommendations = []
    for r in rows:
        mv   = r["model_version"]
        band = r["conf_band"]
        n    = r["total"]
        pos_n = r["positive_count"]
        fp_n  = r["fp_count"]
        pos_rate = pos_n / n if n else 0.0
        fp_rate  = fp_n  / n if n else 0.0

        cur_rw, cur_mw, cur_thresh, cur_id = _get_active_ranking_config(conn, mv, band, "*")

        if n < min_sample:
            recommendations.append({
                "model_version":     mv,
                "conf_band":         band,
                "total":             n,
                "min_sample":        min_sample,
                "pos_rate":          round(pos_rate * 100, 1),
                "fp_rate":           round(fp_rate  * 100, 1),
                "current_rule_weight": cur_rw,
                "current_ml_weight":   cur_mw,
                "current_threshold":   cur_thresh,
                "action":            "insufficient_sample",
                "reason":            f"Only {n} labelled outcomes — {min_sample} required before any adjustment",
                "new_rule_weight":   None,
                "new_ml_weight":     None,
                "new_threshold":     None,
            })
            continue

        new_mw    = cur_mw
        new_rw    = cur_rw
        new_thresh = cur_thresh
        action    = "no_change"
        reason    = "Performance within acceptable bounds — no adjustment needed"

        if pos_rate >= 0.40 and fp_rate < 0.20:
            if cur_mw < 0.80:
                new_mw = round(min(0.80, cur_mw + 0.10), 2)
                new_rw = round(1.0 - new_mw, 2)
                action = "increase_ml"
                reason = (f"Positive rate {round(pos_rate*100,1)}% with acceptable FP rate "
                          f"{round(fp_rate*100,1)}% — increase ML weight +10%")
        elif fp_rate >= 0.30 or (fp_rate >= 0.20 and pos_rate < 0.30):
            if cur_mw > 0.00:
                new_mw     = round(max(0.00, cur_mw - 0.10), 2)
                new_rw     = round(1.0 - new_mw, 2)
                new_thresh = min(75, cur_thresh + 5)
                action     = "decrease_ml"
                reason     = (f"FP rate {round(fp_rate*100,1)}% or low positive rate "
                              f"{round(pos_rate*100,1)}% — reduce ML weight and raise surfacing threshold")

        recommendations.append({
            "model_version":       mv,
            "conf_band":           band,
            "total":               n,
            "min_sample":          min_sample,
            "pos_rate":            round(pos_rate * 100, 1),
            "fp_rate":             round(fp_rate  * 100, 1),
            "current_rule_weight": cur_rw,
            "current_ml_weight":   cur_mw,
            "current_threshold":   cur_thresh,
            "action":              action,
            "reason":              reason,
            "new_rule_weight":     new_rw    if action not in ("no_change", "insufficient_sample") else None,
            "new_ml_weight":       new_mw    if action not in ("no_change", "insufficient_sample") else None,
            "new_threshold":       new_thresh if action not in ("no_change", "insufficient_sample") else None,
        })

    return recommendations


# ─────────────────────────────────────────────────────────────────────────────
# Stage 16 — Experiment / A-B evaluation helpers
# ─────────────────────────────────────────────────────────────────────────────

import hashlib as _hashlib


def _lpr_experiments_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_experiments (
            id                       INTEGER PRIMARY KEY AUTOINCREMENT,
            name                     TEXT    NOT NULL,
            status                   TEXT    NOT NULL DEFAULT 'active',
            experiment_type          TEXT    NOT NULL DEFAULT 'blend',
            champion_model_version   TEXT,
            challenger_model_version TEXT,
            champion_rule_weight     REAL    NOT NULL DEFAULT 0.40,
            champion_ml_weight       REAL    NOT NULL DEFAULT 0.60,
            challenger_rule_weight   REAL    NOT NULL DEFAULT 0.30,
            challenger_ml_weight     REAL    NOT NULL DEFAULT 0.70,
            champion_threshold       INTEGER NOT NULL DEFAULT 30,
            challenger_threshold     INTEGER NOT NULL DEFAULT 35,
            traffic_split_pct        INTEGER NOT NULL DEFAULT 20,
            start_at                 TEXT    NOT NULL,
            end_at                   TEXT,
            created_by               INTEGER,
            created_at               TEXT    NOT NULL,
            notes_safe               TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_experiment_assignments (
            id                      INTEGER PRIMARY KEY AUTOINCREMENT,
            experiment_id           INTEGER NOT NULL,
            registration_normalised TEXT    NOT NULL,
            arm                     TEXT    NOT NULL,
            assigned_at             TEXT    NOT NULL,
            UNIQUE(experiment_id, registration_normalised)
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_experiment_results (
            id                      INTEGER PRIMARY KEY AUTOINCREMENT,
            experiment_id           INTEGER NOT NULL,
            outcome_id              INTEGER,
            registration_normalised TEXT    NOT NULL,
            arm                     TEXT    NOT NULL,
            model_version           TEXT,
            rule_weight             REAL,
            ml_weight               REAL,
            threshold               INTEGER,
            combined_score_at_time  INTEGER,
            surfaced                INTEGER NOT NULL DEFAULT 1,
            actual_outcome          TEXT,
            recorded_at             TEXT    NOT NULL
        )
    """)
    conn.commit()


def _get_active_experiment(conn):
    """Return the single active experiment row, or None."""
    return conn.execute("""
        SELECT * FROM lpr_experiments
        WHERE status = 'active'
        ORDER BY start_at DESC
        LIMIT 1
    """).fetchone()


def _get_or_assign_experiment_arm(conn, reg, exp):
    """
    Return the stable arm ('champion' or 'challenger') for this plate in this
    experiment.  Uses a deterministic hash so the same plate always gets the
    same arm; the first call persists the assignment.
    """
    existing = conn.execute("""
        SELECT arm FROM lpr_experiment_assignments
        WHERE experiment_id=? AND registration_normalised=?
    """, (exp["id"], reg)).fetchone()
    if existing:
        return existing["arm"]

    h   = int(_hashlib.md5(f"{exp['id']}:{reg}".encode()).hexdigest()[:4], 16)
    arm = "challenger" if (h % 100) < exp["traffic_split_pct"] else "champion"

    conn.execute("""
        INSERT OR IGNORE INTO lpr_experiment_assignments
            (experiment_id, registration_normalised, arm, assigned_at)
        VALUES (?,?,?,?)
    """, (exp["id"], reg, arm, now_ts()))
    return arm


def _record_experiment_result(conn, exp_id, outcome_id, reg, arm,
                               model_ver, rw, mw, threshold,
                               combined_score, actual_outcome):
    """Insert one row into lpr_experiment_results."""
    conn.execute("""
        INSERT INTO lpr_experiment_results
            (experiment_id, outcome_id, registration_normalised, arm,
             model_version, rule_weight, ml_weight, threshold,
             combined_score_at_time, surfaced, actual_outcome, recorded_at)
        VALUES (?,?,?,?,?,?,?,?,?,1,?,?)
    """, (exp_id, outcome_id, reg, arm, model_ver, rw, mw, threshold,
          combined_score, actual_outcome, now_ts()))


def _recompute_combined_patrol_scores(conn):
    """
    For every plate in lpr_patrol_intelligence, blend rule + ML using:
      1. The active experiment arm config if the plate is enrolled in an
         experiment, or
      2. The adaptive ranking config from lpr_ranking_config, or
      3. A 40/60 default fallback.
    Stamps blend weights, ranking_config_id, experiment_id, experiment_arm.
    """
    _lpr_prediction_scores_ensure(conn)
    _lpr_ranking_config_ensure(conn)
    _lpr_experiments_ensure(conn)

    active_exp = _get_active_experiment(conn)

    pi_rows = conn.execute("""
        SELECT p.registration_normalised, p.confidence_score, p.matched_job_id,
               p.recommended_patrol_priority
        FROM lpr_patrol_intelligence p
    """).fetchall()

    now = now_ts()
    for pi in pi_rows:
        reg        = pi["registration_normalised"]
        rule_score = pi["confidence_score"]
        pri_band   = pi["recommended_patrol_priority"] or "low"

        ps = conn.execute(
            "SELECT ml_confidence_score, model_version, prediction_window "
            "FROM lpr_prediction_scores WHERE registration_normalised=?", (reg,)
        ).fetchone()

        ml_score  = ps["ml_confidence_score"] if ps else None
        model_ver = ps["model_version"]        if ps else "unscored"
        pred_win  = ps["prediction_window"]    if ps else "72h"

        exp_id  = None
        exp_arm = None
        cfg_id  = None

        if active_exp:
            arm = _get_or_assign_experiment_arm(conn, reg, active_exp)
            exp_id  = active_exp["id"]
            exp_arm = arm
            if arm == "challenger":
                rw = active_exp["challenger_rule_weight"]
                mw = active_exp["challenger_ml_weight"]
            else:
                rw = active_exp["champion_rule_weight"]
                mw = active_exp["champion_ml_weight"]
        else:
            conf_band = _score_to_band(rule_score)
            rw, mw, _threshold, cfg_id = _get_active_ranking_config(
                conn, model_ver, conf_band, pri_band)

        if ml_score is not None:
            combined = round(rw * rule_score + mw * ml_score)
        else:
            combined  = rule_score
            rw, mw    = 1.0, 0.0

        conn.execute("""
            INSERT INTO lpr_prediction_scores
                (registration_normalised, matched_job_id, rule_confidence_score,
                 ml_confidence_score, combined_score, prediction_window,
                 model_version, last_scored_at,
                 blend_rule_weight, blend_ml_weight, ranking_config_id,
                 experiment_id, experiment_arm)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(registration_normalised) DO UPDATE SET
                matched_job_id        = excluded.matched_job_id,
                rule_confidence_score = excluded.rule_confidence_score,
                combined_score        = excluded.combined_score,
                last_scored_at        = excluded.last_scored_at,
                blend_rule_weight     = excluded.blend_rule_weight,
                blend_ml_weight       = excluded.blend_ml_weight,
                ranking_config_id     = excluded.ranking_config_id,
                experiment_id         = excluded.experiment_id,
                experiment_arm        = excluded.experiment_arm
        """, (reg, pi["matched_job_id"], rule_score, ml_score, combined,
              pred_win, model_ver, now, rw, mw, cfg_id, exp_id, exp_arm))
    conn.commit()


# ── Admin: training data export ────────────────────────────────────────────────

@app.get("/admin/lpr/patrol/export.csv")
@admin_required
def admin_lpr_patrol_export():
    """
    Export a clean tabular CSV for Create ML training.
    Features: only safe operational data (no customer/finance fields).
    Targets:
      seen_again_72h   — 1 if any two sightings within 72 h in the 30-day window
      outcome_label    — real field outcome (confirmed_present, no_locate, etc.)
                         empty string if no outcome has been recorded yet
      hours_to_outcome — hours between most-recent sighting and outcome recording;
                         empty if no outcome recorded
    """
    from datetime import timedelta as _td_exp
    import io as _io
    import csv as _csv

    cutoff = (datetime.utcnow() - _td_exp(days=30)).strftime("%Y-%m-%dT%H:%M:%S")
    conn   = db()
    _patrol_intelligence_ensure(conn)
    _lpr_prediction_scores_ensure(conn)
    _lpr_prediction_outcomes_ensure(conn)
    _lpr_sightings_ensure_table(conn)

    pi_rows = conn.execute("""
        SELECT p.registration_normalised,
               p.repeat_count_30d, p.distinct_agent_count,
               p.watchlist_hit, p.result_type,
               p.likely_day_bucket, p.likely_time_window,
               p.likely_zone, p.confidence_score
        FROM lpr_patrol_intelligence p
    """).fetchall()

    sighting_times = {}
    raw = conn.execute("""
        SELECT registration_normalised, created_at
        FROM lpr_sightings
        WHERE created_at >= ? AND registration_normalised IS NOT NULL
        ORDER BY registration_normalised, created_at
    """, (cutoff,)).fetchall()

    # Most-recent outcome per plate (newest recorded_at wins)
    outcome_rows = conn.execute("""
        SELECT o.registration_normalised, o.actual_outcome, o.recorded_at,
               MAX(s.created_at) AS last_sighting_at
        FROM lpr_prediction_outcomes o
        LEFT JOIN lpr_sightings s
               ON s.registration_normalised = o.registration_normalised
        GROUP BY o.registration_normalised
        HAVING o.recorded_at = MAX(o.recorded_at)
    """).fetchall()
    conn.close()

    for r in raw:
        reg = r["registration_normalised"]
        sighting_times.setdefault(reg, []).append(r["created_at"][:19])

    outcome_map = {}
    for r in outcome_rows:
        reg = r["registration_normalised"]
        hours_to = ""
        if r["last_sighting_at"] and r["recorded_at"]:
            try:
                t1 = datetime.fromisoformat(r["last_sighting_at"][:19])
                t2 = datetime.fromisoformat(r["recorded_at"][:19])
                hours_to = round((t2 - t1).total_seconds() / 3600, 1)
            except Exception:
                pass
        outcome_map[reg] = (r["actual_outcome"] or "", hours_to)

    output = _io.StringIO()
    writer = _csv.DictWriter(output, fieldnames=[
        "registration_normalised",
        "repeat_count_30d",
        "distinct_agent_count",
        "is_watchlist",
        "result_type_allocated",
        "result_type_restricted",
        "result_type_conflict",
        "day_bucket_weekday",
        "day_bucket_weekend",
        "day_bucket_both",
        "time_window_morning",
        "time_window_afternoon",
        "time_window_evening",
        "time_window_night",
        "has_gps_cluster",
        "rule_confidence_score",
        "seen_again_72h",
        "outcome_label",
        "hours_to_outcome",
    ])
    writer.writeheader()

    for pi in pi_rows:
        reg     = pi["registration_normalised"]
        rt      = pi["result_type"] or "no_match"
        db_val  = pi["likely_day_bucket"] or "unknown"
        tw_val  = pi["likely_time_window"] or "mixed"
        has_gps = 1 if pi["likely_zone"] else 0

        times   = sighting_times.get(reg, [])
        s72h    = 0
        if len(times) >= 2:
            parsed = []
            for t in times:
                try:
                    parsed.append(datetime.fromisoformat(t))
                except Exception:
                    pass
            parsed.sort()
            for i in range(len(parsed) - 1):
                diff_h = (parsed[i + 1] - parsed[i]).total_seconds() / 3600
                if diff_h <= 72:
                    s72h = 1
                    break

        outcome_label, hours_to = outcome_map.get(reg, ("", ""))

        writer.writerow({
            "registration_normalised": reg,
            "repeat_count_30d":        pi["repeat_count_30d"],
            "distinct_agent_count":    pi["distinct_agent_count"],
            "is_watchlist":            1 if pi["watchlist_hit"] else 0,
            "result_type_allocated":   1 if rt == "allocated_match"  else 0,
            "result_type_restricted":  1 if rt == "restricted_match" else 0,
            "result_type_conflict":    1 if rt == "conflict"         else 0,
            "day_bucket_weekday":      1 if db_val == "weekday"  else 0,
            "day_bucket_weekend":      1 if db_val == "weekend"  else 0,
            "day_bucket_both":         1 if db_val == "both"     else 0,
            "time_window_morning":     1 if tw_val == "morning"   else 0,
            "time_window_afternoon":   1 if tw_val == "afternoon" else 0,
            "time_window_evening":     1 if tw_val == "evening"   else 0,
            "time_window_night":       1 if tw_val == "night"     else 0,
            "has_gps_cluster":         has_gps,
            "rule_confidence_score":   pi["confidence_score"],
            "seen_again_72h":          s72h,
            "outcome_label":           outcome_label,
            "hours_to_outcome":        hours_to,
        })

    csv_bytes = output.getvalue().encode("utf-8")
    resp = make_response(csv_bytes)
    resp.headers["Content-Type"]        = "text/csv; charset=utf-8"
    resp.headers["Content-Disposition"] = "attachment; filename=lpr_patrol_training.csv"
    return resp


# ── Mobile: receive ML scores from iOS ────────────────────────────────────────

@app.post("/m/api/lpr/patrol/scores")
@mobile_login_required
def m_api_lpr_patrol_scores():
    """
    iOS app posts Core ML inference results after running the bundled model.
    Payload: {model_version, prediction_window, scores: [{registration, ml_score}]}
    Accepts only authenticated mobile sessions.  No customer data touched.
    """
    data          = request.get_json(silent=True) or {}
    model_version = (data.get("model_version") or "v1.0").strip()[:40]
    pred_window   = (data.get("prediction_window") or "72h").strip()[:10]
    scores        = data.get("scores") or []

    if not isinstance(scores, list):
        return jsonify({"ok": False, "error": "scores must be a list"}), 400

    conn = db()
    _lpr_prediction_scores_ensure(conn)
    now  = now_ts()
    n    = 0
    for item in scores[:200]:
        if not isinstance(item, dict):
            continue
        reg      = (item.get("registration") or "").strip().upper()
        ml_score = item.get("ml_score")
        if not reg or ml_score is None:
            continue
        try:
            ml_score = max(0, min(100, int(float(ml_score))))
        except (ValueError, TypeError):
            continue

        existing = conn.execute(
            "SELECT rule_confidence_score FROM lpr_prediction_scores "
            "WHERE registration_normalised=?", (reg,)
        ).fetchone()

        if existing:
            rule_score = existing["rule_confidence_score"]
        else:
            pi = conn.execute(
                "SELECT confidence_score, matched_job_id "
                "FROM lpr_patrol_intelligence WHERE registration_normalised=?", (reg,)
            ).fetchone()
            rule_score = pi["confidence_score"] if pi else 0

        combined = round(0.40 * rule_score + 0.60 * ml_score)

        conn.execute("""
            INSERT INTO lpr_prediction_scores
                (registration_normalised, rule_confidence_score, ml_confidence_score,
                 combined_score, prediction_window, model_version, last_scored_at)
            VALUES (?,?,?,?,?,?,?)
            ON CONFLICT(registration_normalised) DO UPDATE SET
                ml_confidence_score = excluded.ml_confidence_score,
                combined_score      = excluded.combined_score,
                prediction_window   = excluded.prediction_window,
                model_version       = excluded.model_version,
                last_scored_at      = excluded.last_scored_at
        """, (reg, rule_score, ml_score, combined, pred_window, model_version, now))
        n += 1

    conn.commit()
    conn.close()
    return jsonify({"ok": True, "updated": n}), 200


# ── Admin: patrol opportunities ────────────────────────────────────────────────

@app.get("/admin/lpr/patrol")
@admin_required
def admin_lpr_patrol():
    conn = db()
    _patrol_intelligence_ensure(conn)
    _lpr_prediction_scores_ensure(conn)

    f_priority   = request.args.get("priority",   "")
    f_watchlist  = request.args.get("watchlist",  "")
    f_min_conf   = request.args.get("min_conf",   "")
    f_result     = request.args.get("result",     "")
    f_day_bucket = request.args.get("day_bucket", "")

    conditions = []
    params     = []

    if f_priority:
        conditions.append("p.recommended_patrol_priority = ?")
        params.append(f_priority)
    if f_watchlist == "1":
        conditions.append("p.watchlist_hit = 1")
    if f_min_conf:
        try:
            conditions.append("COALESCE(ps.combined_score, p.confidence_score) >= ?")
            params.append(int(f_min_conf))
        except ValueError:
            pass
    if f_result:
        conditions.append("p.result_type = ?")
        params.append(f_result)
    if f_day_bucket:
        conditions.append("p.likely_day_bucket = ?")
        params.append(f_day_bucket)

    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""

    rows = conn.execute(f"""
        SELECT p.*,
               ps.ml_confidence_score,
               ps.combined_score,
               ps.model_version,
               ps.prediction_window,
               ps.last_scored_at AS ml_scored_at
        FROM lpr_patrol_intelligence p
        LEFT JOIN lpr_prediction_scores ps
               ON ps.registration_normalised = p.registration_normalised
        {where}
        ORDER BY
            CASE p.recommended_patrol_priority
                WHEN 'urgent' THEN 1 WHEN 'high' THEN 2
                WHEN 'medium' THEN 3 ELSE 4 END,
            COALESCE(ps.combined_score, p.confidence_score) DESC
        LIMIT 200
    """, params).fetchall()

    total = conn.execute("SELECT COUNT(*) AS n FROM lpr_patrol_intelligence").fetchone()["n"]
    last_run = conn.execute(
        "SELECT MAX(last_computed_at) AS t FROM lpr_patrol_intelligence"
    ).fetchone()["t"]
    ml_model_ver = conn.execute(
        "SELECT model_version FROM lpr_prediction_scores "
        "WHERE model_version IS NOT NULL AND model_version != 'unscored' "
        "ORDER BY last_scored_at DESC LIMIT 1"
    ).fetchone()
    conn.close()

    intel = []
    for r in rows:
        zone = None
        if r["likely_zone"]:
            try:
                zone = json.loads(r["likely_zone"])
            except Exception:
                pass
        factors = []
        if r["explanation"]:
            try:
                factors = json.loads(r["explanation"])
            except Exception:
                pass
        ml_score = r["ml_confidence_score"]
        combined = r["combined_score"] if r["combined_score"] is not None else r["confidence_score"]
        intel.append({
            "id":            r["id"],
            "reg":           r["registration_normalised"],
            "job_id":        r["matched_job_id"],
            "repeat":        r["repeat_count_30d"],
            "agents":        r["distinct_agent_count"],
            "zone":          zone,
            "day_bucket":    r["likely_day_bucket"] or "unknown",
            "time_window":   r["likely_time_window"] or "mixed",
            "rule_score":    r["confidence_score"],
            "ml_score":      ml_score,
            "combined":      combined,
            "model_version": r["model_version"] or "unscored",
            "pred_window":   r["prediction_window"] or "72h",
            "ml_scored_at":  r["ml_scored_at"],
            "priority":      r["recommended_patrol_priority"] or "low",
            "action":        r["recommended_action"] or "",
            "factors":       factors,
            "watchlist":     bool(r["watchlist_hit"]),
            "result_type":   r["result_type"] or "no_match",
            "computed_at":   r["last_computed_at"],
        })

    return render_template(
        "lpr_patrol.html",
        intel=intel, total=total, last_run=last_run,
        f_priority=f_priority, f_watchlist=f_watchlist,
        f_min_conf=f_min_conf, f_result=f_result, f_day_bucket=f_day_bucket,
        ml_model_ver=ml_model_ver["model_version"] if ml_model_ver else None,
        outcome_vocab=LPR_OUTCOME_VOCAB,
    )


@app.post("/admin/lpr/patrol/recompute")
@admin_required
def admin_lpr_patrol_recompute():
    conn = db()
    _patrol_intelligence_ensure(conn)
    _lpr_prediction_scores_ensure(conn)
    _lpr_sightings_ensure_table(conn)
    _recompute_patrol_intelligence(conn)
    _recompute_combined_patrol_scores(conn)
    conn.close()
    flash("Patrol intelligence and combined scores recomputed from the last 30 days.", "success")
    return redirect(url_for("admin_lpr_patrol"))


# ── Mobile: patrol list ────────────────────────────────────────────────────────

@app.get("/m/lpr/patrol")
@mobile_login_required
def m_lpr_patrol():
    conn = db()
    _patrol_intelligence_ensure(conn)
    _lpr_prediction_scores_ensure(conn)
    rows = conn.execute("""
        SELECT p.registration_normalised, p.repeat_count_30d, p.distinct_agent_count,
               p.likely_zone, p.likely_day_bucket, p.likely_time_window,
               p.confidence_score, p.recommended_patrol_priority,
               p.recommended_action, p.explanation, p.watchlist_hit, p.result_type,
               p.last_computed_at,
               ps.ml_confidence_score, ps.combined_score,
               ps.model_version, ps.prediction_window
        FROM lpr_patrol_intelligence p
        LEFT JOIN lpr_prediction_scores ps
               ON ps.registration_normalised = p.registration_normalised
        ORDER BY
            CASE p.recommended_patrol_priority
                WHEN 'urgent' THEN 1 WHEN 'high' THEN 2
                WHEN 'medium' THEN 3 ELSE 4 END,
            COALESCE(ps.combined_score, p.confidence_score) DESC
        LIMIT 100
    """).fetchall()
    conn.close()

    items = []
    for r in rows:
        zone = None
        if r["likely_zone"]:
            try:
                zone = json.loads(r["likely_zone"])
            except Exception:
                pass
        factors = []
        if r["explanation"]:
            try:
                factors = json.loads(r["explanation"])
            except Exception:
                pass
        ml_score = r["ml_confidence_score"]
        combined = r["combined_score"] if r["combined_score"] is not None else r["confidence_score"]
        items.append({
            "reg":           r["registration_normalised"],
            "repeat":        r["repeat_count_30d"],
            "agents":        r["distinct_agent_count"],
            "zone":          zone,
            "day_bucket":    r["likely_day_bucket"] or "unknown",
            "time_window":   r["likely_time_window"] or "mixed",
            "rule_score":    r["confidence_score"],
            "ml_score":      ml_score,
            "combined":      combined,
            "model_version": r["model_version"] or "unscored",
            "pred_window":   r["prediction_window"] or "72h",
            "priority":      r["recommended_patrol_priority"] or "low",
            "action":        r["recommended_action"] or "",
            "factors":       factors,
            "watchlist":     bool(r["watchlist_hit"]),
            "result_type":   r["result_type"] or "no_match",
        })
    return render_template("mobile/lpr_patrol.html", items=items,
                           outcome_vocab=LPR_OUTCOME_VOCAB)


@app.get("/m/api/lpr/patrol")
@mobile_login_required
def m_api_lpr_patrol():
    conn = db()
    _patrol_intelligence_ensure(conn)
    _lpr_prediction_scores_ensure(conn)
    rows = conn.execute("""
        SELECT p.registration_normalised, p.repeat_count_30d, p.distinct_agent_count,
               p.likely_zone, p.likely_day_bucket, p.likely_time_window,
               p.confidence_score, p.recommended_patrol_priority,
               p.recommended_action, p.watchlist_hit, p.result_type,
               p.last_computed_at,
               ps.ml_confidence_score, ps.combined_score,
               ps.model_version, ps.prediction_window
        FROM lpr_patrol_intelligence p
        LEFT JOIN lpr_prediction_scores ps
               ON ps.registration_normalised = p.registration_normalised
        ORDER BY
            CASE p.recommended_patrol_priority
                WHEN 'urgent' THEN 1 WHEN 'high' THEN 2
                WHEN 'medium' THEN 3 ELSE 4 END,
            COALESCE(ps.combined_score, p.confidence_score) DESC
        LIMIT 100
    """).fetchall()
    conn.close()

    items = []
    for r in rows:
        zone = None
        if r["likely_zone"]:
            try:
                zone = json.loads(r["likely_zone"])
            except Exception:
                pass
        ml_score = r["ml_confidence_score"]
        combined = r["combined_score"] if r["combined_score"] is not None else r["confidence_score"]
        items.append({
            "registration":    r["registration_normalised"],
            "repeat_count":    r["repeat_count_30d"],
            "agent_count":     r["distinct_agent_count"],
            "zone":            zone,
            "day_bucket":      r["likely_day_bucket"],
            "time_window":     r["likely_time_window"],
            "confidence":      r["confidence_score"],
            "ml_score":        ml_score,
            "combined_score":  combined,
            "model_version":   r["model_version"] or "unscored",
            "pred_window":     r["prediction_window"] or "72h",
            "priority":        r["recommended_patrol_priority"],
            "action":          r["recommended_action"],
            "watchlist_hit":   bool(r["watchlist_hit"]),
            "result_type":     r["result_type"],
            "computed_at":     r["last_computed_at"],
        })
    return jsonify({"count": len(items), "items": items}), 200


# ─────────────────────────────────────────────────────────────────────────────
# Stage 14 — Closed-loop learning: outcome capture, evaluation, training labels
# ─────────────────────────────────────────────────────────────────────────────

# ── Outcome vocabulary ────────────────────────────────────────────────────────
# Labels are operational and objective — no customer or finance data.

LPR_OUTCOME_VOCAB = [
    ("confirmed_present",    "Confirmed present",    "Plate physically located in predicted area"),
    ("repeat_area_confirmed","Repeat area confirmed", "Plate resighted in same cluster zone"),
    ("followup_required",    "Follow-up required",   "Flagged for active recovery action"),
    ("restricted_only",      "Restricted only",      "Access restricted; plate not recoverable at this time"),
    ("no_locate",            "No locate",            "Patrol conducted but plate not found"),
    ("false_positive",       "False positive",       "Prediction was incorrect for this plate"),
    ("recovery_progressed",  "Recovery progressed",  "Active recovery process initiated"),
    ("recovery_completed",   "Recovery completed",   "Successful recovery completed"),
]

LPR_POSITIVE_OUTCOMES = frozenset({
    "confirmed_present", "repeat_area_confirmed",
    "recovery_progressed", "recovery_completed", "followup_required",
})

_LPR_OUTCOME_CODES = frozenset(v[0] for v in LPR_OUTCOME_VOCAB)


def _lpr_prediction_outcomes_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_prediction_outcomes (
            id                      INTEGER PRIMARY KEY AUTOINCREMENT,
            registration_normalised TEXT    NOT NULL,
            matched_job_id          INTEGER,
            source_type             TEXT    NOT NULL DEFAULT 'patrol',
            source_id               INTEGER,
            rule_score              INTEGER,
            ml_score                INTEGER,
            combined_score          INTEGER,
            prediction_window       TEXT,
            model_version           TEXT,
            recommended_action      TEXT,
            actual_outcome          TEXT    NOT NULL,
            outcome_confidence      INTEGER DEFAULT 80,
            recorded_by             INTEGER,
            recorded_at             TEXT    NOT NULL,
            notes_safe              TEXT
        )
    """)
    conn.commit()


# ── Admin: record patrol outcome ───────────────────────────────────────────────

@app.post("/admin/lpr/patrol/outcome")
@admin_required
def admin_lpr_patrol_outcome():
    data    = request.get_json(silent=True) or {}
    reg     = (data.get("registration") or "").strip().upper()
    outcome = (data.get("outcome")      or "").strip()
    notes   = (data.get("notes")        or "").strip()[:500]
    try:
        outcome_conf = max(0, min(100, int(data.get("outcome_confidence", 80))))
    except (ValueError, TypeError):
        outcome_conf = 80

    if not reg or outcome not in _LPR_OUTCOME_CODES:
        return jsonify({"ok": False, "error": "Invalid registration or outcome"}), 400

    conn = db()
    _lpr_prediction_outcomes_ensure(conn)
    _patrol_intelligence_ensure(conn)
    _lpr_prediction_scores_ensure(conn)
    _lpr_experiments_ensure(conn)

    pi = conn.execute("""
        SELECT p.id, p.matched_job_id, p.confidence_score, p.recommended_action,
               ps.ml_confidence_score, ps.combined_score, ps.model_version,
               ps.prediction_window, ps.blend_rule_weight, ps.blend_ml_weight,
               ps.experiment_id, ps.experiment_arm
        FROM lpr_patrol_intelligence p
        LEFT JOIN lpr_prediction_scores ps
               ON ps.registration_normalised = p.registration_normalised
        WHERE p.registration_normalised = ?
    """, (reg,)).fetchone()

    now = now_ts()
    uid = session.get("user_id")

    rule_score = pi["confidence_score"]          if pi else None
    ml_score   = pi["ml_confidence_score"]       if pi else None
    combined   = (pi["combined_score"] if pi["combined_score"] is not None
                  else rule_score)               if pi else None
    job_id     = pi["matched_job_id"]            if pi else None
    model_ver  = (pi["model_version"] or "unscored") if pi else "unscored"
    pred_win   = (pi["prediction_window"] or "72h")  if pi else "72h"
    rec_action = (pi["recommended_action"] or "")    if pi else ""
    source_id  = pi["id"]                        if pi else None
    exp_id     = pi["experiment_id"]             if pi else None
    exp_arm    = pi["experiment_arm"]            if pi else None

    cur = conn.execute("""
        INSERT INTO lpr_prediction_outcomes
            (registration_normalised, matched_job_id, source_type, source_id,
             rule_score, ml_score, combined_score, prediction_window,
             model_version, recommended_action, actual_outcome,
             outcome_confidence, recorded_by, recorded_at, notes_safe)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (reg, job_id, "patrol", source_id,
          rule_score, ml_score, combined, pred_win,
          model_ver, rec_action, outcome,
          outcome_conf, uid, now, notes or None))
    outcome_row_id = cur.lastrowid

    if exp_id and exp_arm:
        rw  = pi["blend_rule_weight"]  if pi else None
        mw  = pi["blend_ml_weight"]    if pi else None
        exp_row = conn.execute(
            "SELECT champion_threshold, challenger_threshold FROM lpr_experiments WHERE id=?",
            (exp_id,)
        ).fetchone()
        thr = (exp_row["challenger_threshold"] if exp_arm == "challenger"
               else exp_row["champion_threshold"]) if exp_row else None
        _record_experiment_result(conn, exp_id, outcome_row_id, reg, exp_arm,
                                  model_ver, rw, mw, thr, combined, outcome)

    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


# ── Mobile: record patrol outcome ──────────────────────────────────────────────

@app.post("/m/api/lpr/patrol/outcome")
@mobile_login_required
def m_api_lpr_patrol_outcome():
    data    = request.get_json(silent=True) or {}
    reg     = (data.get("registration") or "").strip().upper()
    outcome = (data.get("outcome")      or "").strip()
    notes   = (data.get("notes")        or "").strip()[:200]

    if not reg or outcome not in _LPR_OUTCOME_CODES:
        return jsonify({"ok": False, "error": "Invalid registration or outcome"}), 400

    conn = db()
    _lpr_prediction_outcomes_ensure(conn)
    _patrol_intelligence_ensure(conn)
    _lpr_prediction_scores_ensure(conn)
    _lpr_experiments_ensure(conn)

    pi = conn.execute("""
        SELECT p.id, p.matched_job_id, p.confidence_score, p.recommended_action,
               ps.ml_confidence_score, ps.combined_score, ps.model_version,
               ps.prediction_window, ps.blend_rule_weight, ps.blend_ml_weight,
               ps.experiment_id, ps.experiment_arm
        FROM lpr_patrol_intelligence p
        LEFT JOIN lpr_prediction_scores ps
               ON ps.registration_normalised = p.registration_normalised
        WHERE p.registration_normalised = ?
    """, (reg,)).fetchone()

    now = now_ts()
    uid = session.get("user_id")

    rule_score = pi["confidence_score"]          if pi else None
    ml_score   = pi["ml_confidence_score"]       if pi else None
    combined   = (pi["combined_score"] if pi["combined_score"] is not None
                  else rule_score)               if pi else None
    job_id     = pi["matched_job_id"]            if pi else None
    model_ver  = (pi["model_version"] or "unscored") if pi else "unscored"
    pred_win   = (pi["prediction_window"] or "72h")  if pi else "72h"
    rec_action = (pi["recommended_action"] or "")    if pi else ""
    source_id  = pi["id"]                        if pi else None
    exp_id     = pi["experiment_id"]             if pi else None
    exp_arm    = pi["experiment_arm"]            if pi else None

    cur = conn.execute("""
        INSERT INTO lpr_prediction_outcomes
            (registration_normalised, matched_job_id, source_type, source_id,
             rule_score, ml_score, combined_score, prediction_window,
             model_version, recommended_action, actual_outcome,
             outcome_confidence, recorded_by, recorded_at, notes_safe)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (reg, job_id, "patrol_mobile", source_id,
          rule_score, ml_score, combined, pred_win,
          model_ver, rec_action, outcome,
          80, uid, now, notes or None))
    outcome_row_id = cur.lastrowid

    if exp_id and exp_arm:
        rw  = pi["blend_rule_weight"] if pi else None
        mw  = pi["blend_ml_weight"]   if pi else None
        exp_row = conn.execute(
            "SELECT champion_threshold, challenger_threshold FROM lpr_experiments WHERE id=?",
            (exp_id,)
        ).fetchone()
        thr = (exp_row["challenger_threshold"] if exp_arm == "challenger"
               else exp_row["champion_threshold"]) if exp_row else None
        _record_experiment_result(conn, exp_id, outcome_row_id, reg, exp_arm,
                                  model_ver, rw, mw, thr, combined, outcome)

    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


# ── Admin: evaluation dashboard ────────────────────────────────────────────────

@app.get("/admin/lpr/evaluation")
@admin_required
def admin_lpr_evaluation():
    conn = db()
    _lpr_prediction_outcomes_ensure(conn)
    _patrol_intelligence_ensure(conn)

    POS = ("'confirmed_present','repeat_area_confirmed',"
           "'recovery_progressed','recovery_completed','followup_required'")

    overall = conn.execute(f"""
        SELECT COUNT(*) AS total,
               SUM(CASE WHEN actual_outcome IN ({POS}) THEN 1 ELSE 0 END) AS positive_count,
               SUM(CASE WHEN actual_outcome = 'false_positive'  THEN 1 ELSE 0 END) AS fp_count,
               SUM(CASE WHEN actual_outcome = 'no_locate'       THEN 1 ELSE 0 END) AS nl_count
        FROM lpr_prediction_outcomes
    """).fetchone()

    by_model_rows = conn.execute(f"""
        SELECT model_version,
               COUNT(*) AS total,
               SUM(CASE WHEN actual_outcome IN ({POS}) THEN 1 ELSE 0 END) AS positive_count,
               SUM(CASE WHEN actual_outcome = 'false_positive'  THEN 1 ELSE 0 END) AS fp_count,
               ROUND(AVG(combined_score), 1) AS avg_combined
        FROM lpr_prediction_outcomes
        GROUP BY model_version
        ORDER BY MAX(recorded_at) DESC
    """).fetchall()

    by_band_rows = conn.execute(f"""
        SELECT
            CASE
                WHEN combined_score >= 75 THEN 'urgent'
                WHEN combined_score >= 50 THEN 'high'
                WHEN combined_score >= 30 THEN 'medium'
                ELSE 'low'
            END AS band,
            COUNT(*) AS total,
            SUM(CASE WHEN actual_outcome IN ({POS}) THEN 1 ELSE 0 END) AS positive_count,
            SUM(CASE WHEN actual_outcome = 'false_positive'  THEN 1 ELSE 0 END) AS fp_count
        FROM lpr_prediction_outcomes
        WHERE combined_score IS NOT NULL
        GROUP BY band
        ORDER BY MIN(combined_score) DESC
    """).fetchall()

    by_priority_rows = conn.execute(f"""
        SELECT p.recommended_patrol_priority AS priority,
               COUNT(*)  AS total,
               SUM(CASE WHEN o.actual_outcome IN ({POS}) THEN 1 ELSE 0 END) AS positive_count,
               SUM(CASE WHEN o.actual_outcome = 'false_positive'  THEN 1 ELSE 0 END) AS fp_count
        FROM lpr_prediction_outcomes o
        JOIN lpr_patrol_intelligence p
          ON p.registration_normalised = o.registration_normalised
        GROUP BY p.recommended_patrol_priority
        ORDER BY CASE p.recommended_patrol_priority
            WHEN 'urgent' THEN 1 WHEN 'high' THEN 2
            WHEN 'medium' THEN 3 ELSE 4 END
    """).fetchall()

    recent_rows = conn.execute("""
        SELECT o.id, o.registration_normalised, o.actual_outcome, o.source_type,
               o.combined_score, o.model_version, o.recorded_at, o.notes_safe,
               u.full_name AS recorder_name
        FROM lpr_prediction_outcomes o
        LEFT JOIN users u ON u.id = o.recorded_by
        ORDER BY o.recorded_at DESC
        LIMIT 30
    """).fetchall()

    # Outcome totals breakdown for bar chart
    outcome_counts = conn.execute("""
        SELECT actual_outcome, COUNT(*) AS n
        FROM lpr_prediction_outcomes
        GROUP BY actual_outcome
        ORDER BY n DESC
    """).fetchall()

    conn.close()

    def _pct(num, den):
        return round(100 * num / den, 1) if den else 0

    total = overall["total"] or 0

    by_model = [{
        "model_version":  r["model_version"] or "unscored",
        "total":          r["total"],
        "positive_count": r["positive_count"],
        "fp_count":       r["fp_count"],
        "avg_combined":   r["avg_combined"] or 0,
        "pos_rate":       _pct(r["positive_count"], r["total"]),
        "fp_rate":        _pct(r["fp_count"],       r["total"]),
    } for r in by_model_rows]

    band_order = {"urgent": 1, "high": 2, "medium": 3, "low": 4}
    by_band = sorted([{
        "band":           r["band"],
        "total":          r["total"],
        "positive_count": r["positive_count"],
        "fp_count":       r["fp_count"],
        "pos_rate":       _pct(r["positive_count"], r["total"]),
        "fp_rate":        _pct(r["fp_count"],       r["total"]),
    } for r in by_band_rows], key=lambda x: band_order.get(x["band"], 5))

    by_priority = [{
        "priority":       r["priority"] or "low",
        "total":          r["total"],
        "positive_count": r["positive_count"],
        "fp_count":       r["fp_count"],
        "pos_rate":       _pct(r["positive_count"], r["total"]),
        "fp_rate":        _pct(r["fp_count"],       r["total"]),
    } for r in by_priority_rows]

    recent = [{
        "id":          r["id"],
        "reg":         r["registration_normalised"],
        "outcome":     r["actual_outcome"],
        "source_type": r["source_type"],
        "combined":    r["combined_score"],
        "model_ver":   r["model_version"] or "unscored",
        "recorder":    r["recorder_name"] or "Unknown",
        "recorded_at": r["recorded_at"],
        "notes":       r["notes_safe"],
    } for r in recent_rows]

    outcome_dist = {r["actual_outcome"]: r["n"] for r in outcome_counts}

    return render_template(
        "lpr_evaluation.html",
        total=total,
        positive_count=overall["positive_count"] or 0,
        fp_count=overall["fp_count"]    or 0,
        nl_count=overall["nl_count"]    or 0,
        pos_rate=_pct(overall["positive_count"] or 0, total),
        fp_rate= _pct(overall["fp_count"]       or 0, total),
        nl_rate= _pct(overall["nl_count"]       or 0, total),
        by_model=by_model,
        by_band=by_band,
        by_priority=by_priority,
        recent=recent,
        outcome_dist=outcome_dist,
        outcome_vocab=LPR_OUTCOME_VOCAB,
    )


# ─────────────────────────────────────────────────────────────────────────────
# Stage 15 — Adaptive ranking controls
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/admin/lpr/ranking")
@admin_required
def admin_lpr_ranking():
    conn = db()
    _lpr_ranking_config_ensure(conn)
    _lpr_prediction_outcomes_ensure(conn)

    # All active configs (most specific first)
    active_configs = conn.execute("""
        SELECT id, model_version, prediction_window, confidence_band, priority_band,
               rule_weight, ml_weight, min_combined_threshold, source, reason,
               effective_from, created_at, created_by
        FROM lpr_ranking_config
        WHERE active=1
        ORDER BY
            (CASE WHEN model_version='*'   THEN 1 ELSE 0 END),
            (CASE WHEN confidence_band='*' THEN 1 ELSE 0 END),
            (CASE WHEN priority_band='*'   THEN 1 ELSE 0 END),
            effective_from DESC
    """).fetchall()

    # Config change history (last 40)
    history = conn.execute("""
        SELECT id, model_version, confidence_band, priority_band,
               rule_weight, ml_weight, min_combined_threshold,
               active, source, reason, effective_from, created_at,
               created_by
        FROM lpr_ranking_config
        ORDER BY created_at DESC
        LIMIT 40
    """).fetchall()

    # Count outcomes available for analysis
    outcome_total = conn.execute(
        "SELECT COUNT(*) FROM lpr_prediction_outcomes"
    ).fetchone()[0]

    conn.close()

    # Compute recommendations (read-only; no side effects)
    conn2 = db()
    recommendations = _run_recalibration(conn2, min_sample=25)
    conn2.close()

    return render_template(
        "lpr_ranking.html",
        active_configs=[dict(r) for r in active_configs],
        history=[dict(r) for r in history],
        recommendations=recommendations,
        outcome_total=outcome_total,
    )


@app.post("/admin/lpr/ranking/calibrate")
@admin_required
def admin_lpr_ranking_calibrate():
    """Auto-apply all actionable recommendations (increase_ml / decrease_ml)."""
    conn = db()
    _lpr_ranking_config_ensure(conn)
    recommendations = _run_recalibration(conn, min_sample=25)

    applied = 0
    now = now_ts()
    uid = session.get("user_id")

    for rec in recommendations:
        if rec["action"] not in ("increase_ml", "decrease_ml"):
            continue
        mv   = rec["model_version"]
        band = rec["conf_band"]
        rw   = rec["new_rule_weight"]
        mw   = rec["new_ml_weight"]
        thr  = rec["new_threshold"]
        reason = rec["reason"]

        # Deactivate any existing configs for this model+band combination
        conn.execute("""
            UPDATE lpr_ranking_config SET active=0
            WHERE model_version=? AND confidence_band=? AND priority_band='*'
        """, (mv, band))

        conn.execute("""
            INSERT INTO lpr_ranking_config
                (model_version, prediction_window, confidence_band, priority_band,
                 rule_weight, ml_weight, min_combined_threshold, active, source,
                 reason, effective_from, created_by, created_at)
            VALUES (?,?,?,?,?,?,?,1,'auto',?,?,?,?)
        """, (mv, "72h", band, "*", rw, mw, thr, reason, now, uid, now))
        applied += 1

    conn.commit()
    conn.close()
    return jsonify({"ok": True, "applied": applied}), 200


@app.post("/admin/lpr/ranking/config")
@admin_required
def admin_lpr_ranking_config():
    """Manual override: create a new active config row."""
    data = request.get_json(silent=True) or {}
    mv   = (data.get("model_version")   or "*").strip()
    band = (data.get("confidence_band") or "*").strip()
    pb   = (data.get("priority_band")   or "*").strip()

    valid_bands = {"*", "urgent", "high", "medium", "low"}
    if band not in valid_bands or pb not in valid_bands:
        return jsonify({"ok": False, "error": "Invalid band value"}), 400

    try:
        rw  = round(max(0.0, min(1.0, float(data.get("rule_weight", 0.40)))), 2)
        mw  = round(max(0.0, min(1.0, float(data.get("ml_weight",   0.60)))), 2)
        thr = max(0, min(100, int(data.get("min_combined_threshold", 30))))
    except (ValueError, TypeError):
        return jsonify({"ok": False, "error": "Invalid numeric values"}), 400

    if abs(rw + mw - 1.0) > 0.01:
        return jsonify({"ok": False, "error": "rule_weight + ml_weight must equal 1.0"}), 400

    reason = (data.get("reason") or "Manual admin override").strip()[:300]
    now = now_ts()
    uid = session.get("user_id")

    conn = db()
    _lpr_ranking_config_ensure(conn)

    # Deactivate previous configs for same scope
    conn.execute("""
        UPDATE lpr_ranking_config SET active=0
        WHERE model_version=? AND confidence_band=? AND priority_band=?
    """, (mv, band, pb))

    conn.execute("""
        INSERT INTO lpr_ranking_config
            (model_version, prediction_window, confidence_band, priority_band,
             rule_weight, ml_weight, min_combined_threshold, active, source,
             reason, effective_from, created_by, created_at)
        VALUES (?,?,?,?,?,?,?,1,'manual',?,?,?,?)
    """, (mv, "72h", band, pb, rw, mw, thr, reason, now, uid, now))
    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


@app.post("/admin/lpr/ranking/rollback")
@admin_required
def admin_lpr_ranking_rollback():
    """Rollback: deactivate the most recent non-default config for a given scope."""
    data = request.get_json(silent=True) or {}
    config_id = data.get("config_id")
    if not config_id:
        return jsonify({"ok": False, "error": "config_id required"}), 400

    conn = db()
    _lpr_ranking_config_ensure(conn)

    row = conn.execute(
        "SELECT * FROM lpr_ranking_config WHERE id=?", (config_id,)
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "Config not found"}), 404
    if row["source"] == "default":
        conn.close()
        return jsonify({"ok": False, "error": "Cannot roll back the system default config"}), 400

    conn.execute("UPDATE lpr_ranking_config SET active=0 WHERE id=?", (config_id,))

    # Re-activate previous config for the same scope (if any)
    prev = conn.execute("""
        SELECT id FROM lpr_ranking_config
        WHERE model_version=? AND confidence_band=? AND priority_band=?
          AND id < ? AND id != ?
        ORDER BY id DESC
        LIMIT 1
    """, (row["model_version"], row["confidence_band"], row["priority_band"],
          config_id, config_id)).fetchone()
    if prev:
        conn.execute("UPDATE lpr_ranking_config SET active=1 WHERE id=?", (prev["id"],))

    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


# ─────────────────────────────────────────────────────────────────────────────
# Stage 17 — Policy engine helpers
# ─────────────────────────────────────────────────────────────────────────────

import json as _json


def _lpr_policy_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_policy_rules (
            id                       INTEGER PRIMARY KEY AUTOINCREMENT,
            name                     TEXT    NOT NULL,
            active                   INTEGER NOT NULL DEFAULT 1,
            scope                    TEXT    NOT NULL DEFAULT 'all',
            min_sample_per_arm       INTEGER NOT NULL DEFAULT 50,
            min_positive_lift        REAL    NOT NULL DEFAULT 5.0,
            max_fp_delta             REAL    NOT NULL DEFAULT 3.0,
            max_no_locate_delta      REAL    NOT NULL DEFAULT 5.0,
            protect_urgent_watchlist INTEGER NOT NULL DEFAULT 1,
            recommended_action       TEXT    NOT NULL DEFAULT 'promote',
            priority                 INTEGER NOT NULL DEFAULT 1,
            created_by               INTEGER,
            created_at               TEXT    NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_policy_decisions (
            id                 INTEGER PRIMARY KEY AUTOINCREMENT,
            experiment_id      INTEGER NOT NULL,
            rule_id            INTEGER,
            recommended_action TEXT    NOT NULL,
            reason_text        TEXT    NOT NULL,
            metrics_snapshot   TEXT    NOT NULL,
            decision_status    TEXT    NOT NULL DEFAULT 'pending',
            decided_by         INTEGER,
            decided_at         TEXT,
            applied_config_id  INTEGER,
            notes_safe         TEXT,
            created_at         TEXT    NOT NULL
        )
    """)
    # Seed default rules if table is empty
    n = conn.execute("SELECT COUNT(*) FROM lpr_policy_rules").fetchone()[0]
    if n == 0:
        now = now_ts()
        rows = [
            ("Promote — primary threshold", 1, "all",
             50, 5.0, 3.0, 5.0, 1, "promote", 1, now),
            ("Stop — high FP or no-locate drift", 1, "all",
             30, -99.0, 8.0, 8.0, 0, "stop", 2, now),
            ("Tighten threshold — moderate drift", 1, "all",
             30, -99.0, 3.0, 5.0, 0, "tighten", 3, now),
        ]
        for r in rows:
            conn.execute("""
                INSERT INTO lpr_policy_rules
                    (name, active, scope, min_sample_per_arm, min_positive_lift,
                     max_fp_delta, max_no_locate_delta, protect_urgent_watchlist,
                     recommended_action, priority, created_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?)
            """, r)
    conn.commit()


def _exp_arm_stats_protected(conn, exp_id, arm):
    """
    Outcome stats restricted to urgent/high-priority or watchlisted plates only.
    Used by the policy engine safeguard check.
    """
    POS = ("'confirmed_present','repeat_area_confirmed',"
           "'recovery_progressed','recovery_completed','followup_required'")
    row = conn.execute(f"""
        SELECT COUNT(*)  AS total,
               SUM(CASE WHEN r.actual_outcome IN ({POS})         THEN 1 ELSE 0 END) AS pos_count,
               SUM(CASE WHEN r.actual_outcome = 'false_positive' THEN 1 ELSE 0 END) AS fp_count
        FROM lpr_experiment_results r
        JOIN lpr_patrol_intelligence p
          ON p.registration_normalised = r.registration_normalised
        WHERE r.experiment_id=? AND r.arm=?
          AND (p.watchlist_hit=1
               OR p.recommended_patrol_priority IN ('urgent','high'))
    """, (exp_id, arm)).fetchone()
    total = row["total"] or 0
    def pct(n): return round(100 * n / total, 1) if total else 0.0
    return {
        "total":     total,
        "pos_count": row["pos_count"] or 0,
        "fp_count":  row["fp_count"]  or 0,
        "pos_rate":  pct(row["pos_count"] or 0),
        "fp_rate":   pct(row["fp_count"]  or 0),
    }


def _evaluate_experiment_policy(conn, exp):
    """
    Run all active policy rules against a single experiment dict.
    Returns a structured evaluation dict — pure read, no side effects.
    """
    _lpr_policy_ensure(conn)

    exp_id = exp["id"]
    cs = _exp_arm_stats(conn, exp_id, "champion")
    hs = _exp_arm_stats(conn, exp_id, "challenger")

    pos_lift = round(hs["pos_rate"] - cs["pos_rate"], 1)
    fp_delta = round(hs["fp_rate"]  - cs["fp_rate"],  1)
    nl_delta = round(hs["nl_rate"]  - cs["nl_rate"],  1)
    min_n    = min(cs["total"], hs["total"])

    prot_cs = _exp_arm_stats_protected(conn, exp_id, "champion")
    prot_hs = _exp_arm_stats_protected(conn, exp_id, "challenger")

    metrics = {
        "champion":           cs,
        "challenger":         hs,
        "pos_lift":           pos_lift,
        "fp_delta":           fp_delta,
        "nl_delta":           nl_delta,
        "min_n":              min_n,
        "protected_champion": prot_cs,
        "protected_challenger": prot_hs,
    }

    rules = conn.execute("""
        SELECT * FROM lpr_policy_rules WHERE active=1 ORDER BY priority ASC
    """).fetchall()

    pending_decision = conn.execute("""
        SELECT id, recommended_action, decision_status, reason_text
        FROM lpr_policy_decisions
        WHERE experiment_id=? AND decision_status='pending'
        ORDER BY id DESC LIMIT 1
    """, (exp_id,)).fetchone()

    last_decision = conn.execute("""
        SELECT id, recommended_action, decision_status, decided_at, notes_safe
        FROM lpr_policy_decisions
        WHERE experiment_id=?
        ORDER BY id DESC LIMIT 1
    """, (exp_id,)).fetchone()

    fired_rules   = []
    recommendation = None
    can_promote   = False
    can_stop      = False

    for rule in rules:
        action  = rule["recommended_action"]
        n_ok    = min_n >= rule["min_sample_per_arm"]
        lift_ok = pos_lift >= rule["min_positive_lift"]
        fp_ok   = fp_delta <= rule["max_fp_delta"]
        nl_ok   = nl_delta <= rule["max_no_locate_delta"]

        if action == "promote":
            urgent_ok     = True
            urgent_reasons = []
            if rule["protect_urgent_watchlist"]:
                p_n_ok = prot_cs["total"] >= 10 and prot_hs["total"] >= 5
                if p_n_ok:
                    p_lift = prot_hs["pos_rate"] - prot_cs["pos_rate"]
                    p_fp   = prot_hs["fp_rate"]  - prot_cs["fp_rate"]
                    if p_fp > rule["max_fp_delta"]:
                        urgent_ok = False
                        urgent_reasons.append(
                            f"Urgent/watchlist FP Δ {p_fp:+.1f}pp exceeds +{rule['max_fp_delta']:.1f}pp")
                    if p_lift < -5.0:
                        urgent_ok = False
                        urgent_reasons.append(
                            f"Urgent/watchlist positive rate dropped {p_lift:.1f}pp")

            fired   = n_ok and lift_ok and fp_ok and nl_ok and urgent_ok
            reasons = []
            if not n_ok:    reasons.append(
                f"Need ≥{rule['min_sample_per_arm']} outcomes per arm (have {min_n})")
            if not lift_ok: reasons.append(
                f"Positive lift {pos_lift:+.1f}pp < required {rule['min_positive_lift']:+.1f}pp")
            if not fp_ok:   reasons.append(
                f"FP Δ {fp_delta:+.1f}pp exceeds max +{rule['max_fp_delta']:.1f}pp")
            if not nl_ok:   reasons.append(
                f"No-locate Δ {nl_delta:+.1f}pp exceeds max +{rule['max_no_locate_delta']:.1f}pp")
            reasons.extend(urgent_reasons)
            if fired:
                reasons.append(
                    f"All promote conditions met: pos Δ={pos_lift:+.1f}pp, "
                    f"FP Δ={fp_delta:+.1f}pp, n={min_n}")

            fired_rules.append({
                "rule": dict(rule), "action": "promote",
                "passed": fired, "reasons": reasons,
            })
            if fired and recommendation is None:
                recommendation = "promote"
                can_promote    = True

        elif action == "stop":
            fp_breach = fp_delta > rule["max_fp_delta"]
            nl_breach = nl_delta > rule["max_no_locate_delta"]
            fired = n_ok and (fp_breach or nl_breach)
            reasons = []
            if fp_breach: reasons.append(
                f"FP Δ {fp_delta:+.1f}pp exceeds stop threshold +{rule['max_fp_delta']:.1f}pp")
            if nl_breach: reasons.append(
                f"No-locate Δ {nl_delta:+.1f}pp exceeds stop threshold +{rule['max_no_locate_delta']:.1f}pp")
            if not n_ok:  reasons.append(
                f"Insufficient sample ({min_n} < {rule['min_sample_per_arm']})")
            if not reasons: reasons.append("FP and no-locate within acceptable bounds")

            fired_rules.append({
                "rule": dict(rule), "action": "stop",
                "passed": fired, "reasons": reasons,
            })
            if fired and recommendation is None:
                recommendation = "stop"
                can_stop = True

        elif action == "tighten":
            fp_mod = fp_delta > rule["max_fp_delta"]
            nl_mod = nl_delta > rule["max_no_locate_delta"]
            fired  = n_ok and (fp_mod or nl_mod) and recommendation is None
            reasons = []
            if fp_mod: reasons.append(
                f"Moderate FP Δ {fp_delta:+.1f}pp — consider raising challenger threshold")
            if nl_mod: reasons.append(
                f"Moderate no-locate Δ {nl_delta:+.1f}pp — consider raising challenger threshold")
            if not reasons: reasons.append("Metrics within tighten bounds")

            fired_rules.append({
                "rule": dict(rule), "action": "tighten",
                "passed": fired, "reasons": reasons,
            })
            if fired and recommendation is None:
                recommendation = "tighten"

    if recommendation is None:
        default_sample = rules[0]["min_sample_per_arm"] if rules else 50
        recommendation = "insufficient_sample" if min_n < default_sample else "continue"

    return {
        "recommendation":  recommendation,
        "rules_fired":     fired_rules,
        "metrics":         metrics,
        "can_promote":     can_promote,
        "can_stop":        can_stop,
        "pending_decision": dict(pending_decision) if pending_decision else None,
        "last_decision":    dict(last_decision)    if last_decision    else None,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Stage 16 — Experiment admin routes
# ─────────────────────────────────────────────────────────────────────────────

def _exp_arm_stats(conn, exp_id, arm):
    POS = ("'confirmed_present','repeat_area_confirmed',"
           "'recovery_progressed','recovery_completed','followup_required'")
    row = conn.execute(f"""
        SELECT COUNT(*) AS total,
               SUM(CASE WHEN actual_outcome IN ({POS})          THEN 1 ELSE 0 END) AS pos_count,
               SUM(CASE WHEN actual_outcome = 'false_positive'  THEN 1 ELSE 0 END) AS fp_count,
               SUM(CASE WHEN actual_outcome = 'no_locate'       THEN 1 ELSE 0 END) AS nl_count,
               SUM(CASE WHEN actual_outcome = 'recovery_progressed'
                            OR actual_outcome = 'recovery_completed' THEN 1 ELSE 0 END) AS rec_count
        FROM lpr_experiment_results
        WHERE experiment_id=? AND arm=?
    """, (exp_id, arm)).fetchone()
    total = row["total"] or 0
    def pct(n): return round(100 * n / total, 1) if total else 0
    return {
        "arm":         arm,
        "total":       total,
        "pos_count":   row["pos_count"]  or 0,
        "fp_count":    row["fp_count"]   or 0,
        "nl_count":    row["nl_count"]   or 0,
        "rec_count":   row["rec_count"]  or 0,
        "pos_rate":    pct(row["pos_count"]  or 0),
        "fp_rate":     pct(row["fp_count"]   or 0),
        "nl_rate":     pct(row["nl_count"]   or 0),
        "rec_rate":    pct(row["rec_count"]  or 0),
    }


@app.get("/admin/lpr/experiments")
@admin_required
def admin_lpr_experiments():
    conn = db()
    _lpr_experiments_ensure(conn)
    _lpr_ranking_config_ensure(conn)
    _lpr_policy_ensure(conn)

    exp_rows = conn.execute("""
        SELECT * FROM lpr_experiments
        ORDER BY CASE status WHEN 'active' THEN 0 ELSE 1 END, created_at DESC
        LIMIT 30
    """).fetchall()

    policy_rules = [dict(r) for r in conn.execute(
        "SELECT * FROM lpr_policy_rules WHERE active=1 ORDER BY priority"
    ).fetchall()]

    experiments = []
    for e in exp_rows:
        exp_dict = dict(e)
        champion_stats    = _exp_arm_stats(conn, e["id"], "champion")
        challenger_stats  = _exp_arm_stats(conn, e["id"], "challenger")
        assignments_total = conn.execute(
            "SELECT COUNT(*) FROM lpr_experiment_assignments WHERE experiment_id=?",
            (e["id"],)
        ).fetchone()[0]
        exp_dict["champion_stats"]    = champion_stats
        exp_dict["challenger_stats"]  = challenger_stats
        exp_dict["assignments_total"] = assignments_total
        min_n = min(champion_stats["total"], challenger_stats["total"])
        exp_dict["has_enough_data"]   = min_n >= 30
        exp_dict["min_arm_n"]         = min_n

        if e["status"] in ("active", "stopped"):
            exp_dict["policy"] = _evaluate_experiment_policy(conn, exp_dict)
        else:
            exp_dict["policy"] = None

        recent_decisions = conn.execute("""
            SELECT id, recommended_action, decision_status, decided_at, notes_safe, created_at
            FROM lpr_policy_decisions
            WHERE experiment_id=?
            ORDER BY id DESC LIMIT 5
        """, (e["id"],)).fetchall()
        exp_dict["recent_decisions"] = [dict(d) for d in recent_decisions]

        experiments.append(exp_dict)

    active_ranking = conn.execute("""
        SELECT rule_weight, ml_weight, min_combined_threshold, model_version
        FROM lpr_ranking_config WHERE active=1
        ORDER BY (CASE WHEN model_version='*' THEN 1 ELSE 0 END), effective_from DESC
        LIMIT 1
    """).fetchone()
    conn.close()

    return render_template(
        "lpr_experiments.html",
        experiments=experiments,
        active_ranking=dict(active_ranking) if active_ranking else None,
        policy_rules=policy_rules,
    )


@app.post("/admin/lpr/experiments/create")
@admin_required
def admin_lpr_experiments_create():
    data = request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()[:200]
    if not name:
        return jsonify({"ok": False, "error": "Experiment name required"}), 400

    try:
        c_rw  = round(max(0.0, min(1.0, float(data.get("champion_rule_weight",   0.40)))), 2)
        c_mw  = round(max(0.0, min(1.0, float(data.get("champion_ml_weight",     0.60)))), 2)
        ch_rw = round(max(0.0, min(1.0, float(data.get("challenger_rule_weight", 0.30)))), 2)
        ch_mw = round(max(0.0, min(1.0, float(data.get("challenger_ml_weight",   0.70)))), 2)
        c_thr  = max(0, min(100, int(data.get("champion_threshold",   30))))
        ch_thr = max(0, min(100, int(data.get("challenger_threshold", 35))))
        split  = max(5,  min(50,  int(data.get("traffic_split_pct",   20))))
    except (ValueError, TypeError):
        return jsonify({"ok": False, "error": "Invalid numeric values"}), 400

    for rw, mw in ((c_rw, c_mw), (ch_rw, ch_mw)):
        if abs(rw + mw - 1.0) > 0.01:
            return jsonify({"ok": False,
                            "error": "Each arm's rule+ML weights must sum to 1.0"}), 400

    c_mv  = (data.get("champion_model_version")   or "").strip() or None
    ch_mv = (data.get("challenger_model_version")  or "").strip() or None
    notes = (data.get("notes_safe") or "").strip()[:300]
    now   = now_ts()
    uid   = session.get("user_id")

    conn = db()
    _lpr_experiments_ensure(conn)

    existing_active = conn.execute(
        "SELECT COUNT(*) FROM lpr_experiments WHERE status='active'"
    ).fetchone()[0]
    if existing_active:
        conn.close()
        return jsonify({"ok": False,
                        "error": "An experiment is already active — stop it before creating a new one"}), 409

    conn.execute("""
        INSERT INTO lpr_experiments
            (name, status, experiment_type,
             champion_model_version, challenger_model_version,
             champion_rule_weight,   champion_ml_weight,
             challenger_rule_weight, challenger_ml_weight,
             champion_threshold, challenger_threshold,
             traffic_split_pct, start_at, created_by, created_at, notes_safe)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (name, "active", "blend",
          c_mv, ch_mv, c_rw, c_mw, ch_rw, ch_mw,
          c_thr, ch_thr, split, now, uid, now, notes or None))
    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


@app.post("/admin/lpr/experiments/stop")
@admin_required
def admin_lpr_experiments_stop():
    data = request.get_json(silent=True) or {}
    exp_id = data.get("experiment_id")
    if not exp_id:
        return jsonify({"ok": False, "error": "experiment_id required"}), 400

    conn = db()
    _lpr_experiments_ensure(conn)
    conn.execute("""
        UPDATE lpr_experiments SET status='stopped', end_at=?
        WHERE id=? AND status='active'
    """, (now_ts(), exp_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


@app.post("/admin/lpr/experiments/promote")
@admin_required
def admin_lpr_experiments_promote():
    """
    Promote challenger: create a new active ranking config using challenger
    weights, end the experiment, leave an audit trail.
    """
    data = request.get_json(silent=True) or {}
    exp_id = data.get("experiment_id")
    if not exp_id:
        return jsonify({"ok": False, "error": "experiment_id required"}), 400

    conn = db()
    _lpr_experiments_ensure(conn)
    _lpr_ranking_config_ensure(conn)

    exp = conn.execute("SELECT * FROM lpr_experiments WHERE id=?", (exp_id,)).fetchone()
    if not exp:
        conn.close()
        return jsonify({"ok": False, "error": "Experiment not found"}), 404
    if exp["status"] not in ("active", "stopped"):
        conn.close()
        return jsonify({"ok": False, "error": "Experiment already promoted or archived"}), 400

    now = now_ts()
    uid = session.get("user_id")

    conn.execute("""
        UPDATE lpr_ranking_config SET active=0
        WHERE model_version='*' AND confidence_band='*' AND priority_band='*'
    """)
    conn.execute("""
        INSERT INTO lpr_ranking_config
            (model_version, prediction_window, confidence_band, priority_band,
             rule_weight, ml_weight, min_combined_threshold, active, source,
             reason, effective_from, created_by, created_at)
        VALUES (?,?,?,?,?,?,?,1,'experiment',?,?,?,?)
    """, (
        exp["challenger_model_version"] or "*",
        "72h", "*", "*",
        exp["challenger_rule_weight"],
        exp["challenger_ml_weight"],
        exp["challenger_threshold"],
        f"Promoted from experiment #{exp_id}: {exp['name']}",
        now, uid, now,
    ))

    conn.execute("""
        UPDATE lpr_experiments SET status='promoted', end_at=?
        WHERE id=?
    """, (now, exp_id))

    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


@app.post("/admin/lpr/experiments/policy/decide")
@admin_required
def admin_lpr_experiments_policy_decide():
    """
    Record a human decision on a policy recommendation.
    decision: 'approved_promote' | 'approved_stop' | 'rejected' | 'deferred'
    If approved_promote → promotion is executed immediately.
    If approved_stop    → experiment is stopped immediately.
    All decisions are written to lpr_policy_decisions for audit.
    """
    data       = request.get_json(silent=True) or {}
    exp_id     = data.get("experiment_id")
    decision   = data.get("decision")
    notes      = (data.get("notes") or "").strip()[:300]
    VALID_DECISIONS = ("approved_promote", "approved_stop", "rejected", "deferred")

    if not exp_id:
        return jsonify({"ok": False, "error": "experiment_id required"}), 400
    if decision not in VALID_DECISIONS:
        return jsonify({"ok": False,
                        "error": f"decision must be one of: {', '.join(VALID_DECISIONS)}"}), 400

    conn = db()
    _lpr_experiments_ensure(conn)
    _lpr_policy_ensure(conn)
    _lpr_ranking_config_ensure(conn)

    exp = conn.execute("SELECT * FROM lpr_experiments WHERE id=?", (exp_id,)).fetchone()
    if not exp:
        conn.close()
        return jsonify({"ok": False, "error": "Experiment not found"}), 404
    if exp["status"] not in ("active", "stopped"):
        conn.close()
        return jsonify({"ok": False,
                        "error": "Policy decisions only apply to active or stopped experiments"}), 400

    uid = session.get("user_id")
    now = now_ts()

    # Evaluate policy for the snapshot
    policy_eval = _evaluate_experiment_policy(conn, dict(exp))
    m = policy_eval["metrics"]
    metrics_snap = _json.dumps({
        "min_n":                  m["min_n"],
        "pos_lift":               m["pos_lift"],
        "fp_delta":               m["fp_delta"],
        "nl_delta":               m["nl_delta"],
        "champion_pos_rate":      m["champion"]["pos_rate"],
        "challenger_pos_rate":    m["challenger"]["pos_rate"],
        "champion_fp_rate":       m["champion"]["fp_rate"],
        "challenger_fp_rate":     m["challenger"]["fp_rate"],
        "champion_total":         m["champion"]["total"],
        "challenger_total":       m["challenger"]["total"],
        "policy_recommendation":  policy_eval["recommendation"],
    })

    action_label_map = {
        "approved_promote": "promote",
        "approved_stop":    "stop",
        "rejected":         policy_eval["recommendation"],
        "deferred":         policy_eval["recommendation"],
    }
    status_map = {
        "approved_promote": "approved",
        "approved_stop":    "approved",
        "rejected":         "rejected",
        "deferred":         "deferred",
    }

    # Build plain-English reason from fired rules
    reason_parts = [r["reasons"][0] for r in policy_eval["rules_fired"] if r["passed"]]
    reason_text  = "; ".join(reason_parts) if reason_parts else (
        f"Manual {decision.replace('_', ' ')} — {policy_eval['recommendation']}")

    cur = conn.execute("""
        INSERT INTO lpr_policy_decisions
            (experiment_id, rule_id, recommended_action, reason_text,
             metrics_snapshot, decision_status, decided_by, decided_at,
             notes_safe, created_at)
        VALUES (?,?,?,?,?,?,?,?,?,?)
    """, (
        exp_id, None,
        action_label_map[decision],
        reason_text,
        metrics_snap,
        status_map[decision],
        uid, now, notes or None, now,
    ))
    decision_id = cur.lastrowid

    applied_cfg_id = None

    if decision == "approved_promote":
        conn.execute("""
            UPDATE lpr_ranking_config SET active=0
            WHERE model_version='*' AND confidence_band='*' AND priority_band='*'
        """)
        cfg_cur = conn.execute("""
            INSERT INTO lpr_ranking_config
                (model_version, prediction_window, confidence_band, priority_band,
                 rule_weight, ml_weight, min_combined_threshold, active, source,
                 reason, effective_from, created_by, created_at)
            VALUES (?,?,?,?,?,?,?,1,'policy_approved',?,?,?,?)
        """, (
            exp["challenger_model_version"] or "*",
            "72h", "*", "*",
            exp["challenger_rule_weight"],
            exp["challenger_ml_weight"],
            exp["challenger_threshold"],
            f"Policy-approved promotion from experiment #{exp_id}: {exp['name']}",
            now, uid, now,
        ))
        applied_cfg_id = cfg_cur.lastrowid
        conn.execute("UPDATE lpr_experiments SET status='promoted', end_at=? WHERE id=?",
                     (now, exp_id))
        conn.execute("UPDATE lpr_policy_decisions SET applied_config_id=? WHERE id=?",
                     (applied_cfg_id, decision_id))

    elif decision == "approved_stop":
        conn.execute(
            "UPDATE lpr_experiments SET status='stopped', end_at=? WHERE id=? AND status='active'",
            (now, exp_id))

    conn.commit()
    conn.close()
    return jsonify({"ok": True, "decision_id": decision_id,
                    "applied_config_id": applied_cfg_id}), 200


@app.post("/admin/lpr/experiments/archive")
@admin_required
def admin_lpr_experiments_archive():
    data = request.get_json(silent=True) or {}
    exp_id = data.get("experiment_id")
    if not exp_id:
        return jsonify({"ok": False, "error": "experiment_id required"}), 400

    conn = db()
    _lpr_experiments_ensure(conn)
    conn.execute("""
        UPDATE lpr_experiments SET status='archived'
        WHERE id=? AND status NOT IN ('active')
    """, (exp_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


# ─────────────────────────────────────────────────────────────────────────────
# Stage 18 — Controlled Automation
# ─────────────────────────────────────────────────────────────────────────────

from datetime import datetime as _dt, timedelta as _td


def _lpr_automation_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_automation_settings (
            id                          INTEGER PRIMARY KEY AUTOINCREMENT,
            scope                       TEXT    NOT NULL DEFAULT 'global',
            active                      INTEGER NOT NULL DEFAULT 0,
            allow_auto_tighten          INTEGER NOT NULL DEFAULT 1,
            allow_auto_band_suppression INTEGER NOT NULL DEFAULT 0,
            require_manual_for_promote  INTEGER NOT NULL DEFAULT 1,
            require_manual_for_stop     INTEGER NOT NULL DEFAULT 1,
            cooldown_days               INTEGER NOT NULL DEFAULT 7,
            max_threshold_step          INTEGER NOT NULL DEFAULT 5,
            min_sample_per_arm          INTEGER NOT NULL DEFAULT 50,
            created_by                  INTEGER,
            created_at                  TEXT    NOT NULL,
            updated_at                  TEXT    NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_automation_actions (
            id                    INTEGER PRIMARY KEY AUTOINCREMENT,
            policy_decision_id    INTEGER,
            experiment_id         INTEGER,
            action_type           TEXT    NOT NULL,
            scope                 TEXT    NOT NULL DEFAULT 'global',
            before_config_id      INTEGER,
            after_config_id       INTEGER,
            before_values_json    TEXT,
            after_values_json     TEXT,
            trigger_rule_id       INTEGER,
            sample_size_json      TEXT,
            status                TEXT    NOT NULL DEFAULT 'applied',
            applied_at            TEXT    NOT NULL,
            applied_by            TEXT    NOT NULL DEFAULT 'auto',
            rollback_of_action_id INTEGER,
            notes_safe            TEXT
        )
    """)
    n = conn.execute("SELECT COUNT(*) FROM lpr_automation_settings").fetchone()[0]
    if n == 0:
        now = now_ts()
        conn.execute("""
            INSERT INTO lpr_automation_settings
                (scope, active, allow_auto_tighten, allow_auto_band_suppression,
                 require_manual_for_promote, require_manual_for_stop,
                 cooldown_days, max_threshold_step, min_sample_per_arm,
                 created_at, updated_at)
            VALUES ('global', 0, 1, 0, 1, 1, 7, 5, 50, ?, ?)
        """, (now, now))
    conn.commit()


def _get_automation_settings(conn):
    row = conn.execute(
        "SELECT * FROM lpr_automation_settings ORDER BY id DESC LIMIT 1"
    ).fetchone()
    return dict(row) if row else None


def _automation_cooldown_ok(conn, scope, cooldown_days):
    """
    Return (ok: bool, last_applied_at: str|None).
    ok=True means no non-rollback auto-action has been applied within cooldown_days.
    """
    cutoff = (_dt.utcnow() - _td(days=int(cooldown_days))).strftime('%Y-%m-%dT%H:%M:%S')
    row = conn.execute("""
        SELECT applied_at FROM lpr_automation_actions
        WHERE scope=? AND status IN ('applied','monitoring','review_required')
          AND action_type != 'rollback'
          AND applied_at > ?
        ORDER BY applied_at DESC LIMIT 1
    """, (scope, cutoff)).fetchone()
    return (row is None), (row["applied_at"] if row else None)


def _dry_run_check(settings, policy_eval):
    """
    Return True if this evaluation WOULD trigger auto-apply if automation were active.
    Used to show "would have acted" dry-run indicators on the automation screen.
    """
    rec = policy_eval.get("recommendation")
    m   = policy_eval.get("metrics", {})
    if rec == "tighten" and settings.get("allow_auto_tighten"):
        return m.get("min_n", 0) >= settings.get("min_sample_per_arm", 50)
    return False


def _try_auto_apply(conn, exp, policy_eval, actor="auto"):
    """
    Evaluate whether the policy recommendation for exp is eligible for automation.
    If automation is active and all guardrails pass, apply the change and log it.
    Returns a result dict (never raises).
    """
    settings = _get_automation_settings(conn)
    if not settings:
        return {"applied": False, "reason": "No automation settings found"}

    dry_eligible = _dry_run_check(settings, policy_eval)

    if not settings["active"]:
        return {"applied": False, "reason": "Automation disabled",
                "dry_run_eligible": dry_eligible}

    # Stage 19: check scope pause before doing anything
    try:
        _lpr_stage19_ensure(conn)
        if _is_scope_paused(conn, "global"):
            return {"applied": False,
                    "reason": "Automation paused pending review — resolve open review tasks first",
                    "dry_run_eligible": dry_eligible}
    except Exception:
        pass

    rec = policy_eval.get("recommendation")
    m   = policy_eval.get("metrics", {})

    # ── Auto-tighten path ────────────────────────────────────────────────────
    if rec == "tighten" and settings["allow_auto_tighten"]:
        min_n = settings["min_sample_per_arm"]
        if m.get("min_n", 0) < min_n:
            return {"applied": False,
                    "reason": f"Insufficient sample ({m.get('min_n',0)} < {min_n})",
                    "dry_run_eligible": dry_eligible}

        cooldown_ok, last_at = _automation_cooldown_ok(conn, "global", settings["cooldown_days"])
        if not cooldown_ok:
            return {"applied": False,
                    "reason": f"Cooldown active — last auto-change on {last_at[:10]}",
                    "dry_run_eligible": dry_eligible}

        pc = m.get("protected_champion", {})
        ph = m.get("protected_challenger", {})
        if pc.get("total", 0) >= 10 and ph.get("total", 0) >= 5:
            prot_fp = ph.get("fp_rate", 0) - pc.get("fp_rate", 0)
            if prot_fp > 5.0:
                return {"applied": False,
                        "reason": f"Urgent/watchlist FP Δ {prot_fp:+.1f}pp — blocked",
                        "dry_run_eligible": False}

        active_cfg = conn.execute("""
            SELECT * FROM lpr_ranking_config WHERE active=1
            ORDER BY effective_from DESC LIMIT 1
        """).fetchone()
        if not active_cfg:
            return {"applied": False, "reason": "No active ranking config",
                    "dry_run_eligible": dry_eligible}

        old_thr = active_cfg["min_combined_threshold"]
        new_thr = min(old_thr + settings["max_threshold_step"], 100)
        if new_thr <= old_thr:
            return {"applied": False, "reason": "Already at maximum threshold",
                    "dry_run_eligible": False}

        now         = now_ts()
        before_json = _json.dumps({
            "min_combined_threshold": old_thr,
            "rule_weight":            active_cfg["rule_weight"],
            "ml_weight":              active_cfg["ml_weight"],
        })
        after_json  = _json.dumps({
            "min_combined_threshold": new_thr,
            "rule_weight":            active_cfg["rule_weight"],
            "ml_weight":              active_cfg["ml_weight"],
        })
        sample_json = _json.dumps({
            "champion":   m.get("champion",   {}).get("total", 0),
            "challenger": m.get("challenger", {}).get("total", 0),
            "min_n":      m.get("min_n", 0),
        })

        conn.execute("UPDATE lpr_ranking_config SET active=0 WHERE id=?", (active_cfg["id"],))
        cfg_cur = conn.execute("""
            INSERT INTO lpr_ranking_config
                (model_version, prediction_window, confidence_band, priority_band,
                 rule_weight, ml_weight, min_combined_threshold, active, source,
                 reason, effective_from, created_by, created_at)
            VALUES (?,?,?,?,?,?,?,1,'auto_tighten',?,?,NULL,?)
        """, (
            active_cfg["model_version"] or "*",
            active_cfg["prediction_window"] or "72h",
            active_cfg["confidence_band"]   or "*",
            active_cfg["priority_band"]     or "*",
            active_cfg["rule_weight"],
            active_cfg["ml_weight"],
            new_thr,
            f"Auto-tighten from exp #{exp['id']}: threshold {old_thr} → {new_thr}",
            now, now,
        ))
        new_cfg_id = cfg_cur.lastrowid

        fired_rule = next((r for r in policy_eval.get("rules_fired", []) if r["passed"]), None)
        rule_id    = fired_rule["rule"]["id"] if fired_rule else None

        act_cur = conn.execute("""
            INSERT INTO lpr_automation_actions
                (policy_decision_id, experiment_id, action_type, scope,
                 before_config_id, after_config_id, before_values_json, after_values_json,
                 trigger_rule_id, sample_size_json, status, applied_at, applied_by, notes_safe)
            VALUES (NULL,?,?,?,?,?,?,?,?,?,'monitoring',?,?,'')
        """, (
            exp["id"], "tighten_threshold", "global",
            active_cfg["id"], new_cfg_id,
            before_json, after_json,
            rule_id, sample_json, now, actor,
        ))
        action_id = act_cur.lastrowid
        conn.commit()

        # Stage 19: notify admins of successful auto-tighten
        try:
            _notify_admins(
                "LPR automation raised patrol threshold",
                f"Automated threshold tightened from {old_thr} to {new_thr}. "
                "Outcomes will be monitored — rollback available if performance dips.",
                "automation_tighten",
                {"action_id": action_id, "old_threshold": old_thr, "new_threshold": new_thr},
            )
        except Exception:
            pass

        return {
            "applied":       True,
            "action_type":   "tighten_threshold",
            "action_id":     action_id,
            "old_threshold": old_thr,
            "new_threshold": new_thr,
            "before_cfg_id": active_cfg["id"],
            "after_cfg_id":  new_cfg_id,
        }

    # ── Manual-only paths ────────────────────────────────────────────────────
    if rec == "promote":
        return {"applied": False, "reason": "Promotion requires manual approval",
                "dry_run_eligible": False}
    if rec == "stop":
        return {"applied": False, "reason": "Stop requires manual approval",
                "dry_run_eligible": False}

    return {"applied": False, "reason": f"No automation action for '{rec}'",
            "dry_run_eligible": False}


def _check_automation_monitoring(conn):
    """
    For all 'monitoring' tighten_threshold actions, check whether post-action
    outcomes show degraded performance.  Flag as 'review_required' if:
      - positive rate dropped > 3pp vs champion arm baseline, or
      - FP rate rose > 5pp vs champion arm baseline.
    Returns list of action IDs that were newly flagged.
    """
    actions = conn.execute("""
        SELECT * FROM lpr_automation_actions
        WHERE status='monitoring' AND action_type='tighten_threshold'
    """).fetchall()

    flagged = []
    POS = ("'confirmed_present','repeat_area_confirmed',"
           "'recovery_progressed','recovery_completed','followup_required'")

    for a in actions:
        if not a["experiment_id"]:
            continue
        post = conn.execute(f"""
            SELECT COUNT(*) AS n,
                   SUM(CASE WHEN actual_outcome IN ({POS}) THEN 1 ELSE 0 END) AS pos,
                   SUM(CASE WHEN actual_outcome='false_positive' THEN 1 ELSE 0 END) AS fp
            FROM lpr_prediction_outcomes
            WHERE created_at > ?
        """, (a["applied_at"],)).fetchone()

        n_post = post["n"] or 0
        if n_post < 10:
            continue

        pos_rate_post = round(100 * (post["pos"] or 0) / n_post, 1)
        fp_rate_post  = round(100 * (post["fp"]  or 0) / n_post, 1)

        baseline = _exp_arm_stats(conn, a["experiment_id"], "champion")
        pos_drop = baseline["pos_rate"] - pos_rate_post
        fp_rise  = fp_rate_post - baseline["fp_rate"]

        if pos_drop > 3.0 or fp_rise > 5.0:
            note = (f"Post-change: pos {pos_rate_post}% (Δ{-pos_drop:+.1f}pp), "
                    f"FP {fp_rate_post}% (Δ{fp_rise:+.1f}pp), n={n_post}")
            conn.execute(
                "UPDATE lpr_automation_actions SET status='review_required', notes_safe=? WHERE id=?",
                (note, a["id"])
            )
            flagged.append(a["id"])

            # Stage 19: create review task, pause scope, notify
            try:
                _lpr_stage19_ensure(conn)
                scope = a["scope"] or "global"
                _create_review_task(conn, a["id"], scope,
                                    "Post-change performance may have degraded — " + note)
                paused = _pause_scope(conn, scope, a["id"],
                                      "Auto-paused: review_required on action #" + str(a["id"]))
                conn.commit()
                _notify_admins(
                    "LPR automation review required",
                    "A recent automated threshold change needs review. "
                    "Auto-tighten has been paused pending office sign-off.",
                    "automation_review",
                    {"action_id": a["id"], "scope": scope},
                )
                if paused:
                    _notify_admins(
                        "LPR automation paused pending review",
                        "Automated threshold changes are paused for this scope "
                        "until the open review task is resolved.",
                        "automation_paused",
                        {"scope": scope},
                    )
            except Exception:
                pass

    if flagged:
        conn.commit()
    return flagged


# ── Automation admin routes ──────────────────────────────────────────────────

@app.get("/admin/lpr/automation")
@admin_required
def admin_lpr_automation():
    conn = db()
    _lpr_automation_ensure(conn)
    _lpr_stage19_ensure(conn)
    _lpr_experiments_ensure(conn)
    _lpr_policy_ensure(conn)
    _lpr_ranking_config_ensure(conn)

    _check_automation_monitoring(conn)

    settings = _get_automation_settings(conn)
    cooldown_ok, last_at = _automation_cooldown_ok(conn, "global", settings["cooldown_days"])

    recent_actions = [dict(r) for r in conn.execute("""
        SELECT a.*, e.name AS exp_name
        FROM lpr_automation_actions a
        LEFT JOIN lpr_experiments e ON e.id = a.experiment_id
        ORDER BY a.id DESC LIMIT 20
    """).fetchall()]

    for aa in recent_actions:
        try:
            aa["before_values"] = _json.loads(aa["before_values_json"] or "{}")
            aa["after_values"]  = _json.loads(aa["after_values_json"]  or "{}")
            aa["sample_size"]   = _json.loads(aa["sample_size_json"]   or "{}")
        except Exception:
            aa["before_values"] = {}
            aa["after_values"]  = {}
            aa["sample_size"]   = {}

    active_exp = conn.execute(
        "SELECT * FROM lpr_experiments WHERE status='active' LIMIT 1"
    ).fetchone()
    dry_run_result = None
    if active_exp:
        exp_dict    = dict(active_exp)
        policy_eval = _evaluate_experiment_policy(conn, exp_dict)
        dry_run_result = {
            "exp_name":       active_exp["name"],
            "exp_id":         active_exp["id"],
            "recommendation": policy_eval["recommendation"],
            "would_act":      _dry_run_check(settings, policy_eval),
            "metrics":        policy_eval["metrics"],
        }

    active_ranking = conn.execute("""
        SELECT id, rule_weight, ml_weight, min_combined_threshold, source, effective_from
        FROM lpr_ranking_config WHERE active=1
        ORDER BY effective_from DESC LIMIT 1
    """).fetchone()

    # Stage 19: open reviews + active pauses
    open_reviews = [dict(r) for r in conn.execute("""
        SELECT rv.*, aa.action_type, aa.before_values_json, aa.after_values_json,
               e.name AS exp_name
        FROM lpr_automation_reviews rv
        LEFT JOIN lpr_automation_actions aa ON aa.id = rv.action_id
        LEFT JOIN lpr_experiments e ON e.id = aa.experiment_id
        WHERE rv.status='open'
        ORDER BY rv.id DESC
        LIMIT 20
    """).fetchall()]
    for rv in open_reviews:
        try:
            rv["before_values"] = _json.loads(rv["before_values_json"] or "{}")
            rv["after_values"]  = _json.loads(rv["after_values_json"]  or "{}")
        except Exception:
            rv["before_values"] = {}
            rv["after_values"]  = {}

    paused_scopes = [dict(r) for r in conn.execute("""
        SELECT * FROM lpr_control_state
        WHERE control_type='paused' AND active=1
        ORDER BY effective_from DESC
    """).fetchall()]

    conn.close()
    return render_template(
        "lpr_automation.html",
        settings=settings,
        cooldown_ok=cooldown_ok,
        last_auto_at=last_at,
        recent_actions=recent_actions,
        dry_run_result=dry_run_result,
        active_ranking=dict(active_ranking) if active_ranking else None,
        open_reviews=open_reviews,
        paused_scopes=paused_scopes,
    )


@app.post("/admin/lpr/automation/settings")
@admin_required
def admin_lpr_automation_settings():
    data = request.get_json(silent=True) or {}
    conn = db()
    _lpr_automation_ensure(conn)

    try:
        active           = 1 if data.get("active")           else 0
        auto_tighten     = 1 if data.get("allow_auto_tighten") else 0
        band_suppress    = 1 if data.get("allow_auto_band_suppression") else 0
        manual_promote   = 1 if data.get("require_manual_for_promote", True) else 0
        manual_stop      = 1 if data.get("require_manual_for_stop",    True) else 0
        cooldown_days    = max(1, min(90, int(data.get("cooldown_days",   7))))
        max_step         = max(1, min(20, int(data.get("max_threshold_step", 5))))
        min_sample       = max(10, min(500, int(data.get("min_sample_per_arm", 50))))
    except (ValueError, TypeError):
        conn.close()
        return jsonify({"ok": False, "error": "Invalid numeric values"}), 400

    if not manual_promote:
        conn.close()
        return jsonify({"ok": False,
                        "error": "require_manual_for_promote cannot be disabled"}), 400
    if not manual_stop:
        conn.close()
        return jsonify({"ok": False,
                        "error": "require_manual_for_stop cannot be disabled"}), 400

    now = now_ts()
    uid = session.get("user_id")
    conn.execute("""
        UPDATE lpr_automation_settings
        SET active=?, allow_auto_tighten=?, allow_auto_band_suppression=?,
            require_manual_for_promote=?, require_manual_for_stop=?,
            cooldown_days=?, max_threshold_step=?, min_sample_per_arm=?,
            created_by=?, updated_at=?
        WHERE id=(SELECT MAX(id) FROM lpr_automation_settings)
    """, (active, auto_tighten, band_suppress,
          manual_promote, manual_stop,
          cooldown_days, max_step, min_sample,
          uid, now))
    conn.commit()
    conn.close()
    return jsonify({"ok": True}), 200


@app.post("/admin/lpr/automation/run")
@admin_required
def admin_lpr_automation_run():
    """
    Trigger the automation engine against the current active experiment.
    If automation is enabled and guardrails pass, applies the eligible change.
    """
    conn = db()
    _lpr_automation_ensure(conn)
    _lpr_experiments_ensure(conn)
    _lpr_policy_ensure(conn)
    _lpr_ranking_config_ensure(conn)

    active_exp = conn.execute(
        "SELECT * FROM lpr_experiments WHERE status='active' LIMIT 1"
    ).fetchone()
    if not active_exp:
        conn.close()
        return jsonify({"ok": False, "error": "No active experiment to evaluate"}), 400

    exp_dict    = dict(active_exp)
    policy_eval = _evaluate_experiment_policy(conn, exp_dict)
    result      = _try_auto_apply(conn, exp_dict, policy_eval,
                                  actor=str(session.get("user_id") or "auto"))
    conn.close()
    return jsonify({"ok": True, "result": result}), 200


@app.post("/admin/lpr/automation/rollback")
@admin_required
def admin_lpr_automation_rollback():
    """
    Roll back an automation action by re-activating the before_config_id.
    Writes a new action row with action_type='rollback'.
    """
    data      = request.get_json(silent=True) or {}
    action_id = data.get("action_id")
    notes     = (data.get("notes") or "").strip()[:300]
    if not action_id:
        return jsonify({"ok": False, "error": "action_id required"}), 400

    conn = db()
    _lpr_automation_ensure(conn)
    _lpr_ranking_config_ensure(conn)

    original = conn.execute(
        "SELECT * FROM lpr_automation_actions WHERE id=?", (action_id,)
    ).fetchone()
    if not original:
        conn.close()
        return jsonify({"ok": False, "error": "Action not found"}), 404
    if original["status"] == "rolled_back":
        conn.close()
        return jsonify({"ok": False, "error": "Action already rolled back"}), 400
    if not original["before_config_id"]:
        conn.close()
        return jsonify({"ok": False, "error": "No before_config_id to restore"}), 400

    # Restore the before config
    before_cfg = conn.execute(
        "SELECT * FROM lpr_ranking_config WHERE id=?", (original["before_config_id"],)
    ).fetchone()
    if not before_cfg:
        conn.close()
        return jsonify({"ok": False, "error": "Before config no longer exists"}), 400

    now = now_ts()
    uid = session.get("user_id")

    conn.execute("UPDATE lpr_ranking_config SET active=0 WHERE active=1")
    conn.execute("UPDATE lpr_ranking_config SET active=1 WHERE id=?",
                 (original["before_config_id"],))
    conn.execute("UPDATE lpr_automation_actions SET status='rolled_back' WHERE id=?",
                 (action_id,))

    conn.execute("""
        INSERT INTO lpr_automation_actions
            (policy_decision_id, experiment_id, action_type, scope,
             before_config_id, after_config_id, before_values_json, after_values_json,
             trigger_rule_id, sample_size_json, status, applied_at, applied_by,
             rollback_of_action_id, notes_safe)
        VALUES (NULL, ?, 'rollback', 'global', ?, ?, ?, ?, NULL, NULL, 'applied', ?, ?, ?, ?)
    """, (
        original["experiment_id"],
        original["after_config_id"], original["before_config_id"],
        original["after_values_json"], original["before_values_json"],
        now, str(uid or "manual"), action_id, notes or None,
    ))

    conn.commit()
    conn.close()

    # Stage 19: notify admins of rollback
    try:
        _notify_admins(
            "LPR automation change rolled back",
            "An automated threshold change has been rolled back. "
            "The previous ranking configuration is now active.",
            "automation_rollback",
            {"rolled_back_action_id": action_id},
        )
    except Exception:
        pass

    return jsonify({"ok": True}), 200


# ─────────────────────────────────────────────────────────────────────────────
# Stage 19 — Notification and control automation
# ─────────────────────────────────────────────────────────────────────────────

def _lpr_stage19_ensure(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_automation_reviews (
            id               INTEGER PRIMARY KEY AUTOINCREMENT,
            action_id        INTEGER NOT NULL,
            scope_key        TEXT    NOT NULL DEFAULT 'global',
            review_reason    TEXT    NOT NULL,
            status           TEXT    NOT NULL DEFAULT 'open',
            assigned_to      INTEGER,
            created_at       TEXT    NOT NULL,
            resolved_at      TEXT,
            resolution_notes TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS lpr_control_state (
            id                INTEGER PRIMARY KEY AUTOINCREMENT,
            scope_key         TEXT    NOT NULL DEFAULT 'global',
            control_type      TEXT    NOT NULL DEFAULT 'paused',
            active            INTEGER NOT NULL DEFAULT 1,
            effective_from    TEXT    NOT NULL,
            effective_to      TEXT,
            trigger_action_id INTEGER,
            reason_text       TEXT    NOT NULL,
            created_by        TEXT    NOT NULL DEFAULT 'auto'
        )
    """)
    conn.commit()


def _is_scope_paused(conn, scope_key):
    """Return True if scope_key has an active pause control in lpr_control_state."""
    row = conn.execute("""
        SELECT id FROM lpr_control_state
        WHERE scope_key=? AND control_type='paused' AND active=1
        LIMIT 1
    """, (scope_key,)).fetchone()
    return row is not None


def _pause_scope(conn, scope_key, trigger_action_id, reason):
    """
    Create a paused control_state record if not already paused.
    Returns True if a new pause was created.
    """
    if _is_scope_paused(conn, scope_key):
        return False
    conn.execute("""
        INSERT INTO lpr_control_state
            (scope_key, control_type, active, effective_from, trigger_action_id, reason_text, created_by)
        VALUES (?, 'paused', 1, ?, ?, ?, 'auto')
    """, (scope_key, now_ts(), trigger_action_id, reason))
    conn.commit()
    return True


def _resume_scope(conn, scope_key, user_id=None):
    """
    Deactivate all active pause records for scope_key.
    Returns the number of records deactivated.
    """
    now = now_ts()
    conn.execute("""
        UPDATE lpr_control_state
        SET active=0, effective_to=?
        WHERE scope_key=? AND control_type='paused' AND active=1
    """, (now, scope_key))
    count = conn.total_changes
    conn.commit()
    return count


def _create_review_task(conn, action_id, scope_key, reason):
    """
    Create an open review task for an automation action.
    Returns the new review task id.
    """
    cur = conn.execute("""
        INSERT INTO lpr_automation_reviews
            (action_id, scope_key, review_reason, status, created_at)
        VALUES (?, ?, ?, 'open', ?)
    """, (action_id, scope_key, reason, now_ts()))
    conn.commit()
    return cur.lastrowid


def _open_review_count(conn, scope_key="global"):
    """Return number of open review tasks for a scope."""
    return conn.execute("""
        SELECT COUNT(*) FROM lpr_automation_reviews
        WHERE scope_key=? AND status='open'
    """, (scope_key,)).fetchone()[0]


# ── Stage 19 routes ──────────────────────────────────────────────────────────

@app.post("/admin/lpr/automation/resolve-review")
@admin_required
def admin_lpr_automation_resolve_review():
    """
    Resolve an open review task.
    If no open reviews remain for the scope, automatically resume the scope.
    """
    data      = request.get_json(silent=True) or {}
    review_id = data.get("review_id")
    notes     = (data.get("notes") or "").strip()[:300]
    if not review_id:
        return jsonify({"ok": False, "error": "review_id required"}), 400

    conn = db()
    _lpr_automation_ensure(conn)
    _lpr_stage19_ensure(conn)

    review = conn.execute(
        "SELECT * FROM lpr_automation_reviews WHERE id=?", (review_id,)
    ).fetchone()
    if not review:
        conn.close()
        return jsonify({"ok": False, "error": "Review not found"}), 404
    if review["status"] != "open":
        conn.close()
        return jsonify({"ok": False, "error": "Review already resolved"}), 400

    now = now_ts()
    uid = session.get("user_id")
    conn.execute("""
        UPDATE lpr_automation_reviews
        SET status='resolved', resolved_at=?, resolution_notes=?, assigned_to=?
        WHERE id=?
    """, (now, notes or None, uid, review_id))
    conn.commit()

    scope_key    = review["scope_key"]
    open_reviews = _open_review_count(conn, scope_key)
    auto_resumed = False
    if open_reviews == 0:
        auto_resumed = _resume_scope(conn, scope_key, uid) > 0
        if auto_resumed:
            _notify_admins(
                "LPR automation scope resumed",
                "All review tasks resolved — automation scope resumed and ready.",
                "automation_resumed",
                {"scope_key": scope_key},
            )

    conn.close()
    return jsonify({"ok": True, "auto_resumed": auto_resumed}), 200


@app.post("/admin/lpr/automation/resume-scope")
@admin_required
def admin_lpr_automation_resume_scope():
    """Manually resume a paused automation scope."""
    data      = request.get_json(silent=True) or {}
    scope_key = (data.get("scope_key") or "global").strip()
    notes     = (data.get("notes") or "").strip()[:200]

    conn = db()
    _lpr_automation_ensure(conn)
    _lpr_stage19_ensure(conn)

    uid     = session.get("user_id")
    resumed = _resume_scope(conn, scope_key, uid)
    conn.close()

    if resumed:
        _notify_admins(
            "LPR automation scope resumed",
            f"Automation scope manually resumed by office. {notes}".strip(),
            "automation_resumed",
            {"scope_key": scope_key},
        )
    return jsonify({"ok": True, "resumed": resumed > 0}), 200


# ─────────────────────────────────────────────────────────────────────────────
# Internal Messaging System
# ─────────────────────────────────────────────────────────────────────────────

def _ensure_msg_tables(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS conversations (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            type       TEXT    NOT NULL DEFAULT 'direct',
            job_id     INTEGER,
            subject    TEXT,
            created_at TEXT    NOT NULL,
            updated_at TEXT    NOT NULL,
            FOREIGN KEY(job_id) REFERENCES jobs(id)
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS conversation_participants (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER NOT NULL,
            user_id         INTEGER NOT NULL,
            joined_at       TEXT    NOT NULL,
            UNIQUE(conversation_id, user_id),
            FOREIGN KEY(conversation_id) REFERENCES conversations(id),
            FOREIGN KEY(user_id)         REFERENCES users(id)
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS messages (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER NOT NULL,
            sender_id       INTEGER NOT NULL,
            body            TEXT    NOT NULL,
            created_at      TEXT    NOT NULL,
            is_deleted      INTEGER NOT NULL DEFAULT 0,
            FOREIGN KEY(conversation_id) REFERENCES conversations(id),
            FOREIGN KEY(sender_id)       REFERENCES users(id)
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS message_reads (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            message_id INTEGER NOT NULL,
            user_id    INTEGER NOT NULL,
            read_at    TEXT    NOT NULL,
            UNIQUE(message_id, user_id),
            FOREIGN KEY(message_id) REFERENCES messages(id),
            FOREIGN KEY(user_id)    REFERENCES users(id)
        )
    """)
    conn.commit()


def _get_or_create_direct_conv(conn, uid_a, uid_b, job_id=None):
    """Return (conv_id, created) for a direct conversation between two users.
    If job_id is provided, look for a job-linked conversation involving both."""
    conv_type = 'job' if job_id else 'direct'
    if job_id:
        row = conn.execute("""
            SELECT c.id FROM conversations c
            JOIN conversation_participants pa ON pa.conversation_id = c.id AND pa.user_id = ?
            JOIN conversation_participants pb ON pb.conversation_id = c.id AND pb.user_id = ?
            WHERE c.type = 'job' AND c.job_id = ?
            LIMIT 1
        """, (uid_a, uid_b, job_id)).fetchone()
    else:
        row = conn.execute("""
            SELECT c.id FROM conversations c
            JOIN conversation_participants pa ON pa.conversation_id = c.id AND pa.user_id = ?
            JOIN conversation_participants pb ON pb.conversation_id = c.id AND pb.user_id = ?
            WHERE c.type = 'direct'
            AND (SELECT COUNT(*) FROM conversation_participants cp WHERE cp.conversation_id = c.id) = 2
            LIMIT 1
        """, (uid_a, uid_b)).fetchone()
    if row:
        return row["id"], False
    ts = now_ts()
    cur = conn.execute(
        "INSERT INTO conversations (type, job_id, created_at, updated_at) VALUES (?,?,?,?)",
        (conv_type, job_id, ts, ts)
    )
    conv_id = cur.lastrowid
    conn.execute("INSERT INTO conversation_participants (conversation_id, user_id, joined_at) VALUES (?,?,?)", (conv_id, uid_a, ts))
    conn.execute("INSERT INTO conversation_participants (conversation_id, user_id, joined_at) VALUES (?,?,?)", (conv_id, uid_b, ts))
    conn.commit()
    return conv_id, True


def _get_unread_count(conn, user_id):
    row = conn.execute("""
        SELECT COUNT(*) c
        FROM messages m
        JOIN conversation_participants cp ON cp.conversation_id = m.conversation_id AND cp.user_id = ?
        WHERE m.sender_id != ?
          AND m.is_deleted = 0
          AND NOT EXISTS (
              SELECT 1 FROM message_reads mr WHERE mr.message_id = m.id AND mr.user_id = ?
          )
    """, (user_id, user_id, user_id)).fetchone()
    return row["c"] if row else 0


def _get_conv_list(conn, user_id):
    """Return list of conversations for a user with last message + unread count."""
    rows = conn.execute("""
        SELECT c.id, c.type, c.job_id, c.subject, c.updated_at,
               j.display_ref AS job_ref
        FROM conversations c
        JOIN conversation_participants cp ON cp.conversation_id = c.id AND cp.user_id = ?
        LEFT JOIN jobs j ON j.id = c.job_id
        ORDER BY c.updated_at DESC
    """, (user_id,)).fetchall()

    result = []
    for conv in rows:
        cid = conv["id"]
        # Last message
        last_msg = conn.execute("""
            SELECT m.body, m.created_at, u.full_name sender_name
            FROM messages m JOIN users u ON u.id = m.sender_id
            WHERE m.conversation_id = ? AND m.is_deleted = 0
            ORDER BY m.created_at DESC LIMIT 1
        """, (cid,)).fetchone()
        # Unread count
        unread = conn.execute("""
            SELECT COUNT(*) c FROM messages m
            WHERE m.conversation_id = ? AND m.sender_id != ? AND m.is_deleted = 0
              AND NOT EXISTS (SELECT 1 FROM message_reads mr WHERE mr.message_id=m.id AND mr.user_id=?)
        """, (cid, user_id, user_id)).fetchone()["c"]
        # Other participants
        others = conn.execute("""
            SELECT u.id, u.full_name FROM conversation_participants cp
            JOIN users u ON u.id = cp.user_id
            WHERE cp.conversation_id = ? AND cp.user_id != ?
        """, (cid, user_id)).fetchall()
        result.append({
            "id": cid, "type": conv["type"], "job_id": conv["job_id"],
            "job_ref": conv["job_ref"], "subject": conv["subject"],
            "updated_at": conv["updated_at"],
            "last_msg": dict(last_msg) if last_msg else None,
            "unread": unread,
            "others": [dict(o) for o in others],
        })
    return result


def _mark_conv_read(conn, conv_id, user_id):
    ts = now_ts()
    unread_msgs = conn.execute("""
        SELECT m.id FROM messages m
        WHERE m.conversation_id = ? AND m.sender_id != ? AND m.is_deleted = 0
          AND NOT EXISTS (SELECT 1 FROM message_reads mr WHERE mr.message_id=m.id AND mr.user_id=?)
    """, (conv_id, user_id, user_id)).fetchall()
    for msg in unread_msgs:
        try:
            conn.execute("INSERT OR IGNORE INTO message_reads (message_id, user_id, read_at) VALUES (?,?,?)",
                         (msg["id"], user_id, ts))
        except Exception:
            pass
    conn.commit()


def _post_message(conn, conv_id, sender_id, body):
    ts = now_ts()
    cur = conn.execute(
        "INSERT INTO messages (conversation_id, sender_id, body, created_at) VALUES (?,?,?,?)",
        (conv_id, sender_id, body.strip(), ts)
    )
    msg_id = cur.lastrowid
    conn.execute("UPDATE conversations SET updated_at=? WHERE id=?", (ts, conv_id))
    # Mark as read for sender
    conn.execute("INSERT OR IGNORE INTO message_reads (message_id, user_id, read_at) VALUES (?,?,?)",
                 (msg_id, sender_id, ts))
    conn.commit()
    return msg_id


# ── Unread count API (shared desktop + mobile) ────────────────────────────────

@app.get("/api/messages/unread-count")
@login_required
def api_messages_unread_count():
    uid = session.get("user_id")
    conn = db()
    _ensure_msg_tables(conn)
    count = _get_unread_count(conn, uid)
    conn.close()
    return jsonify({"count": count})


# ── Desktop messaging routes ──────────────────────────────────────────────────

@app.get("/messages")
@login_required
def messages_list():
    uid = session.get("user_id")
    conn = db()
    _ensure_msg_tables(conn)
    convs = _get_conv_list(conn, uid)
    users = conn.execute(
        "SELECT id, full_name FROM users WHERE active=1 AND id != ? ORDER BY full_name",
        (uid,)
    ).fetchall()
    conn.close()
    return render_template("messages.html", convs=convs, users=users)


@app.get("/messages/<int:conv_id>")
@login_required
def message_thread(conv_id):
    uid = session.get("user_id")
    conn = db()
    _ensure_msg_tables(conn)
    # Verify participant
    part = conn.execute(
        "SELECT 1 FROM conversation_participants WHERE conversation_id=? AND user_id=?",
        (conv_id, uid)
    ).fetchone()
    if not part:
        conn.close()
        flash("Conversation not found.", "warning")
        return redirect(url_for("messages_list"))
    # Mark read
    _mark_conv_read(conn, conv_id, uid)
    # Conversation meta
    conv = conn.execute("""
        SELECT c.*, j.display_ref job_ref, j.id as jid
        FROM conversations c LEFT JOIN jobs j ON j.id = c.job_id
        WHERE c.id=?
    """, (conv_id,)).fetchone()
    # Messages
    msgs = conn.execute("""
        SELECT m.*, u.full_name sender_name
        FROM messages m JOIN users u ON u.id = m.sender_id
        WHERE m.conversation_id=? AND m.is_deleted=0
        ORDER BY m.created_at ASC
    """, (conv_id,)).fetchall()
    # Participants
    participants = conn.execute("""
        SELECT u.id, u.full_name FROM conversation_participants cp
        JOIN users u ON u.id = cp.user_id
        WHERE cp.conversation_id=?
    """, (conv_id,)).fetchall()
    # Full conv list for sidebar
    convs = _get_conv_list(conn, uid)
    users = conn.execute(
        "SELECT id, full_name FROM users WHERE active=1 AND id != ? ORDER BY full_name",
        (uid,)
    ).fetchall()
    conn.close()
    return render_template("message_thread.html",
                           conv=conv, msgs=msgs, participants=participants,
                           convs=convs, users=users, conv_id=conv_id)


@app.post("/messages/new")
@login_required
def messages_new():
    uid = session.get("user_id")
    recipient_ids = request.form.getlist("recipient_ids", type=int)
    if not recipient_ids:
        rid = request.form.get("recipient_id", type=int)
        if rid:
            recipient_ids = [rid]
    body = (request.form.get("body") or "").strip()
    job_id = request.form.get("job_id", type=int) or None

    if not recipient_ids or not body:
        flash("Recipient and message are required.", "warning")
        return redirect(url_for("messages_list"))

    conn = db()
    _ensure_msg_tables(conn)

    last_conv_id = None
    for rid in recipient_ids:
        if rid == uid:
            continue
        conv_id, _ = _get_or_create_direct_conv(conn, uid, rid, job_id=job_id)
        _post_message(conn, conv_id, uid, body)
        last_conv_id = conv_id

    conn.close()

    if last_conv_id and len(recipient_ids) == 1:
        return redirect(url_for("message_thread", conv_id=last_conv_id))
    flash(f"Message sent to {len(recipient_ids)} recipient{'s' if len(recipient_ids) != 1 else ''}.", "success")
    return redirect(url_for("messages_list"))


@app.post("/messages/<int:conv_id>/reply")
@login_required
def messages_reply(conv_id):
    uid = session.get("user_id")
    body = (request.form.get("body") or "").strip()
    if not body:
        return redirect(url_for("message_thread", conv_id=conv_id))
    conn = db()
    _ensure_msg_tables(conn)
    part = conn.execute(
        "SELECT 1 FROM conversation_participants WHERE conversation_id=? AND user_id=?",
        (conv_id, uid)
    ).fetchone()
    if part:
        _post_message(conn, conv_id, uid, body)
    conn.close()
    return redirect(url_for("message_thread", conv_id=conv_id))


@app.post("/messages/<int:conv_id>/read")
@login_required
def messages_mark_read(conv_id):
    uid = session.get("user_id")
    conn = db()
    _ensure_msg_tables(conn)
    _mark_conv_read(conn, conv_id, uid)
    conn.close()
    return jsonify({"ok": True})


# ── Mobile messaging routes ────────────────────────────────────────────────────

@app.get("/m/messages")
@mobile_login_required
def m_messages_list():
    uid = session.get("user_id")
    conn = db()
    _ensure_msg_tables(conn)
    convs = _get_conv_list(conn, uid)
    users = conn.execute(
        "SELECT id, full_name FROM users WHERE active=1 AND id != ? ORDER BY full_name",
        (uid,)
    ).fetchall()
    conn.close()
    return render_template("mobile/messages.html", convs=convs, users=users)


@app.get("/m/messages/<int:conv_id>")
@mobile_login_required
def m_message_thread(conv_id):
    uid = session.get("user_id")
    conn = db()
    _ensure_msg_tables(conn)
    part = conn.execute(
        "SELECT 1 FROM conversation_participants WHERE conversation_id=? AND user_id=?",
        (conv_id, uid)
    ).fetchone()
    if not part:
        conn.close()
        return redirect(url_for("m_messages_list"))
    _mark_conv_read(conn, conv_id, uid)
    conv = conn.execute("""
        SELECT c.*, j.display_ref job_ref
        FROM conversations c LEFT JOIN jobs j ON j.id = c.job_id
        WHERE c.id=?
    """, (conv_id,)).fetchone()
    msgs = conn.execute("""
        SELECT m.*, u.full_name sender_name
        FROM messages m JOIN users u ON u.id = m.sender_id
        WHERE m.conversation_id=? AND m.is_deleted=0
        ORDER BY m.created_at ASC
    """, (conv_id,)).fetchall()
    participants = conn.execute("""
        SELECT u.id, u.full_name FROM conversation_participants cp
        JOIN users u ON u.id = cp.user_id WHERE cp.conversation_id=?
    """, (conv_id,)).fetchall()
    conn.close()
    return render_template("mobile/message_thread.html",
                           conv=conv, msgs=msgs, participants=participants,
                           conv_id=conv_id)


@app.post("/m/messages/new")
@mobile_login_required
def m_messages_new():
    uid = session.get("user_id")
    recipient_ids = request.form.getlist("recipient_ids", type=int)
    if not recipient_ids:
        rid = request.form.get("recipient_id", type=int)
        if rid:
            recipient_ids = [rid]
    body = (request.form.get("body") or "").strip()
    job_id = request.form.get("job_id", type=int) or None

    if not recipient_ids or not body:
        return redirect(url_for("m_messages_list"))

    conn = db()
    _ensure_msg_tables(conn)
    last_conv_id = None
    for rid in recipient_ids:
        if rid == uid:
            continue
        conv_id, _ = _get_or_create_direct_conv(conn, uid, rid, job_id=job_id)
        _post_message(conn, conv_id, uid, body)
        last_conv_id = conv_id
    conn.close()
    if last_conv_id and len(recipient_ids) == 1:
        return redirect(url_for("m_message_thread", conv_id=last_conv_id))
    return redirect(url_for("m_messages_list"))


@app.post("/m/messages/<int:conv_id>/reply")
@mobile_login_required
def m_messages_reply(conv_id):
    uid = session.get("user_id")
    body = (request.form.get("body") or "").strip()
    if not body:
        return redirect(url_for("m_message_thread", conv_id=conv_id))
    conn = db()
    _ensure_msg_tables(conn)
    part = conn.execute(
        "SELECT 1 FROM conversation_participants WHERE conversation_id=? AND user_id=?",
        (conv_id, uid)
    ).fetchone()
    if part:
        _post_message(conn, conv_id, uid, body)
    conn.close()
    return redirect(url_for("m_message_thread", conv_id=conv_id))


# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000)
